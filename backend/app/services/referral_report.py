"""
Referral Report Generation Service

Generates comprehensive Word reports for referrals with:
1. Landscape Orientation (Letter)
2. Narrow Margins (0.5")
3. Grouped Data Columns for better readability
4. Professional headers and section coloring

Extracts data from AWS database with all CCU and payor information.
"""

import io
from datetime import datetime
from typing import List, Dict, Any
from sqlalchemy.orm import Session, joinedload
from app.models import Lead, CCU, Agency, MCO, LeadComment
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def get_referrals_sent(db: Session) -> List[Lead]:
    """Get all sent referrals"""
    return db.query(Lead).options(
        joinedload(Lead.agency),
        joinedload(Lead.ccu),
        joinedload(Lead.mco),
        joinedload(Lead.lead_comments)
    ).filter(
        Lead.active_client == True,
        Lead.authorization_received == False,
        Lead.last_contact_status != "Not Approved",
        Lead.deleted_at == None
    ).order_by(Lead.created_at.desc()).all()


def get_referrals_confirmed(db: Session) -> List[Lead]:
    """Get all confirmed referrals"""
    return db.query(Lead).options(
        joinedload(Lead.agency),
        joinedload(Lead.ccu),
        joinedload(Lead.mco),
        joinedload(Lead.lead_comments)
    ).filter(
        Lead.active_client == True,
        Lead.authorization_received == True,
        Lead.deleted_at == None
    ).order_by(Lead.created_at.desc()).all()


def get_referrals_rejected(db: Session) -> List[Lead]:
    """Get all rejected referrals"""
    return db.query(Lead).options(
        joinedload(Lead.agency),
        joinedload(Lead.ccu),
        joinedload(Lead.mco),
        joinedload(Lead.lead_comments)
    ).filter(
        Lead.active_client == True,
        Lead.last_contact_status == "Not Approved",
        Lead.deleted_at == None
    ).order_by(Lead.created_at.desc()).all()


def format_date(dt) -> str:
    if dt is None: return ""
    if hasattr(dt, 'strftime'):
        return dt.strftime('%m/%d/%Y')
    return str(dt)


def get_latest_comment(lead: Lead) -> str:
    if not lead.lead_comments: return ""
    sorted_comments = sorted(lead.lead_comments, key=lambda c: c.created_at, reverse=True)
    return sorted_comments[0].content if sorted_comments else ""


def generate_referral_report_docx(db: Session) -> bytes:
    """
    Generate professional Word report in Landscape format.
    Group columns for readability.
    """
    doc = Document()
    
    # Set to Landscape
    section = doc.sections[0]
    new_width, new_height = section.page_height, section.page_width
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = new_width
    section.page_height = new_height
    
    # Narrow Margins (0.5")
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    
    # Title
    title = doc.add_heading('AWS Referral Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # Column Headers (Grouped for readability)
    headers = [
        'ID', 
        'Client Name', 
        'Contact / SSN / Medicaid',
        'DOB / Age',
        'Address',
        'Emergency Contact',
        'Referral Info',
        'Status / SOC / Priority',
        'Payor Details',
        'CCU Details',
        'Metadata / Latest Comment'
    ]
    
    def add_table_section(doc, section_title, leads, color):
        heading = doc.add_heading(section_title, level=1)
        heading_run = heading.runs[0]
        heading_run.font.color.rgb = RGBColor(*color)
        
        if not leads:
            doc.add_paragraph(f"No {section_title.lower()} records found.")
            return

        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'
        
        # Format Headers
        hdr_cells = table.rows[0].cells
        for i, h in enumerate(headers):
            hdr_cells[i].text = h
            for paragraph in hdr_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(8)
                    run.font.color.rgb = RGBColor(255, 255, 255)
            # Set background color for header
            tc = hdr_cells[i]._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:fill'), '%02x%02x%02x' % color)
            tcPr.append(shd)

        for lead in leads:
            row_cells = table.add_row().cells
            
            # ID
            row_cells[0].text = str(lead.id)
            
            # Client Name
            row_cells[1].text = f"{lead.first_name} {lead.last_name}"
            
            # Contact / SSN / Medicaid
            contact_parts = []
            if lead.phone: contact_parts.append(f"P: {lead.phone}")
            if lead.email: contact_parts.append(f"E: {lead.email}")
            if lead.ssn: contact_parts.append(f"SSN: {lead.ssn}")
            if lead.medicaid_no: contact_parts.append(f"Medicaid: {lead.medicaid_no}")
            row_cells[2].text = "\n".join(contact_parts)
            
            # DOB / Age
            row_cells[3].text = f"DOB: {format_date(lead.dob)}\nAge: {lead.age or ''}"
            
            # Address
            addr_parts = [lead.street or ""]
            city_state = []
            if lead.city: city_state.append(lead.city)
            if lead.state: city_state.append(lead.state)
            if city_state: addr_parts.append(", ".join(city_state))
            if lead.zip_code: addr_parts.append(lead.zip_code)
            row_cells[4].text = "\n".join(addr_parts)
            
            # Emergency Contact
            e_parts = []
            if lead.e_contact_name: e_parts.append(lead.e_contact_name)
            if lead.e_contact_relation: e_parts.append(f"({lead.e_contact_relation})")
            if lead.e_contact_phone: e_parts.append(lead.e_contact_phone)
            row_cells[5].text = "\n".join(e_parts)
            
            # Referral Info
            ref_parts = [f"Source: {lead.source or ''}", f"Type: {lead.referral_type or 'Regular'}", f"Staff: {lead.staff_name or ''}"]
            row_cells[6].text = "\n".join(ref_parts)
            
            # Status / SOC / Priority
            stat_parts = [f"Status: {lead.last_contact_status or ''}", f"Auth: {'Yes' if lead.authorization_received else 'No'}", f"SOC: {format_date(lead.soc_date)}", f"Priority: {lead.priority or ''}"]
            row_cells[7].text = "\n".join(stat_parts)
            
            # Payor Details
            pay_parts = []
            if lead.agency:
                pay_parts.append(lead.agency.name)
                pay_parts.append(lead.agency.phone or "")
                pay_parts.append(lead.agency.email or "")
            row_cells[8].text = "\n".join(pay_parts)
            
            # CCU Details
            ccu_parts = []
            if lead.ccu:
                ccu_parts.append(f"Name: {lead.ccu.name}")
                ccu_parts.append(f"Coord: {lead.ccu.care_coordinator_name or ''}")
                ccu_parts.append(f"Phone: {lead.ccu.phone or ''}")
                ccu_parts.append(f"Email: {lead.ccu.email or ''}")
            row_cells[9].text = "\n".join(ccu_parts)
            
            # Metadata / Latest Comment
            meta_parts = [f"Created: {format_date(lead.created_at)} by {lead.created_by or ''}", f"Comment: {get_latest_comment(lead)}"]
            row_cells[10].text = "\n".join(meta_parts)
            
            # Format row font size
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(7)
        
        doc.add_page_break()

    # Data
    sent = get_referrals_sent(db)
    confirmed = get_referrals_confirmed(db)
    rejected = get_referrals_rejected(db)
    
    # Sections
    add_table_section(doc, 'REFERRALS SENT', sent, (0, 102, 204))     # Blue
    add_table_section(doc, 'REFERRALS CONFIRMED', confirmed, (0, 153, 51)) # Green
    add_table_section(doc, 'REFERRALS REJECTED', rejected, (204, 0, 0))   # Red
    
    # Summary Sheet at the end
    doc.add_heading('Summary', level=1)
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Light Grid Accent 1'
    
    data = [
        ('Referrals Sent', len(sent)),
        ('Referrals Confirmed', len(confirmed)),
        ('Referrals Rejected', len(rejected)),
        ('Total Records', len(sent) + len(confirmed) + len(rejected))
    ]
    
    for i, (label, value) in enumerate(data):
        cells = table.rows[i].cells
        cells[0].text = str(label)
        cells[1].text = str(value)
        cells[0].paragraphs[0].runs[0].font.bold = True

    # Save
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def get_report_statistics(db: Session) -> Dict[str, int]:
    sent_count = db.query(Lead).filter(Lead.active_client == True, Lead.authorization_received == False, Lead.last_contact_status != "Not Approved", Lead.deleted_at == None).count()
    confirmed_count = db.query(Lead).filter(Lead.active_client == True, Lead.authorization_received == True, Lead.deleted_at == None).count()
    rejected_count = db.query(Lead).filter(Lead.active_client == True, Lead.last_contact_status == "Not Approved", Lead.deleted_at == None).count()
    return {'sent': sent_count, 'confirmed': confirmed_count, 'rejected': rejected_count, 'total': sent_count + confirmed_count + rejected_count}
