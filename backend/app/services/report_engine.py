"""
Generic Report Engine
=====================
Data-driven Excel & Word report generator.

Any "Top by" report (CCU, Source, Status, Staff, MCO …) builds a `report_config`
dict and passes it here. The engine handles all formatting.

report_config shape:
{
    "title": str,                     # e.g. "Top 5 CCU Report"
    "generated_at": str,              # ISO timestamp string
    "sections": [
        {
            "heading": str,           # e.g. "#1 — Community Care North (12 referrals)"
            "sheet_name": str,        # Excel sheet name (≤31 chars)
            "summary_row": dict,      # e.g. {"CCU Name": "North", "Phone": "555-..."}
            "detail_columns": list,   # ordered list of column names for detail rows
            "detail_rows": list[dict] # each dict key matches a column in detail_columns
        },
        ...
    ]
}
"""

import io
from datetime import datetime
from typing import Any, Dict, List

# ── Excel ─────────────────────────────────────────────────────────────────────
try:
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False

# ── Word ──────────────────────────────────────────────────────────────────────
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    WORD_AVAILABLE = True
except ImportError:
    WORD_AVAILABLE = False


# ── Colour palette ────────────────────────────────────────────────────────────
RANK_COLOURS_HEX = [
    "1A6B8A",  # #1 deep teal
    "2C8C99",  # #2 teal
    "3CA5AA",  # #3 medium teal
    "59B976",  # #4 green
    "86C96A",  # #5 light green
    "B5D96A",  # #6
    "D9C46A",  # #7
    "D9906A",  # #8
    "D96A6A",  # #9
    "A855F7",  # #10
]

RANK_COLOURS_RGB = [
    (26, 107, 138),
    (44, 140, 153),
    (60, 165, 170),
    (89, 185, 118),
    (134, 201, 106),
    (181, 217, 106),
    (217, 196, 106),
    (217, 144, 106),
    (217, 106, 106),
    (168, 85, 247),
]

SUMMARY_BG_HEX = "D9EEF2"    # light teal for summary rows
HEADER_TEXT_HEX = "FFFFFF"    # white text on colour headers


def _get_rank_colour_hex(rank: int) -> str:
    idx = min(rank - 1, len(RANK_COLOURS_HEX) - 1)
    return RANK_COLOURS_HEX[idx]


def _get_rank_colour_rgb(rank: int):
    idx = min(rank - 1, len(RANK_COLOURS_RGB) - 1)
    return RANK_COLOURS_RGB[idx]


# ─────────────────────────────────────────────────────────────────────────────
# Excel generation
# ─────────────────────────────────────────────────────────────────────────────

def generate_excel(report_config: Dict[str, Any]) -> bytes:
    """
    Generate an Excel workbook from report_config.
    Returns raw bytes suitable for HTTP download or st.download_button.
    """
    if not EXCEL_AVAILABLE:
        raise RuntimeError("openpyxl is not installed. Run: pip install openpyxl")

    wb = openpyxl.Workbook()
    # Remove default empty sheet
    wb.remove(wb.active)

    # ── Cover sheet ──────────────────────────────────────────────────────────
    cover = wb.create_sheet("Summary", 0)
    cover["A1"] = report_config.get("title", "Report")
    cover["A1"].font = Font(bold=True, size=16)
    cover["A2"] = f"Generated: {report_config.get('generated_at', datetime.now().strftime('%Y-%m-%d %H:%M:%S'))}"
    cover["A2"].font = Font(size=11, italic=True)

    cover["A4"] = "Rank"
    cover["B4"] = "Heading"
    cover["C4"] = "Records"
    for col in ["A4", "B4", "C4"]:
        cover[col].font = Font(bold=True, color=HEADER_TEXT_HEX)
        cover[col].fill = PatternFill("solid", fgColor="1A6B8A")
        cover[col].alignment = Alignment(horizontal="center")

    for i, section in enumerate(report_config.get("sections", []), start=1):
        row = 4 + i
        cover.cell(row=row, column=1, value=i)
        cover.cell(row=row, column=2, value=section.get("heading", ""))
        cover.cell(row=row, column=3, value=len(section.get("detail_rows", [])))
        # Alternating row fill
        if i % 2 == 0:
            for c in range(1, 4):
                cover.cell(row=row, column=c).fill = PatternFill("solid", fgColor="EAF4F6")

    cover.column_dimensions["A"].width = 8
    cover.column_dimensions["B"].width = 48
    cover.column_dimensions["C"].width = 12

    # ── One sheet per section ─────────────────────────────────────────────────
    for rank, section in enumerate(report_config.get("sections", []), start=1):
        sheet_name = _safe_sheet_name(section.get("sheet_name", f"#{rank}"), rank)
        ws = wb.create_sheet(sheet_name)

        colour_hex = _get_rank_colour_hex(rank)

        # Section heading row
        heading_text = section.get("heading", sheet_name)
        ws["A1"] = heading_text
        ws["A1"].font = Font(bold=True, size=13, color=HEADER_TEXT_HEX)
        ws["A1"].fill = PatternFill("solid", fgColor=colour_hex)
        ws["A1"].alignment = Alignment(horizontal="left", vertical="center")

        # Merge heading across all columns (we'll know how wide after writing detail)
        detail_columns = section.get("detail_columns", [])
        summary_row = section.get("summary_row", {})
        detail_rows = section.get("detail_rows", [])

        # All columns = summary keys + detail columns (union, ordered)
        all_summary_keys = list(summary_row.keys())

        # ── Summary row (row 3) ───────────────────────────────────────────────
        ws.append([])  # row 2 blank
        ws.append(all_summary_keys)  # header row for summary (row 3)
        for col_idx, key in enumerate(all_summary_keys, start=1):
            cell = ws.cell(row=3, column=col_idx)
            cell.font = Font(bold=True, color=HEADER_TEXT_HEX)
            cell.fill = PatternFill("solid", fgColor=colour_hex)
            cell.alignment = Alignment(horizontal="center", wrap_text=True)

        ws.append([summary_row.get(k, "") for k in all_summary_keys])  # row 4
        for col_idx in range(1, len(all_summary_keys) + 1):
            cell = ws.cell(row=4, column=col_idx)
            cell.fill = PatternFill("solid", fgColor=SUMMARY_BG_HEX)
            cell.font = Font(bold=True, size=10)
            cell.alignment = Alignment(horizontal="left", wrap_text=True)

        ws.append([])  # row 5 blank

        # ── Detail rows ───────────────────────────────────────────────────────
        if detail_columns:
            ws.append(detail_columns)  # row 6 — column headers
            header_row_idx = ws.max_row
            for col_idx, col_name in enumerate(detail_columns, start=1):
                cell = ws.cell(row=header_row_idx, column=col_idx)
                cell.font = Font(bold=True, color=HEADER_TEXT_HEX, size=9)
                cell.fill = PatternFill("solid", fgColor="00506B")
                cell.alignment = Alignment(horizontal="center", wrap_text=True)

            for row_num, row_data in enumerate(detail_rows):
                row_values = [str(row_data.get(col, "") or "") for col in detail_columns]
                ws.append(row_values)
                data_row_idx = ws.max_row
                row_fill = "F0F9FA" if row_num % 2 == 0 else "FFFFFF"
                for col_idx in range(1, len(detail_columns) + 1):
                    cell = ws.cell(row=data_row_idx, column=col_idx)
                    cell.fill = PatternFill("solid", fgColor=row_fill)
                    cell.font = Font(size=9)
                    cell.alignment = Alignment(wrap_text=True, vertical="top")

        # Merge heading row A1 across all used columns
        max_col = max(len(all_summary_keys), len(detail_columns), 1)
        if max_col > 1:
            ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=max_col)

        # Auto-width columns
        for col_idx in range(1, max_col + 1):
            col_letter = get_column_letter(col_idx)
            max_len = 12
            for row in ws.iter_rows(min_col=col_idx, max_col=col_idx):
                for cell in row:
                    try:
                        val_len = len(str(cell.value or ""))
                        if val_len > max_len:
                            max_len = val_len
                    except Exception:
                        pass
            ws.column_dimensions[col_letter].width = min(max_len + 2, 45)

        ws.row_dimensions[1].height = 22
        ws.freeze_panes = "A7"

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf.getvalue()


def _safe_sheet_name(name: str, rank: int) -> str:
    """Excel sheet names must be ≤31 chars and have no special chars."""
    safe = name.replace("/", "-").replace("\\", "-").replace("*", "").replace("?", "").replace("[", "").replace("]", "").replace(":", "")
    truncated = safe[:28]
    if not truncated.strip():
        truncated = f"Section {rank}"
    return truncated


# ─────────────────────────────────────────────────────────────────────────────
# Word generation
# ─────────────────────────────────────────────────────────────────────────────

def generate_word(report_config: Dict[str, Any]) -> bytes:
    """
    Generate a Word (.docx) document from report_config.
    Landscape orientation, one headed table per section.
    """
    if not WORD_AVAILABLE:
        raise RuntimeError("python-docx is not installed. Run: pip install python-docx")

    doc = Document()

    # Landscape
    section = doc.sections[0]
    new_width, new_height = section.page_height, section.page_width
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = new_width
    section.page_height = new_height
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    # Title
    title_para = doc.add_heading(report_config.get("title", "Report"), 0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Generated: {report_config.get('generated_at', datetime.now().strftime('%Y-%m-%d %H:%M:%S'))}").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    for rank, section_cfg in enumerate(report_config.get("sections", []), start=1):
        colour_rgb = _get_rank_colour_rgb(rank)

        # Section heading
        heading = doc.add_heading(section_cfg.get("heading", f"Section {rank}"), level=1)
        if heading.runs:
            heading.runs[0].font.color.rgb = RGBColor(*colour_rgb)

        detail_columns = section_cfg.get("detail_columns", [])
        detail_rows = section_cfg.get("detail_rows", [])

        # Summary block
        summary_row = section_cfg.get("summary_row", {})
        if summary_row:
            summary_para = doc.add_paragraph()
            for key, val in summary_row.items():
                run = summary_para.add_run(f"{key}: {val}   ")
                run.bold = True
                run.font.size = Pt(9)

        if not detail_rows:
            doc.add_paragraph("No records found.")
            doc.add_page_break()
            continue

        # Detail table
        if detail_columns:
            table = doc.add_table(rows=1, cols=len(detail_columns))
            table.style = "Table Grid"

            hdr_cells = table.rows[0].cells
            for i, col_name in enumerate(detail_columns):
                hdr_cells[i].text = col_name
                for para in hdr_cells[i].paragraphs:
                    for run in para.runs:
                        run.font.bold = True
                        run.font.size = Pt(8)
                        run.font.color.rgb = RGBColor(255, 255, 255)
                tc = hdr_cells[i]._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:fill"), "%02x%02x%02x" % colour_rgb)
                tcPr.append(shd)

            for row_data in detail_rows:
                row_cells = table.add_row().cells
                for i, col_name in enumerate(detail_columns):
                    row_cells[i].text = str(row_data.get(col_name, "") or "")
                    for para in row_cells[i].paragraphs:
                        for run in para.runs:
                            run.font.size = Pt(7)

        doc.add_page_break()

    # Summary table at end
    doc.add_heading("Summary", level=1)
    sections_data = report_config.get("sections", [])
    if sections_data:
        sum_table = doc.add_table(rows=len(sections_data) + 1, cols=3)
        sum_table.style = "Light Grid Accent 1"
        headers = ["Rank", "Name", "Records"]
        for i, h in enumerate(headers):
            sum_table.rows[0].cells[i].text = h
            sum_table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
        for rank, section_cfg in enumerate(sections_data, start=1):
            row = sum_table.rows[rank]
            row.cells[0].text = str(rank)
            row.cells[1].text = section_cfg.get("heading", "")
            row.cells[2].text = str(len(section_cfg.get("detail_rows", [])))

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()
