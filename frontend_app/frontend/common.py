"""
Common utilities, CSS styles, and shared functions for Lead Manager.
"""
import sys
from pathlib import Path

# Add backend to Python path
backend_path = Path(__file__).parent.parent.parent / "backend"
sys.path.insert(0, str(backend_path))

import streamlit as st
import pandas as pd
from app.db import SessionLocal, engine
from app.utils.activity_logger import utc_to_local
from app.crud import crud_session_tokens # crud_leads moved to local import
from streamlit.components.v1 import html
import os
import io

def get_logo_path(filename="icon1.png"):
    """Find the specified logo file in multiple possible locations"""
    # Base path is the project root (lead-manager/)
    base_path = Path(__file__).parent.parent.parent
    possible_paths = [
        str(base_path / filename),       # Project Root (Absolute)
        filename,                         # Relative to CWD
        f"frontend_app/{filename}",       # Docker relative
        f"/app/{filename}"                # Docker Absolute
    ]
    for path in possible_paths:
        if os.path.exists(path):
            return path
    return filename # Fallback


# Token-based session management (secure, database-backed)
def get_session_token():
    """Get the session token from URL query params"""
    query_params = st.query_params
    return query_params.get("token", None)

def set_session_token(token: str):
    """Store the session token in URL query params for persistence"""
    st.query_params["token"] = token

def clear_session_token():
    """Remove the session token from URL query params"""
    if "token" in st.query_params:
        del st.query_params["token"]


# Caregiver Type Constants
CAREGIVER_TYPES = ["None", "FHCA", "PHCA", "HCA"]

# Global CSS styles (SafeLife UI theme)
GLOBAL_CSS = """
<style>
    /* Import Montserrat from Google Fonts */
    @import url('https://fonts.googleapis.com/css2?family=Montserrat:wght@300;400;600;700;800&display=swap');

    html, body, .stApp  {
        font-family: 'Montserrat', -apple-system, BlinkMacSystemFont, 'Segoe UI',
                     Arial, 'Source Sans Pro', system-ui, sans-serif;
        background-color: #FFFFFF !important;
        color: #111827 !important;
    }

    @keyframes fadeIn {
        from { opacity: 0; }
        to { opacity: 1; }
    }

    /* Hide Streamlit header (Deploy button, Rerun status) but KEEP the sidebar toggle */
    header { 
        visibility: hidden !important; 
        background: transparent !important;
    }
    
    [data-testid="stHeader"] { 
        visibility: hidden !important; 
        background: transparent !important;
    }

    footer { visibility: hidden !important; }
    .stDeployButton { display: none !important; }
    #MainMenu { visibility: hidden !important; }

    /* Force the sidebar collapse/expand buttons to be visible */
    [data-testid="stHeader"] button,
    [data-testid="stSidebar"] button[aria-label="Collapse sidebar"],
    button[data-testid="stExpandSidebarButton"] {
        visibility: visible !important;
    }

    /* Suppress transient Streamlit errors and warnings highly aggressively */
    /* We use a delay of 0.8s which is enough to hide all typical developmental "re-rendering" glitches */
    /* EXCEPT for our custom toasts which we want to see immediately */
    [data-testid="stNotification"], 
    .stException, 
    [data-testid="stFormWarning"],
    [data-testid="stStatusWidget"] {
        /* Immediate visibility for critical alerts */
        opacity: 1 !important;
        pointer-events: auto !important;
    }

    /* PREMIUM INDICATOR STYLING */
    
    /* Make Alert boxes much more professional and modern */
    [data-testid="stNotification"] .stAlert {
        font-size: 1.1rem !important;
        font-weight: 500 !important;
        padding: 1rem 1.25rem !important;
        border-radius: 12px !important;
        border: 1px solid rgba(0,0,0,0.05) !important;
        box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.05), 0 2px 4px -1px rgba(0, 0, 0, 0.03) !important;
        background-color: #f0fdf4 !important; /* Soft emerald green tint for success alert */
    }

    [data-testid="stNotification"] .stAlert:has(div[data-testid="stIcon"] span:contains("✅")) {
        background-color: #f0fdf4 !important;
        border-left: 5px solid #10b981 !important;
    }

    /* Make Toasts (Pop-ups) prominent and professional with premium emerald theme */
    [data-testid="stToast"] {
        min-width: 520px !important;
        max-width: 700px !important;
        padding: 1.25rem 2rem !important;
        border-radius: 16px !important;
        font-size: 1.25rem !important;
        font-weight: 700 !important; /* Force bold */
        box-shadow: 0 25px 50px -12px rgba(0, 0, 0, 0.25) !important;
        border-left: 12px solid !important;
        background: white !important;
        color: #111827 !important;
        z-index: 2147483647 !important; /* Max Z-Index for toasts too */
        animation: toastSlideIn 0.3s cubic-bezier(0.175, 0.885, 0.32, 1.275) forwards !important;
    }

    [data-testid="stToast"] p, [data-testid="stToast"] span, [data-testid="stToast"] div {
        font-weight: 700 !important; /* Thick bold text */
    }

    @keyframes toastSlideIn {
        from { transform: translateX(100%); opacity: 0; }
        to { transform: translateX(0); opacity: 1; }
    }

    /* Toast color coding based on icon */
    [data-testid="stToast"]:has(div[data-testid="stIcon"] span:contains("✅")) {
        border-left-color: #10b981 !important; /* Success emerald Green */
    }
    [data-testid="stToast"]:has(div[data-testid="stIcon"] span:contains("❌")) {
        border-left-color: #ef4444 !important; /* Error Red */
    }
    [data-testid="stToast"]:has(div[data-testid="stIcon"] span:contains("⏳")) {
        border-left-color: #f59e0b !important; /* Warning Amber */
    }

    /* Style the confirmation "Double Check" boxes if they are inside buttons/dialogs */
    .stButton > button {
        transition: all 0.2s cubic-bezier(0.4, 0, 0.2, 1) !important;
        color: #FFFFFF !important;
        font-weight: 700 !important;
        white-space: nowrap !important;
        width: 100% !important;
        min-height: 48px !important;
        padding: 0.5rem 1rem !important;
        display: flex !important;
        align-items: center !important;
        justify-content: center !important;
    }
    
    /* Force primary buttons to be Deep Blue (Chicago/SafeLife branding) */
    .stButton > button[kind="primary"] {
        background-color: #00506b !important;
        border: none !important;
    }
    
    /* Ultra-strong override for all button text containers */
    .stButton button p, .stButton button span, .stButton button label {
        color: #FFFFFF !important;
        font-weight: 700 !important;
        text-overflow: ellipsis;
        overflow: hidden;
    }
    
    .stButton > button:active {
        transform: scale(0.98) !important;
    }

    /* --- PROFESSIONAL MODAL SYSTEM (NATIVE ST.DIALOG) --- */
    
    /* Target the Streamlit Dialog container to match our theme */
    div[data-testid="stDialog"] div[role="dialog"] {
        border-radius: 14px !important;
        overflow: hidden !important;
        padding: 0 !important;
        border: none !important;
        box-shadow: 0 25px 80px rgba(0,0,0,0.55) !important;
    }

    /* Fix missing close icon in modals (X logo) */
    div[data-testid="stDialog"] button[aria-label="Close"],
    .stDialog button[aria-label="Close"] {
        background: transparent !important;
        color: #3CA5AA !important;
        border: none !important;
        box-shadow: none !important;
        padding: 0 !important;
        min-height: unset !important;
        width: 32px !important;
        height: 32px !important;
        display: flex !important;
        align-items: center !important;
        justify-content: center !important;
        border-radius: 50% !important;
        z-index: 1000 !important;
        margin: 8px !important;
    }
    
    div[data-testid="stDialog"] button[aria-label="Close"] svg,
    div[data-testid="stDialog"] button[aria-label="Close"] svg path,
    .stDialog button[aria-label="Close"] svg,
    .stDialog button[aria-label="Close"] svg path {
        fill: #3CA5AA !important;
        stroke: #3CA5AA !important;
        stroke-width: 0.5px !important;
        width: 18px !important;
        height: 18px !important;
    }
    
    div[data-testid="stDialog"] button[aria-label="Close"]:hover {
        background-color: rgba(60, 165, 170, 0.15) !important;
    }

    /* Professional header style for use inside dialogs */
    .modal-dialog-header {
      background: #1f6f73;
      color: #fff;
      text-align: center;
      padding: 24px 20px;
      font-weight: 900;
      font-size: 22px;
      letter-spacing: 1px;
      margin: -1rem -1rem 1rem -1rem; /* Negative margin to fill the dialog top */
    }

    .modal-icon {
      font-size: 52px;
      margin-bottom: 10px;
      display: block;
      font-weight: bold;
    }

    .modal-body-content {
      padding: 10px 15px 25px 15px;
      text-align: center;
      font-weight: 700;
      font-size: 18px;
      color: #1f2937;
      line-height: 1.6;
    }
    
    .modal-body-content b, .modal-body-content strong {
        font-weight: 900 !important;
    }

    /* Prevent flashing of unstyled content during reloads */
    .stApp {
        animation: smoothStart 0.4s ease-out;
    }
    
    @keyframes smoothStart {
        from { opacity: 0; transform: translateY(5px); }
        to { opacity: 1; transform: translateY(0); }
    }

    /* Primary brand colors */
    :root {
        --safelife-deep-blue: #00506b;
        --safelife-green: #59B976;
        --safelife-aqua: #3CA5AA;
        --safelife-soft-gray: #F9F9F9;
        --safelife-blue-light: #B5E8F7;
        --safelife-blue-extra-light: #DFF8FF;
        --required-star-pink: #3CA5AA;
    }
    
    /* Square alerts/info boxes */
    .stAlert {
        border-radius: 0px !important;
    }

    /* Authentication container - separate from main app */
    .auth-container {
        background: #FFFFFF;
        padding: 2.5rem;
        border-radius: 1rem;
        box-shadow: 0 8px 24px rgba(0, 0, 0, 0.12);
        max-width: 500px;
        margin: 2rem auto;
    }

    /* Required field star indicator */
    .required-star {
        color: var(--required-star-pink);
        font-weight: 700;
        margin-left: 0.2rem;
    }

    /* Form labels - bold and black */
    .stTextInput label, 
    .stPasswordInput label,
    .stSelectbox label,
    .stTextArea label,
    .stNumberInput label,
    .stDateInput label,
    label {
        font-weight: 700 !important;
        color: #000000 !important;
        font-size: 0.95rem !important;
    }

    /* Main page headers */
    .main-header {
        font-size: 2.4rem;
        font-weight: 800;
        color: #FFFFFF;
        margin-bottom: 1.5rem;
        letter-spacing: 0.04em;
        text-transform: uppercase;
        text-align: center;
        background-color: var(--safelife-aqua);
        padding: 1.5rem 2rem;
        border-radius: 0.75rem;
    }
    
    /* Fix Expanders (Leads, Payors, CCUs) - Ensure headers are NEVER black */
    [data-testid="stExpander"], .stExpander {
        border: 1px solid var(--safelife-aqua) !important;
        border-radius: 0.75rem !important;
        background-color: #FFFFFF !important;
        transition: all 0.3s ease !important;
    }
    
    /* Header/Summary of the expander */
    [data-testid="stExpander"] > details > summary,
    [data-testid="stExpander"] > div:first-child,
    .stExpander > div:first-child {
        background-color: #F9FAFB !important;
        color: #111827 !important;
        border-radius: 0.75rem 0.75rem 0 0 !important;
    }

    /* Ensure NO BLACK background on hover or selection */
    [data-testid="stExpander"] > details > summary:hover,
    [data-testid="stExpander"] > details[open] > summary {
        background-color: var(--safelife-blue-extra-light) !important;
        color: var(--safelife-deep-blue) !important;
    }

    /* Icon and Text colors inside expander headers */
    [data-testid="stExpander"] svg,
    [data-testid="stExpander"] p,
    [data-testid="stExpander"] span {
        color: #111827 !important;
        fill: #111827 !important;
    }
    
    /* Force Aqua for the border even when open */
    [data-testid="stExpander"] > details {
        border: none !important;
    }

    .main-header-signup,
    .main-header-forgot {
        color: #000000 !important;
    }
    
    .main-header-signup *,
    .main-header-forgot * {
        color: #000000 !important;
    }

    .main-header .main-subtitle {
        display: block;
        margin-top: 0.35rem;
        font-size: 1rem;
        font-weight: 600;
        text-transform: none;
        letter-spacing: 0.02em;
    }

    /* Force Lead Pipeline Analytics heading to white - multiple selectors for maximum override */
    div[style*="linear-gradient"] h1,
    div[style*="linear-gradient"] h2,
    div[style*="linear-gradient"] span,
    .stMarkdown div[style*="linear-gradient"] h1,
    .stMarkdown div[style*="linear-gradient"] h2 {
        color: #FFFFFF !important;
        -webkit-text-fill-color: #FFFFFF !important;
    }
    
    /* Override any Streamlit default h2 styling in markdown blocks */
    [data-testid="stMarkdownContainer"] h2 {
        color: inherit !important;
    }

    /* Stat cards */
    .stat-card {
        background: #FFFFFF;
        padding: 1.3rem 1.5rem;
        border-radius: 0.75rem;
        border-left: 5px solid var(--safelife-deep-blue);
        box-shadow: 0 4px 14px rgba(0, 0, 0, 0.04);
    }

    .stat-number {
        font-size: 2.1rem;
        font-weight: 700;
        color: var(--safelife-deep-blue);
    }

    .stat-label {
        font-size: 0.8rem;
        color: #6B7280;
        text-transform: uppercase;
        letter-spacing: 0.08em;
    }
    
    /* Make all subheaders bold and black */
    h2, h3, .stSubheader, .stSubheader > div, .stSubheader > div > div {
        font-weight: 700 !important;
        color: #000000 !important;
    }
    
    /* Ensure all subheader text is black */
    h2 *, h3 *, .stSubheader *, .stSubheader > div *, .stSubheader > div > div * {
        color: #000000 !important;
    }

    /* Buttons - aqua background with white text */
    button, .stButton > button, .stForm button, .stForm button[type="submit"] {
        font-family: 'Montserrat', sans-serif;
        font-weight: 700;
        border-radius: 10px;
        border: none;
        padding: 0.5rem 1.0rem !important;
        background: var(--safelife-aqua) !important;
        color: #FFFFFF !important;
        box-shadow: 0 3px 10px rgba(60, 165, 170, 0.35);
        font-size: 0.95rem !important;
        letter-spacing: 0.02em;
        white-space: nowrap !important;
    }
    
    /* Ensure all button text is white */
    button *, .stButton > button *, .stForm button *, .stForm button[type="submit"] *, button p, .stButton > button p, button span, button div {
        color: #FFFFFF !important;
    }

    /* Primary buttons (active/selected state) - deep blue background with white text */
    .stButton > button[kind="primary"],
    .stForm button[kind="primary"] {
        background: var(--safelife-deep-blue) !important;
        color: #FFFFFF !important;
        border: none !important;
    }
    
    /* Ensure primary button text stays white */
    .stButton > button[kind="primary"] *,
    .stForm button[kind="primary"] * {
        color: #FFFFFF !important;
    }

    /* Secondary buttons - aqua background with white text */
    .stButton > button[kind="secondary"],
    .stForm button[kind="secondary"] {
        background: var(--safelife-aqua) !important;
        color: #FFFFFF !important;
        border: none !important;
        box-shadow: 0 3px 10px rgba(60, 165, 170, 0.35);
    }

    /* GLOBAL RADIO BUTTON STYLING (Professional Overrides) */
    
    /* 1. Base circle styling (Solid Aqua for unselected) */
    [data-testid="stRadio"] label[data-baseweb="radio"] > div:first-child {
        border-color: var(--safelife-aqua) !important;
        background-color: var(--safelife-aqua) !important;
        box-shadow: none !important;
    }
    
    /* 2. Unselected State (Force solid Aqua) */
    [data-testid="stRadio"] label[data-baseweb="radio"]:not(:has(input:checked)) > div:first-child {
        background-color: var(--safelife-aqua) !important;
        border-color: var(--safelife-aqua) !important;
    }
    
    /* 3. Selected/Active State (Force solid Deep Blue) */
    [data-testid="stRadio"] label[data-baseweb="radio"]:has(input:checked) > div:first-child {
        background-color: var(--safelife-deep-blue) !important;
        border-color: var(--safelife-deep-blue) !important;
    }
    
    /* 4. Ensure internal dot doesn't create weird artifacts */
    [data-testid="stRadio"] label[data-baseweb="radio"] > div:first-child > div {
        background-color: transparent !important;
    }
    [data-testid="stRadio"] label[data-baseweb="radio"]:has(input:checked) > div:first-child > div {
        background-color: #FFFFFF !important; /* Small white dot in center of blue circle for active */
        width: 4px !important;
        height: 4px !important;
    }

    /* Fix for radio button option labels - make them black and bold */
    div[data-testid="stRadio"] label,
    div[data-testid="stRadio"] label p,
    div[data-testid="stRadio"] label div,
    div[data-testid="stRadio"] label span {
        color: #111827 !important;
        font-weight: 700 !important;
        -webkit-text-fill-color: #111827 !important;
        opacity: 1 !important;
    }

    /* Form fields + selectboxes - WHITE backgrounds with BLACK text (Professional Look) */
    .stTextInput input,
    .stPasswordInput input,
    .stTextArea textarea,
    .stNumberInput input,
    .stDateInput input,
    .stTimeInput input,
    .stSelectbox div[data-baseweb="select"] > div,
    .stSelectbox select,
    div[data-baseweb="base-input"],
    div[data-baseweb="input"],
    .stTextInput > div,
    .stNumberInput > div,
    .stTextArea > div,
    input[type="text"],
    input[type="password"],
    input[type="number"],
    input[type="email"],
    input[type="date"],
    textarea,
    select {
        border-radius: 0.6rem !important;
        border: 1px solid var(--safelife-aqua) !important;
        background-color: #FFFFFF !important;
        color: #111827 !important;
        box-shadow: none !important;
        font-weight: 600 !important;
        transition: all 0.2s ease !important;
    }
    
    /* PROFESSIONAL INPUT FOCUS - AQUA BORDER (Replaces default black) */
    .stTextInput input:focus,
    .stPasswordInput input:focus,
    .stTextArea textarea:focus,
    .stNumberInput input:focus,
    .stDateInput input:focus,
    .stSelectbox div[data-baseweb="select"]:focus,
    .stSelectbox div[data-baseweb="select"]:focus-within,
    [data-baseweb="input"] > div:focus-within {
        border-color: var(--safelife-aqua) !important;
        box-shadow: 0 0 0 2px rgba(60, 165, 170, 0.25) !important;
        outline: none !important;
    }
    
    /* Target the focus state of the wrapper too to prevent black outline */
    div[data-baseweb="input"]:focus-within,
    div[data-baseweb="base-input"]:focus-within,
    .stTextInput > div:focus-within,
    .stNumberInput > div:focus-within,
    .stTextArea > div:focus-within,
    [data-baseweb="select"]:focus-within {
        border-color: var(--safelife-aqua) !important;
        box-shadow: 0 0 0 2px rgba(60, 165, 170, 0.25) !important;
        outline: none !important;
    }
    
    /* Global fix for Streamlit's default focus ring (black/gray) */
    *:focus, *:focus-visible, *:active {
        outline: none !important;
        border-color: var(--safelife-aqua) !important;
    }
    
    /* Force specific streamlit internal div borders to Aqua */
    [data-baseweb="input"] {
        border: 1px solid var(--safelife-aqua) !important;
    }
    
    /* Ensure text visibility in input fields on focus */
    input:focus, textarea:focus {
        color: #111827 !important;
        -webkit-text-fill-color: #111827 !important;
        border-color: var(--safelife-aqua) !important;
    }
    
    /* Fix Number Input Buttons (+/-) icon colors to be white */
    [data-testid="stNumberInputStepUp"],
    [data-testid="stNumberInputStepDown"],
    .stNumberInput button,
    .stNumberInput div[role="button"] {
        color: #FFFFFF !important;
    }
    [data-testid="stNumberInputStepUp"] svg,
    [data-testid="stNumberInputStepDown"] svg,
    [data-testid="stNumberInputStepUp"] svg path,
    [data-testid="stNumberInputStepDown"] svg path,
    .stNumberInput button svg,
    .stNumberInput button svg path,
    .stNumberInput div[role="button"] svg,
    .stNumberInput div[role="button"] svg path {
        fill: #FFFFFF !important;
        stroke: #FFFFFF !important;
        color: #FFFFFF !important;
    }
    
    /* Ensure no double borders or inner borders */
    [data-baseweb="base-input"] input,
    [data-baseweb="input"] input {
        border: none !important;
    }
    
    /* Input placeholder text - subtle gray */
    .stTextInput input::placeholder,
    .stPasswordInput input::placeholder,
    .stTextArea textarea::placeholder,
    input::placeholder,
    textarea::placeholder {
        color: #9CA3AF !important;
    }
    
    /* Date picker calendar - white background with aqua header */
    .stDateInput div[data-baseweb="popover"],
    [data-baseweb="popover"],
    div[data-baseweb="calendar"] {
        background-color: #FFFFFF !important;
        background: #FFFFFF !important;
        border-radius: 8px !important;
        color: #111827 !important;
    }
    
    /* Target the calendar days and grid */
    [data-baseweb="calendar"] div,
    [data-baseweb="calendar"] button,
    [data-baseweb="calendar"] *:not(.aqua-header) {
        background-color: #FFFFFF !important;
        color: #111827 !important;
    }
    
    /* Ensure the header navigation stays aqua/professional */
    div[data-baseweb="calendar"] > div:first-child,
    div[data-baseweb="calendar"] > div:first-child * {
        background-color: var(--safelife-aqua) !important;
        color: white !important;
    }
    
    /* Fix for the black squares in the grid */
    [data-baseweb="calendar"] [aria-roledescription="button"] {
        background-color: #FFFFFF !important;
        border: 1px solid #F3F4F6 !important;
    }
    
    /* Selected day - deep blue background */
    [data-baseweb="calendar"] [aria-selected="true"] {
        background-color: var(--safelife-deep-blue) !important;
        color: #FFFFFF !important;
    }
    
    /* Dropdown menus - white background with high contrast */
    .stSelectbox ul,
    .stSelectbox li,
    [role="listbox"],
    [role="option"] {
        background-color: #FFFFFF !important;
        color: #111827 !important;
    }

    /* Dataframes / tables */
    .stDataFrame, .stTable {
        background-color: #FFFFFF;
        border-radius: 0.75rem;
        box-shadow: 0 2px 8px rgba(15, 23, 42, 0.05);
    }

    /* Sidebar */
    section[data-testid="stSidebar"] {
        background-color: #FFFFFF;
        border-right: 1px solid #E5E7EB;
    }
    
    /* Sidebar navigation links - black text */
    section[data-testid="stSidebar"] a,
    section[data-testid="stSidebar"] button,
    section[data-testid="stSidebar"] label,
    section[data-testid="stSidebar"] p,
    section[data-testid="stSidebar"] span,
    section[data-testid="stSidebar"] * {
        color: #000000 !important;
    }

    /* General titles/subheaders - all headings black and bold */
    h1, h2, h3, h4, h5, h6 {
        font-family: 'Montserrat', sans-serif;
        color: #000000 !important;
        font-weight: 700 !important;
    }
    
    /* Tabs Styling */
    .stTabs [data-baseweb="tab-list"] {
        gap: 8px;
        background-color: transparent;
    }
    
    .stTabs [data-baseweb="tab"] {
        height: 50px;
        background-color: #FFFFFF;
        border-radius: 10px 10px 0 0;
        gap: 2px;
        padding: 10px 20px;
        font-weight: 700 !important;
        color: #6B7280;
        border: 1px solid #E5E7EB;
        border-bottom: none;
    }
    
    .stTabs [aria-selected="true"] {
        background-color: var(--safelife-deep-blue) !important;
        color: #FFFFFF !important;
        border: none;
    }
    
    /* Force white text for specific headers (with sidebar override) */
    .white-header-text, .white-header-text *, 
    section[data-testid="stSidebar"] .white-header-text,
    section[data-testid="stSidebar"] .white-header-text * {
        color: #FFFFFF !important;
        -webkit-text-fill-color: #FFFFFF !important;
    }

    /* Hide "Press Enter to submit" hint */
    div[data-testid="InputInstructions"] > span:nth-child(1) {
        display: none;
    }
    
    [data-testid="InputInstructions"] {
        display: none !important;
    }

    /* Call Status Tag Styles */
    .call-status-tag {
        display: inline-block;
        padding: 4px 12px;
        border-radius: 6px;
        color: white !important;
        font-weight: 700;
        font-size: 0.78rem;
        text-transform: uppercase;
        margin-right: 5px;
        line-height: 1.4;
    }
    .call-status-not-called { background-color: #EF4444 !important; }
    .call-status-pending    { background-color: #EAB308 !important; color: #1a1a1a !important; }
    .call-status-called     { background-color: #22C55E !important; }

    /* Color Tag Styles */
    .color-tag-dot {
        display: inline-block;
        width: 12px;
        height: 12px;
        border-radius: 50%;
        margin-right: 8px;
        border: 1px solid rgba(0,0,0,0.1);
        vertical-align: middle;
    }
    
    .ct-Red    { background-color: #EF4444 !important; }
    .ct-Orange { background-color: #F97316 !important; }
    .ct-Yellow { background-color: #FACC15 !important; }
    .ct-Green  { background-color: #22C55E !important; }
    .ct-Blue   { background-color: #3B82F6 !important; }
    .ct-Purple { background-color: #A855F7 !important; }
    .ct-Pink   { background-color: #EC4899 !important; }
    .ct-Black  { background-color: #1F2937 !important; }
    .ct-Brown  { background-color: #78350F !important; }
    .ct-Grey   { background-color: #94A3B8 !important; }

    /* Referral Tag Styles */
    .referral-tag {
        display: inline-block;
        padding: 2px 10px;
        border-radius: 4px;
        color: white !important;
        font-weight: 700;
        font-size: 0.75rem;
        text-transform: uppercase;
        margin-right: 5px;
    }
    .referral-sent { background-color: #0066CC !important; }
    .referral-confirmed { background-color: #009933 !important; }
    .referral-rejected { background-color: #CC0000 !important; }
    .referral-assessment { background-color: #f59e0b !important; color: white !important; }
    .referral-refused { background-color: #6366f1 !important; }


    /* Plotly Modebar Styling - Pill/Capsule shape for icons (RHS look) */
    .modebar-btn {
        border-radius: 20px !important;
        padding: 4px 10px !important;
        margin: 2px 2px !important;
        background-color: #00506b !important;
        transition: all 0.2s ease !important;
        display: inline-flex !important;
        align-items: center !important;
        justify-content: center !important;
        height: auto !important;
    }
    
    .modebar-btn:hover {
        background-color: #3CA5AA !important;
        box-shadow: 0 2px 8px rgba(0,0,0,0.25) !important;
    }
    
    .modebar-btn path {
        fill: #FFFFFF !important;
    }

    .modebar-group {
        background-color: transparent !important;
        padding: 5px !important;
        display: inline-flex !important;
        align-items: center !important;
    }
    
    /* Ensure Plotly chart areas don't clip the modebar */
    .plot-container .modebar {
        padding-top: 5px !important;
        padding-right: 5px !important;
    }
    
    /* Local Time and Relative Time styles */
    .local-time {
        display: inline;
    }
    .local-time.badge {
        display: inline-block;
        background-color: #3CA5AA;
        padding: 2px 8px;
        border-radius: 12px;
        font-size: 0.75rem;
        font-weight: bold;
        color: #FFFFFF !important;
    }
    /* --- NEW REFERRAL CARD UI (USER BLUEPRINT) --- */
    .referral-card-header {
      border: 2px solid #35A7C7 !important;
      background: #EAF7FF !important;
      border-radius: 10px !important;
      padding: 0 !important;
      margin: 8px 0 !important;
    }

    /* Target Streamlit/BaseWeb select control via marker presence */
    div[data-testid="stHorizontalBlock"]:has(.status-marker) div[data-baseweb="select"] {
        width: 155px !important;
        max-width: 155px !important;
        margin-left: auto !important;
    }
    
    div[data-testid="stHorizontalBlock"]:has(.status-marker) div[data-baseweb="select"] > div {
        border-radius: 6px !important;
        font-weight: 800 !important;
        border: 2px solid transparent !important;
        height: 48px !important;
        min-height: 48px !important;
        padding-bottom: 15px !important;
        padding-top: 1px !important;
        font-size: 0.85rem !important;
        width: 155px !important;
        max-width: 155px !important;
        display: flex !important;
        align-items: flex-start !important;
        justify-content: center !important;
    }
    
    /* Target the text inside the selectbox directly for uppercase */
    div[data-testid="stHorizontalBlock"]:has(.status-marker) div[data-baseweb="select"] [data-testid="stMarkdownContainer"] p,
    div[data-testid="stHorizontalBlock"]:has(.status-marker) div[data-baseweb="select"] span {
        text-transform: uppercase !important;
        letter-spacing: 0.5px !important;
        line-height: 1.2 !important;
        font-weight: 800 !important;
    }

    /* Color states (background + border) */
    div[data-testid="stHorizontalBlock"]:has(.status-not-called) div[data-baseweb="select"] > div {
      background: #FFEBEE !important;
      border-color: #FFCDD2 !important;
    }
    div[data-testid="stHorizontalBlock"]:has(.status-pending) div[data-baseweb="select"] > div {
      background: #FFFDE7 !important;
      border-color: #FFF59D !important;
    }
    div[data-testid="stHorizontalBlock"]:has(.status-called) div[data-baseweb="select"] > div {
      background: #E8F5E9 !important;
      border-color: #C8E6C9 !important;
    }

    /* Font colors depending on background */
    div[data-testid="stHorizontalBlock"]:has(.status-not-called) div[data-baseweb="select"] * {
        color: #D32F2F !important;
        fill: #D32F2F !important;
    }
    div[data-testid="stHorizontalBlock"]:has(.status-pending) div[data-baseweb="select"] * {
        color: #F57F17 !important;
        fill: #F57F17 !important;
    }
    div[data-testid="stHorizontalBlock"]:has(.status-called) div[data-baseweb="select"] * {
        color: #2E7D32 !important;
        fill: #2E7D32 !important;
    }
    div[data-testid="stDialog"] button[aria-label="Close"] svg path,
    button[data-testid="stDialogCloseButton"] svg path,
    .stDialog button svg path {
        fill: #3CA5AA !important;
        stroke: #3CA5AA !important;
        stroke-width: 1.8px !important;
    }

    /* Custom Attachment Icons */
    .att-icon-btn {
        background-color: transparent !important;
        border: 1px solid #E5E7EB !important;
        border-radius: 8px !important;
        height: 38px !important;
        width: 38px !important;
        padding: 0 !important;
        display: flex !important;
        align-items: center !important;
        justify-content: center !important;
        background-repeat: no-repeat !important;
        background-position: center !important;
        background-size: 20px !important;
        color: transparent !important; /* Hide original emoji/text */
    }
    
    .att-icon-btn:hover {
        background-color: #F9FAFB !important;
        border-color: #3CA5AA !important;
    }

    /* Specific Icon classes */
    div.stButton > button.btn-view-att {
        background-image: url('data:image/png;base64,iVBORw0KGgoAAAANSUhEUgAAAoAAAAKACAIAAACDr150AACAAElEQVR4nOz9CbwcVZn/j59zqnq7+80eskISwk7CFiAhYECQMAwC6iACioMrrui4zIyKjgu4jI6o/P4qo/gVZMcZGHbCHtkCgbBJQhJCQsh6l+57e6tzzv9V5zl1unq5IcBtAsnnHb3c20t1dXV1PefZPo+vymXGBROcc66V5oJrQ/in+cmY5lwwprWWWgtzC4vuYu53d7vWmjEmhFBKuV+EEHQ7bdb9Et+ae10WUb9l+rP+1XEv7sW92/4e1X+t3oF79c7c53fmvc07Vtu/5fgT6ar+Vj7f131dpZS7Pb7DNY93e1LzjuLWp+bemi286eO8beKHyxpKzWKGs3oj0Ua5UtI8WtBzGdNCNH6r8b2PHzJ3F1nl+K64P+t3t+FBrH9K/Lk7x70NGa7X3cmOZEN24tfd/nu3zTtzr97mfX5Xfxe2zRt9v9tPzQbdpf6NbudNvG5DU9rQ+sbtUQ0N73VvwRmd7d+rml+Gehg9wO22MEZUaB4aVtZo+RDfN/OLcL+7F2349tw7cTTc8jb2uP6u7V9lNORtOD/eLbzukXwr924bfApxmnc0GnoD2/m6zft837bd2P5N7cDvwjCy/Qf2jV5FX3fLcaOy7ee+oS1v5+sO171xR3Ybz31DJ3BDap4ihOCBDHjFxLKawIK5JTS39a73NhYgr/uAoXjTTwQAAADeBG+bxXHbd4Hh0PlldcY/vi5gNjDtwtb0GF3js8epfvrrrxTcMqomoFETMKnx4uvv2sZybNuv+ya2/G68d9tH8q2AT+GdcE42j+a97o46zjvqu/DOfL/DeMbWbGFHnbFviO03VW+F+jR8iNSSDojHuGKavOGGnui2/xzeXYQHDAAA7yJw9X4T8ECryMk1Bc8MhwwAAABoFm6xIpjSXHOmjF+sbZVz+IAo1Ny8WA0AAICdA5iJ7aeSpZVSCs6VjUrX1lsBAAAAoBkY0xv9oU0IWjNmDDEz4WjNEJQGAAAAhhufIgVRwpwrpkzjkTLZYE8jKwwAAAA0AaE5N1ZWM21/NYXQgjPBrBMMAAAAgGFGaKVMoJmH/wudYWUKsTiztVgwwAAAAMDwIyjNqyt1WVHOl5vgNEcxGwAAADD8+CbkrIzN1Sb7y43Xa3qTOEMzNQAAANAMfK51aG2Nyys0lyYEzZlW5oa4fCUAAAAAhgufccGYEiS4YcLPnMqwTOC5ZjoEAAAAAIYFobXkFHLWpgY69Het/ws9TwAAAKBJCGpAUib/y2xDkqnAovthfAEAAIAmwKVUjEka0W+ln00tlr2bcavFAUsMAAAADB+Ccyq24tqmgTVjijGlw5+2PAvWFwAAABhefJPj1WYQYegBCxbNJdaMax71JsEGAwAAAMOJ0Fqa2YM8coCNMJbxfDWzMtE7eicBAACAnQ1hfVxWGYGkaRgwN3dSLTSHDQYAAACGE66UYpTuZTxeasVdMRaCzwAAAMBwI0Lby7jSnCLNiiLRWmmmtNaKbDOi0AAAAMCwIjjJcDAtTCCairAEjSNUVIfFGOwvAAAAMKxQEZbp8yUrq0MfWFFEmtN0JAYxLAAAAGB4CT1eTcLPJIUVjSO0U5HMPSjCAgAAAIYXYWLPWggaQMg497iZz8CNCxzpcMADBgAAAIYTqoKuJHlNRzDjZhS/Mc+aC4GRDAAAAMDwYgbyK2ki0MLcQqbWOL3a5oABAAAAMLwIrRSPtCdJjcOoQlMGGOXPAAAAQFMQRuhKaCu4YaQnmam7osKsHb1/AAAAwE6JIK9XCBHNQRJWAYuKsNCGBAAAADQBwU0TsNKkduUpMsPhH0YJS0EJCwAAABh+fDODQZH6lbmFm8GEbhQhfF8AAABg+BHKubehE6w5U5zGIWlnenXsJwAAAACGAUHiG3buILcDkXT4l7O4PPYTAAAAAMOA0IJppoWwRc+mG5hzO56QItCkwgEPGAAAABg2BDfmVhnJZ6aY1GYQQ/jTFGZpSbFpGGAAAABgGBE2A6w1N+pXtuqKi/B/odn1SJwDIWgAAABgGBE8KrgiP9fM5deuCEvD8QUAAACagGCc6yjla27RVgiaGoSjcUgAAAAAGEZ4oKSRo6RRwJQK5qYOGgpYAAAAQLMw4huaUeUV09IoYWluirCY1loFphqLGZVKAAAAAAwPgnNbh2VcXhGpP5sirPAXD33AAAAAwLAjmFI8dHZJ/5mjCAsAAAB4GxCMkwaWm71v+5GYVeaAAQYAAACGHx4oRdpXSkdD+aNqLNRgAQAAAE1CUJjZKmFpmkRISlhGCAtFWAAAAEATEFbnmVK+3KhecW2KsMz9KMICAAAAmoAZR2gywKYIyyhh6dAC2yIsCHEAAAAATYD6gDUpYWmtIwloTRbZZoVRiQUAAAAMK1wqxbk2eV5hhCeNt6utMqWtiAYAAADAsCKMu8s156bviCvr7dr5g0qjCAsAAAAYfoSImn2p6UgwzUJzLOxYQhRhAQAAAE1AaKa1NiZYUbpXM80Vs2XRSP8CAAAAzUBQ6XOsCMsUX2kaEUyRZxeUBgAAAMDw4Fu9Saq/on4kTmlh8ytpVHKEoAEAAIDhRJhGX0EzgLXWsSIsFf5H2dkM8IABAACAYcQ3kWdl6q2c+hXjTFhBaO5Fj4QHDAAAAAwbQnHFTbevnQps5hKSFEc0qh8AAAAAw4zgSpsqLG6qr3TF0eUMbi8AAADQJLiUklMVlruJ8sEYRwgAAAA0DWHajbjWZgohY4ppxZTJCiuttcI4QgAAAKAJ+JT/NVMIBdOm+Cr6Z7qCUYQFAAAADD/ClllprsnFNR3Abhwhuo8AAACAZiCiNl9XhEWzCRXj1iDHfgIAAABgePAZ97g1wjxSwuKcKVOXJYyFRjEWAAAAMMyYNiSj/ByVWhk5LApLa62sEpaGEwwAAAAMIz6LepDIzSU1LM5dEZbxiZlGERYAAAAwjAhFltd2IYU+MLUFa1hcAAAAoGn43BZhmaFIldJnUwm9g/cNAAAA2GnxhaZoMylBM6YFzQSmSDQ9CFVYAAAAwPAimHV4wx/c/NNUfWWVsBRjzkcGAAAAwPDgGylKxnlMAYsxLoR1ieH7AgAAAE1AaK4411YJS0dFWIoxmF4AAACgaQjb5cs0F/FgNMLOAAAAQBPxGRemCssOAxbM01qT98tRhAUAAAA0B0GSz8pIQHPGFVNmOKEbR4giLAAAAGD4oSIszSM3lzdQwgIAAADAMCM001Zv0nq5mjGtoMMBAAAANBPBnaElyWcbikbYGQAAAGgiPrdzFuw4Qq3JJEMJCwAAAGgiwplfzcwYQq4VlLAAAACAJuMz23HEXCyaQwkLAAAAaDJCUx+SNkVYCkpYAAAAwNsBKWGZkmdupxBqM5UQYWcAAACgeXCpFCcJDq3iAed4ERYAAAAAhhdh+o6Y0opSwEYM2lRhoQgLAAAAaBqCcR5aW66jIqzwpxCCm1HBKMICAAAAmoFQPPRuzSB+Sv6aIiwNISwAAACgiQgjxEE/bAzaziTUmjFEngEAAICmwKVUnCvq/uVWibJWCQsAAAAAw4uZwq9FpISlGdcaSlgAAABAk/E118JOYBDMDmKAEhYAAADQXIT5yZkpwjI5XyhhAQAAAE3HZ1R4xTTnImoDZm5KPwAAAACagW9akBRnUSG05pqZ/2AcIQAAANA0RGhvtVA6coQFirAAAACApuOHHq/xeY37a0LRKMICAAAAmoygwmduZxEyrrVRwuLRnH4AAAAADD+C22orbicxcApG25JoAAAAADQDUwXNNfm9oSOshfldmekMKMICAAAAmoLg2uPKuLycMSVIBVppLTWKsAAAAIBm4WtjfhnXjAvSgDaq0IKhCAsAAABoGjT014wDNn4u18rMYmAowgIAAACaBw1jMK1HJhVMJjeaSQgLDAAAADQFn6YvkMNrAs7cZIM115yjCAsAAABoDoKb9l/TBhw1IGnNtDKVWEYRC0VYAAAAwHDjM87JvgphlLCMu8s5lLAAAACAJiLMEH5ThmU8Xm7GEYZ+L4qwAAAAgKYhOHX+Mm2dXW1LoBF2BgAAAJoHN1IbmhShGRlg4w87iwwAAACAYUcY1WfupJ/Nr8rEoCWUsAAAAIAm4RsRaBa1HCnOuK284qEphhcMAAAANANfcy2YcYJtL7A2AWijzIEiLAAAAKA5iGj0IKV8hRHioI5gWF8AAACgWXCpVFWUmYqwGM0IRvwZAAAAaAqCPF3NmaJfhDZiWEoxjCMEAAAAmoXvWn6F7QM2etACSlgAAABAExGK29IrMsOkhKU14s8AAABAE/E5DSOMxhGS1aUiLFhgAAAAoEkYKUoWub+mEFqzqB6a86glGLYYAAAAGE6E+andBH7GGYqwAAAAgGYjdOTiWi1oI8AhhODGNsP3BQAAAJoBecBREVZU+awVirAAAACAJiIowkxKWJri0KYWC2FnAAAAoHn48TgzN4MZTBFWlRIWAtEAAADA8EJV0MwVYdFkBhRhAQAAAE3FFGFZMUo7kB9FWAAAAECzEaHvyxXNRDLVWIpxpqCEBQAAADQTwa3wlUn8ah31IkVtwQAAAABoAlSEFTUjmUEMnKLSuro4CwAAAADDh6DeIxV1AWvGlCnCMv9DERYAAADQFHwzhMH+n24yXrAgkwvfFwAAAGgGQnOmQy+YR1LQNBYJRVgAAABAExHUAsxZpfwqtL8cRVgAAABAE6koYVmnVzMzIDj0iqGEBQAAADQJW4SlaSChslLQWiuFIiwAAACgafhm+JHWNI7QOLqcmyKsaDLSjt5DAAAAYCdEKG4afrV1c6GEBQAAALwNCK5j4wjjSlgIOwMAAABNw4+mDvIqJSyjhsUYirAAAACApkDjCIVmnPqOFNNmEKEZR2j+wRsGAAAAhh3fCD/rqABLCdOCxKk32NVlAQAAAGBYEbbsWTNlxaDJIodmeEfvGwAAALDTIijCrJkm35cUOcL/cmUegOAzAAAAMPz4zNRaRS3AZJI5Z9IoYcXHFAIAAABg2BBc0/QF04lkBCiNEpaJSBtBLBRhAQAAAMOOz6L2XyhhAQAAAG8bRgnLDEOySlhMcc4wjhAAAABoKiYEbQLPnAyvGcqgEXYGAAAAmgmXUlamDjLn+2rO4AQDAAAAzUIYI2uVsBRz0WijhIVxhAAAAEBz8DWn4Uem5yg0wdyAIiwAAACgiQitNeehFQ6Nr+k5igWiAQAAANAUSAnLjOTn3ChQuiIsKGEBAAAAzaJ6HKHgTJuyaCY5gxIWAAAA0CxoHKHze5kWtiVYMShhAQAAAM2ClLCYmcBv3Fwd/ocLFGEBAAAATUQorjlXVgC6SgkLCWAAAACgWQjuRChJFBpKWAAAAEDz8U0OWNhQM2dMC1OEZSLSFYksBKIBAACA4URY35cpzZRJBGsoYQEAAADNxo0j5NE4Qo4iLAAAAKDZCDOD3yhBayq6UlYVC6YXAAAAaBrC9B1xkp7ULCrC4gg7AwAAAE3EZ6b2Koo+c84EDeiPjyNEIBoAAAAYXmqVsBSKsAAAAIDmEylhMW3dXChhAQAAAM1HKB6aYOfmxpSwYHoBAACAZuEzRfJXJv8bRZrJIMMCAwAAAE2CSymrQs3W960qwgIAAADA8CK4wRVhMc5QhAUAAAA0G9/2/kadSCQIjSIsAAAAoKkIzRllfyMlLCMNDSUsAAAAoJkIVhlHaG6gSYRQwgIAAACaCZdKMRo+SKAICwAAAGg+wrQbVYqwtEARFgAAANB0fGZ1NxiUsAAAAIC3DRH6u9SGRGZYQwkLAAAAaDom4swUaUGbyfwsPpsBAAAAAM2ASyk5Nf+6OiyGIiwAAACgufjG0goTgdYiGsvAQp/Y3s650FojGQwAAAAMI35oWZVmTIvQwio7lSE0vaTLYRqUdMw9BgAAAMBbRihTg8UZd1LQVIQV/ocjDwwAAAA0BUFurmbKTOYXVH8V3qCVyQJbswwAAACAYYRLLY23G0ENSFwyzTj3EH0GAAAAmoEIXd9Y35HmOvSGTQ2WUcLSUMICAAAAhh3fRJ4Z404OmnNuhxNqDSUsAAAAoCkIxRnnynQhUdGzMg4x5X7h+AIAAABNwWdK00R+mgpMmCn9Gs4vAAAA0CS4lLJq6ALGEQIAAADNxyfTa+SgGddC29izMkMKhdZaCChhAQAAAMOMb7UnK6VXxvXFOEIAAACgmYQuL2V/lVImB2zGESrEnwHY8VBppI5qJIdxUzr6n/09fq9m9LPq8VU7UHle7E8AwBtDuC8RF5H+s2aao/cXgO2i1rC92S9OaPKszatsmYaEcsM2tlxvp+v3h6JZblOcMa5N2z/9EruXcRsRq9xSGwzj0U8e+xMA8MbggZLG8HIreYUiLACaiTOrNbfTt+3tWfdGw7/tFFIy1DV7xY0gbcPnUtKKRWkqFm0KAPCG8I3h5WYkgy3DMk4wirAA2C7iBtU5rEPdWP+n81adORva73zze1jjEFf84LgKrW1INKaXHia4VpocYvdgc8GgshEaWuq51TsA4A3hmwWsMgEnCkGHXzIUYQHwpqmxwYRSyj3AWV9h8j7b+pZpVg7KSqpASiUDrVkgA86YlIoLZ0l5EASe59Hfvu9zxoRnCf+MzPlQe6tN0YddAmhXi2l+p6lo1QsIRrZYU+zMxMt4zV5HjjXDOFMAhoSXtfS0TQmZL6kyE/iZCzHVr5QBAJpprrmZ2xnzDiOnlswt3U6mscEWtM7n8729PZs2berr7e/t7cnmstn+XG9/z0B2oL8/WywV84ODuYEBGcjBwYFiscQZK5SKgvFASi6MayoE5zyfz6fTaSFEEARtbW1CCN/3k8lkpiXT0daeTCY7ujo72zs6ujq7O7vaO9q7urq7urpGjBjR1dmVackM9R6llNZ2xkx47EIQWWkystzdyhmi0gBsB77J8/AolaPdVwpKWADUUBs6Jgscc3CNNQz9wrjRLRaLPVu39vT2btiwYf369WvXrl23bt3mLVu2bNk8kBvIZrMDA9mgLMulUqACprjmKvzJNPcEuZaCceEJik9zT7BIpJ1rGhpq87WRC6tNEklpEnnXKggkp0dwlvATnPNEMpFKpNKt6c6OzrbWtpGjRo4aNWrCxAlTp+w+dfepo0aN7urqymTSiUSi5ggopQKluHO9BWdOSV67AhLykLWOvOOaHDPW9AAQXEplcjlRQSOKsAAYolSqJnFLP2sc3CAItmzZsuqllS+veXnFypfWvrL2lTVr1r/2Wm9vb29Pbzkok8H2fN/zRDKRNFFiLlj4fxF+4wQXWnCPWZvm6VgtVNzVrq2ZqrNqnAy0rbTiUiutFMWracUgVVkpVQ4CaVBK+cJvbW0bOXLU6NGjWtoykydOnhr+m7LHtD2mTJ3a3d1dY5JjcfXYK0dzXLQNaIe7JWI7DwAguJTh6liZgJGLHbHwe8s4QxEW2IWozXTW3Ru5uSJ+bykor335lZWrVr700ksvPP/CiuXL16xZs6Vna1AOyqUSE9xP+Ck/6Rvi3yYa9lmpjaqkXiuERpRXpViHLqLm9RY6/qzK78x6w6HB16F7rXn4KwkABKFFLpXLoVXmyqwVEolkKtnZMWK3ibvtPnX3PWfMmDlz5h7Tp02ePDmdTsf3oRyUGdc+Tyizz7SQF6zK8WVRWhiXFAC4VMoGjRh3xRKuxQDxIrCLEK+cathZ6zxdrfWmTZteemn5kseffHLpEy+99NK6ta/2bN1KX6J0Op1IJBLJBKVn45a7qiuX/ny975XLNLtbhvo+1hi5ev84ts1KhRQ1ILrRZ1YMz+SWNdOCHsq1DIKylKVSeXBwsFwOUolU18iuSRMnzZw5c5/99pk9++Dp06eNHj0m4fu0TRm+X1tTLbho2OCESwoAPFCBMb00Ftikmyj/K1CEBd71vG45blynIh5ellJ6nkdVysTmzRufXvr0k0uXPvX0U8uefmbdujXlUsAYTyYT6VRaxB5c0+RTZe3iPUjbFtdw/i+rdWpf18Q2bI6KbdUlZs0jKw2/UaK4cg/pwmvBmGTKxMM4557WulgqlktBUC4prZLJ9JgxY/bdd99ZB86afdBB++23324TdnMvpwwmYSzsysN1Opk4dWWBwStDyFE4DXYFeKDKXEe9EEqH33guaLnuBvTDAIN3HfakpUqkRmIR8VAzOan0lHiEef369S+88MLDDz/82GOPvfTS8k0bNw8MDCSSfltLezLte8zTofVQ5p8rPLI6UxWrTxZlu6uQ4nc1TPrW31v/9PoXqvqFxQJecVunrSRPA/vNrVIXucuUrrZTTCUrB+VCPl8slTKZTFd319TJU2cfctBR847aa++9Jk+e7BqIAyW5sgeZXtBOQeWVl4fxBbsOlSIse31CERZ491PlZdoMa2QFnWmMNenGTUIQBC+88MK99977yCOPPP3006tXry4Wi8lEsqW1NZVIcLLOmjOhuPa4MDXA2o90rGIW0dhjMsC2BthFnofI47LYjm0jrVtzb81zhzoaLHr06zre9U4zXRgEzU0L1+gmPu0anZktHlFaBrJcLpe55oOFguB8/Ljx++y37+FzDj/66PkHHHhgKpWirdFyxx7z2EcCDxjsUnAVKBKEtrEoikG5IqxwqYsiLPDuI27kyAQKxhW3uscUF/U8L+7sLn36qUcfeXjx4sXL/74i29/PGMtkUql02hM+BaXj24+Uqpxr3diq0U4wE15yxVTxL1SNN1y/kW0b4DdHJAfNNS26SddKKbKIDQPX8Xddbx/j0XthCI0sU8VCKT9Y0Fp3dHRMnz798MMPn3fUUbMOmj161KjomawsA2HCDsoFqE1MWuCaA3Z2TBV0eGEyTkC4vGVcMS7MQCTSpTNydFzgywDeNQzl4dmiKs69yN9dvXr1/Q88cM+iRU89uXTd+leDcjkd2tyU7yeFEEpKqZXWlXRMfGsufBoTRa6IJEdxJNuXY61djVtJzjfnzimPq1a5EHr9u6gyybbdJyaU0eiJ5P0rrRXTVIBtux6oVIrVvrXYq1RiYdY218lE10TFqRXZqGhxT3iMsVKpUCgWC4VCwk9OnDJ59qxZCxYsOOKII6ZMnUrPtOsbczSiFqaYoFhNKh0uMtgp4IGWJrnjtCiVcXyZu67AAIN3MvXJ0YaSy0qFa0sjZxGy5pU19997/92L7l68ePHGDRuE8DLpEE+I0OIyqkJkLr5qQ6PV3io5t1HhEgV3aaansPLItry4ylZRJxLTWpowLM3/k0rJQBq3XEqpjDWKcrHR2yR/XSnl+z6VNQkhAmO3dKTlrpQ0SnYq8vw5/SmE8DxfCC5M3zEzW/N8T4S2VZhaZW6epGMZYmolcp1DMc/ZlkwzVjeJsMZyu/ulDIxpFUEgSyHFIAjGjh131FFHHXvsgrlz506eMoUeWRWdrn7JirwHTDDYKeCBkjyWgOFUhKXtKhdFWOBdRMX0Uj4l6v9xHUSvvvrqQw89dNNNNz3x5BObN25mjLW1tPoJ35Zg2ZRtJWRNlpIJwYaI+sZC0JqZILep6rXzCRzkz4Xep5RBEJRK5SAoUwegn0h4CS+TziSTyZaWlq6urpaWls7OztbW1vaOjnBZkEq1GznJlpaWdCbtCS+ZTGqmPRH+U8ZWhQabs1KxlM1mgyDI5/OlcqlYLA0ODhQLxf5s/0BuoLevr7enZ3BwsFAo5PP5YrFYKpVCV1hp3/d8308YVRAKINM+6/Dokc/saa7sEj1WvBVXpqzvPK5eA0VpaPPRCI9rpYJA5fN5pdS4ceNmzZ79voXvO/LIuZMnTaItkE9M6wznB7PI26/xyyGzBd6NkBBHzcmKIizwrmEoD5ikLYQRhsxls/fce+8t//d/Dzz44LpX1nIu2tpa06kUF8L5W5XG36hiikQkSO5Rm7iQnUygqFPWBlqjCiRJwwlCP88YKmlsbbEYunpM8EQi4ft+W1vb6JGjxo4fN2HChLFjx44ZM2bcuHEjR45sa2trj6gXgBwuisXiwMBAf39/Lpfr6+vr7e3duHHj2rWvrFv7ymuvvbbu1dd6tm4dHBwslctaKuF7CT+RTPqJhC+4r7nUkg5N+PZco4RuNN9p258X/WJj7WaZlB8sDgwOBjKYMGHivHnzTjrppAULFnR2dtIjw2uUENwOK2dOVrNG4qOhAgkA72R4IJXLsJj8rxbUWWHSMVDCAu9Aai738d9rXN6lS5fecsstN9988/IXl2ulQicynXYdR/b0piixJsVEHok8U0hXMxF9PUJDLDg9XAk3I4i+IIEsl0vlclAuB4EsScmCtpa2kSNHjZ+w25QpU3bbbbcZM2ZMMYwePTqTybzeeKLXt2Rv6Fix2IjDhqhA9uf6Nm3atGbNKy+vXrPypZdWvbx6fcirW0OrnBdc+6HvnUok/FQq5XmeE8GuT403VMpUWgluVjzGlgob43YP00ywYqFQLJSTyeSkSZMWLDj2pJMWHn7E4cJ8mvSpUeOTrWaLfoEHDN6lGCUsktvhLmUVKaqTDA5OaPDOJjQDmpKv1vRu3LDh7kWLrr/hhiWPPtbX19fW1uZsBhkMHQWcWbVElIuxkjss3O+VLwIz3phNR5ZKpXx+UGmdSiXaO7rGjRs7ffq0Pffca8b0PaZMnjpqzJgRI0a0tLTU73B8QG989G/DL9rrfge3MZO4VoeyTuErqvvi9XUe2Wx28+ZNmzdvWrly9fLly1csX7569csbN27o6e0NykFoiZOpZDJpa56jML6xjF5ddriCirouTJyBxq9xE+LWnvAF96SShUJ+YGCgtbVt/wP3P/nkfzzppIWTJk1mFBCXkjLEZjWkaw6Ce1NDvToA7xx4oKXQPKYaoJgZBswxjhC8kxgqyEk21Y9EEB9//PGbb775xhtufHn1Kj+RbGttTaVSzuWNPY2Ra2utrOs9JSUNpUz+U5j6ZE166Z7nk9xxoZAvlwLNZKalddzYcfvsu8+sA2dNm7bHfvsfsPvUqV60Jw4pZXwMfkN/sf6Wt/iN24ZktI7m98Y0KN2zyNGPdxNVkc/nV65cuWzZsheef+GZZ595/vkXtmzeVCqVfE8kEr5vwuwmWiy0ouA0havjzUtcc1oG0fLG3BtJc1H1l5RSeNz3/aAscwP9pVIwfvz4E0884dRTT587by7NmwoCaV5o6CVL/ZhiAN5hmCIsZqdwG9E7FGGBdxb16T3yYql8gVzeLVu23HnnnVddffUTS5YMDg5m0ulMKkPNtxSUrrJ80YiDyCA4M2WFoIxVVlp4piVGkEM2OJAPAtnd3Tll6uQZ0/eaNXv/OXPmTN19j5EjRzlDFVfUqnnR+jxljZ/6Nh5QZuPnUTeRE2d2IliszlF29VDuz3K5vGnTphUrVjy+5LGnnnhq+Yq/r123LpfLpUy9mO/71HRdiTpQ3jZ8LRlZx9h4hki3hNskr1JacSY8j2kmCvn8wECupaVt/wP2Pf0Dpy9cePL48eNDM2wKtTwhKqMs3LurPsIIUIN3IKEBFrXrRBRhgR3GUNHUuHYVPZKMwcqVK6+44oq//vWvy5cvT6dSba1t1KXj3Di3zQaSyPZcNwXMppCKLIQwF/RCIV8oDJalyqRbJk+cdOihhx5++OH77Lv3nnvu2dLaFtthZeTkbAtr/L00NLo73gwoZsuo6P0LRn2GFU2ryj3x1ltjrXVUp6a1X+3r9/T0Ll/+92XLnn744UeXLHls/auvFkvFVDLVkmlPJn2zLFHWwmpps+ssqlc3ugMimiwc/k9V7CitaTzPC4JgMJ8rB8HkyVNOOukfzj7rrL323ruBrpYdxbiD1jcAbDfWjagoYVGbAIqwwA5EV8UsK6ZLM6mku+4/+uijV1999S03/9+GDRsymQylWkPvSnNyiLTSrDrkW1PwXNM3TP08+Xy+UChIGYwZM3b/Aw849NBDjzrqqFmzZrW2trqNGJfOurgN067ukUPGgXfgd6pO1aKR8rOVrqyaIRgLlVU2Fi2J4nORN27cuOSxxx9avPjhRx5+/tnn+rL96WQq05JJpVKe8JjL2FPhG/fCI8oV00bZk1IApMTHtV0vmD0KZFkIP5lIZnPZYqE4atSo40844cwzzzxszmH0uuUgEB4TxpS7oYfbnpoMwA7EeMC6MnNUV6JB0SNwyoK3EV1z0acaZWN+6RJfKpXuv//+P/7xj/fff19+IN/Z0ZFIJlVUA1TZDl1/mZOZssSEqqyVEZxLpQqFQrFQTGfSY8aOOeTQQxcsWHDQ7IOmTZ/mznyXyqXdiGk/OcXJd4yP+1aIiXmx+LuIVkWVFHJlzoWOh9/jU6QGBgaWL1/+0EMP3X/vfcuWLduyZYtWuq2jNZVMW+OtlDQzYCgC7o5XXA9aVxJkVvLTMxQKhf7+/raO9iOPPPKMM8447rjjMpkMqXnXzGyOnVEVuWkAdji8rKXQ5PXa/JpxfGMXqnfvpQS8q6jxF22SUmulpe8nGGN9/f033njj9dde++ijj0opW1tbU8m6AqvKBCT6g0cVzZXHKNPZK0wTcH6wUCwWk6nk1Km7H330/OOOO+6wOXM6OjrsI6NRemRRhmp13Z7xAQ3MM+PDpa/4xmz/dnvAPBruVClr0rz+KfVHhkQ8lFLxnuY1q9c88OADd9111+NLHt2wfqOUQSqdzqQz1iZa6WfOK4pi3CyhmI7WN9qErZUbAkFxaRn092c5ZwcddPBHP3r2Ke9/f6tJEJAZdtJjbhwFe6dlBMAuDBVh0RlOFwWMIwQ7gPp6mfD6KpWX8E1+seeaa665/I+XP//C8+lkqr293aUGazciSD+RGVWqSvmViWqasKYQWqlCMZ/L5VKp1F577TN33rwTTjhh/wP27+rsou24uQvOkWpgovDViFHz8dEoBe1S8Sa/Kyp6ZGsfffjx226/9dHHHlnz8lrGVVdHt+ebbL0kH4BXtKqpbJpHotvVjoFLK3DO+/v7SqXCvvsdcPY555x66qmjzLwH+igjj9ytebRLEtfs/I49jGBXo04JC+MIwdtOvfWlUUWMsd7eviuu+POVf77i2WefbWlpaW9vd4OJ4gMAaoutzLCBCG4uweGWy0GxVChx7k2bvsfxJ7z3mGOOOeTQOel0mh5Xo33YbFtbZdRj5Uhv4uXeULK5STnR+g8x/hIuRh3vbtq8eeMjDz9266033//AQ6+88orvi/bWjkTSk0pTERa3SijC/C241kNFCOzrcjmQzefzhZkz9/zwR84866yznBnmRnElquHmLmuAhRTYgXBSkaWVpUsDV8YRoggLNI2Gl+wgCChu2d/f/5e//OV3v/v98hf/3tbamsm0mkkDDeqeaq6hMWkle3uxVKTm3QkTJixYsGDhSQvnzpvn4sy0Bo03KdVoSAxVV/VWL9zxmmMWVSbXdt5WBYxrdo1iV7p6NJGLIfNGx/lN7up2vqFtVqJFlXThvtUIlq1b98odd9x5+213PPrII1t7t2QymXQq6XtJY4SVS5Dpus/FlUm7/jTjf6tSsVgqyUlTJp999llnnnnm2LFjGTNzD7nYxpRDXOXA24xtQ6ppWkcRFmgq9ddoco9o4E8ul7vuuuv++7LLnn3muUxLpq21hUqs3OMbtnjau6I2FOoKKpVKg4ODrW2tB8464P2nnHbMe46eMmUqPTKQAdNW1Uo4WxZNGGz2ex9CbLLO+Mefth0br54a1EAWahubaYY3PIQyF5XAaSnJHltLrJV+9rlnb7vt1ltvueXFF/8elIPW1nbP59qI9WnTFsyrhclqdj5SM5PCE5x5+cLgQG5w9z2mfeSsMz/ykY+MHj26Zr0FJxjsWKI+YB0VN5I6KzXp4dQEw0fDNCpBkUnOeaFQuOqqq37/u98/88yyTDrT0dFpAs6Bu2IOJbDg1DW0lVjihWK+VCxOmDjp5JNPfv+p7z/ooIPoYS4paAv+3ZheallhrGFw8nVjvHV+ap1nWv2UmiPDqmfaV8lmvfHvHhWNG8lHHR9TMZQOV+M8d/wdxdQ54s95vXhAZRvx2uroGNEm7epfSknLL9Lbuve+e/56442L7r6np2dLOt2STqe1lkwL4QmtdFUUhFrO7KdAsWpl5xEL7Qk/NzCQzWZ33333M8/88D+fd97IkTYoXa/zhasceJupjCN0ZzPnDEVYoHnELQ2FIn3f11rffPPNl1xyyZIlSzLJdIvxeqmoSqtK8249zgAIHv6Tqtzf159IJA86+OAPfPD0E088acyYMfH8bsUIUYkvo/qeYTq9o7LhSGur1su3GtTRy9V0ywyFlLJYKgUyCErlQEoqMK6YC/POPSFEwvc8L+UnkqnkNlQY43LQ9aJdsbS0a9phrCpI1jhEUC9YNvSDaWOKBlZpq/BR2afwzRi9yRUrVlx//XU33Xzz8889l0z4He1dFHCmFVt1DbZrPWOVZQOTJLPiCZHPFwcGs3vO3PMznzn/Qx/6p0wmEwRBjbYXrnbgbaZOCQtFWOAtUdXjso1oa1zDedGiRZdccsl9993n+357e7vHRaMpmdvKxXLOi8VibmBg5MiukxYuPPX0046cO8/3EnGXN3q61VuMOm2cDjG9RmOTsc24cdUuspiupTMrQ9na0LIWi33ZfpoMuGXzli1bNvf29Gaz2d7e8Gc2mw2UDErlvv4+KWUhXygHZSpDE55nBBuVoMZbLrykn0gkWtKZkJZMZ0dnR0dHd3d3V1f36FEjR40ePX78+JEjR7a3t7e0tNTvDG22kWA1Z7ECMRYpXLDXl7DWjYLqNa/rSpKrNhl3iLPZ3L33Lrr+uusW3X1PPp9v72hPJpKuQ6w2uG3XVJVWJZck9hMil8sNDA4eethhX/riBQsXnugi0lGlS2UUZc36CYBmUKeEZatCUIQF3jK6yp7pWPumM70vvvjir371q2uvvbZULHW2txlL6NXEBmNyyk78wbpBnucppXK5XLFYnD59+kkn/8MZZ5yxt5EndDW3biNxh+mNvY/tyJsqmmNgHFOvka1VSm3ctHHDhg0bN2xcG7Fu3bq+vr4BEyMdGBiQUgblIJCBs1rk1psRw1wwM4rXrpftmANSd7QvYaZC0fHVTvnRbCRpZu63tLS0tbV1dnWOHTduzJgxEyZOmDx58rhx43cbN37cbuO7u7rq36ARW9bC9G/ZvTKjkRU1GtW3bgthhTPs2oZOAFtIxUj1k1dJnVX+qAhwVTZIzb7056OPPnrlFf/vtttuW79+fUdHR2um3Wh9S8YjYTImaLoSYw1yB3Yp5omB3IDnead/8AOf//wXpk+fFksMi0pU2/YhwwiDJsKlUu4LAyUs8JZxVje8RNf7jVJaLcmenp7f/u63f/jvP2zetLm7u9MTnpFEouylmbXOQ/simBuSF58WyMj09vX1Cc4PnDXrn874pxMXLhw3bpwxdVIp6g2lhK5glabPN3xFHaqmV2updcVQel7VoqG3t7ent2fNy2tWr169cuXKtWvXbtywcf1r63t6evKD+VKpJEKL5nm++ed5XAgy24JXAlJRT21tKrlKRJOGCkW/qEq+l0yIdVvJLmutA0KG/5cy8D0/mUy2d7R3d48YN27cpEkTZ8yYsedeM6ftMa2rq6uzszMuMOnaoxmjmYy1RUyNLhdxJ5iKpLzYfCSCTpUoeF/dYBaPOdNyasWKFdddd+2NN96w4sXlqXSqo71Ds/D9mEizFztVTB7YrhZo2WA/u0QioZTa0rN57JhxZ37krHM//rHx48aHixcluWfEwJ2SGn/d4AgAbx5e1tLT3Mm0RRUotFa13gYMMNgG1dHgKhHD6BGMRcPvhBCDg4O33HLLL37xi6VLl47s7s5kWoKgVNV/41KzdlKgkdFgtrjZ88JNZ7NZxtkRRxx57rnnLjxpIdkJI35EI/Co0ZNF8wVcUdGQdUPbKNytmafr8tbxg5DL5VavXv3i8hf//vwLy19c/sKLf39t/Wt9fX3lUilccyQTyUSSc55MJZPC577nWU+RVxy+6GUavnpVeZcpzao8JjrGpLrZOOnaqBSLTDKFwcvlcmiQg0AplUynOjs7u7u7p02bNnOvvfaaOXOvPWdOmzG9q9pFJm3OikaYSxJXiTDHS9NoErON/Ve3Ybld5pUcQXWAgVOXEWO+bRDfctVV1/zxj39Y/uKLyWSytbWVc89lgIV9CfpYKRHgjLpd2AjBiqVSb2/ftGnTvvilL51zzjk0Ysvol4bWWnAbl3aF1/HacgDeOrwsA8+eZzVKWFZ0DgYYvEmiC6zpIbIx57vuuuv7//H9Z599NtOSyaQyUgbxs4vr0Od1v1uLUrFDWik9ODiYSCTmzJlz3ifOO+697/V9XxnNLCZCn5LXvHYV2+vD1Ng/N2wnHhsfGBhYt3bt8hUrnln2zFNPP/XSSy/19vT2Z/vLxZLneclUKpFIkKkw7qnxQekarioBJ7tbQ4yLcD4uBZq5sAsKERnghse73mtn1TfGdRlV/OBHFUk6NK5BsVQsFcvloOz7XntrR/fIrunTZ8yaNXv2rFnTp0/fbcJuNP2CkFIqE/wV8bJtW6Bc+SjsHvL69ubqNQareMA170JRal1ZddJstv+vf73xT5dfvnTpU4lkqqOt3Rhp6zxEYXJXss0iAxwdWqE94RUKxf5sdv78+d/4xjcPP2IOdQx7NNuKV5LW8IDBsMNlEJhJZNyunrXJNDFFDUk7evfAu4AhC5S4naNAl/VVq1b9+Mc/vu6667hi3SO6nYazswdRytIGC8mNldH8A81ULptN+Kn5R8//+D//83sWvIdeJDyBPc8GYGNOt5E0V9XdNB6zQsKNA6c1b4rFZA7pxlwu98Lfn39iyZJlS59Z9twza9e8srVnK2Ms9G8TiWQy6eYQOLNtj0TM0214rMjoDlFIHPNfSaCxkacef2T8ReNGPf4LjxxDm3MyHwCLGqjCoym0x31qwA3KqhwUg7IqlUpM667u7ql77D5z5sxZs2bNOXzOXjP3SqVS7nVrZyHEfN3YnpPrro1zSw9W7jNqIEfiDoUVoq4q4guC8q233vL731722GOPaa1JLi3u9NfuR7WIB+1tX2+vSPinn376V7/61SlTprhWJeeh80qGDhdGMDxwJWlmtoYSFnhzDKWT7GbjDAwMXHbZZZdeeumG1zZ0dnTU1wPXjCRwKz+qyFI66OvP+l7ihOOP/9SnP334EYfX11hVNsXdeKLXN8DRUyp/0j7Hw8v5fP65555bunTpY48/tuypp1etXj04OMiU9pOJVCqR8pMsCiDV70z1jjVuI274SCftVHOcaYBEjZF2YpaMN64Xqwli1QyEaNisTA+kZiHzuwnHKiaVZEyXSiVKJTPN29s7pkydNPugQw4//LADZ83aa+be8RcKAikEr7mGxBuJ7QsZ9Qz3GcUNcN3yjvJkts2aBhv6QpC61t13333Z735/1913e77X2dHGmScDyQR1V3oUjq5qZtZMCman0Qglpezt6xs3bvynPvXp8847r7W1VRp/2quKSONKCIYNakOir1rlpEcRFnhD1LhuzvQyxu6+++4f/ehHjz/6WGdXZybTIoOg3ksjI0mukPPPuAi3luvPeQnxngXHfvrTn5k3b17N9PWhT8sGIehaLYjo3FY2wWwG8ZsNBkquX7d+yZLH77/vvieeeHLV6lX9vX1+IpFJp5OplAksCxUuXZWJiluh6W1/WYYKC0frDV4TfHY52rjMdc2h41UmJZrYN4RaSK1yBauMJWAN+3U1lS8rwStCPVTL7HEv2jKTgSoU8/lCXnAxevSYPffc8+CDDz7ksEP333//CRMm0MMqwy24q+GuP2INQtBDfbJmacU0ozWW0LEqraAc3Hrbrb/97f/32COPMO53dobesJLk1OuaaAQlPaIyfcVFeM7m84VcduDgQw/5yle+esIJxw8133DoPQRge3HTkFhU20EnNNW/oAgLDElc5ahSKxTZDM75q6+uv+iiH1199dVa6a7OTorJxt0g5lovtTIh50gzydRAFwqDUqojjpx7/vmfOfa494amV1JgJrSUtQaJOkesPeKxzduoYfR418vEXPOJK/TNFwpPP/XUgw8++OBDDz795FMbN2/ijKfT6ZaWllQy6Wx/9O5NtYQyFWM8KmN0xVDx4HNsqF/jhigzjE8ps/6I2oc5i/03dtApI8xcMtNVSpsZflxUDm+jNlnX5Rpt0H0K1uLGbo7uY9xeDLQ12PQIFfqEnLJWXuheSp4v5PP5wXI5YJxPnDjhkEMPW7Bgwdwjj5w+Yzo9SWmppBLC45Xcqo5kNGLWN0pxNwi200hEEvGIenV1ZJKltGqmUgY33HDDpZdeuuzpZalksqW1RUrj69pjThVh5j2aDqmoQo9prjzP54z39vdqxT70Tx/6l69+dfKUKTXLvuqVUPXOA7DdkBJWZTYIirDAGyNmXVxaTmt9ww03XnTRj15a8dKI7u7QLwmCuLyDK9+1TTxMauZJprngCe4Xyvn8YGH27IPOP//8Exe+z/cT0ug2cyOlb+3SkO0utqclMi7RfOvodZ2Ag1eZjvfqE088cddddz3+2OPr1q3d2tPjCa+9rc1PJsjs0GyluA/EKTCr6YJuYplUNFEjqh7bs+hoNXBkbRBYWy9NWUE625hKjqNteDUPJGNgfjLjUnJ6gjBhcHpuJKgT6WzH2oWtI+u6m2L613b3yd910//I8JsrQmiEmFCh9dKUorK5KqM5RVsXPDSBxUIxO5DlXIwaMWr/WfudcMIJ848+esaMaZx55i0E5qByz6PVkqg5TtuyZG5uFO0ed6qf9qNXSpI3PDg4ePVVV1/6//1mxYrl7W1tmXRrTL6DaaY4dcrFlh7mPVKRF+fa27J1y4SJu/3L1755xhn/5Pu+c4VxVQTDgg1BV053+uJaJZkdu29gx1A3r30oOcFYVFdXuoyeeeaZCy+88P5770smEplMi8kasnh+tKbLk0pbueaez4ulUrY3u/c+e3/uC59//6mnplOp0O+VVCbDTIiaZMqH3vlYn0gs5qzJbsUn8Kxfv/7ee+9dtGjR4sWLX3nlFc55a6YlmUz64aWXu7h0fFdrxBqtxEToX3vmF66j7tv4kYwnZmuEouxCQUQGWESVt5oFUpZLpaAclKXimkljRhOJhGke9nxajHA7IJ+6em0fUSAZ58ITRvvJT/qJZCIpjKHTTgaF2aYlE4HVoiI/aSxZRYiC1Uhm6GhEfh2hD+oizPaTNc55uRTki/mgVOweMXL2QQcfc8wxxx9//IwZM+LzixpqbzV4Dc5rpdZYTPEj2mc7hjgaJr1169Yrrvjzb3/7/1u/fkNXZ6cnvLjOWryjmtHQaNqcyS5zwQuFfL5YPvro+d/97nf322+/+uKDaKUCDxi8YawSllVldeNIrQAsirDAkFRStnSx8z2t9R/+8IeLL75484ZNnV2dbroRPX6okQZSSpHwZCB7e3onT5z00Y999Nx//viIESPqFfNjUpIN1wMV3f/4K5LH47azadOmxYsX33rrrQ/e/8BrGzcE5XIqlc60ZHzhkWflrGztC9TmVivtu8b5ZO467p5JmRxKVDof1MmBVTp/ONdSlsvlYqlULpcDGXi+n8lkRo8ePXbM2FGjR40fP37UqFEjR47MZDKtra2dXV0tmYwQwvd9z/OCICiVSvl8PpfL9fb29vT0ZLO5np6ta9as2bBxQ8/mrZu3bBoYHFRKJ3yRSqWTqbQfrjKEivaDRKoUZXiV+RcLfVXc4rhYT403H9lE65pyu027SDGUyqXc4KAuy1GjRx122JwFxx47/5j5e82caY9VJJ9CH671iBuVqfOKMLXxermIXHZe89HHhbTWrl37m9/86sorrujt7Rs9erTgnpSBi/BRAMNz50z0lswg4dBN37K5Z/SYUV++4MvnnfcpzxNBEJjNUihbMeY5IRFcKsH245Swqi5uVoYtXi6Is2qXoUYfoX5hH29icW7lipdWfO/C7/7f/93S0d5OSkPxq2d1ApLbKQvmFillLpdraW09/fTTvvDFL1IHCMX66gbjD3kaxgusKHXpCsHoOZs2bXr4kb/dedudjz3++IoVKwTn7e3tnu+ZIGR47TdCi8YhtCFNUQm8Ml3jtdf2LDk549o3G0uQ88p4Mbe1cqk8mB8syyCdTHV1dY0bN27PvWbus88+e0zbY+KEiWPGjOnu7s5kMm/6o8z2Z3t7ezdu2vDKK6+8+MKLzzy3bPnylza89tpAbqBQLKYz6fbWtnAJILiSyoqQVSvh2fh2dIyjk6NxxVnVEYg1JZsKT05hf61VOSgPDAyUS8GYsWNmz5p9/Pve9573HEOfu9X3MGOPOKsko4dWmZamcDqusSVqOrlJMIQq25c++eTFP75o0aJFHvfb2ttNvDow78w0edOIC/PuFXXD0atL6Xl+WZay2exxxx73r//27wcccIBzheEBgzcNLyspdGyKKqOcGd/G1wzs3DQ0wKxRr5HTlbzyyisv/M53tmzZMqJ7RDx4Gz95IgeFLuPhH57nFQqFfD4/d97cb3zjm3MOnxOvOK0XgIztXdXtuiqPR3rMjFQJgyB44P77b7nllkWLFq1atUpw0drSks5keLj0lLZey5p5HTVBRTW/jA11Pa3p6qkurGXxJDerttbC85jW+WKhXCgprUaNHrXffvvNPvig2QfO2meffSZOmlTfy2SHzFf3TA/lo8cf1nD8Q6lUXLVy9bPPPvv4kseXLl26YvmKrVu3CCHSqXQqnbaWWCnjhEYpXhsWE9x0J7Ko6alm2VHzi1s8mao8o65XaXjSIlwXicGB/ODgQFkG48fvNv+o+SeffPIxxxzd0dlJDwrKQWiFKShfa9rc2ySLG+lnueWBPSG0Defxqr7hm2++6ScX/+SpZU91d45IpvwgKHPmCU5ZBJv2VrY+mptUvzHR5svQ09vT3t71xS994bOfPT+RSJArXC3vFQlYasZRmwW2CQ+kJGNbV4TF4hEzGOBdCvuJxzxKVl1M6xzfdevWfetb3/rf//1fSqDS7Y06Ppkt+zGb9BKeLKutW7dM3X3ql7/05TPO/LC5lpVNz6i/PcMP6iQkrZ3yfUE+0PLlL95xx5033njD8889X8wXWlrb0ukUiyls8Hglrv0KKOZE+DUVGLPYoyoHpyYoWtMJzaonT1ghEa1L5XJ/tt8X3uSpUw4/bM6xxx138CEHT5gwwbUduwPrdq9mmMRQ2pn1D4gELEkVyj1Gc15JChQKhXXr1j388MOLFt3zxJIlq19ezTVra2tL+H5oPo1ol8tJVwYnR9H2immuiw/XZhxso5WtCTNCj1IzrsIFnCeEKBbLuVwumUxMmjT5PQvec8oppxxy6KGJRMJOg9DaE577EGLhEO2K7Dh3+s8VMUtS+WAV3Y9KHcDg4OCll/760l9f2tvb2z2iW3CPdC4FlQtUqr2p69jUx5lJF37SL5eCrb1bjp5/9Pe//4P9jSvsZjdF6WGYXLBdcFePEK12OcYR7uJEbkTthd5dZym5yxi75pprvve9761fv35EVzfdXt/3Ur1lcyXVrKd/S2tL+1lnnfXZ8z+7224Tokorn9Vf0But/2psT7y6anBw8K677rrhhhsffeSRzZs2JVN+OpP2uE/mubIFTopbPGoqjXKYwpb6cmWqfeskmRo6fENpgZEbGgRBLpcTQowZM2b+0fNPeN/75hx62OixY9zDZBCwaAzDUJny+Da30XM85J7EUqrxZYF7wIYNGxYvXnz7Lbfdd/99m7Zs9jyvrTXjewkpKRxrm7yEaT/SwhwnUw831DWi0s9drwJWFdS3et2ex6UK8oPFYrGYTqcPOPCAfzzl5Pe//zQasBEbWBRf27GozUzZJUEkZV9dBifNI0mGRVBMgd778uUvXnLJJVdddTWTqnvECCNpGlh3nSrqtDKniBfa1phNFYL1ZfsymdavfOVfPv/5z1GsJXY8K2VisRpzAGrhUirbaVAZ+cmhhLWL01Dcyl1VPc9bs2bND3/4wxuuvyFhhtyFJlBYPUkW0yDUJhpsPAMtBCdVrEKhcNTR87/5zW8eeuih8UqrhioZ295JF1RkjD311FM33njjbbfdtnLlykIh397W2tLSYsOH0Qat4CX5tyI0IWT2rLwUlTSb0mDj94lw/4WuMcP1oeZ6GUt6R1QblUqlDj/yiNNOO+29733vqFGj6DFu4bvtb9ZbiT8N9SHWHNia4vB1a9fddvtt/3PjX5csWVIoFNraW9OppFLV4lzVH1V1QzZndbH3+LsIt0M+om6wjqGwealUyOZyTPNx48Ydd/x7T//AB+bNnUuHlDLEnqARUMLN2IgLupjX9KLSNxFLCEibX4s+Ajp57rjjjl/98pLFixen0+m2trZyucR4eLpqJWxWTnlKmDbi2Bgn47gXstmBBce996KLfjB9+nQlJUUYKMmiw++EDUK4ioc38TmCnRgqwrI5FbdSgxLWLk4Dof/IoRFC3H777d/4xtdXr3p5zOjRTOvQgaPi1eorr4mnhMs7JZnnhR5Eb2/vtGnTLrjggtM/+AHf96UMmOmXiWpnGgjx13fNOjUMshnZbPbee++95uqrH3rwoa1bt3Z0dFDVkhnzYMynzatEp7SNqNNr1EbLw6utrcURLPYVqD8aNS6+04mkxxeLxXw+P2bsmPedeOIZZ5xxyCGH0OPdBKG49W3gyMabf1yTzDB9DWtj12bZ7dqB6KiWy+XHH33s2uuvu+2WWzdu3Nja2ppKJ5Vs8Ik0rsaqq9qr8t3dKIrYPCjn0ZLOs6mP44VCoS/b39rSMmvW7FPef8r7Tnzf5EmTq0qmzVxEszhwQle2Lj1WxkAKIh7j8XZoFuXXle8nisXiVVdddckvf7lq5aru7m7PE1IqW69AMiVCMiWi5ufKOSOEt3nz5tFjR37rW9/5yEc+Yj7isonlVA4OPGAwFDxQgbBSQeb/QkdKWBxFWLsM1b2VNv9bpRpl0nX+4ODgxRdfdOmlv0n4ifa2jnK5TDFBq5wWFQ07xSttBH59L5HN9XHuf+xj537xi18YM3Zs5Phy8zwVqQHHhCF0Yw+Y3GVyhl5++eUbrr/+uuuuf+7ZZ4UQHR0dqWRK60CqeHMnBRKjAv/6AubYuU0+eiDL2kwiYLymq6hG0qoCLQuklIlEeCnP5XK777H7Rz7ykQ988EOTJk10e14vRVJTM+yqhuI1XG96BF69B9xYIzOqpbId3Sau4KrHV69a9Ze//OXqq69Zu/bl1pa2VCodyHL4sUWp+trcswl8mEVDlAx2ctMVnW7XzmQ/J7Jzto7KlTPRgs8TMpC5gVy+kB8zeszChQv/6YwzjjzySNr56sBvpdDPLafciWlHfFTm/9uoHw3rEtzjgr/66qs/vvji6667jjOeackoGZjJ1DTLiro0oxWmomxFuFZLJBKD+YFsNnf2OR/7zne+NXLkyHK57Hue8XqjHXOtU7iOghg8UGXOBItqbripgkYR1i5OvG3GaDCphO+/8MILX/nKVx584IGRI0eISA1Yxw1mzHOu1EZp1dOzdc6cwy+88HtHHHmEvWjGamoM9IdyLaA1Brgm2vzEE0/86fI/3XLL/23csJFaYys9IUY2n9WkGRv59G6hIDxPyqBcKmumCvlCW1u75wtqx7OPHMIDdq1WUkrfFJH19fbtueeeZ59z9ofPPJNamUkCrKaWqrIP1bnKoab5NreQtnrT7q1RrN7zww9608aNf7z8v//f5Ve8vOblzq7OVDIRlJUX9ezGLhHWKAvO4tcQresNcKS6Gb1FJ2dl1FZMT3LkxNKqgHu8XAwGBvsz6bY5hx9+9jlnH3/88S0tLXT83UF2u2ElXkxkJRKwjDTHYrtGO0YRAN+ql9/1rW99+9lnl40dPVZrprS06pWaxvuzaDt2TAWFowXnGzZt2mfffX75y18efPDBDUdC1Rf0gV0cUsKqDCYTFXk5FGHtOsTaN2xuzhqn0LSYq9IVV/zlO9/5976e3u7uEUEQ0HAYKxNMLgFnHrPajzQvzvf8bDbrJ7zPfObTn//Cl9ra2qQMKEfmqNGSjM7DyrU8rqWQy+X+53/+5+abb3rogYd6e/s6u0KXt2ERcrxHKN4P415JMOZxoQUPgmAgPxgEQcpPTt196jHHHLN06dJly5al02ltspWxUQFV8pnxtxAEQX8uO3bM2HM/9rGzzzln7LixNS7v9oRt46hYZMBFCbY1a+mtUbUmcEbJOH3S6Fj6vtWy+OMf/3jllVe+tuG17o4uz/Nq9MJiLm0UetAVZed4YzSryEpzbUci2PJoGolau2+Rtfa80OnM5nKlIJg9a9Y555xz+umnt7e3D3HAq7umhOaKMyZDj0OFn795i1qZs16EZjh8s57v9fT0/PznP7vsd5dJJbu7u8olqZm0A//DhYLQlbdnFNrMB+T7fjaXTSQT3/3u9z760Y/G6hus5XWibHVxDfQq7aLYKuhIAIvbU4kpF7TBEK6djzqxSXuxigKR9vayDJJ+Il8o/PAHP7j0N79JJBItmZY6UUYVhXaZVRCKZJO3btk6++BZF373u3PnztNaKsXMJVtzXtO4Ui0o6FpoYqZ306ZNN9xww5///OfnnnmWc97a2kqCkXYLVcasrmO0Ol8b+jKeUErmBwq5wVx7a9ve++57xJFH/OPJp0ycNPEnP/npjTfc4JlVBck6RC3x7hJp4+uBLDMdes8DAznf8086+R++9KUv7b333pWhEdtRYzXUJ1JDsy/PtZpWrFLDa4U0aClmIhArV6789a9/dfVVf8kPFrpM9Tt94iYC4VF/uDmVOBOmxZHqybXt2BWcl8PjoxKJhFYsZsJ5dDqxuABLTTyGMhxSSS7EQDanpJyx18wzzzzz9NNPHz9+vJu8FG+njsq9aOFlR0ZFB5RHQWpNNd0k1Uln3aJFi77/vf948sknO7u7ksmkCqTLkkQRbU7nc/QGwu0U8sVCqXD22edceOGFnZ2dQRB4XDDBXd120z5G8O6DirBM5YgQrEYJx5T8mf57zbelvwve3VQp8ZvLbbgsM1VOL7749699/ev33XPfmNGjtSl9adj3YqSFtZThFdX3vVx20E8mPv7xj//LV7/a1t5uwnFOcJ8xKoiK2kmd3j95SMqoE7kaq7Vr117x5ytuuP76FS+taGttCx1TQzzquM0Foq1hpvsFF6VSsS/0yxPTp087/oTjT3zfifvus3drW/uG1zaee+65jzzy8KiRI50vpbVVOqTkptLao3i4SW8XC6V8oThv/lFf/cpXjzpq3tACXkP+8k5c11YrnTDWuN3rb39b/MMf/PDBBx5sMcR86LioQPXCwpwBoSFPJqUMtmzanM60tLe3mqFSkmYzcF67ZHE22Ja5USFyVA/ItB4s5AcGBnbffY9TTzv1jDM/PH2PaS7yb1eKNvxbNSnDav0Ze87IiaUYeFSQTSmPvt6+H130o8svvzzpJ9ItmXK5bKqjlWAeOS22sqDio2shfM/jGzZuOviQg374wx8ceugcGQTchOu5e2UeTRRpVH4Idh14WUthbS1F2kzYzRVhaRjgnZDYDAVTMhOpJdgLUuQB3HD9df/2b/+6YeOGUSNGS6lo0B5twY0rsJfM0P0NhJdQUvb29x5yyKH/+m//Nn/+fBeFi2JwypWVuoKrmI6VNpKQninOYs88s+zKK6/8643/8+qra1syrW1tbTVKTzUVyE4GgdqCohm92shPeUFQHsgNFsuFcePGHTX/6JP/8R+PXfAep/K45LHHzzvvE2vXvjJ61KjAFCozVw2mq5p96dLPhdff3zti5MgvffnLHzv33ITv2wEHNQuCN10/VRuiePs8YCcXau/iblXGdRSaIOOklLrs97//6U9/1tvT09HRaeMfsZYbF3CmgyjMwk5KmUqnTzvttFGjRl573XUrXlzOtGppbU9n0mbgk3KR7So/mD5UK+hht2nPP2POiqVSfzY7YkTXKe9//1lnnX3QQQe50oF6cTFX3mD2laYzclaJAdX2u99///3fu/C7S554oqur0/OE8cOFjrUhU5gwtqrUvu9v3bol3ZL+/vd/eM4554TrNSWdRge9fmWEJocB3kWhcYTxr41RwopWae/opTp4szQYMxuVp8hAJhKJbDb77//+b3/+05/a2ttb0ulSqcy4p7kZAKtMNi1quYmuKczz/P5sP9Psc5//3AVf+Uo6nXb1R9W1qVWzglytlpRaCKsf+cgjD//pT5ffestt2f5ce1urn/SkVKxa5r6+mqla81K7CqmBgYF8Pj9u3LgDD5x90skLjz32uIkTJ9DltVwupVLpO++84/zPnD8wMNAeOuuy4ZkelSZp3/fK5XJvX//7Tjj+e//xvekzZihltBs9jzfamZ2Ahtlrs1QKP9xVq1b96zf/9eabbho9erTv+4EMhFmTRUOV6PnWUSRzWC6XEsnkf/7i58cde9yD9z9w7XXXPvjAQ2tffTWTSbW3tUXGTMcD+FVV1lEbkDKCkUYZRAtPMM3KQXFgYLClpfW49x778Y+fN3fuXNp/UqqqFvKkpnURzV2OyVc10pHO5/M//+l//tcvf+H7Xmtbm5RlqszipJ3FuSvKo+3TGqUclLds2fLJT3zqwu9d2NraGq/ZrshV7iSnCXgzkBCHMpNtqmRcUIS1M1MpujWF71GuT2qV8Pynn376i1/84hNPPDF65ChzDVJim/EPrRUXbMvmnv0O2PcHP/jh/PlHM6aDQNJgYOfj2ioUJ3Dpakmton34sDvvvPN3v/vd3xY/GJRlJtOaSFjP0hbn6HiHEmucM+Whm849XiwUc9nBZCq57z77/cPJJx173LH77bcfPYRahjhjfiJx0//+76c++SnBeUtra5QjDs23c7+cBVJKJRKJvr6+TEvL177+tU9+8pM0iYguqbvICM+a2AMlhrXWl132u4t+dHFvT0/XiE5tRFe0FjaHZSbvksgjWSnPZ+WS6uvv/dlPf3buP/8zY2zVqlX33HPPTTfd9MijjxQGC22trelMWklVn0T3jJoGM6MPJal+2/0y5yHjni9KxdLgYN5P+EfNn//xj3/8+OOPd2syV5msXb0fk+H3gHtRstpW4ddIrdGnfN999/3rN7+x7OlnxowdKbgfOrWcdqrBgXLrhk1bNh155JG//K9LZuw5Q0rJohoKCinFq/zehk8QvKOwHrC2/gWPZOjp1IcS1s5JZAeNfq65Xrpq5+uvv/5r//K13t6ekSNGUj0LnRt2DHtMJYPOlkQikcvmCsXC2R8959vf/nZ3d7cMJBeuSqtqUIzJZdDQc6a5lkr6ns8YK5WLt912+x/++7LFDz0YlFl3d5cwu2TjkOYCbuQ0qCSYu4BNvB3IZm2ZLOSLA7nBCRMnHnf8caeccsrcuXNJVZgG/nATjzS+rH/jjTd++pOfSiT8TCbFVHjxVkobzf1GWlEJsXXz1v32P+CnP/vpIYccUiu12GAS/85DzCAZnQCtTHkm6ScrmsD/1FNLv/H1b/5t8UMjR4/QkruAfE0S11mmUqmUzWW/feGFX/riF92rLHliyc033XznbXc898LzqVSqrbXV8zw6E2y8mbvMAqcsLg2TFlF3k0tUl8vl/mw2mfQOO2zOJz75qZP+4SQK+MbkmhXn0RhBJlhVl3BleBct85SWTDPP87ds2fz97//gT3/6Uyadbmtto1Z4CgjF3WutbLWXEQWRvb19Y8eO/8lPf3LiiSdW0hkRuLrusthxhDQHVZAEn4YS1s5PRWKJMWkW+Fqpiy+++D9/9p/pVMoGkE2np2slsXoELvwrtO95vb2948aN//d///YHP/TBSBhBcG6zblF+117T3EaMHiE30cjy3Xfd8etfX/q3v/3N9/2O9vbwcmXifvF4sp0EYA0/QW6wMo1Q4Z4GUvb19QnB995n31NPff/JJ//DHtOm0+sGgayoLZpsnO/7f/3rXz9x3ifS6UQqnZZB4HGfU7IwJs9EZc+cs2KxnC8MfOCDH/rRjy7u6uqsKbbalb4msXmEla4uLmXg+4mB3MB3vvPtP/z3f7d1tKUT6UBSYzerkYMOF1ShDZdC8L6+7Hf/47vnf/b8fH4wlUqShtRrr712++23X3/99UufeDKXy3V2dqZSKWeGK5V1jMmoEkrEunzIhVVaJxK+0rqnt4dzPn/e/I99/NwTTjghmUzSpiqLNq2H0quqMcP0NkkF/Qf/8f1XX13XPaLbvBSFo5VpKRZRgEY7Fa6En8hls1Kpr37taxdccEHod8uAEi7ELnP+gCq4VIEpKGCRSIudLhIfX4oz412KjoV9a5WWI9kHFUg/4ff29Fzw5QuuufbasaNHC+FJRc1ppuqq7sJktOzDK15vb8/Rxxxz0UUXTZ8+QylJ2ny2et4mMoxLwSuqQ6QhRL7Rbbfd+pvf/OahBx70PK+9vcP5OnHDFtVY8VhJEg99cRWuDYyeh5cv5LO53Iju7hNOOGHhwpMWHLugpaXFrAYk5+R1mRp/uoaa93vzzTd/7KMfS6USmUxaSW3eqd2ruC8bXnATfqFYVFJ9/Rtf//znvxArK4tNi9plcOV7PNaw5FIJtMS5/I+XX/jd72gz8IBHgbSKELQxeSp0EGVohwLWn+3/+S9+/tGPfrRcLgnhUaifPvcHH3jwmmuuvuP2OzZv3tza2prJZCgvaxVzKS4S7YIgLY/q/ITpZBOasWxfthSUjjrqqE984pMLT1pIPqtT7a4sK7lT6oq5wjaZbac5UXfcqpWrvv71f7nrrjtGjOgWXkIGpEytualfVa702xyd8OCEb01u3LzxjA9/+Mc//kl3d3e5XDZJlkqUsb6jHezc8LIq00yY6ONHEdbOjwvqBkGQTCRWrFjxqU99asljj48aNSpe/1KjVGyH5XGWSib7+/t9P3HBV7/yuc99zvf9GkXACk4IwQSAeeSGPvjQQ//1i1/cddedjOkRXaM8Xwdlacw9Z0MrJMdEIrg2zaCFYr5ULO8xfdopp5zywQ9+cObMmeETzZXVtSNH1dqC7vE876677jrn7HM415lMSkoluE8z8oRXCSQqrbjmiaS3acuWSZMm//KS/zpq3vwgkELwqpkEu1wHSUyzpUZoIyp3SiQSDz/8yGc/+5l1a1/p7u4qFaWpHHbFXNpkM0QU1w3XeIP5wq9+8+vTTzvNjtc1n5kwQuGMsZUrXrrqqquuv/6GVatWtrS0pNNpOyA4OkEU1azE3IaaPWZRnXwulyuVS/OOmvflL19w9NFHu4al6hItt3B1gqBMVxcuUvJbSnnJf/3ioot/5Ht+a1t7uWwk3tzrauosZpUJVJx5nti8efPsgw759a9/tffee9cIZuEyu6vBA6W4ze1VqjhRhLVzsA0PWEcm6p577vns+Z/duGFjd2d3mWR+uYhbr7iqVHhJ8VhP79YZ02f87Kc/n3vUvNCuMs09j2srYipiL65MMlBrLlU54acYY/fec++vfnXJ4sX/f/a+A06KIm2/QvekzRlMJFEEUUFPz5zuFBTJIgZABSUKAmIACQomFJQkioKInyAIkjGDmP8GFJWgggQXWDbvzszuznRX1f83laZndr37Pm+X4M17HA67E3q6q+tNz/s8XxCb+JI83CkCLgPMWf0YjNIDO2rNIFYVEQAQDAYZY61bt+7T77ZOnToLlSFRu3aqwqkBFiaayhjj777b0uWGG2yb+pI8xKYQiUFRpCZQxQxOJGnByCwuKj6nfbv5C+Y3bdrUtsMYG4pUEEEHFWGD0lQd+xYDWeIAacMw9u/ff/ddd3355ee5OY0ZI2qeWDYOgGL84QjpiL8N2/br//M/V155pVbHips8LjxcuG7N2kWvLdqxYzvGRnJysmxnKGkNVcJj4mJwfENka+PKRJGICvG2B2V2IOiH0OjYsePw4cPbtWtXi6lbDQ4DoP13DGebQB4QKuKDzZs+um/M6L1792RnZ9u2zRQgQQhvMQSBgFJE8n7CKDRNs7yyIjU19bnnnuvUqZOo+ojhq0Rj+L/NoGDtkSXCBAjrr2XcMUQcHGBEUeZFrrJo+kIAXnpp3pRHp9i25UvyUhLj6mqL1iEMqE3LK8q7du/2xBNP5eXlWZalWfudVhsoK+QCn3vu2bVr1hJCsjKznNpBWjUByMIiipO1FXuT0NatCYcwxhdeeOFtt93WsWNHl8vl3EBrTwnrd8AYHzp0qHu37nv27E5JSY58X6BlARDXQeKaOeoMFJcUXnLxpS8vWJCXl6enmWOwEeLBf1cGXLdpwkjpBAnFBi4uLh48ePD777+fk5UtUj01OC6XIX8J4jETDAaDyUmpK1a+dWbbM50+GMQygdfU1Lz77tuLX3998+aPwzVWWma6yzBtjq9DnBgVMknspirSQJaFGWKQiCIyxtiy7WAgkJya0rVbt8GDBrVs2VIIM0OMgJJljh4AlCLDzslmGIkvpQLHgQMH7rvvvvXr1udkZytSMGdlmzdN5BsixrvE1aEq27YeePDBkSNHO2FZMSXoBHPWX904F7RSn9Z3UAKE9dew6KgFYJp7x7ZtwzBCodDESZNenDs3MyPTNI2wZfHaKqpbgpcBjJE/GMDYuP+B+4cOHSp8YcxQYyzjo5NIcs+ePXPnzln+5opA0J+akgKBwUTegkT3EEkcMVAts8guCVAkdoiGAqFQKBgMZmRkXHnVVf1u73PJJZcICKttR4OAP/K+4oeVlZW9et343fdbszOzaqqrODMEiHGdYmaGUyoWFRVf1+m6eS/OS0pOjpwxhAFKMAn+O2Oa11QGXtXV1cOGDVvx5vKcnBzKTYlZRfu0oicMEAgGg01OafbWqrcaN27sHE4TF1dniqLL8Mmnn7z6yoL33n8/GKhOS0vHBlIvQVSgtbW6uaQSVQ9EV4IBwzDC4XB5RXnjxifcfOvNdw3o36jRCZLODGMgibdYNMbSHlEhsxiNknSGw+GpU6fOnfM8QkiM/DpB8tRBSc3zdvmlSkpLBtw1YMqUxwTKTBJHQ81JGDOUfLSvbsLq3wQTlsh65VJJgLD+Gub0iIIekHHCC8MwCgoODRky5IP3PmjcqDGHJDPFGwUcBM9AYYwjvy0rKT/plJNnzJp52WWXRTmwatWHxUfrrLeoqHDB/AULFrxSWHgwIz3b5TJtwkF/gPAipP5E6OSmdHAXQi57HqqpqcnKzu7UqdOtt95y1tlni+FjQqgzAogrU0fJm5Q/6Ne379q1q0888eRwKFybHUnN8kZS7cLCw7163zxz5kyv10tsirCao9I7MHCQeSUyYGXMIbOvaTcIIffee++iVxfl5uQ4Wp5Qt0dVlMiwgQ8XFHbo2HHRa4tEUFXn5iNaCeKHW779ZuHCV9esWVNZUZGakubxemxCVLtWhXS8/wK1JAMUQleR/JUy5uJiwKVlpaeccvLgIUP79evnFNcCDt4sB4E01GN4Yu1pXN4777wz5r4xhwsKMjIyLNuGenRKF5aE940cFeVS2Kjg8MFrOnScM3sOLymFsYGhqFQpBPnRuZYJOyIm5oCjUqlCuzwhR3g8Wm0IZZRenw9LEkJNw9ixY+eAAXdu2/ZTblaeUGIHcntCwv1oHkYKbAQMCMHhwoIbOnV9ZtrTjRo31iTMzlWhP1pvjn5/5ZLFS+bNe3Hv3r0pyammyyA2iXJZcgCrqBgCPrsZ2Y41uEd1oKuqqvyBwMknn9zrpl639enTrGnT2pr2TnMej/6hiAYemfTIM08/3bhxnk0oipmScgK+Iou/uLi4d+/eM2bNdLlMvrHiuDdP2B9alGE5yrIiPNC9I+59hRfzw+Gwy+VyAqnVihUpLCwuLr5nxPBHH300SnLiWNVi4kdMGQHOUQUA+Pnn7QtfefWtFSsLDhekp6W73S4Ow+MYKBoDaXaUyeU4OqCRZyGIwqGaQDBwRps2Q4cN69q1q9vtjsKkNUjK+fI/KPns3bt3yODBn3/2eW5ens2nhDFGzhtRxW0yQjBNs6ikuGXL016e/3Kb1m1sO4yhAVBMMzixA/9VTaohRX8g4vwECOs4tzjqf91F27Rx49Bhw0pKitNSUiyLi9tH4i2CJB+9GOxAQqXc5TKqq0PBqsDIUSMffHCsQDsL5qPaZJAa1UUIWbp06bx5L/z0wzaPx+1L8lGin8a5LCmFMcO8TGmVi8ohFmjVmpqa5s2b33Rz7969e5988snOIuS/FjaoTee7YvmKAf37C4He2q5UnyJsoNLSsm7dus+eM9vj8ehOZGIH/HOmTyyHHZDRo0YLH8xb/hRKUmd91QCX4YgssOLSkrlz5/bu3bu2D44zXoxhHBwH9vz22/Lly95YsvTX3bvS0lKTvMm2bSt4XYzCtBB5A3Llye61yMiDwaBlk7+d/7d77rmnY8eOWmBDXH6nA66z1iLusmAwOGH8+Pkvz8/MzDRNk1ALQozEGB4PTQR5CFTxCjZgeXl5RkbmzFlzrr32GgVoQEDUuNXHHclrl7AjY1qOEOjxN75ClRxhAoR1/FgcNzJlMvgX6FOE4JLXF4++fwyzCZfmjcG5QC1J6BDAqampSU5OfvzJJ3r06MH7bvH0PbXbvR9+8OH0adM+//xzn88nPiV28aixJKX1xxxlXMT1hauC1VVVVW3ObNP75pt79eqVk5NTJ8aqrvw1/lSI1bt169bON9xgW7bH46n9fE35izGuDPivuuqqBQsW+Hw+kY3pFC2x/v+EOT2TeDB0yJClS5dmZWcQS0KIIYhWI2SXFEYWnmGaa9evO7NNG40Ndl6CeKpI7vpENlxw6OCi1xYtW7ps96+/+VKSkx2K/aIhG+PHBUyaA7QUm0bkUyorKyGEV/3j6tGjR5977rnaDUv+yCjvlUP6yXE7iDvrlfkLJk+eHAqFkpI9xAYYG1qvkwq6U45BA5gCwgxsBKuDFgGTJ0+5667+fwjLSthfyyChVKmfOPfWBAjr+LN4/BHPKyM7RSTENmbOfG7yI1OSkrym6eYE4OCPMMOMUdNlFBcXtzyt1QsvvHD22WfH4q1iKoe63bvz553PTH1mzZrVCKKUlJTa1IOxGYz2xBL5DAGuqg5aVvi0lq1uue3WXr16ZWVlxQn81WYq+KMMWIcFfr+/Z4+e323ZkpaWVuewh/gKLpfrcFHhxRdf9Prri9PT0zV3v/rKiVvgT5qTThlCaFv27X37vf3u243yGoWtEK+14ZjeP9+MEELl5eVntm27avWq5ORkCRMQRWMm9Lu0rr9mzwC8YkwFU9WBA/tfW/Q/byx948DvB3xJPo/bLXSHRQs/2qVmVPEdANUBEVSAkX+UV1R6fd6bb7552LBhugbjmFZycJIolFZcSPr5Z5/fO3zEnr2/ZWfn2DZhjEDEAMBi+AoCJQTBDwchZNlWRWVgzP1jHnrwQcqb2XF0nolF+BczPGHiBDX3K1S5BFEN0E2I2uFnwo5NqyM7jNzWkQs6buzYqU9NzchMxxhRyng+7Bg0ghJ1LOY1IYKFhYU3dO6yYMErLVq04C6Qk1IKZTolo6BnNCsqKqZPm37viHu3fv99Wlqa1+vVhFYOLCtQGYM6WImEjexT1dXV/kCgxaktxoy578knn7ro4ot9Pp9+k9p14NrwnJieNzexAz4yadLKt97Ky80VE1O1V7LwtaVlZe3atVu0aGF2do6jNgD1bptY/3/C4vQiGWOmaV79j398/dXXe/fudZkmlaqRQKkA88IbbwZ7PJ5fd+2qqan55zX/lFdEkn9IoSWh0yyxVaJjxud8Rb6bnp5xyaWXdu3Sxevx7vx5R2FRkWGabpeLMiJZJ7lJWDLUvEPyH5D3mJOSkgBjn3326do1a21it27d2uPxMM6iCbHIhiWgC8D43VKMGzRt2rTTDZ327N3z00/bInEqFLgHgW1GYlYKQEn/yzWdoNvtev/d90pLSq66+h9C89G5yCUtVwL791cxJUeo2IXkBLn4HQCJ8P94sTjf45yOuOee4UsWL87NzWHMjtzkALOosK8sViMBgcHQsuzKior7xtz34EPjEIIKtipUWygHagHKifaEP1u9etVTTz71048/paenu1wujVuJWzDRmorAp0aOkBqmGQ6Hy0srW51x+oC77+7Zs2dqamq03PfHGKs/OgNOxi6M8YYNG/r16SveU2xkzrxWpylV1VWNGzd+c/nyZs2aiVw/RsEpYX/WapcrxJrcv29fj+49CgoKfD6vTSxOhgWdo2iURp7pdrtLSkoWvvZqp+s7SR8MlQCvE4BeW7SYSdJIbESW6IH8/S+/PP+NxW8UlxRx6mZAbAd+XpIeAInJVnhpTqJBIim1aYSqw5WByjPbtr1vzH1dOncBAOgJeKoUjVDtMoxiHaeUPDb5sWefnZ6dncUYpMxC0JAiJ/zzIzcgb0lTwhCOWEHB4W7de8yZMzspKcmyLFFkigaaf2nlj/8q4w5Yk8gkQFjHq0Ubt8Js2zZNs7S0dNiwYevXrc/NybJsi2cIAAEMAaISpgq0UptogpqmOfXpp2+5+eY6u1DOmvP3330/bfq09evWuUxXSkqKZVlOFqpasuqMx/6UP8aGCSNZr99/yilN7rijf5++t2VmioKzhZChJ4CdyVNt7xtXPHcCryCEpaWlV1xxRWFBQVpaugZOIwTlEIjywIQQQsmqNSvPO/d8Z188EXHWi8VBhQEEQnL/m6++6dGzG4xEdlhpE3HC70g0KAvLGONAIJCTm/vOu+80atSIC3gghwMW6jGShCy69mM9NKFUyHz9+ssvc+fOWblyVWVlRUZGBoKGiCPjiMdFHqJGcJkeY3O7XMFgsDoc6tChw9iHx51xeismyidC/Fi/tjYRpgInLpg/f+zYBxHGST6f5PzlqTOvSFGO7xaEdZFvhRAqKi65+OKL582bd+KJJ8ZRvYqcHyYy4OPf8MQJE50gLCb2n0gu7NRLT2TAx65pVl49TSvYefbt23/77f0+/OC9RnmN+c8xv4eR1mjV11SUYcvLy3JzsxcsWHBDpxssKxzx00hy9wjhBjH9y8PzgqeefOr+MWN27tiRnZ3tcrmEe+M5AUKyvB1XuRWpCcDYCBM7EAhmZGb1H3DXtGnT/nnNP0XVmk83QbktyRnIaEGvzu8O69r7RDDxyKSJH77/YVZ2pm1TQTcJoZRPl3KIAODINlc85bHHOnfuonPfRNWnviwuA+arNOJmCSEnnXxSWmra6pVrklOSdUkCiGSWRWlQvV7vwQMHS8vLOnXqJGbVoIrGoPR88u+oAbVa+F9IZd45OTnXduh49VVXh0Khn37a5vcHXF6XlGSIrA4cFZeA0bfTTZBIROtyeTzubdt+XL1yTXV1ddszz/R5vUwRbDq/uHPxi/VMiH3ueee1bNnyo42bA0G/1+PlCl1Q1bCjRy0p4RhMTUndsX37Z59/euWVV2VmZkowNmAIIFmeUVPJiWL08WtSjtBJypdgwjqWrXYFDkYRpIABSggzDOOHrT/cdfeAPbv3ZGSkh8MWZ8NncYzFfCJIsueXlpSedXbbeS+/3LJlS17ywhwaQlm0qMsMA9s2Wbr0jWenT9/z2570jFQIsFO5gbt5LuTBEAPx48iMMZuQqqqq9LS0Lt26Dh48uEWLFirrxSjKMxUPdv3fwE8c3tfG2Pz448039erp8fgQkiLEgoxBKCMpVQajtKSkZ69ez8+dwwMIcQxHe6nXeYH/EqZF/QCgtk1N07x3xPBFr76Wk5tlhUl01JbPA+lzwCAI+P3z58/vdMMNMdB9KVBEFY46KioIWCxkOrIAuIYSA6Io/dmnn8ycMevDje9jZKampdnEhhRgw1CzAOoDQByfGuAgLyMUrq6sCJx9zjmjRo+64YYbdNMkDp1Qi53GNgzzu+++u+uuAfv27s9ITwtbYSTDYuFPKYOcMivyOiRG/ir8FSef3GT+/JfPOeccgcZwIMAch/lXWST/bQYtRjDTumowwYR13JhixXNqwIibfMu3W2677bbi4qL0tHRRB45KBArIOxP0PTaAzMDm4cOHr+3Y4fk5c3I4WEk3nLhB7tIiP/nxxx8fmzLl3Xff9fq8yb5kzmYFeLmXxfIciBIZ5Tz0kNfTIIQ44Pcbptnphk6DBw1qe9ZZdTPg83dwkoH8i7XnLD6LsrMYxayqqr6uY8dff/k5OSWFKV04HmQirqEg4Nm4vLyi1Rmnr123Pi0tjVKi9XCO9GqP87jRH2tu9j//3nV+ndqt2SP0fZW+kBjfRQhVVlZ07tR558870tLSI94lso6hgCXJBQsAwigYCObm5W3ctDE9Pd3ZFlGKmrFDbiwqHRhFLakbJRJIAtkoWbd29cwZM7/55luvz+fz+sTYbSz0JaayIhIUwhhEkbumsrIyFK7p1r37/WPGnHZ6K61T6ajHRIeUxLcnxDZNMz8//64BA7744vPs7CzCB5Aka3vkiKnqaCPehYmEBWVlZZmZGa8uWnTBBRdYvE6DAKRQKDFKZu2EBz5ODU+cOBHEMGFpHULgHABPOOBjxFjsBA+TAKfIA2JHctnPPvm0b98+FRUVqSkpimWFAVFhkzUOsSlEIi2AYGlJ6e133P78C3PT0tJsyzJM06HeKsSFjPKy8udmzBw1auQvO3/Oys5CCMoeLU+QIYqW/kTRTs+TY55XVvorglXV1/zzmunTp9199915eXm2HRaVRoSc8bx4F+aoOv9h8Rk4yo7ORBljY+7zc5csXpyZlU1sJpuGDHOmLaYnP22LYBPPf+XlFi1OpZT8Qdm8gS+lcIRQ0TdKJJEGYkCn79U8D3GRRxymTMVbdRdF45zukb61o1dKV5h955xz9ptvrmCUaXoyQcrGgcFIxB8et+fAgQMQoiuuvEJzlilkX8xoOtCZIYxuXExRQCvYEyTcDbdqdcaNN/Y6pekp27Zvy8/Pd7lMwzApFUPzWgCJxby7QmETQlxul8/r/WHrd8tXLLcJbdu2rcfjcZz/eAgNhHIuID09vWu3rvv27fv666/TIjFiJLHmQwFAoR2RJKfjUbLP6wsE/G+tXNmmdduWLU/lN7WMKWB0FiqxPx+XhidOmAgcXRO+9PktErsLHu3jTBiQm2xsmUyXoAmNeN/33nnv9jv7VldVpyanWralsMQCxiyRJUJJD2Ns21ZFeeX9Dz4w5bHHTGwIBRugyBAE2AohtGLFihHDR6xdvdZwmclJSSCyV2LeXVPyvcJVxFbsGIPYgIGqYEVF5bnt2j/+xBMPPvTQiSeeSClllCJs1Br1UW8mutTgf4VAjnItqST499/3jR41CnEgs9jT5Z4JGW9pyz50cXHxuPEP9+jeMx7ecgQzQjXHqmNfUZ3gUoeKsEmcUnEdkcNgLav9W6ef1g/qnKU+Ujc4c8DLgRjUOfGkkzxez7q161JSU8TgWcT7KV407pEAZczlcm359ttrrr02Ly+PUktRMnPXBVXvVoIVaoHnnSmsY2ejlLpM19lnn9O9e/f09NRtP+08fPiw2+1Rw+LIOaqu+Jx131kqePl8vnDYev+99z/evLlJkybNmjfn70z4y6nzePT7UMrcLtf113eybfLeO++ZLuxyeVS3RwxGCRAsL55jEHm+221ZNW+9tbJZ8+Zt2rQhlESWc2RBy3s6sUsfpxbLhCWVUJUaNUgAUo45i6anavKB8tTHwHjVylVDhw1l1PJ6kxkFepuLuXYwcvNijGpC1YSQKY89cccddwgBe4yh6I9qnPOOHTuenvr06jWrTMNITUnT3s5xOI59HnGRV/4UbKBQOOyvCLQ8reXAQYP63Hqb4TIjR0oBRGIEOEq/rKrNQAu+iXf+t4U1zfUhXK847JEjR76y4OXcnFzblkVp2feFROxTBsbFpcWXXX7Z0qXLeatbch0ceYckHSGNatpD3rHm7NxQiNgLs207Pz+/oKDg0KFDBw8eLCwsrKioqKmpEa1H0zQ9Hk9aWlp2TnbjRnmNG5/Y+ISImY5Wgh6qjuusHxWeOyYV9KVCAWO0d++bN278MDM9nRAmewGyvswXOWclKK+ouP76TvMXLHBcJiFaEAOHqGvNUKXazCdvmU4aoZiaExFY/v7f586du2TJ4srKyoyMTA4Wi4qOCA1oMeEkC4SQF4yBWNDMHwgABHvffNOYMQ80btxYDxHoWEJ9LgEAC/kTCOFLL7085dFHKaU+X5JtWwhLQUzqIPgQN4mBjZpQdU04PHXq0/369SXERghL1q2oclJioz7OLF6OsI5nJBzwsWHiJo4bdaAMEGq7DHPFihX3DB3m4hYOh1VWp9NKqsvCCKJgMJCUnDRr1vPXduigJhoREFRCnOA+GAzOmzdvzpxZ/spAWmqaHuOJOyTndC9P3yhCmFK7rKw8Kzu7/50D+g/on52dAyiwqYWxoSO8+lpXTuQzQujrr7/q3q2bx+1WLJxRCLf8TASJZVNA1m/YcOaZbVWmcuTyhzrVEvVXYIAZyBBHXun3796967NPP/vxhx/379v3e35+RUVFMBCwVeOcT+1I+A5/aeQ9DRP5vEnp6RmNGzdu2bJl6zPbnH/++a1atRJ8UrpGreU0js4NLprBnBeZsciC3PHzjuuvu45Ytmm6gM5mBaQPSY9tuozS4rLFbyz5xz/+IdwPFDgD+G9jNaadbnSpqo4bgAKfxTCPV77+6qtnn5u28cNNDIC01NTImua5KYsyZgEaDRMlYhkAgrFh23ZpeWmL5i1HjRzZq/dNgjtdT7SrKEtCHDRx9McfbR4yZGhxcWFaWiohkgNY1z8UupJfXIRD4XB1deiJJ5+4s/8dlhUWtF/CElv08WhqDpjBKKZdLVioWMsTDvhoWZ3Mi9Ff8nzWxMYbbywdNnRISnIyxiajBCAh+i3RoQKBwodrKYTIHwg0btRowYKF7c8917JqDMMlbnFCJJvu5o8+euSRR7/99puszAzOW0lieaagowoOFa8G46qFqKKy0uUyu3fvMWTokFatzhAgZ4wMqPgD63Gj0JVnnQ3feuut7737dm5OnmURiICjpcpkZRuxosOF4yeMH3P//WruKBbN1vDeKJqAqvaBmAAVMVNVVdVHEfv4k48/3rN3d6gmbBqGaZoYY9M0DcOQdXUImOOtlI4UJ5mwqG0Ry7Ysy7KJ7fX6Tjut5SWXXnr5FZdffNHFXq9XpNRHq/AOorV12bQQpYsZM2dOHD8hJztbVOMpR4OqeW3IEHGZ5uHDRZdfccXSpUuVWCEDDkqZ/xMgmOl6i7OJTiV3x+rVq6Y9M33H9m0+X5Lb7aGUiHRbcGRR/t+IH0YAqbxa5PQuwwwEg4Fg4J//vOaBhx5s37491/knCCPgkLTU0YCY19++bXu/fn3zf/89NTXNJhYHgkMhmChH5gSOnzKEMaWkpKzsiSeeHDRoILGJQGDEYtMSdtxYDBOWInsDcSCshAM+Fiw6cSRHFhilxDTMJW8sGXXvSLfbY3D9cI42Es1FQWurk2fgMs2S0uJWrVoveOWVU09tyXdhJIq3PPE1CwsPT5361P8s+h/GgECcalCJMBhFfsQEA6ZpVlX5A8HqK668avR9oy668CJZ9uRCb3JAiWkV3fpZS/FCT5s2de/RPTU52TTMKB+2qrLyB0ZFZdnpp7d6++31SckpsQBDpgqVDYspjWHtB4By3KwoLezevfuNN9546623du362QrTZF+SN8ljYFPLT2mk9x+8Nb93+RJAInXkaW44HPb7/aFQyO31nHrqqV26dOnVq1fz5s21Gz4aN7hYmtzR8gVGKQ2Hw106d9n6/fcZGRlcNFqiUgSwmXFRapvYfn/w9ddfv/baa51CSf+bjwNRpw8l5aSaWGJRjxgJQyPnDeOqYHDhKwtnz55VWFSUySvSlmWroV15JUV2DplipVR5NjJAZWXQZZoDBw0ads89qampdthCBoIwSvOir6Pwwfv37e9z223btm3Lzs4Ohy2MZAOQyXMlPogCyVRIysoqHp00edjwYVYkwMVxA/FHqcefsP+z8RJ07PSDhAuCBBPW0bc6lBL435R3fk1sLFy48IEHxvg8Poww5cBmDvtEgFGu7iuxlJQRhGBRYclll1/20ksv5TVqJIZ9I1E1lR3fZcuWPfHYY3v27MvKyoCQEQKxgQTXkCCuVWhSeRhctD+SWIdqQn5/sMWpzcaMeaDXTTdBKBQ+GOAibkgmNAyJLa9e15TwSYJ5o0ePHh9+8GFWZiaU1EgxwjecRwn6A4FFixZ17NhRDFk6T3J9HtYfWBxdKKNUlD23b98+/6WXV65eVVBQkJKU7E3yYGgAxixiMRoJbnTxv05uMmEqq6ZAFPoZ4v1OohV4CCHVNTVV1VWNGjXq0aPHyJEjGzVqVEsU68gU4rVDZBqNZRjGx5s/vqlXLy9ntxA0yJE1I1r8QFCHk0Cw+pJLLlm+fLlzZkwVAf4dP6PuBQP1QH11MRQguNDVNLkBANq7d+/sWTMWv/6GTcKpKWkarS/iAqaQX5Bn7Dwv5jcMZAjiUCjk9wfant124oQJV//jHxFfSwiOPdtSo8K2TNNVUFDQr0/fL7/8Mi+vEaEWbzEgMZFFkVjQUMz+ARTZAUpLyh9//IkhQ4bYtu2kWU043ePI4uUIFfghIUd4TFiMA444MCrCYVF5Xrhw4f333+fzJJmmK7J3QK5uJttVlEGpx0Ap8Ho8BYcLrrry6nkvv5SRkSG2XTGRiZBx4ED+Y49NefON5W6PJynJJ+5n8YGClU/hUFhkK4CQcR5LjDEf5az0eDy9et80YsQIgT0R7Vjnt4ibFK5H0yCaLz7/vFvXbsnJyWp+M/LV+AKOfmwwWH35lZctXrxYH5Wjffh/AH/9CZPlYVXJ1zC37du3z549e92atdXVNampqQJ4FUkKWSSX5ThYPpAjO9lRi5OoixtXVV8pJorWbNjhcKi8ouzEE0+a9+LLF158oZI64OjrI2hODhnNXzZo0KBly5bkZjey7DCHG2MmNKrFSyLXGvgDVcveXHbppZfxvUuLBP4vDv5fkZxorkxZJxZ4CJFkb9y48bnpz3722Wc+bhaxBcdW5EaTd0iUjc7xXpEworysAmN8x513jBo9OjMz0yms6bzBxZIoLCzsP2DAp5s/zsrKEik+EnMJYjiaT7Xzl0WiSduilYHKaU9Pu/3OO8STndt1whkfF4YnTJoIpedVRQzByqog0InWwtEyKXCmhhdFr5ICRgk1DePVha8+8MADKSlJGGNCbCS5rpTMt6hV8STD7XIfPlx0ww03vLxgfmpqqvCvYr+DEK1YsWLwoMGfffpZWno6lywl/GILD4p06xYqzj8mae4hsWllZeVFF144Y+bMfrf3S0lJiYGc1JouqueTE6uQ+ujkydt++ikpKSnaTJFEmiKJYGHbTk5KmjHjuUY8SnCCw8Vhqj/1f7Bq3hdqcmHDMCoqKqZNmzb2wYe2fPut22N6vF4x/SJ0iGWhk6n5MejkWoyZKXL6YDllFP2b1aa0E+VWny+psOjwF1982blL15SUZAei+Mjd5rrD5fzhGWecsW7NOssOcXkDrJpjUGMAIcIBfyBUXdO5c2egpsb1jfJvLp2+vHVc59izxyBCEd8uujMtWrTo3qNHbm7e1u+/P1xw2ON2Y4QjrlfpM4nt0zlgLb6Xbdkut9vlNj75+OMPP/ggNzevVatWoumjP1jOEPDPSklJ6dK5847tO77bsiWZz/EjiOU0g2Yq5JxuhDPMuN3m2xs2ZGfnnHvuubomHzOjmNi9j23D4ydNhNFJMgABlTxEMH587Wgf6n+dQdnwhUw250Xfl3vfRYtG3ntvSnISRgavdymXKSG9DKjumulyFRUWd+rU6fm5c5OSkgix+W6OMMa7d+2+f8x906dND4VCaelpAu0ZHQmvtR0LvC3m/qGkpCQ7J2fChIlTHnvspFNOVl3kqGZLXABej0ml3uP0R+zcsWPSxIkej1tMw/KfI020CxhE2CgtKx00aOBNvXtrhcSGXthxfIRQJXkIoffee+/uu+9eueItAxspKSl8voUD1RGMgtuYYnOI7fnG+eA6WsLRfD7ew6iXRI4tJTVl3759Z7Y9s23btk4KpwY6G39kcXRm2dnZ/kr/+x+8n5KSqk6FJCrRp9E0zV2/7upwXcecnBwpdQA1Tca//rB/lQHLNR/LPKYnhk3TbH9u+86dOweC/h9/2FpTU+P1eqN8qzB+pkuvUj6ty5KTfAcPHly1alV+fn7btm3T09NjGwo8KuADdW63u1OnTnv37f3qq69SU1NlKcMx7Cyfz2950zAghmvXrslr1Pjcc9vr9FpwbmvMZCIbPmYNT5w4ETpp/xJMWMeMOcun4mYlhLhMc8mSN0aNujclNQUjJOHqDCk2FeZ8KUa4pLik5403zn3xBY/Hw90kMQwXpXThwgXDhgz79tuvMzOyeO5lQ6hZO/TmRJnc+8QMDzVNo7KiAkB8Z/8Bs2bPvvTSS3Ufq64hJeiE6daXxVVfEUIzpj/36eefJiX5JEhFTl/q/QeGw+HcvLynn346JSUFxBJTNNyWFOMdISC8zen3+ydNmjRu3LiKsnLhPwSYRx4GrJ2bxWRUcdRX+oOcH6rdVZ15s+OOxlVVVWeddfZFF190FG/zaOKuvk7zFs03rN8Q8Ac47bGGNukjp1w4pCI5NfnKK69Uubt4r5g6cF0fpv6uKwNWbxEfmemzRwhJz0jv2LHjRRf9/Zdffvn5559dpst0GVxWAcRFMLHHAAklbo/b6/N9+cX/W7tubWpq6plnnimyXrlaBZgaShmozp07l5SWfLTpo+Rkn5j3ZVBHlLIojSBgHBftcpvr16/Py2vUvn07njRz+hLVPDo6rGcJ+9+ZVENyXB7J4ZZgwjrqxgScmTeABFrKZZrLl60YPnyY1+M2DZNQChFFwBAuU1Z8he9kzGW6ioqK+vbt+9yM5wzDEJJHhmH++suvw+8ZNmfObIhAamoaIRy8Fe0XyuElvq9BB2MDrKquqvQHLrviiudmzOjXr19KSoqmd6hzhTTQ4nHKviKEDhcWjn/4YZvYnCKXiaYZUBsnA5FvXVZeNnDgwE6dOv3BNHODLG9nps4oMwxj69att99++7q1azPS0rWElONgJVWUEKQTAo5xNWR9tkU3V03jREvQdfb/6gROGwaurKzs0rVL+/bt/yiEOjLmRJYJpsaAP/D+B+9GIirRVAHOFBMySiHGBYcO9ezR3evziQgGIgFKrp+rWfuUijMvMAenNGl64429crKyt2zZUlBQ4PV6TdOMFd+Mj4rkuDBlSclJ/kr/2rXrtv74Q5s2bXJzcyXjlb731As7XNuBEnvzx5s9bo+YPgKCxoaLEMsJdySY3ZBpmmvXrD3hxJPOOeccKqYPQJTJq068XsKOBcMTJ04SAyISRCgnJxNyhEfZFEuAZCGglLoMY83qtUOHDTExdLk8amRMDLHoHh5VYxWopLh0yLAhU595WiiruFwuCOHChQuGDBryw49bs7IyETSBIutBiEFBZQWAszbHq6aQAVpSXHLiSSdOnDjpscceP+mkk2jEb0fVf+PakEcmdBO5wtKly5YuXZKRnsHpvCRNZtS3QWaF7YzM9KnPTE1NTdOIMFRXali/x+Z8gDFetWpV/zv75+//PTsrG0T3WYFuEw1rIZks+j/ImaCJwjVjzOYWqqmpqqqqqY78LxQKhcNhy7Js247Lw/7oKojzFgqF3B5z7LhxOTm50Q7r0bjN42q2AIBmzZqtXrW2ujpUp8QQA8Dr8eTn55/W8rSzzj6bEKoUheutpFFnHKNNVKTP+9t513fqVBMKffftlqqqaq/Xo0nZ6j4AKGYCqOnCPp/npx9/WLnyLQRBu3btDdOorUXNALviyitTUlPfefsdhKBpGPxfUhIAAIAASURBVGqKXgUkQMotAgbEuNr6DetPOeWUtmedFUmspbxY7cpHYic/hgxPnDRRQgl0yTkBwjra5mgcQoErdhnGBx98cPfdd5mG4fG6bJvwMpPoBMlylKw/M4qwUVJSds/w4ZOnTLZtmxDL5XIfOHDg3hH3PPfcswgjAcXiXSchASTcElIKMIo5DzHDMAPBAICg9823znl+9qWXXS5VXxzSuX+UaTVQ6KY/V+x34x8eX1x82OV2C/ZNzlOBFGpbcoP07duva7duGp4NG1jpy9kIFH3f2bNnj7nvPoywz+ezLEv9FiAk+n9IU+sLmFuUrwoDQmggEAgGApRSj9ebk5d7WqvTz2l3zjnt27U966zTTmt52mmn5zVu5PP5CCHhcPhfdIhF2YADocMlJSX3P/hAly5ddfp71ENt8emEkLT0tKLCok8+/sTr88aloSoZhZYdLq+ovLFXL4SxGp6sNx9TZwQTNwNGCMnMzOzQoUP79u337t29Z88ehJDpMjk6Kqb9ryEHEmfN5ci8Xl/YCm94e8PX/+/rNm3aND4hyl6pP5EQcsEFFyQnJX34wQemyzCwGcluY0lmdXmA94PZhrc3nH76Ga1atYpismILLYls6pgyaDMbcQpfCSKEfB6OCgBsonV/RK1OlRsxp//pZ5/3vulGBKHL9FA+rYh4JI6QkCgXE0O8NEVpcUnZ6DFjxo8fr6ddN2xYN3bs2L17fuPpDicbVirgkvpdYD2QED1lopZLCPH7A2e0PmPcw2OvuaZDndQNR4VCWXNPfvXVV5073ZCcnCRjSCCYhCFjVLhnUXhfu25tqzPO0PCrhlCPiRFmjuKNIzZu3LjZs2bl5uYxXsB0nCuhuic8H9Xsc5QRTvwbqqqqYoykp2eeffY55513bus2bVqedtrJJ52Umppa5zFMGD/+hbkvCPF2R+8AOUd9IIJFRcVNmjQZM2ZMn759dCW/fs/GnzLGNQohIRZCxi+//Nzh2msZBZyv2+mAJWMGwqiionLDhvXtzz1PhlZR8o6GolKJWe28syvqE1Y4NGvW7FmzZgaDwfS0dK5vJF4RnT5Tx+7AyCGIMSouLk5NTR01evSgQYMNwxDUsNoHiwBu7vPPT5w4PjkpVRFlq6/K2Tj5fCKDkXOCQ+GQZZHFSxZffvnllmULtvM4S2zmx44hGRUDKdclGdAh+EPCnYQ1gDll7Z2eWHjf77//rn//OznfE584imxTghsIRa4cFZQTTOQxJWXS+4bDYcMwAn7/2Ice6NunT3FRUW5uI6EwGMmxmOgyCpw1VIyVQJQxTcMsrygjlI0cOXLdunXXXNOBEFs4s9qNxiNeI4lmeGvWrK2pqcHY4FmFlDoWrVSxKVdVVV122WVntD6DUiLaw0zL8dT7McV6X2GjRo2aMWNGnsP7OhIs/XSR7zIIBQEFKSouTk5O7nnjjfPnL9y4adPSN5eNHTeua9eubVq3TklJEYGFNJtYlkUpLS8v/+DDD03T/BcoM4xxRWXlrbfdun7Dhj59+1BHQFCX0sYRNS30xSMn+/TTW1199VV+f0BIODsh3LJ4AHFNdc26dWv1D2HDf4WYeg9kgsUscl8Y5qjRo9dt2HDFFZeXlBZZdhgbDpFOhW5lADh0EgEj1A7b2dlZELBxY8f27Nlz586dup0sDCFkWdbgIUMmPfJIRXmFallEGxwQKF4vGEnK3aYHAtC/f/+vv/7aNA3RmHAefMOdnIT9CUPAiVzg1R0RsSVAWEfStN917oOUUNM0t2/f3q9f38ry8iSfj1GGkCG3JKYQm5BQEW8zcLigaPR9940fP55xBbfNmz/q0uWGuXPnpqWleTw+27K5l0J8xhECgLlmPmKS8T3ilQm1q6qDxcUlF1100eo1qx8aNzYlNdW2LcwFBI866yz/dJm0BQKBjRs/FGgmGLsxCRPZwy233iI3Z40KbogDUx5VX0SM8diHHnrxxRcbNWoUJ9MbW7cHPB4ilJGq6qrCwsKMtMyHx49ft379Cy+80LVb1yZNmpic1j/ibW1bvAohJBrDCEsJwm+/+Xbn9h3Jycl6xlR/UR1hF5UU33333c8///zJJ58ksuT/2/lQagRReYv6dHZQNQ7EggS33nobxggqecq4u4Mx5na7P/hgUyAQwJizjXLeFVVaaECTF1FxcIn7wrbtNq3bvLF02YwZM7web1FRoRAWY4DDkplUUnFST4s+SuRCIJSXl/vZJx93uu761157TciASmkUAASCcsiQYRMfmXS44KCkToqeCl4+4dUvBqhlh70eX1Uw0K9f3x9//NHpzmMSrYQdG4agap+IEW9RjebsSbS2AHjCGsji8w8ICSXYwHv37Onbt2/BoUNpacmRDBi7OARSygbwRICKjR8jo6ioZOToex9+eILg9J88ecotN92yY/vPebmNOTSJIgz5ho2lemzk2lPedJCVZ4RARXlFWmr6E1OfWPHW6rPPPlvs+Ejw8UDdnD4ap0iemGiLdMuWb3Zs/8mX5KGUAC73Gwc5CQQDbdq2vvCii5yM0A2xnqUYhVZ+4B83a9asOc/PPqFRHqAxnQVRZxY5ET/tkR8Y2PRXBtNS00eNGvX2u++MHj26efPmhBBBzSEatwghwzBiXKZD5ee999+zLEs7KgePN6SMYoxKS0svvfTSCRMmaN4lZxwQDar+hZfVuLEG4CwR9Bfim3FNLfr3Cy9q166dPxBJ+5jjHHKL+KfkFN+uX3du3bpV8XzpxsIRWaAxZULI+XAIhPi2PrevXbu+a5duxcUloVAII0wiR08UgMwpaiLKNggCTGyQmZltk/DIe4f3798/Pz9fIhxVWYVSOnzE8IfGjS8vL2ec6lXUeATiEgIs/mBsMGCnpqYUFxb173+neB8ZkzHBUfpv+UoSduQMCX0PzttHNdc+p3c7aqjI/zZzqsrLn0R8rXHw0MHb77h97549ycnJVlgwhlIOyxKkzPKWxshwuz2FRYWj7hs1adKj3DNt6dWz54znprk8JtcZtXn6hyDECl/Jy55Q+QtgI8wsyyorK+t5443vvPfeXQPuFkG9bkdF5U3B0eEIh9FzJR3G22+/jSASehJO3h+gVJJsy77u+uuTkpLisKkNgQuDCkMl0u4NGzY8NuWxvJw8jhZ3ItSgbPlE3DLjYgyR1wUDgRt73rhhw4aHJ4xv1LixQDvHEYbE0HpovBwnbbJta8uWb30+nxZscCZJCMFQTU1aWvrTTz/tdrv1xh030wW01McfeVntj2n9Z8CqtiFZJiilXq+30/XXh2ospwCCunzigKFNyLvvvKNPjjzJ7AgtUAcymimFR2rb9qktWy5YuHDu3LknND6hpLgE4EgWy6jzhSAeZc0hVxibaWlpa1a/dcMNnZavWKFL3GINWMR+6KGHRo8eXVxUTEWRjBGoSAyldjKjEDHLslNT0/b8tvvOO+8oKioWxNoMSjFQFo2tEnaUDVHJGCyhDYp6ECaUGI6AMZUBOd2DSHf8fv+ggQO/+25LZkYmjcrri9uG6laQqKoePHBw6LB7JkyYCAB4ad6LN/bo+c0332RmZmGMeegt7jx+qzKkplzEFgw5HTT2+4OGYTz77Ix5L710yikn8/okEEmSo3B6NG9apnM7ENnp/AH/xg83JSenIsjHM7hfVr+HhNJQKJSbk9Ota7d4ht6GyYBFGYlXKfDOnTuHDhlsRGIXJDyuEqRiukEMId9wTVxZWWGY5vTnnn3+xblNmzezbVu4cCcSxxkxxOkpAZ5t79+/v6CgwO12O1wpUJPFDCNcVlbevXu3008/3Tnu4uwm1JEBx3pZ7aGdDUimr0n9nNXoLSBChBu6dEnPyNCFd+dhi+6Lz+fbtOmjqqqqGBzZkd26VB1YHpWmFO19y82r16656ZbewUAgGKgCKMrCJZRCoCqtQ6EHqyjvsrJyDxcUDB509wMP3F9eXs5L0LYQFqOEPDxh/LAR9xSVlPBSFhD1Z3GpqQwFuZwwI1lZWV9/9fXAgXcHAkFO901lF5qxIxajJOxfG5IXnbcj5M3FtMJrwhrc1JZMZd+X/8O27XvvvXfzRx/l5eTZtm1gQ2kM8rsMUSbRRgBBdPjw4cFDhkyZMjk/P39A/zsfePB+SklGRqZtEcB4VK647BhkmhVAswshhIPBYMuWLZe9+Wbffv042IoKVQAN5XWgdo+m6SlMAMDnn32+f//vpmkqTxOl0RQLN1gVvPCii5o2baZJphrIop6Jf1BFRcXgwYODwYDH4+EaO2JPBPEPKPO43eXlFRdedOGq1atu6t1bYKriuDWcb14b8gbVr/bv/z3gD2ifrRNCwexuE2IY5uVXXC6q2eKo/69lADnyIg4AORrq9YQUqR0KMMaaNG3y94v+Xl5RYQj0b2yoihA2Xa59+/d+9dXXWpkDHKkKdO0jd1J2iALSCSee8Pzzzy989dWmTZsWFxfp5B6IcFGHMpJthHHmb2yFLZ/Xl5WZOe/FF7p07vzdd98ahkk5UkDwS0957LHBQ4YcPHiAR9hcKJivNFX94SAPjCzbysvN3bTxwxHDh0VcuKi+cLJhABM7/DFhKCoGDHhHUIKwWAKEdQQMAtlQFQV/ITKIMR47duybS5fl5Tbi874wLsOQo7oMYGQcOnT49n63P/7k4xs2bOjapcvaNauzsyKJr00sJDXLkZS4YlpZVYTActLUtiiCeNasWef97TzLComqV/3urfViGlssCPY2f/RRqKZGaOXGsUMDIYpM2T+vuVa0QhlrqOw9LjNDCE2ZMuW7b7ekZ2SIGRW1NWMFfhFgNw6JKirs0+fWxUuWtGrVSgwHO3NT8Adkk9ET4qhhFhQUhPk7qB9qZxC5lJYVTktPadqsmeIs+5exlIzzGJ8vFyeX/Vv7zzPguDhG97A7duhIOEzJ+TRVkQYGQlXBqk8//QQqDk5Hin6kLbZcFE2Fr72249r160YMH1FZGaisrOSjQaJUI/VCpQodxVDuu5G/rXAkhd318y83du85e84cAyFsYF5zBoTQJx5/fODAgQcPHoxE57JEz+uX6iQxihBAlh3Ozc1ZsWL5ww8/LPsO8oqCmKJ9wo6SyUavAmFhAHWHMAHCOgLG1PwoFRuHgfH0Z6a9+Pzc3Nxc2fuBUHGUEV4WhBy3TA3DOHSooG/fvs88++zECRPuvKPfoYKDaRlZOsMRNyNlSPV6RX0q+tmc+xmWV5Q88ND957Rrx4eGXY6e1DFmmn8D46qqqi+//NLj8TgRnnqtIghD4VCjExr//e8XqBiCNsTokdMLirrxBx988OrCVzMyMsIhC0Gn8itRGvC2OMjyirJRo0dNf/Y5QaBRq9oP/20M5KxFBwMBKmK1KCBMtqQj+7VNvd7ktLQ08TrB7vhHfl05OInWgSpdEogw6fIBbIg+elwGLOySSy456aSTampqnE8DsvkS+TIul+uLzz8Ph8PRov1RXbzO5rqs2RCSkZH5yOTJixe/3rp168OFhyizIu6ZqJiMSQAar1FRfbdSApJTk21CH35obN/bbtu3b1/kDhWNXgimP/ds/zvvKikpEX1igaYEYipJMGTx8ItQkpeb98LzLzz77LN8jFg2yeM0wRJ2VMwJworOsSdAWEfAmJqeEUxUovy4ZMmSxx9/PDs3S9SloQyJxZOx1IVlECOztKy0a7eudw2468Ye3WfNmpGamuLz+ohty9squoNF9UdjIDycuba8vPyySy8fOHAwbw1iNYsJj267t05jjgz4tz17fv11l3DAcdAqyBPJcCjUtm3bJk2bOPqFDUJJ7ywOl5eXT370UZfLwBjr+c/Ytg5gNJJ8BAIVDz740LiHxxNiSzaV2GsUlxHWGQTHYrvUb5WileQwleRLTHGExUDo6gysFY4tmoaKtrFhGAKJHXHDcvHW6h/X00kVh89JOUjTpk3bntW2qqra8f5MqRZFbhGv17tjx459+/cBrfd4VFdvXEWar0dICSXEuvqf/1y1evXQYcPCIdtfGTAMLPEBYnlAGpWx4kPhCENCGMZGTm72hrfXd7mh84oVb2KMDYwsO0QpfWbaM506dS4qLHa5XbwSLXXcHZhELHKpnJycJ5544vXXX8cYi566855KZFlHyxCTyqFAlW6o0n1PuN76sLqGOuLZoxizCTFN85OPPxk5cqTP6xMCL1HaAaXIJt4SIlBdXdO27Vlt27a9+eabP/n040a5jXgQTQUvpVa00ruz6hmzqHOFIOgPJiUnT54yxeVy8cYVaBiSqPoxJ153yzffBPx+zo6r+BshdPBPIcuyr7zyCq3UpJSM63mvYY4REYTQ/Pnzt3y3xef1Co+lGIv46AekotFvmmZxSeHw4SPuGzNGvEpPVzurr3VmhP/iSLxen+RIgiwqYwmpYH8wsBHw+8tKy9Q9LhUJnd8i+qVifb8YON62bdvi119f/ubyffv2Yk79qAqdDQJtExdKp+lXXnkVpcRB6aUSBQ4kNA2zvLz8++++V9m5KOgefYs2j1RX1rattPS0xx9/8rX/ea1Jk5OLi4uA5P0V+Eqm7lzVNIpuyCwrK6u0tPjuAQNGjxpZUel3uby2bSOMZ8+Zc/Ell1SWV3J3TgW6XtcwENOUecjrdY8ZM2bjxk2maRIH10eC6/Aomhj7lSidBBNW/VtdQx26VChSW8Ilfnfs2HnXXQMQAqZpUCISGY5VBiDa3OKvJtROSvYZyHh2+rPVNVVZWVmCCwkCTg8d1TsFDlyGFupnQhzN5TJsag8aNLjtWWcLTgYB31D34bF1N8blW19/9TWUpbYYxyyeY9u2LynpkksvUeFGFGfYEC1t4eN/++23+fPnZ6RnWIrgwhE0EZFVulyukpLi22679eEJEzkoBkjkI/gzHXdnSSMtLY2j7fj+y+dSmIRIA8qoaZrl5RXbt2+TrQlHj7SOiEQ9JhzXU1FRMXDQwKuvvnrQwEF3DRhwxeVXPProo5ZlyWhGpcv1uoPrYkAkd2SMXXLxxYK93OnsIXCIDUD41VdfiUiIOcoGR93kNVJEv6IrbNvWlVdd/c57797Z/87KiopQqMYwsShS8BdRoCgrgW7uQk6K53JlZWcvXPhK1y6df/hhq8vlIsROSk6aN29e06bNBV5aVrAVLEFE2xBCy7ZcpmlgPGTIoB07dghyD/mkP5bMSlhDW4z+RoIJq/5NpaAqs5D/F/EpYHJw5XBBQf/+dxYWFyX5fBQQiARFBlNXQDCkCAANRggR2/75l5+Tknwut2nbNoWCxxkjCICjAq3K13zqlFHRFRZpjd8faNOm1ZAhQ8U/dQLZEBRH/7k5XWw4HN62fbvLNJmYznLsHeIJVVWB008/7dQWLZVPkwXZ+jXnh/L0d8GhAwc9bo8odSvodeRiYz52grFRVlbarn27p6Y+I0tMspn3J/c+Z5GzceMTXC5TEkoDAbSDjmoHBIitWbNGRFhAAfmcrVad+kYxxiByqu+44/Z5L84DAKSkpPh8vuqq6kcnPTJeIXq0L6/XLYI5rnjkQYtTTz3jjDOqqqqiDWD53SVOxTSMbdu2iclpmUEcM1uWCFP4JaaCkI0XgUl6euYz06a/9tprWdmZRUVFECKbUIFVllkwR8VCVUGJ3M0MEkKys3K3/bSta5fOCxe8YhgmgLDxCSfMe/mlzIysYDDIiTgoY0S4bhL5QCEVRWxCfD5PaWnJnXfeeejQITEcHD3IhB0NczJhcZ4UxBJMWPVpGhJFRTwqt0SxR1BGIUKBQHDEiOG/7fo1NzOHUYiRiSCGAKvRV/4OcoJeKAPjSDrlNvngKVfkBbxnD+X9pOczKYjkRDxNREoxCQjQdU11eOiQ4ULTFznysPqnOKoPc3avDx06VFxU5PK6CCBO0T7pSTCqqQmdd+55gpVC7cX1D8zRMQHGePfu3UvfWJyRkSGgWDJzYUKaGRKGIKThcAhjY+pTTycnJ1Oq9Q/+/GE5Ghm0efOmmZlZVphAWdMSYy2IV+gpY3ZWZtaHGzd+/OknDmIHnfsyCZwWsCaFKUMIrV69+v333j+tZctkX5I3ye3z+VJSk5s0PWnBggWfbN6IMRYEVfWdP0HntyOEeH3e9uedW11TrSo6Mc/klB2+/fv27d+3P+KA6Z+PaRrC5IycSGq4MBKnzZJKIR2uv27N2vVdu3crLCoklm2aJiWK1QsgRnFkhwCyeCX6FYTa6RnpAIBRo0bddVf/wsLDAIAzzzxz3ssvQYZD1SEsYiNImJ7NYxgjE3Pay4z0zN2/7Bp098BgMBiD9ThmQpb/KkMC/x4FYXFcUAKEVW8m3ZkY6qD8DwMOWgkE4ZgxozdsWJeZmS0mSZgUFdDpG+CkzaJyLTRuZbGZ/zE4v5USc4/RLVVjv1A1GPjBYAOVV5ZddvllXbp2jXYrj+0L7awz7969u6S01OVSuWa0ysCfQyL5ULv27UX4KLj3WDxLdD0cD1N8WwCABQsWFBcVi/6rtEjKImsYkedgo9LvHztubLtz24tE7T8viTtxW+kZGaefdrrQIowd2iaRxUcRHyWnE8ZPKCktNTizdBS2JSGxRDeiBR6woODQzJmz0tMzKBEqW4hJ/jSMMZ7z/Isy43TQjNTT2QXOkyMOqX379sjAxCZxp0s8zTBwcXHJr7/8EtdHPxZM35vOLoPE83OhhZNOOmn+/AUzZ81MTU8pLy/HBpJoMkAhksx3iGEpuyD2a0oRNDIyMlasWNGjW/fPPv0MAHDJJZfMnD0rWFMTtiyMMKVArEAHXWnkxNh2ODs366PNG8dPGI8QooSoBiQ7ihNc/7WGmGQ2VIoiQDIDJwKi+jFR8JM5kXP8DhDON/nkU48vWrQoL69RKFwjOl6Cl0z3fEWKAhThnNAugopDWHBcCTg1f3MapScSo0hMzhw4y+GEkEFDBpkuUzn5Yz3Scm6pO3furK6q4tJGkMYDiJhlh1PTUs5pdzaEoOFU9pzp74EDB9asXp2RkaHppuXkjoQ6McpoWUnpJRdfcvfAgcK3ab3YegGFURq5plf946qwZSEMHadL6rbz8iZJSkr+cev39w4fHqyqMkzTti1Fp8UkdEn4YT4TVVhYOGDA3b/+stPn8wi+QwfoiiUlJX20efM3X3+jHGSMa6kvg1JMF0EI27Ztm56Wbqlj1nqD8iogHA6Ht2/fLl5Ym/vsKJoey2bqztQ/Z4wJsQRKab9+t69cterSyy4rKirk7GZYCqxA+R8VmmMefUv6gNzsvN927bq5902zZs20Lbtb925TJk8pKS0jlHHQO4uiqpE4n4g3cUKZmVnzX5o3/6WXRTTGdPyaIOg4soagpCVT1SemOBsSl6FeTE1VipYOL/FzP0ls0zDeXLZs6lNT83LyNF7KEeCLMUcmM16NnVE3sRr5U3KC6vn6cXTGX19JxhBGlRWVF5z/90svvUzP8Bzj3teJyIUQ7ty5E+Ho/Lpz449sLlY4r1Gjpk2byHamrMHVcwlap78QwnXr1u3ft5/PgYj5S6W1oCaCiE28Hte48Q8r4SY5WfufYyygQ+/9qquvbnxC4+rqKoSR1hnkhDqCtgVSwnJzs997993bbrll165fTdOFBABXDKHz5YUQMk3z+63f97rxxv/3xRfp6emE2HJngJTLXTCuYotrqoNvvrlMN2Ubwts5+VWaNGly0oknhcNhxewGVKVdmmkY27dvhwA69Y+PhYUdM9YMYNyUs66027bdsuVpy95c9sgjj4TCoUBlwDTd6jpGSWd5EwtzeBrCGFuWlZqSbLrwpIkTbutz2/79+wYOGnjf6DGFhUWcexLzs0TVxi458HglHGdlZU2YNGHjpk2myyS2DRSHSW3mn4Q1nCHIkA5eRaMiAcKqTxNxuhgOAUy0cjlsxPzss89Gjro3OSlZTuUq8jHGdEEIqlQJQR75xkwZiuISw0IdiEoADpYFbA33YlHENfcZNmO09823eD1eJ3XfsXmz1eY0sCxr165dbpfcm3QtV1FwIKvGbtniNLfbyxTPNqvv2prOXBFClm0vX77c7XZRIsuDcpBadOVZxFdVVlb26Nnzggsu0DzM9UVkoc8MpbRJkyadOnWqrKwUrI0AAq2GpJ6GrbCdlpb2+eefdep0/eQpU3755RebHxJWpP87d+589NFHe/e66eedP6dnpAo2U6HPoiBdfNyN0KSkpA/ef7+0tNQpM1W/ppc6IcTlcrU6o5WoscuAAkTBbpRSl8u1e/du0Sd2cpPV+1H951abcQUAwKW+CcZ4xL2j1qxZfdY5Zx48lB9xlRF3S6MASfkA8dY+jaxAwiDCOTk5H77/3nXXdVi5cuWEiRP63Hrr7/t/52TRUM2BY0W5BVXd3mCM3TNs2O7duzUgS8olqdD2aJ+qv74Z4j9y3FRWSYEQeRfp2nHRIzyWLTqwwfMiQpmJcf7+34cNHRK2wqlJKYRQw3AxRiTAROV1DpYJGfKrgqFE2VBR3AKEU94wiEAobLlMDIHIEgT+UuFeIxeY2CF6yilNrr+uo04oAQCgAUDC9WFM9c+0WiIqLi4+8PvvHPEb0+rTe7FFalqf2YpXfm2ITedMbX2ZeDcBU/rh++9/+uFHl9ujWL31xor4jsZs287MzLx74GCgShlqMdTDUcVRdgwePHjVW6tC4RDvhTPBdMZVcPQFRoyQ9NSU6qrqmc8+t2jBwtatz2javJnX6wsGqn77bdfu3bvKK8q8Xq9A58naAVP/l+stkkW53e78/PyPNn3UvUd3EVg0XMFffLuz2p61bMkbCo8op7yYwvF6PJ7ff/+9oKCgWbNm0pccq1tWnKRVNHzkOy2ltP25f1u5atXTT019/vkXEK5OS00nPG52IOyIwjbzUjMDYSuclp7qr6wcePdd33/33dPTnykpLXv3nXcan9DIsq1IcC9YPhjSZTbbIsm+5IKCgqFDhi5fvjwpKYkra0W2FQaclADH6Gn8a1gsCAtIMFwChFVfFjPjyWtJCIJQKDx8xPC9+/ZkpKURKlQFiWzXQgpZFOPjuFHlG8KoYgyPU3mLCGMEIAoGq3NycnjvkapdCvI+qfRhEOCacNVll1+WlZ0dR21/TF5oTWsMdByTn5+v5h3VkxzTOHwHMVvwASQGo3xA9ZsJOb3+pk0fBYNBPhPFVB9HbF6EMWpgs6Kiolu3bm3atJH+TA1z19+xCGx9JCBo0aLFwEEDy8rKMEY0WjyQ1ReJgYeIEuBxeTIzMwm1tnz/7RtLFs+f/9LSpW9s3fq9bdvpaZkYG4RavIYp1hkV4ZweaoUwEnzYNtm0aXNMJFffpnNrPozUApuGPPnKM2jvhTGu8Ffu2rWrvoeS69/iqLLiDpifWNvnS5r4yKOvvPJK8+bNK/0Vgi5NPRMJsnP5ZryNgIABGHK7vRlpqTOem963b5/77x9z+RWXl5eXR/w6JIoRRMtmAhgJxuzszOwvv/hizJj7WWQJRRYtgxrtlfC+DW5IAeugBjNy4pWEHGH9WBxThACMTJgw4d133snJzhZaCwCpZS5Lx5LHqjawJc6XMAYoiWw94XC4qKhoyNBhffr0qagoQxgyoeygtmghf8YoNQ3XdddfLxHC9YcDahBTjhc4hk337d1XVV2FJKUtcB65KKImJfmaNG0KgBSFbSDvK7yOZVmff/6FYRjaA2m9I1GBCFthj8fTp28f2cMGciSF1Uf66zgkABQh18BBA9u3b19eUYkxigndhGKtGIzCsm0BIXK7PalpaenpGampKR6vmzFmWRbv80q8vQrdINMU4bKcj3w+39dffRXw+0Xu2xCFaKeXOvHEE5OTU7Q4bhxlGODt/99++63O5LJ+j6q+LLYKHf0ugq+DEPu6TtevWrX62g4dy8vLbWILhJ0yGXg5OGYESyjIa9Ro44cbhw4ZctFFFzVqdIIVtjCPpZQciD45kVVLiZXXKHfRa6/OnjnLMExKBZEeEwB+x6nT9ADHIk/A8WuSWzgK+RFIR65vxZ+QONH/kTndGyUUY/zWirdefGFuTm62ZRF553DuIgapaNeIvU7eWix+gAFGueYjQbHb7SkvLzddxryXXho3buyPP/4oxeel0pLYdeXgUk24pnnz5u3atVPi4SyuI3VsGZRwYk6pKDf3/AMHbNuKi0z0TmHbdnZWZm5uru55g/pmnHeOkeTn52/f/pPP59PMUEojn0GGMDKCweD5559/Trt2ol2nfw3rzSvInVdQmDHGUlNTpz49zev1WVZYQKxEIotk1i0AtLIljpABAYaChJQvCYRxJMECWqYFSEljOcIGJTkMQ5RQt9v9++/7fv3lVx1cNgQQWltOTk52VlYoHNIqjOI5UTgbAwcOHHDecces663LoLOWwxkqMSEkJzdv/vwFjz/+hBW2qqurDdPgo/80iu1w9D740DewLTsvL/fggfyXXnrJ5TK5YwaKXcvxKQCLZWNZVk5W9pNPPv7J5o8xxpwMRJPHAQ3kFgd5bPIEHL+GhGJGLRCWJF46dnfnY9ui978qShJKDdP4+edfRt83OsmXxFk4gMqIMJ/+EgJGFPDfADW7q99Qm24GU2YfKjhw8UWXvr3h3V69epWWlm7fvt3tcVPCNNkvk6SwDEEUClnt2p+dmZkpnXS9eqb6N6YAbPJPxAoKCjjpZgxiXP9t23ZObuOszCyouq1Mg8vr66AcNYOtW7cWF5cIGXxdA9cZMELYsqwOHTpw1wvkTLeMuOrrnmKaSUMk5YSQc89tP336NMuyCLPFiIu4zmLegSlQPVRVXMYIZUTAZSUfKaq9MBRBMS9pc/oeappGdU3NF//vS9BgdIYOshGWmZmZk5NjqUp+lCVOV0ooKywsrF32OKYXucPiK9J8jk6QcgwcNHjFytVNmzQ7dOgANrB0ohJXJUAjGpwVCbJs23b7vAySoqJCA8t+jSN2lP/loFBICTBN0yZ0+IgR+fn5hmESqrnT9QsTGXCDGBJ7vAKqc8fLl7RIOnhJ4/gKJI8JizoGJqEiCMLq6uoH7x9TFfAnJ6dgZBq8Ugz4uB7iHToKFWc0Q3IgzLHFaOwPJ8UlgUDAX1k1cuToZcuXtWh5KqX00MFDhw8e8rg9VMr7i3YjBRBJalkGzvvb3zXB2bF+ZfWdTiOnSCjVHziQ7+SQittnbdvOTEvz+ryC4aQhOoHO8ub/+/JLSqiQCVLXKFo1JcROSU25/MrLY3Tf9Bhr/Vg0HdERlW3bPXr2nDBpUllZpSZakYB31XOgVIURgAociCy6AM1s4uh0yI+AkmWaIRKJznEkX0bou2+3NFwkp10spdTj8eTm5tqWJTwOcirwq6ZAWVlZbRT0Mb3IHRY7oeScpIKEkPPP/9vKlStvufm24qLiUCgkFpgMxMWVkR0DxEsaBtcLRwLq7Ehk5bvyv6n4FNOMeOisjIwDBw6MGjkqVBOCEMTNK6t/Uv4n4YDrzcR4tki/ZF6RAGH9hxaTrartAyE09cmpmzZvysjMsG1LFR5ka06oT6n7iIkURQiFOt5W7v4YQ78/cMIJJ77y6sIJEycibITDIYTQrt27Kv1+DqiOPEvLqYv2D6E0yedr3759nd3lY9F0WUafKsZKS0sxMuJ2VFEWYIASm5x44olAkE8x1hDhujhvwqtt375dKMo54i2o6SOqqqratGkjGtLKK+tNtp6Px4nXEyiewYMHT5s2DTBWVVUlQhYS8cIUaL8lB0wBpFhKNglYHyIMsqggNIvkWYxzKgEmuMap+CqUUtPl+uXXX6uCVQ03K+HMC/Py8hihcTmufowNo6SkRDgncJxPUTIQc+Timmbn5sx98YXHHn8cQhgIBBDCAESZUCTcQ0baTHPBaMyB42zogk30zFgWyczMeO/9dyZPmSxFJ3U5J3plkZIPOP5O6bFpiAr+vAQTVv1ZDL4RMEKJYRjL31w+a87sjPRM24qyP0ZHdqFg4qZi9EbyaaiUSeUBROzglX7/ZZdfseKtldddd52QNME4Eufu2LFDEDtD4XSBklgAEX8QDofz8vKaNGki9mjgIAw6Rk0KCTCtPltTU+P3+8WXlU9xHD+lFCJ28imnaBA5UhDq+v2S4mCKi4vz8/NdLleULYGJMU0qNOZCoZr27dt73B5Ni6Gnkeu3JA5ADGweqaHe/v37L31zWdNmTQ8fPmxZFi9FRmEETBHvsCjkFTAtogTEF4nkuUBdBSELABy0Ji6XcajgYH5+vrM5Ur8GHU2E3NzcyA7FaFx1WsEbkd/vr66urg0tPu4Mas8n2TM0MosMHjL4zeVvNm3arKSkhCPvRMuDqoqbImKTI4gwDk7lqIToToTYYzCxaUZGxqxZM5e/+SafDLZ1GydKqZXoAderIUFkCOTgV4IJqx4sZs6PMtMwP/3kk4fGPpSanMxlD7SQi1IH5CgKyDRdJe8J8B2QMSL2dMaY6TJD4VBJaemQIUOXvLG4WbOmgowXqcB29+7dshaqFH9FjVEEVaFQqHnz5j6fzwmZOaavcrQHLK2qqirgDwi1gzg8lHApDACvzwcd4T1sgElGcTAHDhwoLS0VDbbYhAwImQ2McZvWrZ2eQOPn6vF44rI9pycmhFzw97+v37Bh+IjhAIDSkmIEoGEYCGMkmAlVc5hvq0h1qDE3w3QZVVWB0vJSLpqIAFNVModfMwzT7/cXHDrkRNTXrzm/YEZGhihvxLoT+TSEUE1NjQDENQQA/gib1EwT9R1dDELItqzzL7hg3Yb1PW/sWVRUJMcZ9J4NJXFtVAWMxVwzx/kEUQgbEOVlhgDKyMgYN3bcp59+ahgu0YKE0XZPogdcz4bUpQKxIKwEE1Y9mLg3ykpLH3zooYqycrfbLdGakVsLaT4fEe7y7Alp0IOQMoJyNJ8ZhlFaUpaakvbygvmTHnnEMExKCObjsOIWsgnJzz+gmpFQ33IquoKWZbVo0cLJ73Gsb0/KWenuZnV1dZWQcHGkfYo/jzFGEMTpaWka/hbVlq+nI3Iycx08eNDv9xumobNbUYJmvPBAbOrzJTVr3sJZHz4CJxw6TOTB6enpk6dMWbt+3c239LaJXVhYWFZWWhMOA9UG5mgfIeTOVesICQQChYWFJcVl55xz1tBhQ4XP+//sfQeclMX5/5T33Xa3d3t7jaKAICpFjMSusUclUWPvDRU7NuxRMLbEgtgFNfa/GhuIFDXmZ0GjsaMoShBE6vVedt93Zv6fnWdmdnYPjYlXFtwnfsix7O69Zd552vf7feQgEGyutaq0U5pIJNeuXWf4bt11IhmZvVApmBp7zNP8oixObYc0c8E3gEX+Q4ZVoQJGfOu4HEGlnTFWWhp/4MEH77nnHsdxG+obJOuMp+XeVflNSaH92O9RpTYDGSGhQLixsfGyyy5trG/QZTyM0+/MZ8DdaSCMInkeyK745ccR/u9mxyuEkEcfe2zR51/EYkWe36mSDCL1MTCTYz6ZdhhKJRhZag2cC4IdQkhNTc0OO+zw7PPPHXbooYwxpcJo+YOW5uba2hpXKkIYH2yA0HADhw0bZtzAhhhaJRMJ0+HLqEJjxZ8LuIF4aTx9Xub6dNMB2JdrzZo1yUQS0nG7nAA5C2OsqKho0KBBdhTba1fbeCZTtxwzZsw9994/d97Lky67ZNttxxKEqqur11WtWVe1Zu3aNWvXVK9bU71m9dqqdevaWts23XTQscce89DDf509Z945Z58bCoY8nwkdpdsMbIKJYBwm4mGLIffzLaO3reF08Xg86Aay1rCd9ycSiWQyiTLlS7vrkHrVLNCx6lVZ5wKhFef8+BNPeOqZp0ZvPbq6upo6MJ7EVxOhoWVgxkP/WBRq06kRY340Gv30008emD4DmsE9pPWdNylFaeAhlhIWlsSJDW6DzhGztRra29pmz3opEol4jDnUgeBd5qc89e962gmIzMHIBcQJwkyqOnNHZhjt7e2nTTh9ypRrCwoiIESgxFoRMqp7TU1Nba1tahitRgLJe8kxF5yzYDA4YMAAu1aZCa/IaYNLmkgmOzs7zTka9T6I+RnjxMWF0WjPhRf2vl9TU2Nv/QipdoLM14jPvFhJrLSsNIv12ztXO0OMQsLBASswctTWI0dtPeniScu+/XbZsmWrV6+uqq5uaWn2fT8UDJeUxAYOHLj58OFDhw4tLy+HsysoLCiJx+vq61CQC45NZVcPe0j9WV/f2O2nlhHT6EUbjUbdgMs4A+WTrFAS3JJxwL18zbvZrDzTsJLs0jqclOd522+//QsvvHjppEkvvTw7Go06IL8qWGoPRwzqFnqgsMJmZeavamS2fTGFEPFY6bPPP3/K6adKcT2RhwT1kDkCcVATFUoLmmNMBEeIYJQPeX6GgXv49ttlq9esAVS5CkwRNDaJUkZQJWICGjTqXxFmzKcObWhsKi4qvn3atBNOPAHJ8qC99QgLhNLS0pxIdBBqS1fKCTayz8A5DwaDJSUl5vBy3/vaWzD86Xke59x13awmnzllSmkgEOhKS+2J02xpabHrvSodVMVvzHxWWhIPhUIqQur2X/+fLH1/ZSpECU2FYdJCodDIUaNGjhr1Ix+XQP3UeYUjkcKiQgnuQwyDonB6BxcCEUxaW1t61M8Zb5q6uYQgSVPNagPDe7jgPsueGbyhW1bEbCt8ua7r+368tPSvjz46/M8333bbLZFIJBQKypYBklE9R4JgrTCP1oOIMDMr1ZdDUBuOhKqq1n65aNFee+8Na1jOodKwlb66FhudES0qo0C5CoSF8wWHbjAhRGeiE1i5xAAUld9UyvWA1UUG9oxVS9N1A+vWVW3zqzEzX3rphBNP8H0fuEzIkh60H8iW5hZJRtJ0FJQxsQ2e1VAoZIEpNoBNSu+taQ40FMRsREla7QsJh7rggLMgJ917PGDt7e3IguBiNZlZsY+5EAUFhfqfUJ+UQ9U2jfX/NMIKitK+7zNpnHF4RZova5uMUuJI4jXGOFoYVR7QaIwj4xJSGXZnZ0e3lygz7qBeAK7rAipbdLnFCgnh+77MgO12eDceVR+aXdSxgw+pXZWyK668/IknnywqLm5sbHKdgBAQKwk9UjzDrC8xbX1LeYkoxbSWlhZhANNC9ASn4BduBAuKLASdHDaJ8iCsn29w9fpV9isIRQAQYTpW2kemKz+KEJJy1VzyXEVDQ8P48ac+++zzo0eP8n2fUsdOt2zXq8qziYTvCXk3keEhYJCj0IGtcU7CTBPOYbPOV70iBBOcK0oM0sMFsJnZQGx8rAlQeggN6+tegHV8WlxIHqobcPU/9xICa72maUMZ4CwiJRrgB0IJOGapROnIHyjsA0Jv8XLillyf2De4XDgzhLjvd3/SmZHg6i8mhOCMm46MFpiZtcut67xRZhE409RlkaD3cePGzZkzZ/fd91i9eo28nwQIxZgwJX+GsK0Vk+GSMdO7EKDdqeM65eVl5tcogdwc3zU2NJOicpig1H+pKJkr3X4NwhJGhjd/2X+qWVpIbNNBm241ckRba6sllgT/Co00oywDpBvhOE5HR0djY9PFl0y66+67SkpKuM8IpQYHYe3kkqmJ7aA4vVWlvTyIKcuBqaFQBBnYEsr1WFZkni2CbE6jMWEqe/pagJan5QO6Nsx+vtnfQzGhSM/RwHoAvgHlIuRoupRWQejjq/0DCRDKLOPrZaZxO5pjKs8YaXUXDdnU4pXdf7Q2CMtgqx3HMS1o3XfXfwpEKBJcQLfb/pKNzLLaLkjz2wgmHmNDhgx54sknzjjzzKqq2s6OpEMDnCPBKaSvZkq4rTuNCWwIlEAdDilCV//+/UeN3hpgBGp/gT9Q3gt3mxGhCs9GWjW1t8lIk6BUskGQSjU2wqXcQ2YrAyOEzj33HMd1kwlP7tggbqXGm8v6M9EFVOE6webm1lAodN99919xxRVQWYIERQPksLDqSRCXwv7IOctwP9j4YGgyA0pFfongeEPQk7MYJuYVU7dHAHayO5224EBWw6wnMiFKCFO9GyDR6rn0WlOqI9GpWTRdIom+MPyfLKPOqbWUMqgQyvFpuVr9fkOH6/Zj1gs9/VfwGKb4j0CiS/4puEoHu/0wcsoymVfIQOEQRg4hnu+FgsHbpt42bdodBQXRpqZmTGCLAKYZS/0ngILG08tT7hBcqR/KllZr42mnnl5UVCTbyepJU2RVlHcH3WaKYS9RubCJWUpYWLGzrYntOb1l547B40EpZZzvvsce11xzTSKRlMxflp6vgFXGCj7Vcdy62rohQ4a88OILRx19tO97hMiZRbAPWltkxtQUYYkX2QcAv0Lp0iGoQ0tkjQ6ocv5WWnG62twd15ECPVK9TYsUpH0wxsmk58kWYC+kPsFA0EQ38vdzBaSQ+RohpE1Sls0rOZWN2cClLN8prAwHjr+jo4MgglVwoSVT02VMHpbYgp44QbXmtb/v7Ow00YAGjUpHAs5BpFLkYDCYZvH0ddDTo5ZufhssJ8Ku4wKZe/xppzz/wvNjxmyTTCQxpiD5rwajKb028MEMrhbgAJKe197Wtnbt2gkTzjxl/ClyWrnuo6haRHoUSt5+vjmKaqYoFEInY/qWIpE5PjyHNpGctYyUS/7Q0tLCTbSOuGbpEaMASB3a2NC43fbbP/jQgwM3Gej7SUIco2WV7nOqRw5+UIBbQDo71KWqUp3u0wBrBMsiksTdMMXSx3iDuJVZMOaA60oINAcBT03FUldDluD8js7Orh/sifZkcUkMNCJUiCDnEyA1G044Dm1qbPSSHih1wDT13PHBWZel64EJJbmBE4lEY2Mzpa5efqr/q94mPWFxrNgmH/fIocrf2NnRIdXfTPlZaAk56P4yA4PHIueCnh4yUxVTeCrd8EgmkyNHjdhn330WLvxUovGJwJYqDdYJVSqWop7nd3Z2Yoyj0eiwzTc/4fjjTz7lFFDd0KRiQxfOo6C70xxszwaH7ppQij7yrwDqEemcKn/x/5PZ9U9KaXNz89y5c6EhqCpGiKtCpXSmlDrNTc1bbbXlk089EY+X+n6C0gB8k3agiqmXlRcilEYlOg5VrVBEhZLaUvULqQibyg47OhIGtYLTPjpHzZZZMCBYOTctiRWSDKN0Eozk/NREY2OjTc/tXrqz/VVl5WWSaqlqE6kLzhHoRcn77tTV19fV1fXr32/DysPSDk9e5cbGxoaGeqpnTkDMZ9jnsEbLKyp6ON1UmXprWxtjTA7AgN/O5KRjruU/WcANgQNGFjS6x44qVyzrSZG1LhYIBObNm3/71KmF0QLOPWgmciXQI3X4JHDBcUhTU8voUWMmXzuZ+ayiomLz4ZvDhE1ZpCNaZRpreXWUL0F3ozkizewSQPMjQDhQ2AdVZYJeT94B/0SD/QJEH79ctGj58u8iBUHOJO9LcMU3gqamEJ7nRQqC0+68Ix4v9TzPcQCrLKyZgSLrGdN0Jv3rkCgqLpaDARgl1GQqcg9VjGTf9zo6WlWPBz6ENrBuWSAQjEQiTU2e0OG+vfuADkNTU1NWob4bMyG7YFtZWek6jhTLBdF8aJoSQRjG3BFuc2Pzd999V9mv0kQD3XIMPW3mUIH2tnrV6pamZurgLPgPXF7Oues4AwYMyPpsdx+SKik1NDR4fmcgEIXXEIfOO4ZZMoLjcDgMJehflGXR4iVX3vnmm39fetml1KGpyJBzSmQ4rlWHZcWeECKY7AhfcOEFu+++u/k+GEJsPUFG2tZE7XlH0D2mlPwNt1KCoTEioBqBtPhnfgrVf2FZe9Ann3za3t6mOihKL1dKZKQyCUYpbm1tPe7YE7fddlvpfR2UbgcIrKP7zMcs408pzIGLiqLhcFhwtB68j4BpSF5DQxMcoPVn7iZntreDkwqGgqlztAmLFtgNaKANDQ0/1N3sBrPwL5WVlZGCAsY9xQpTXTgYakWoQzo62pcuXbp+OHcOmx3qIYS+Xbasvb0tlXR2WSqgN1JYWFhZWWnflG40fWeVElNzczNjXIKw4ACIxPQSnZcLKUMR6vbAK7ctQxkbtHqqq6vHjz+5prq6sLBQwjOVoraafATarYgT4tTVNUw444yDDj7I8zyghkO8jhWyAevoR/2Z14LuXiMCOvPCAGNlc54b5BXPDHby1/0/mL3Vwpa0cOFCOUyeqC6uoNBhIwQx5icSibKyslNPGy/XvWbxInteQrpRp7c5YaBJwHwROJUBRwoicnvS9CPdGxIyO/S85Jo1ay1QqR3G5rRvEPpKhsPhgoIC3/ezILvGDQshamtr7USte48EWwczcJOBpfG453mIAJYUnAFMv5ANad//8stFWTDj3DcTLsAxf/nFF57nyfHSKItXjTH2fC8Wiw0YOKCHzlEPA1CHVF9fL7SCq7zgWAkIyaCWcxYOh8OhkM1w3YihWLoxi21GGSGpqOjiiy7+ctGieEncT907rLYBbLrEgjFGKW1qbNxhhx0uv/IKIYd3gYGAlpagRobdt6Es4A3LiHnikI2vwFmVfrOCN86l3L1mfCdM4V22bFkqr5U9GGADwQ9yxpHb0tq62267DdlsM1nnV31dTQ/IKDV3JY2ACSGH7RcWFhUVp/xBenYK1hNO1NeuWrVKlmqRHn5i15RyztL+VW/EkUikJBbTo8LTmw78DOISVVVVtkvuZviVCViEqKyo7D9gQDLpEdDx5nDdYQJr6hiDgeDChQsTiYQZ2LBBeAJz2VOrN5H4+KOPXDcgOMpqqAstDjpgwADIgHvC1aWl0ORfq6uqBBeGuScfJ4o1EI9zHovFQrpG0r1HkoOmtWzT11ymv3Tq1Ntfnv1yZUWln/QghcUoPf9botVST1CiM1lSEr9t6tRoYRSGCquvhQkOPyhdmbfuNKI0r0CZWC5oKLQRhaajZg+U78/fjP9g9uA5jHFdXV1tXa0UECCKNaQqO4BqIcwXO+20kxDInhKoN8GMDFh/rTCYFF2ITuVewWBw4MCBqS9R8131RqkmBXHq4JUrV1g1W57j0z11gK/OG65PWXm5zoBV11W9GQiqhKxbtw6qcD1R+BUaRM44dxxnxMgRqYNBFGnuAMEKoM05D4VDi75Y9N3y5XZOmZuX2jbbla5evfrzzz8Ph8Nq6GzmxQSfN3LUKIc6PVSCTv8maWvXrSPEwXr8rSJ5ISlaIA+mtKwUbn02rXljND3pV20RnPuu6z7++JM33/yXstK453mYYE0WMPu3ALgiJqQz0X7TTX/ZeuutJeORGN2A/Fbfm0awRpYDyY8rqRQOURJI7+aVsH662TuREKK5qSnR2em6ruLemRwXg3QDdwN04MCB8qPYkmg2lz3jm7Ub5laVFbjaHAsxeMgQz/Ol/2VI8bql8If81a7rLP12aTKRSEt/ACoj902DyoQQFRUVAjEtgmj/OwK1r7Vr19bX1/8IweZnHYhFCRBC7LjjjkIIEFPGumLHBSISiO44TnNL899ffz0za8zpC54lj/Xhhx82t7RIcB8XmcN3CSGYEp/72/162x5M7gH8L39dR0fHurVrYRQH1jU6ye5T69xnyVhJSdfi8wZRePivzBQGtCiZ8Dmj1P3gww8mX/PHaGGhIu0qfQ4MgGZMQKZTYIrq6utOn3DmoYcfIqF21OBO+vrMfnFGkNauVQMJJdxBiouCKgrJB0T/lWUN4Ukkk0BmN5uvDts5xLCUkHA4YoOJYPdQpWGALFo9ORA9htxXTQkDgQSEhg4dzJgP3CQjmIW0ELHrBr5f8X1Nba3W8kOWj8/FB0+riKjtAy5dv379HOpqQhcWKGOfdV23tra2rq6uh4DHaV61tLHbjo3HSjzPM+0bOYBVETaEEKFQeM6cOclkEmeAiHPyamciaeHSvTL/FYkKRF01PoUQXsIrLira9tfb9mCWKdLz/+vr66tqqh3HEZkio0KPmkCYVFZWZogsbqQZsI1+kOJJ3KXO6tVrLpx4QTKZCIYCevHzdEog55xCXFhbW3vEkUdcc801vu+bL9ywkAobjREOikymqomMEhbsy3klrP/O1qMxJFAmj0hi31TyikBvUrHusOYHyahVbuLEetJMLI8NIFfeKSVMsPmw4ZFIJJWsAMpLwR1TAZUQiEhG8uLFiy0xWK1okJPRlVJAFYJYndcBAwY4jitF9YQeP4QNIRVY19VV1WbcRbeXoIVWRBFCDB40aOTo0e0d7WZGJDgCgTmWJMtIJPLxxx+/9957hFDGuHYouZhq2MqdkNAvW7bs7QULYEXZFQVIfxHCrW3No0aNHjJkaE/v3XBg69atq6qudl3XvIh0V0DWPwTB1BCiNjjw+X9l6rz0CRJC2tvbzzzjjC+/XBSNRmWPRocdCqamwFeO4zQ0Nvz619tNvW1qSKLVzNLdWKsFOW4EZwW/Gjdr2o2Z78/FnTrXzAYBhUIhRxbNrLQ4PV+PEMI4q6+vM4kvVCAUX0ATjazHA+lBDiI9nkQoRMqgQYMrK/t1dHQSNRVA6OHbxMDtPvv0U/28qeYxHFSfXrD1G0ZWmC8QljvFkCFDpEtgZoCaurDSBVNCE4nE8uXLkEgreHbrIaWJT5xzNxjYZdddEsmk/QSZ0edSnowy5j/yyCOqZqvrGfBtObXZZSmWYIwfe+yxhvr6QCAgOM96A0CRk8nELrvsHAgEIYLsodNREaIMCFpbW1UGLLLegwUXoWBw8ODB9k3fyPK5rj4S1iHn/KKLLnr33XcqKsq9pE8wTWv1GOqQ1Cvv6OgoLSu/6657SkpKDFRi464W5LgRCbYyBWeNvpKEYJy6kQDC2jD4Krlg9grGGMdisUg4Ysbk6SQjjcrxfX/JkiUG5iZA2xbBm7iWQEl/oUVS0lVrLczbf0D/IUOGtLd3yLsJbpmD+wZskOM4ny1cCA9eZi00F2+rKrmoYEGd4+DBg+NA/kHaMev01+y833yzRIBOoRXad6Olsz0h9t1nn6KCqK5hID0SmCJBcOrqs3hJ7NVXX/3www8JIVxLq/ScZsXPMUMkpZR+vXjx448/HpOYc5v9bFY4435BuHCPPfY0L/bQ6Rji3ZIlS8xxmvEMqkeMBOOsOFY8ZMgQ25fk4EX+OdaFd4fgob7uuuuffuqpstIyrZgngWlKSAmqQAwh4fu+z9nt024fMWIrxgB4ZZCeeesbI7qyzNXQMXVb9ThCBcLiCDH1xrz9qNmQV855aWlpZWWlKeKliUaYg+JMKBR65513PM8DJWdZeaYYUayn2wlbK9RKhjPx0ql903XdUaNG+p6fRq3L+jNWfl2Ew+Evv1i0auUq6QyEFS7k4iaV5p6Do5WrMRYvGTx4MAChTQPYTCDigruuu+Sbb+SMNSKQ6Img3tT9uBDbbjt2zDZjOtrbU9uZJKTqLgLMd0YODbS3tk2fPt32vllYgb41BcDMPJ4Z02fU1NRQh2ZJHhmv1tGe2HKr4dttt52M4rGF0+zucEf/3sWLFwfdgEFRpFXPsY8x8jzWf8CAisqK9fSANgoT6zFOKX300cen339/eVk540xOVpDoTjNvBRpVnCCBm1qar7zyynEHjPN9SexWX5vPevvSiBqqqrliioNECEY2CCuvhPVTzYYzQHy6xRZbQMYmhLUdCPmccFFQUPj555+/+cYb4LDTk9cwXHyyvh2NZLAKzHgjIX41dltKqeDQeEayfyo0MxiFwoE1a1Z/+NFHPYRR6gnTWb6a4eK67rDNN/c8T83KUykwNrPzAoHAihUr2lpbbVBbtx+SwV1Thx5x1JGdXjK16SGpDoGE7AGbhNIvjkXnzZn77jv/pBQ6wbnlGLAelgUrkFL66SefvvDCC6UlceYzDcxMoxBSwQclHR1tfzj0sEAwqBDgPYDWFFaPs7Oz89///rdBYGkCGix0QjBNdHYMHTosFApzwXEmZ767jqdvDWfWIXzfp5S+9tprV15xWWFBgVx+hiENXF6uBMIEcwO0ob7h1FNOOe+88xgzarV519v3loridaql0AypH7kG4OaVsP57E5mzzcdu92vf96X2ZDp5NRRfx6FJLzHtjjsSiURWAm1hdtZTfUJK0wOBDiKAKcaOHVtcXOwzpubG6URBaCgW5+KN//tH1n6acfD6zxzhCJurYa7MliO2ACkODK5XCK2zl/IfATdQU1Pz7dJvszeXbjoNu90LLbRx48ZtNmSzRCJpjc3GFnsbBQNBLtifb7qxszOBUXpwZI5sf0IpWasrxBm74YbrW1tbYaqB3QgwB5zoTAzcZMBhhx6WFsrAvNvXi02LX7Zs2ZrVa9IOWLPAlfQYRkkvOWrkCCBQar79Bo/AsrN5i3mIfd93Xfejjz6aMGECJVSS5DlRD7XqxqSukSTDU0Jrqmv33mfv62+4wWg7mwwhRxbhL9YINqwU0+SVSzuToJLTzcIcNPPAYIx/PXZsPB5nvo8JQXqmEfRqpE66KC4ufvfdd+6/f7rjOJxBodJIdhiyr92E0/GSwcVAGiPEpptuutWIEclkEgqwAubWqwQScyYKIpF/vfd+bU0tVFC7PoG2A8mF8nTWToEx3mL45gWRQo37gSQYKTlChFzqtLa0frZwITRQhMGWd9Np2AAfqPz379//kD/8obW5Dfo1SHXyLcgSxyUlsfffe2/6/dOpQw0gIFdMT5jjsvv7xBNP/OP1fxQVFRnhJAvAL+CqtrS0HHzwIZtuOshWj+n24zJJuRBi0ReLWptbYPI//KMSK0YKcBQMBkePHq3+zQpVu/2oetPssNs8AuB9v/zyy5NPPhkhFIqEBRckFYtT/TkIugmI67W1tW+x5Za333F7OBJRCHa8EZboN1CDsAkTjHX4bkBYeSWs/9Hsx4ZzPnzz4dv+6lcdHQmCiRmmKYShshLBRVlp/Oab//zKK6+4rpuKZpHWPTAVae16kQGoW6RScNKMsUAgsOtuu3a0d1LiYJGmIJjGcSgUWLHy+zfffCP1D4x1Vc3NtQzY9nmg6bj58C0qKis8L5EJzVXi8YQSxtlHH30IFAskbL5KN1jWtgVJ8CmnnRovjUu8ktDjU9Mjg1Jv91ksFvvzn//8z3++Z9K4vt0E07UWATQe7jjOl19+OXnKlMJotCvkR20U8nyLi4tPGX8qwraf6ynBAPj+zz791Oc+ocQi0ysEIiE4kUiUxku3GL6l2spEdkt7wzPdVEIio0wlJx25y5cvP+mkk2pra0NBl3OGiSNbWsgeXiRzXzeR6IwURGY8MH3TTQd1hT1vBDHKhm4ESTCOTLvkXkDWC8LKK2H9F2ana6kHJuAefOghPlfNWpJyFAy0JWT9XyDmU0qDbuDss86eOXOm4ziOnMrAGBPcQlwIwXjqjjCW+k9KlRkxK5x6mXm+7//mN78JhgKM+3p3hNFjSJakUmEy5+LluXME5/KvyG5qKuSkRQDPBQds76SpjLPfgE0GbtLW1kqodBIcxmgqxS9CSGFh4eKvFre3tsFQFxXF9MxpQBK82ZAhx594QmNjo0Nd2b4xlQlVr+CIUMehGF904YV1dbWUUjNS12439Kal2xl68mBbW9vFkya1t7UFA0ENkeWga51mCRPS2Nx0xJGHbbHFFswK4CyVzW4lXsuJ2p2dnZ99vjAYDsnngivHpOMqjElnZ+egIYM3HbQpZORZgPMN0tJYEENfVB36poam88+b+P13K4oKo16SEUSoFBLO6nxhjH3OOpPJO+68c8yYX2nuA+rzyC9vtqmiqDB8jmwQVj73/a8tky6ZCjkPOuig4cOHt3e0EUpkmkT0PGAYTJjaLwJBFyF+3jnnXHD+hUuXLoXRAtDZ1V+LKCGUUgeMUikHKBWACKbUCQbDjuPstddev9pm27a2dkqJUtsARyCfN5/5seLiBW8tWLJ0qZ3+pvMh3UHKHQm0rBKcEGLkqFHp4ifKkEjknIfD4X//+99Lv10qtFpBN55GpghRukx6+umnDxo8uDPRqajaCBset3lPtKhw+fJvzz3nvEQiYYqrlg/rniP8iZYO7HRJ/Iorrvjg/X+VxksZ93VniuiZdGqnTnR29O9XeeaZZ9sbfU9oCNtOYsWKFV9//XVBpEC+QpTGeXpVIN/3R4/eGsr7thJWtxxJn5jQTXl4eAWCXAglOpPnTTzv3X++WxKPceaDT81Wj9cIwZaWpquv/uOBBx7oeYZ0lJH4btCXaOMwR8ih8JD9cjkG3gzO6+2YfOMys7cCGensM8+46KKLwqEQR4wIPdBUYC4gNUaMCZkjBR5//JG5c2cfeOBB++yz75htxlRWVoL+bWd7Z31Dw5rVq1evWV1TXbN61aqWltZkMul5XjgcLiwqqigvi5fEB26yychRoz759GNgI6SCAEQ45kTu+AgJN+CsW7fub08/PXnKFEh97G0Um0q3tlx4RrPIpjtsvz2MbpQ5mfbNIq3d0dDY8Nbbb4/ZZpvUKRNkj9L7+UeSxYgFVzpw4MCzzznn8ssu7VdZCf3ptCdQnyK+75fGS1+ZP++CCy645557aSoag+tvmDw9yF7NKszaySuldPLkyY8/9lhlRSVs1ip2xFqoVq0Q0tDYcMGF528+fLjJqHrogM3hEULeeeed+vr6stIy5vv6LHiWOubOO++sJbrQRkAChl6VQBkLiRBy0aSLXpr9Ur/KSs9LEoSNnrtR5SREcMbdgLtqzaoLL7zw/PMvYIzJjQXbz1Ffn1/elGGPMQKgfiW3kXLAXKshbgRLuQ/NrvMkE52HH3b4xx9/XBwrYr4t4SQkmwIThBkXOOUgXd/3mptbKXVKy8rKSktD4TBGuKW1pbW1taW5uaOjgzPOGZPiUAR4IlwwGEPruG5RtIhxn3GGuSKX6TmTHHxRorMzXlr66iuvlVdWAGQD9lyE0kpdcAoYoRwZIGAuJiFk2dJvf7vfvkJwghUqx0SMRJ5oU1PTbrv/5vnnnwNNZtkMFqgHCMH2D+3t7YcecuiiRYsKCws44xlZO6TpPPWD49Kq6uoTTjjxjjvugFq0dhv2NI6eVeoQuu8LnfUbb7zx9qlTY8UxE8pAfxdhnYbJUnBbW9uwzYe+NHt2cXGsa2Wi2w/PXIRjjz32tVdfi8dKuGS+mmUAly71IFD86muvDR06THOiNsgtK1PZDSs1N12qoJReeeWV9993X3l5KWdpYflMeCbhnLmus66q6vAjjpgxY4ZsgJB8vpuzRgDmCHsTASUsga1xhPk797+Y3VyBqxcKRyb/aQp1Hd+TBUnBEWZ6mCdsdZwQRIjjexwhUlISK4xG2lpbvl367aKFXyz64ouVK75vbmxCAkULo7FYLF5aFo/Hi2PFsZJ4LFYULykpjZcWx4rD4UAi2ckZMHMQEhTurnqiEREcRyIFy5cve+aZZ0CRIyMx0n/mYO/B+OBNNtlk9KjRHR0dpgSHBLTE5UQELsKR8BdffLF06bfpYmYPoHKyQKqFhYV/uu5PCCHmZ6PbFCZd3gYv6VeUl/2/J5+YMGFCZ2fC6gen6xB2vb27jjaj0aAbioSQK664Ytrtt5fEStJPusb6YZEGPfnMR5hPufbakpK4fXg9AXcypSNCyIoVKxZ+8lkkHE49Mpn5AFyu9o6OrbYauckmmyIkNlzvaxdpNL5MVp7lxaWUXnPNNffdd19paQlLuV9hF6rMpzhnlNLqmprfHfj7O++8k1JicIIbPCptIzWCFYlDQJtBqAms9jjCfK/+v7aswAX2ux132OnSSy+tb6jHEp2LBNEdHiL/k0gKmLWHBGMCMeJQJxwOFxQWhMPhYDAIXWG4L4yBtBzzfU8isFJ/l/OACVXzxeAADD4Jhj3LNgPnxcXFjz32WFVVlRmdu0HEWwB6CoSCu+y2azLpmRc1nAGxVGSDgm6gpqrmnbcX9Jwch/nV8M2EEMbYLrvscs4559TV1cuSsklt5S1Ww80FJoIxVlFRPnPmC0cdddSqVaukQIeE1el2bFpWGuOfedz26Zs7CzIODQ0NEyZMmDFjRnFxjHcJGZGG2kO1uamp4ZTxp+2zz742nKcnLAOTL8S77767as3qsByzbxcGzKL1/OROO+0YCAQY8/WLG+B+ZaGoVCcIQGacU0KmTp161513lpbEkSBYUM1toDbcRMrC4OaW5u222/6uO+4qLCyUrQ2a40/0L9wIUH5Nhw0CK2scYf7O/e+W1dVjjJ133sSjjzm6tqbGcWgqweAa/ZrKhjGWnthIuaodW5YK7dad+cHgWNPMAgSgZ7s8CBMJsQK16sGf4XDku+++e+qpp6CUZ+dbuRlv2bAaIcRvfvOboqIiz/Oz6sCwbXEh3KDzyiuvQJ6n04WeaqyaI+ScT7rkkp132bmhodF1YYwuU1rc6fOgCBHf45WVle+//8+DDzr49ddfdxzHwLJMTo/NXvwzLCtNh1/hOM5nn312xBFHzJo5s6K8XKTitvT71Z6OAKLFHJc2NjbtvMuuV111lanx9hyWx3wnpSkHM2/+fIfSrPeYFe/7fkFBZNddd9HhCrZwoxuSCYt3pJRNhGA+cxzn/vvuv+H66yoqytJLQVEU06sFFLzb2js3Gzr0wb8+WFZe5vvJrF5GfjPPQSOKaC+UnKHWpZPalDlYhdygLL3i1WXElDq33nrrHnvuUVVdRR2c8h9cEMyQuuywgVCE0vMvbAdjMwKtv2I9tlxv2+k0gsA2rt8Dyh6p38F8PxgMPfboY+vWriPUEuUQAiGcIwQk20ykDw517NhfDxo8KJlM2NcEacVUjkQ0WvTBBx989dVX+lr1CNMnKx6SIwjDt0+7vTgW6+jsVHBWKD5DoKX6PfIWJFlZSUlV1dqTTjjpppv+3N7eDhRhVXZSs64ySsfmXIDV/eMCzFk6SrJqwqDs/PBf/3rIHw5Z9MUX8ZK4l/Tss+gahHV0dJRXlN1++7SCggJAk3U9pG65mPa3wUVYtmzZuwsWRKPRTFqRQt4JIZLJ5KBBg3fYcUdZhDB+esPbsDDSQEKsNN04547rzJg+46qrriopLuFMYFuRUD+mELc5jtPe1l5aWvrwww8PHjRYVjgcG0+Qs4H1L9yUUoHKm8BwOhpTlr9x/6tltay44NFo9OFHH/3tb3+7rmpdIBh0qIs4IKQUuUI9Nhh1lbS1U0B704dMFwIpkMkylDLzn8FJaqE+VFxc/P2KFTPuv59I7IaFdlVLIKe2MftKMsZDodDOO+/c1taWHdfra0UIbW1tmzVrpg1a7rkzsgvRI0eOvONOpS0KEJh04dQaPoNxKgKLRMLRaOG0qVMPPvjg1//xOpXGOGdyC0YmeYd7rXw4FDTEekmC9m5rnBn0Ux3HWfT5F0cfedQlky5BQsSKimGshb260n9KZ0Ao5VzcPm3qFlts4fvAezHV7O7MgLOcBFyx2bNnNzY0AgsgzZbG6ZCqs7Nzl513jUQKUtcqvU1teBuWMHpkEFZx7jrOjOnTJ19zTTxeIsUmhVWlAPVQJZBOHZpIJiOFkfvvu3/EiBHQX0CmhrYhtJZ+sSYZMRlKWMiAsOTOkb9zP8vU1q+ZuJRgxlhJLPbY40+cd975Lc3Nbe0d1A3KAFjobFWloyLTusJP1ueGiZJeUvRNrhKmjI+DoALlnMVisUcfe+yzTz9zHVdBgYw/7+78pntMObDUse299942Tzpj75bDkwKhwNy58+rr602fWwuCdfdBWRsc+OADDzzw5r/c3NzcDN1SfW0lrEaCa+TsBi4I5j5ijJXG4199+dVxxx538vhTPvn0U4dSV/K4uewMp1uearIkzlghXQBWWWUScOorVqy48vIrUm7+9ddjsRghxGcsC0uVddMxxvX19df+acr++4+D7DlzQ+/ma2j3VCil7e3ts2fPDgQCmS5EOh6kUVoO3m///cyJ68WwAexXWfUJuJ/wCufMcZwHps+48qqrIpFIavsgApq5ErCD1ZRSoZQGfN9ngk+7885dd9sVvK+S0EovkLzlqNEpk6dooId8thUVggvdLLRoKRvAss41UxsK1omlzuQCgcA+++yzxRZbLPx84YoVK0jKkxB4cvQ+YvrHKOvarxfQqAc2CPNbf4zzR6R4LOfUJR0dHW+/veD3v/t9cXExIKLBUxCdj+dO+AVMHrMXl5aWvfbqqzW1NcFAUJ0/xlzmBcCzDgVD369cuc0226TSAsYoIahnOoRdUaac87G/Huu67vx5r4ZCQQdEjGVKQqzuvhzlQwlCPueOS13X+eLzRTNnvrhs2bLS0tJNNtkEIgwAuHLGUjuxJK0hAqEWUEyENftBMMYgDjD+7Ouvv7733nuvuPzyN99803GccDisdDr1Vc0gguvbTSmtq6u78KILL550CYhe2StBzwjpkWsInfu3315w1113FRdFdb1dcWHhDhJC2tvbhw0beulllwWDQZ3v5dBy/XGzI2lL0V1pgj7zzDOXX3FFtKAYIS5PVo4oxYggGDWoyiFwrk1NTbfeeuuRRxwB+Lh0vILy6VOuG51y7bVy+JEU5NB1S73I8xyybrD0zoIzeHuc8xEjRhx88MGO43y3/Luq6nW+74eCIUwEIOAsWogWnMysShNpgFWB3l5W1ydrS01Xs7Wb5lwEQ4F169b867339z/ggGi0UKFs4P1dvqEPTajiDAQHqSyzoKBg6b+XfvCvDyKRCANlTemR4N1y5WIvkWjraD/88MORjSju7vPJKk7Ai4yxXXfdFQvx99dfjxYXQSGC6Oa+LiZjNXhSwh4pdQoiEYzxpx9/MmvWrLfferuxobEwWhiJRILBIJE5MYEzFILDjGSku9vwrOpVkfS9ld9//8r8V6dNm3bLzX/5vzf+4VK3uLgY/jVLkcM+EUjWCSGNjY3nnnvulD9dawbY2T6j2zNg2yCxu+GGG75ZvLigoAB1WcNcvqGlpeWoo48eN+53UGDfsHaq9eqigPd94YUXzz9/YjAQpPpmyaqzjLkktR+iH3C1DQ0N11577RlnngG5b8Y9whvM1fjFGvYEIxrfoLpUGUpYpnEockmVYcM0hZtIMwekSE0qN6quqn722b/NmjVryZJvEolO1wHGEegHY0IxbLTQGWTMkI4A1MqR3nkdx3FdNxAIUEKSySQXwrGEC9IpGuhSSw1lLrjjkJramr322veJJ58Mh0NQoUY9LMz0v5qCNDG5577/3vuHHXZYtLDQZyyVFnKEqBKjxwpPKpLJxIuzZm2//fY+Y0SPE+ihM8ogEemu8B3Tpl1//Q2FBQXBQECo9wgsMCdCqqiQtGCWrqM7jpNMJjs7OhLJZEFBwYiRI7fddtutt9566NChgwZvWlZWHgwGu/7q5ubmlStXLl26dPHixR999NHnCxfW19c51JUcNtf3uX1sme5KmJGXsCybW5onTbrkj1f/Ed5MMEG4Z7VBkJakTqVxlCz+8qtx436HZSJuFmGW1HMymZw1+6WxY8eq13UKsUGYqUCbthFjzHXd2bNnn33WWQ6ljuMKgaRcmuQyYC36LsVlGfMxJk3NzRdeeME1kycbDc40Mz4XH968ZZtSwoJkV6GxMNEj8WA8MLEccN5+rnXt0kHYK+Gm7V99uejvr7++4K0Fq1evrq2t6+jo8DxPbjqp/CgQCBZEIkXFReVl5SXxeHl5RSwWI1IMur6hvq6ubs2aNatWrWqor/eSyWAwVBSNOq7j+wwD5Zunh6UYHKUkf/NgMLhubfVxJxx39913C5VOEWha5shjrME3EDcIQDNx3z/siCM++tcHhdEo4MiAW61pNNyhTkN9/ZFHH33Pvff0Gh+j68195ulnLrv0UsmZKTDykwgxGeWaugXUOYgl/SgRZ5wlOjvb2ztcxw0XhIuKoqWlZfFYvKi4KBQKeb7PfL+tva21ra2xvrG2rqa1pTWRTAbclN8NBJzU3WUZbP6sfN06Wk4I7ejowBjdcMMNJ48fD6gruzrd07GLqj8TesUVlz8w/YF4WalpkxrpK+gQNzU17bHnHk89/bTKEUGGMQcW6o9Ypo6pAeUD2Tq1Tl58cebEiedRQoOBgCz7U6jayFH7XNHs5IIhDq6prp14/gXXXXcdNAjsCDuXL0LebMPMZ6beiZDW9ENZmsDG++bdcDebeWZMNiz3IH/1yjXV1VWNTY2NDU1tHa2u6xQWFhUVFZdJKykpgXnpWdbZ2blixYqlS5d++MEHCxYsWLL4606vs6go5jgOY0loLyDTV1az+iVtJ+W2aE1N1cWXXDp58mSIr3tn5/2JxqVr0sRpjAVn0rc99dST5507say0zPd9gZFDaMrtqequXMYcJX1v3vx5o0ePNoznnt6kROZoGkLIOwsWnHvOud9//31FRYXv+6kDgOfOwjmarqrJYFIHTETKxSCKBPaYJ3vBLPUNeiSQVCTF0IyAWR1EzzjCKt3GKFPFwj7I9A+YN9Q3Ddxk4F133bHnXvvYiXJPpFNdI1FDSV/67dIDfru/6Wh2/SwhpL6+/v7p0486+iif+VTLTfT5Kv3pJpTqPgfemeM4Tz/99MUXX+wQGgwGjZOW6pIIkM8y8Ez5YoJJbW3NxAsn/mnKdXD5Nqzye96MYTVTDMFOADkGlnMZEEago6tBd4jqbTBvP9eyNP2zSCOml/MjZjMjzffYEkWJzsRHH3744qzn5897tapqXWlpSSrjTX0qW9YAHmyOCMWivq7hhhtvOvucs7NwN33+bKdbp0ifsbyG9fX1vxs3btXKVYFAAJwQ54Ijrh0wppTW1NYec+wx9913nz1CoMcPOBOTRSn97rvvLr300tdeebW4JBYMBBFnchsmBBmGp8LXZE3TU117rmiD0h/LpEiKqekJ0yJLkySNBtDsb/0GxeNJOy2CPd9vaWree5+9b7n11qFDhyoqi7U4e9q92UXUq//4x/vuube0rIxpnLZ1KVJ/bWtrHTx48PxXX43FYlAh51jkfgF6vX1fiLyffurpCy+4ICjNOmtspnTAF3DBKHFq6+rOPvvMm/58s6nb50iUnLf/1uiUa6/FSqpf3XIJussCYWG8Xsph3v5Xy2IT4TR7lditON7FbEE+43EBWQPfbN7pOu6gIYP33/+A3+y2e31d/eLFX2NMQXMj7VbT+7fcoykKBcOvzJ8/aMjgMVuPyRId7NtnO40L07QuSBkLCgpqqmvefOONSCSC5cI1mHO4xkLwQDD49ddf77n7HgMGDjC+rXvYqxlhQRolYYmTqdvq+348Hv/DH/6AMfnwXx8kE4lQMMhB7VF7YNOI1UdoVojOb+B/oLfCM0ksFiZZIe260E9MPgxKDhIBJiilHR2dGPPzL7zwtqlTy8vLfd+HYkzW4lQrs8vJ/o+Xros3Ah/87bJlf7zyKpAG63qP4MXm5qbTTj/tt/vtp3N3RYo2sylz1szTDek+V1Wcpy6+6KJwOOS6AS6YHNQNsUjGpzgXRAqIjh8//s9/+YvkK8IWbTUUfmBB5i03TYrEG/gH1/UqjvXTzDPvYP5Wdr9Z5Ues2TQKUaUn/1L7ZzvTtdmE8D2qCCnTKqhVjvnVmEcff+zPf76ZYNLR0eFQakgdSgMcITUQiSFMcEFBwUUXXDB/3jzHcTzPM8fZt4RCo/gkLDkLEBU55phj4qVxz/PgohAYbqHiytT7HMfpaG+/88477cZnt5xORlhgv4KxfV+gbck5DwaDf7z6j3979m8jRoysqqn2PM91XSSyayGw/+rzFDLiwEAThP+QxtQDVV/GyFyANqFBRevYeb3kYEwQcRzOeX1Dw6jRo5/527NXXXVVMBiAhMzOqNLXSsuzdGMwnnUXCCEPTJ9eVVUVCAS6krsgykgmveJYyZFHHS2UEpjQFdq+r9P8iKkFnCYgKO87Y8YD5513biDgUKLmGWtxJJ35qFEjqUy3vqH++BNOuOXWWy26kRDCavDrP/MJ0wZhMKwOpUkaisEtsbUIdVG6y3O6e8pgo8PZ8huiK2E/U/0ZZ1AJMzGuhBAqYaXM908ef/JTTz8VjUaTiYQtcmk44HAU4CRCgdBZZ571zjvvBBQYBGc5ld73x1k7izkkxthmQzc79NBDW1tblWiw6XDLRZza7BgriZW89tprIDVlD637mUeVERZYr6hjtihhJunxfX/X3XZ7ee7Lf7ruuoKCgurqapgiYD4iDPtAdW51fUph10EoVosGI566d6kXORYgYQpAcZUUZwF/oFPBmJ/0Eg2N9bHikuuvv37WSzN32WVX3/flLm/Y/yKLdyR0BvyjCpg/1bJKplCl//TTT59/7vnS0lLAHsJFM++B97c0Nx1wwAHDhg2D44QgpEcEVrrVdNEprU3mOM5NN930x6uuKopGCaFqxrHajTNlNDByA25tXe1xxx03bdo0aAvC7p3V/f3xBZm3XDPsc4ZRugelQVgSN4tpvoDRi2ZgRhwhByEmX/zf++5ZOCDGmeu4b7/91sknncyZCAZDgvuIcDlfhXApr0QAtCQYpaS1pSNWUvzc88+PGjXKrkXnSJIhdDYBPbBFXyz6wx/+YLJPbCIaIknU8sXm5ubtd9zhueeeAw5P2q/8jM6ZuWdCI5h/6G6ZXqwENnOIFb77bvkjf33k6Wf+X11tQzQaBfQNF77s72K7v5vlsfQrIMaBUt6Xk9TZEvlZTgTmmt+SPgBwZh2d7e3tbeVl/Y4+9sjTT5+w6aaDJAeGEyJAh/yHTuKnn+yP2A+hH+CHo48++u0334rH44lEQslumPdwNaY8kex4cebMHXbcKRUaUpJau1IdOtcAKgDY1xOVYSQZJ4hwoeK/K6+88t577+1f2S91yxm8yIHcnTLZ15aLBjuOs3bd2uNPOG7atDtNbUBLJGUs3W65R3nrNctUwpJKAUg9JSivhNXrBleY2Jme9fp/YVmJsql6bbbZZpsMGPjSrJluwEWEpDYuqe6r92mpNSU5wK7jNjY2Lnh7wf4H7B+LxYw+g72n96Uz1idI5Cy//v37f7Pkmw8/+EBOYVN9QdjaqUQuMS4ikfCSb5ZsPnyLrbce3Y0hxY8X/TLQMaYyId/CGIvHS/fae6/fH3iQ6zrLli5btWqV53mUOqCKA0OKsjMYrOeXwROq9aHhnHEqhoJZ3shIkCo6E2Otra0tLS39Bww8ZfwpU2+fesghhxYXF0vlLOieZ8HT1nNZfn6FMwv9AC8CoH32Sy/dMe2OkpISw6vJ+KAEOjQ2Ne5/wAHnTZzocx8ThV8TGFnST7liWDP+oXWgV6zC1E255pr77r+voqwccn0pfwQz2YnZB0DwjBKnqrrqyKOOvPPOu0OhIBecYipQeiZ/VgSZr0JvQJbKgInZsbSrtREceXDdBmrrTTWgw3f3XXdPueaasrIyCYDm6QYE0n/IkcSuG6irrR89Zuunn3mmsrIyy2n17argqk2KNHOUfPbppwcdeJAboA51YfyQKsMR+D9MMG1raxuy2WYvz3k5FovB9/SoLodt6y13y7BGwOCaVStXzZk798UXXli06PMOYP1GIqqdb32JqVEjrV0IsDSQLFSSdtLACSc9r62tzff94ljx6NGjDz/88APGjevfvz8sBnDa5vt74bauR7ZTnk5zU/O4ceNWLP8OQqj1fFKO8PKY97dnn9tl550ZY9ShOsXMxe1JZcCpPZXLIhPmUg9VCHH11Vffe8/dFRUVnGUwwTKgc1Id2nUD1dVVBx/8h+kzZgSDASYElfdLYJT7qO+8/Uej11w7BSs2qFoACnWJ06SI3Ck8btRmAxizfvhfLn5GqqHDYEww43znnXfq6Oz4v3/8X2E0YuCzqexJfxC2A8ZYpCCydOm3H3340e9+Nw5EJLL6zfbv6mXDeswT9FYHDBy4ZMmSDz/4IBothgEGCoKEFCga2tvfS+XtPffaq1uS4J8OOs1q2JtXCCZcCO6zWElsu+22O/GkE3fbbbcBAwf6jDU2NjQ1NjHGksmkz3wFd3colYxfkCmkxIH/l2qRBMDznud1dna2trYmksmSkpKddt75xBNPvPLKKy6eNGnbbbeNRqNARO7KsoW//VC5q7sQtlgPqVYta84dOaZz9qyXTGCUiQZP/ey6bm1t7e9///tzzz1XXz2C7NJCX7XLujy44HhtTDyg8R1Kk0nvoosufPDBhyorKpjP0lUlSfvUUVPqFjDBOOeNjY0HH3zwffffHw6HheAE8mNd8egaLeVR0BuWyR6wQFqB0lbCEuph7OtcJ28/37hQLA2u2mmCOnTShRc99vgjlRWVnQkPNO4hsUI4DSzlwk9tfDV1u+6y6/975qloNNqbVNr/bCYJlj2T5Uu/PeTQQ9va2uTpqPabLXABSWfSS744a+YO2+/Qyzl915pEmjeixwCb0msikVi5cuWiLxZ9/vnCxV8tXv7d8oaGhtaW1raOdsH1dC1FGiRSzgw5lLquWxiNlsRimw4atPXo0TvustOYrcf0798fvtPQ2Ozeau8/4IATU3gH2Q7/1/vvH3rooa7jYoxd181CyZkfPM+bOWvW2F+PTd04WW1ffyO0j8w0fSHoQxYwzPdTz1FjY+OZZ5712quvlJdVSuQdz6Daq6KNikwc16muqj7muGOmTbtTkoN9QhwjIo7y0oQbhWHGONZa/0rnTKriAbqwrw8vb91jGaP6AEErEGds0sWTnnzy8bLycu6rofGCYC6hIohzWABcsGAgsG5t1Z577/34E4/bPthY72+AtjS5lIvCvpfa4+656+6rr766sqJCopkkyY7grIpmS2vL1mPGzJo1KxKJ2GisvjFL/gppRX6gk9nvaqivX7duXXV1dVV1dV1tXWNjQ0dnp5dMCoyDjhsKh6LFReWl5SmrKK+srKyoqLA/bkQl+zxyysIGCiHa2toOPeSQRV8sKiwstPXgUJr/yh3Hqa6uPuvss/78l79I74t1EVafSx+e0w9kwBKqzrFIJTkudZYtW3b22Wd/+OGHpSUliCE1WlIOx1J0KiI7Qkytxrq62jPPOuP6G25yHCqRhjSdJikWd942eEsrYYHv1fhSWwkr1xVW8/ZTTJW5BCAkU/+jhHR2dk447fRXXpkfLy3xkgzmIZruoiUNwYmDa6pq9t1vv4cffiTlgzkjOCc0OsyIXyHTpsbGxt//7nerVq4KBYMEY5Y5HYEQwjgPuO66desmXXbpNVdfbQUTGPWdjEMWtM0WRobuACH/3YHJQcICJidllXOz0spezoDtWJAJ7lLnuuuuu/Xmm/v36+/JwriRgbP1OD3fj5XE5s+f379/fyE46J0RaJrgvpbf0MhjAGMLrjp4ijrNBaV04WcLzzrr7CXfLImXlnDGzBUwsyUsHlHqX+vr6yeeP/G6626U6zrlfS08bH4r3ngMM84RjNNB9lOaB2FtLKYEnw0CEyHtqwQXhJLGxsZjjz76k08/KimJe0kO+7xcEALrmS3giB2X1FbX7fvb/R7660PRaKEkruQKNwm2MOCSPvfss+eefW4sFuOC/8hHWltbn3rm6b333ltjbnGO7G7ZEs2ZLHB4T9fjtP9J95n73DWtxwwdC27Wm2++ecxRR0ciEaXyITvbWcxpx3Fq62pvmzp1/PjxOmCCgSKamNy3p5jWVBcZMxaEGnD09ttvn3PW2bU1tUXFxb7vpQJXYVoIGvCMEWN+KqIQpLGp4eJJk6666qrUAhaYUILScqI9PpMqb71p2BeMgPgGAlK7pA9yNZQl3wPeaEzmtRYoRA4UZYw5lK5auerIIw7//vsV4UhEcIkKkeTaVEzP9eyG1OeZ67hr1q3d/4AD/vrQI0VFUcaY0bXoq6TKyqgkiYoLztmxxxyz4O23imMlMAgonQSr6iCnDm1taes/oP/cefP79QOAN01L0mRs6l3Bcb10Xl11g+3UsKtlg+OwpbHTR4/w+tve2vvW19fv99v9Vn7/fTQa5ZpzBSvKRhG2trZut/12zz3/vOu64J6lVC5MzMR9CTQC1yvFT1XrVuftnKtJEs8+++zFF12EEQ0FAz7zU8cvR8xJ1idHKvIDvyp85rW3dV59zdUXXHABhzHtiJhVlx9Jt/EZUbqyGi+KNMGhb3UH89ZtpuVwJGNUUhJhu4AMWGoUb7LpJg/+9ZFYUay9vc0NUM4ZJjLeFlr0ECkFn0Qy0a+y8u+vvXbqqae2t7dLTaXsmQ293Gi0ccVgjutefc3V4UjE85K2erYhikK2UVxcvEJOR/B9pkUNNeg/Y5uzeZW912y0kdJdVSGzNKrsjDmtkiZw79+OHzkLeMWk8pzzyy+/fOnSpUVFRb7vw5j9NHfWUsgKBoPXXDM5FApp+rcq6vQ90VWtFWjhYaUNIxDjDM75lr/cfM7Z5wQDgXA4xPSQFTnYCinCiQxvheCEpvKgjo7EjTfdeMEFF/i+n3o8MczDEJrKkPe+G5sRzebQkEpg8YOgLs4J1EbefpZZ2nRYeWI951f+C6XU9/3Ro0c+/uSTRYXFLS2tjusI1SoGmT9uIJcEO8xj/Soq/v7qqyeffHJnZ4dqG2fqZfamUKU1nQLwvZj5/q9+te0555xbW1OXBWUSevQ1wa7n+fGy+OzZs2699RZHqiKvVwE0U92vz6T9fsQf2043w0/nQBjddWHAX6UG8oxnnnmmvKzMh6oywE0wyjqX5pbG008/fYcdd9DTCYHorIbE9O3Z6TavUu6EIgyTdCPO2MSJE2+86cbioiLQ/iS6H48JpnrYJ1wWx3GSntfS1nLrrbeeccYZMIrKAt9g0fd19rz1iGGAFGreEQYEfR6EtfHYD7CLrXqZFN2Q6NN33n33hBNO8BLJSDgMEo/A8AGlISnnw0HyQQheW1vz+wMPfuCBBwsLCpngpMsK6Z3K53ol+4UQ7e3thx1y6KJFClurc8LUFQASFecwalM0t7TcP/3+ww8/nLEkwY6cBWbvd9CPYTkl7bfe0m5fH9R6rOtxwkp74403jjvu2FAglIYRyI1HDsvUcBQhmpubR4zaavbLc6PRqECIEmK6ZbnAcRVcWKzb1KkyCXiuq6u76OKLn/vbs5UVFQDqlv1qLsXJiBqtkYY0Y9/3k17y5ltuOeGEE+ySUi7f2bx1i2HGQUrVkt7Ig7A2dlO9hkyDnfHVV189dfyplBAQTE73I0hqv0ltHtIb+74XDLrV1bX77rPPgw8/FCsu6XN+sL3XQ4vxX++/f8ThRwQCgbTSUCqaVKO+IKsgEgqOCf7bs8/utNNOZg6uMWuvz4lNfwMyrWSNTC4Oa+yLL7446sgj21rbgoEg4+nZt7pljcy0rmQy8f+eemr3PfYww/kVaStHboJIs3elVrWglC5ZsuS888776MOP+lVWJhIJlcBYb9eIrdQ5uq7b0dGBEJp2x7TDDj/c95OUurobQnLmPPPWUyaBNpD36vlnepB5/sZvFNa1gCp0KzTz2aaUep63//77P/DAjFRInkxigriQCod69g4XSnyKEsf3eXlZ2WuvvXbi8SdUV1dRqclsI3Z7pxBt/yJ7DJTv+zvutNOll11WV1cHkYGEuQi12LEiB/uMuY7jJb0zJkxYvvw7qEXDN+vGKtZuoW/7jRmWNcOgz0vNWaYHbcFoNWHGZjiOs3LlytNOHd/Q0BiJRHzup8lsStsJg8YopbSurm7ixPN332MPIDGDWomSTMW9er7rvdpCjQGEUj+HFu+CBQsOO/TQTz7+pKysLBXYqXEZStNIjZZUWjicuk5jY2NBYeETTz4pva8vRUmx5qHk1j3NW08YKGGl0c4ICaKUsPR+k8+AN3azH3WQ7Jk3Z+748ePDkWAwGPI8jxIqYzIG899MgYQL5jpuXV396DFbP/7444MGDeqaQfZyTgxrFZwoVNFPOuHE+fPnx+NxwGwLzhHB6TREpmWBQKC+vn7U1ls/++yzFRXl5izS8+H0qeSTkp9o+qqpiAdqEjU1NUcdddSXiz4viZV6ngfZodDRjfwIESJ1m5qamn6z227/7+mnAsGASIWA4Ke56v6KPhaiyCq3AJ33kUcemTJ5MhIoFAqZIZ7KdHtPYvBTfwkE3JraupFbjbjnvnvHjBmTWnKEYqVHh/v6/PLWS4Y5Y8LeIpWwuTCgrLxtxNZ1VLt0q9yhzvy58yaef24imYiEC5jPtVRPdlSe8sGuU1/XOGjIoIcffnibbX4l8xUKjjpjFlAvnpT5gRCyZvWa340bV11dDbpX67WUe3BobV3tbrvt/sSTT8aKizjzMAFJJhj4l49Bf4JZ1WYYagiXDMYvNjU1nXDCCe8ueKe8otTzGFZMR4EISSW0XOW10BQoLCx8ec7Lmw8fbmQpkSF193ojvivOIKvZ4fv+lClTpt8/PVpYaKQ07Y8TKJ4rxichhNbWVO+1zz733Htv//79IOBT/WEzGCUf7f0CjBh+mWpkYADAco76ANGat142G+VuELYEE9/3x/3+dw89+HDYDbe3thKqKJdE73uaaYoJpl6SxWKx1atWHnvMsQsWLHAcR6pAiqyycC+fFPzAGBswcMC0O+9AOK2EbL/TqBExn5WXli94++0zz5jQ2tZGqMv0XFYzTBAha3Jj3rqaKrXK0r32I1wISkhra+uZZ5yx4K23y8rKfF9QDFA37W4kyEDIe4QI9ph/y9TbNh8+3OSRuj0mdF2uz+6CgS4LIbisqSxfvvyYY46ZMX1GUTRq0+Ktj8imHpEio4JyxGpraw4/8qhHHn2kf/9+0BcH4RQz+yan+h156zmjU66dAjUPDOBQkONAaWGsPBNpo7eubinlkJg/dNiwocOGzZ0zl/t+IBTkjKvxNen1gE3zIhwOtbe2v/TS7KFDh44YMSKLH9zLS8jGkTLGhg0bFgqH58+bF4lEDLp7PZNkuIgWFy5c+Nk3Xy8ZN+53wUBAp1/WOecfhB+2NLQe4OYY+Zw5lHYmOs8848x5c+fZQy2hOYq18GnqgzJsa2pumDTpktNOPVUikhwQQMaYKwISMOL6IqSzyVEQzDmO8+abb5500kmfL1xYGi/tulWmHxPQhSbE8/z29tZzzj33tttuC4fkiAXdsrHUKPMTjX4pRidfO4XASGt4EOR9FlqUz0jZ5R3wRmzZtxgG3RPqM7bVVlttvvmwufPmqQaqkLuhrMpm9kexECwYdJPJ5IsvvlBaVvrrX2+XpWfUm5yZrGohY2zHHXdcuXLl+++9F4/HNahH/atSQZexBWcsGi389JNPFn355W9/u18kIulYUILOmACB88joLAO2mjBIZjXpyGltaz1t/Glz58yprKyE4fMoDS5Po4MFF8FQsLa29pBD/nDzX27WSyxN9k0LpPTe5MGutWf1KvB0Z0yfcf7EiR3t7SWxEj1XP0OURhG1JfbecRzP833fu+766y+59FJCCBecUqO1kNZ5tpVf8onwxm108pQpaQE+MJzmXqRrInkHvFFbVoxPpDeCwuyIkSO33HKLWTNnc+EHg0GZQWK9zSoXJEnA2Ge+Q2nAdV6a/RJjfPfddzdIqKwEoqdL01lZCPzevfbe6/PPv/jmm28ikUiWD9azaDBnnDEWKy5auHDhRx99vO+++8L0XHUKRGnGKQFCUyfKPxzgb3D6nvq+7zhOU1PTKSef8tqrr4L3VarHev0YeTLBeSAQqKur23333R948MFQKAT9e/lNoLwhaTl6a+pNOTKBMtwqRKId7R2XXXLp1NumRiKRYDBojjbL+yKtreYGAi0trQUFhffed+8xxxwL85iNogjEGfkM+Bdo2OdMDSDUHtfoBuerbb8cW3+or6kXlNJ//P31004f73leYWHUS/oOIVwuGiE3SZX4gKheykvx+vqmU8afcsstt7iu23Xsbk+HdF1PB8AydXV1hx162NeLF4P4sDoqM7kmdUbpMQA1dbWjRo16+OGHt9pqKzPOT8UQKAPg+svcINezZqAALSEADqVrVq8ef8r49957r7y8XK0BgRxKhBz/g/RoAc4ZpSlXveWWW86cOau8spwxn0opcqnFKBCyabS9ea2FsJQ2zLPw1VdfTbro4vfff7+0tFQAac1SKMv6Ci6nbzU2Ng7abPCMB2Zs+6uxjHlSeJwgxDWzOeO89HQl9VKuiL/krQcsPY5QK2Gp0W55JaxfsmXtrYATeeutN884/YzmxgYYw4Cpm/K36fKZHG6KOHQzCMJ19fX77b//vffeGy+N2z64FxbVemGD4IMXL158xOGHN9Q3FBQW2KMa1uOwXdza3FYSL50+Y/ruu++e2dWGyY5y4yQI/1I3SPsm2mRZSunChQvPPuusfy/5d0lJCSR8MNoIGL2AOpHMbISJaG/rKCwsenHWC6NHb73eadO9tQXpYSXQnpZKVzJbF4wjRzZrZ8+efeWVV1atWRcvjYOENThIK3+FmYTIyL22tLZuO3bsvffeM2zYMM1wU/g+oyK6noMwp/8LDfB+EUag4YdMUC9SoamcPpqvPP9CLYs7BPup73l77LHnY48+WhIvaWlrJQ4FAqdaKRhEpuUIfIAzCVRRXvb631856qgjvvrqK1umY72JQvcazjR4EdQ5RowY8dBDDxVGo4nORJbbsGvjqRCBoWg0WldXd/xxxz/2+OOSKGIG9AKAN61p8osyG99u667AxaGUvjzn5aOOOHLZ0m/j8bjp+1rvh60GIRgH4jHqONMfmD569NaeD9JRSpej14dJwEAanCGvncpSUgm953k33njjqaee2lhfX1oOUUWaHmXIVxQp8DdPeW3e0NR42OGHPffcs8OGDfO8pGbJE11zzva++R7wL8qwJxgVdlWQSYVSnAuzzPLWV5YlrQx3HzQ6vlq06MSTjv/++5Wl8bLUHqTGzsoZrkRpVipoq0CE4obGpvLyynvvu3fPPfdkjFuTP7CF9UPdG+l3neVnutHQm/z7a38fP/4UjDA08Exe3jUPBshrfWPDeRMn/unaax3HUaxNCa7Rk11/cU9HRggl/wJCV77v33TTTXffdXdBJBIIBIzaRtfmKGSIXtLzmTfjgQcPOugg3/fkvCCi1KVwr84aBM0uFVsJJdTMOUMIU0pXrFhx8UUXv/raK/0q+yGB5HAFcKWy+qHnERIYdYxS6XJHoiPpeZdcctkll0ySa4kZTJkazo3Wz5LPZ8C/HKOTJ0+GcBTr+ZoYpXngeRrSL9Oy0kfYJlLba9Kr7N9/v/32e++995YvX1ZYWMg5I3Kopaq9GWcqa4uM8cLCwva29uee/1tBQcEOO+xg55pYDxnHiorSbbtM19lB5kXJsGLDhw8fvsUWL700mxIC2bld9rS9BXwkEgm/8cYbny/8fJdddonFYqlsHv4Jp3uZG+sz0nWgQprsC5wcoQYcrV69+owJE5548sl4SUnWVU1fXijRpoIzwpjfmUjeede0ww47QnbZHeOtexPhppJdgTV5SumEM+Y5TmqBvPjii2dMmLB48delJSXgkgkBRgAAZfRsBZDrwth1nKaW5sJo0Z133nnqqacyxqCMlAVFhOAznwH/kg1DYTAN988rYeVNml2ITmtLyX6w6zh1tXUXXXT+yy+/XF5eKZujkPrKrUs2trBGr8jB/5hzr6ml9aSTTr755pttWJb+cnsUSM9uN7aA0Zw5c84680zIg6HPbbJe+6wV6ZM6DY0N/QcMuOXWW8aNGwetcZk0I90y7APZr142Vdrogm5DCM2bO++KKy5fs2p1PF7KBENdprlkKZR5vteZSNx+++3HH3+8lVCa9/e234GRNBwLIjCX2s6u47S3t99www0PPvBAwA0UFhYCkLvr/RVKzkuqpWNcV1f36+22u/ueu0eOHGkvpw1igFXeetPolClTMCYmbVHMktQ6AknSjX9Pydt6LYu/q16UwT8TvLCg4PcHHlRVVf3uO+8WFBTIN3I5rxCryVpYeVQAtRDiRMKhd95d8MH7/9p5513i8ThnDBEiZX7lJ7UmUs+BmbImvTPJch4xYuTcOXOSyWQwGARMTbaIoAwUAEAUDoebm5pmzZrV2NS4/fbbh0IhgHGR1BO0cT4p2SItOtcDR8yYcBza3Nx8KusSywAARzRJREFU/Q3XX/3Hqzs7OoqLi1MBGcL21cv4QiSog72kl0gm77rrrmOPPVYGcCSr6tBrl9EKLjXkUIYUlJCPPv749NNOm/XizJKSkkAgoCBXJGNKujlHwMZ7ntfS3HzYEYc/+NCDUho9G1DWJ+eYt5w1OuXaa2UAb/pYcvqNWjS9D4LIW85ZupArt0loj0Hes//++7e1tb33z3cDgSDsQtjWCMSa2IYFZE1FRcVLlnwzd868IZsN2WLLLZmeO0Q0lBr3ZOJjdsy0/CRjW2655Zgx27z+97+3tbYGLEJnF06nQscGpL355pvvvPPOZpttNmSzISD0YVQYe4Fk1ZvWRaQFwTVkIuVaKCHv/fO9MyecMfOFF+PxEghiul5A81fGuFRw5AKhO+6866ijjoKGeteWRy+cmh2QmbUBnWzG2PTp0y+YeP7q1avLyksZ413J5cikLBK85ThOMplEGF1y6WU33XRTQUHE97k8M7JeVOBGs0Ly9nOMTr52sqkWynUhA3kuGSXS8sslbwbhqX2skDlfKlDbZ999HUr/8Y//o5Q41AGpLKSrcsgIGMnklnFWWFjY3Nz8wgvPBwOBnXfehWDMOAfdD1tZqxc0OmAr9H1/8+Gb77Lrrq+9+mpdfX1htNBLemagk8popUyWwGnMa1G0aMWKFS+88GJdfd3o0aOjRVFTu15P2QCb5k6uQ2t+qN1r9yCExLJRQluaW26++S+XX3bZ2rVrS0tLZZ2em/tms84MCC4UCra0tgQCwQcefODggw823tc+hl7TSss6QTAYmHjO2efcd9+9kUgkHAn7PjNlkez7q26soJQ0NzeXlpXec8+9J510EiGIcyaL6jirX54vQefNNjp5yhSiykpEKEQBEWrXRHkHnLdsTFY6+E9tJpyxXXfbbdCmg155Zb4QKOCGODhUjKwcARt9NcZZIOAGA+78+fO+Wrx4zz33LIhEPN+nVtrUsxodmYuZ4FRYMHDgwP323/9fH/zru+XfFRQUaKANTn8IhIj15B7GWDgUdihd8NZbc+bOjUajI0eOhMypq483gg6qPp/D+lk2eK3rsy9nCSgw0dw5c84488wXX3g+WlgQDqsrhhCXN5lksb+UHkUgUN/QsPnwzR977LHf/OY3kCsb1FVvZodde9JQ0SGEzJw5c8KE0z/+6KOystJU2MXXr6qmPi4BZYTQ+rr63ffY4+GHH95+++1935MnrdS7DERgvdc2b79wwz7nRIXlRGlxIJQHYeWtq6mETSisvBoOJFkXlJJ3Frx17jkT165dW1RcxBknlKrEmQCjBEq48BUMc0QDbm1t3YitRt55911jx46FBhvQe+xf2u25AkfIpoDAi4DAqqurO//88+fMeTleEkNS7le5Ty6wmt2T3rJTP7NULtjR2dna1rbX3ntdeeWVO+60E/C1TKfQnAV8HGuKVu5nwFmvGxzvBx98cMstN7/5xj8oCRQURrxkkhCHOnqCC9RHrFQPIGyEkOqamn1/u+/9993Xr39/mfsS1OtiAzbbW1fFGfQjmpqabrj++oceeigUDP5/9t4EXo6qzPs/51T1dvfsJAQkCyEKhDUCsgTzjiibStBhXph5RccFVMAFx3EZFVf0P+Po6OCC26ufd1TEhIgriw47jlkgEBJIgBASQta7d9/urjrn/6nnOXW6ermXm6Q7QPx9A0nfe3upqttdTz3PeZ7fL9eWLZdCDsl8DdHwyCil+vr7gjC85pqrPvrRj+VyOTKQ8JKdDAi0YAySSljsUc5VFShhgcbEWkG2YkzmlZITiLVrH73qqvetWrHykKnTjJZhPCNrKiJrVWcxz4vOX12dPZ/4xMcvf8fbOXQlnPBbuf3xlzIx7MvGrp/65L9897vf7ezo8HyfR4fj9799XE3yxJs6MDDQ3t5+8Vsufvd73jNv3jwO6lXjN+5iQsuXrMAgH4rk591pawgh1q1f993vfHfpsqX5weGOzjbP8/lCPbnsbaXC6S3hhrhGRorDwwOX/cM/XH/9l9vb26l73CMhxsa2VM2mRt+qEkRd6/vNNy/9/758/dNPP93d3U2bpJMXRzXVY76k8Dyvt6932iGHfP7zn1+yZIm7znCPerHaucHLCBlqbY2A7TW+dSI0dVMEL/KWgpcMlfNR4q0ShoHvp/p6e6+56urlv1p+yCFTjRaBDu1Qm52XtSdBraMTU2hC31OlUnl4KH/pZZd94l8+MW3qNG5oUrI6NMWnsaYljRXnvMoeccxQSv3g+9//zGeuKxWLnZ2dYRhyGFYerQPT49zpmLeH55HCMOzt650+fcaSi5dc/vbL5x05z40wJSY+rZOPPHCWPmMeBndZYasTteJWHHo3bNjw/R/8YOkvb965Y1dPd7fv+9oEwih7ZRWfHBKxTdOSVvRlIV/0U/JjH/v4le99nzEmDAPP810F5QCfWNxb111YbNmy9fovfemnP/1pJp3mX3e9WmoyAFthuCDYvWfPa05/zVe/+tWk+eYBsxsBBweshJXw7YhyXxV3xPB6GQIwGIX4io3fQKEOfc8Pg/Azn/70t7/zrVw2l06nOUhrI5QTeCGVXUEW5ezn4Hnejp0758yZ82//9m+Lzj6bz4+s3U+iDcatqSbP8fv2nqzPgJNE8VIIT6kHH3zwA9d8YMMTT7B9IZXHlaYdtp60iUomhysuO5dKpcGhwQkTJr7pTW+85O/+7pRTTuFnDoIy3SEx/Rz30La0Macme6uZLEp0X1YCjNHG820b2spVq37+s5/dsuyW3Tt39nR3e6lUqVRWSioVX0zZoTNSTjaxdrLQxkjlqR07ts+ePfs/vvGNs85axDFvNDPmVlBXN7ZXGtzqLIRYunTZZz796c2bN0+ZPJlV0hp2hCUbtVKpVD6fL5VL7/jHf/zkJz/Jzlqulay+Uo0zJxgDGeUoxurn0zuFAjBrscZRGQEYjEEyfYxOslJ6Sv3i5zd99KMfLRQKXV1ddF5TScVJvtKjAMyXfNpT6aGhAa31+9531Yc/ci0tpwWe8lwrYOsah2tOl5rwfX/3nj3//E8fvennP5s4cRI3WHFjDWtQugTOlakrl7G0sjg4OJjNZV9z+umXXXrpWWed1TNhgkuIOQ41VJVqKIFSo/LRMGC/oF7VGL87vuKJriE8xbWHwcHBu+6662c//dl99907PJTv6mj3fT/g7FDai3MrH6AldY/w5bpmtW8+XHv27Hnjm970xeu/cOiMmQk7qSrdkmb/MhvuoKt36DA0HCw3PbP5s9d95tZbb81lsul02lljVc31Juz3+deqPG/37l3Tpk777Oc/d/HFF1NCH3p1jvoAjBMZam3HS+I5EOeqjfcTeEEaNuxwgHl49UNXvveKxx5bO23qdB4ApfGMZDAVNq0VWodCedGZrnd336mnn3b99dcfd9xxHBgk6XXIugDcimTR3XYCTz/4/ve/8Pkv9Pb2Tpk2SQc1i4MVktHRHYFyuZzP54UQR8478oILL1yyZMlRRx3lXqvih5h4hvpwO84tH634Obb7BRuxGCW9uOb/1FNPLVu27De3/vqRNY8oT+VyOd/3azbM6ZdJKbTRSkoaQRJSamr79fr7+9va2j704Q+99/3v9xOalBQFQ/bBbekaeFIGvMamiWd8b/jPG7Y+t3XShIm8dlDfe1+fyBZLxf7+gfMuOO/zn/v8nDlz6svOGDQCe4vUoaZlKZo0N4oFsIR1mUMTFngBGiZeFMACz0s9v23bR//pI8uWLZ88ZZLv+WHIzX12ECO+M3UBGqllKIVU0tuzZ093d/eHr/3wFVdeyYVBj3qqJb0V3Uu3sFHLru8aqYX01Jo1az77mevuuuuubC6Ty2XDIJSyMihcI3VUE/D44zM0PBCGprOja+HCk88777zXnHHa3Lnz3ENYK7i+4Xb/M+Cqanv8fbcs7VI3jrv33X/fH37/hz8/8OCuXbtyuVx7e3vS9ajyezJkxEjnBmGcZ4s2Wnq+KpdLAwNDx59wwuc+97nXvOY1SenvymbEobgVcxb1DolRzI9rzvfff/8Xv/DFe+6+u6urK5fLlYNAUsNBspJRH3qllP2DA52dnR/60IeuvPJKrnC4RveGS8UIwGA8yFCHtiuEl6OkkDpex+F74J0ExqRmtMN9OwjCVCpVLpe+9IUv/ed/fjOTyeZyOR2G9IC4aillpU1a2OxSKVUqlfoH+l93zjmfue66V73yla7FtGHUae6O2OfUPBGggzDai6Bc/uGPfvi1r/779u3be3p6orRJB1R7bfzqyRM67ZHU2hRLxXxh2IRi6iHTFhy7YNHZi84888xZs2d1dnS6DXAJdL1OU91rsJ9Q/ZGv2SmdbKhMXi4MDg5u2rTprrvuuu/eex9+aM1z27am/FRnZyenvPXXOrwYH6eWsWkQr8sr40mvv7+vu6f7yve+/13veldHR0cYBNJTUrR2sHuUva4aMBNCPLN58398/eu/uOkXxWKhq6srDIwQpn5BOlF2NlFspjJGX1/f6Wee8cUvfnHBggU1ZWecIcH+IMth4ElVWZ6jJR1j7Pgm3l5gbEbLgDm9idIi4Skllt689OMf/2caOuoS2oTJ5CduB5ZRwmmc/IuUctfu3RMmTvjgBz94xRVXWB/AeFXYiRInPQ2b1CAd95bFAmCuO3r9uvWf/vSn//THO9PpTFt7Voe1Dx2t2Cs5Q5RCeVHWOFIqlEZKxWKxs7tr7pwjjz/+uIULX33sscfOmTsnl8vVH2Eda3aOkfcnC+B8g2Uuau5WKBQ2btz46KOPrl4VsXHjk319fSnfb8vl0pkMS4Nx+5sRDbt5KyGYXS14mno4P1woFM47/9xPfOKTxxxzLEc+Tynu0KoUhN0AW1UGnGyMG/U3Oc7s333HalQpr1Qs/vCHP/z3f//ajp3PT+yZJJW0atXR200l5TWSL6d1KKQcGBjo6OiM3oFXXsGOHTWXRMlfBzJgsLdwE5ZI2GJpKZU2sXog3klgn7D5o6DGYa0931+//rEPXPOBBx58YOqkaWRFF7qzrhsTTj6eVuz8YqnYN9B/5hlnfOpTn3o1dRSXy+WkfqENvlZCppnv0ng4x55YOZfSRi9ftvT6L31lw8bHJ06Y5PspLiCPrQLtklr+UnlCURE71MHQYL5QKPi+397RPnXatLlz586fP/+YY46ZN2/ehAkTJk2a1NbWts+7EB29Pb3btj//7OZnH3t07cOPrHnyySe3bds20D8gtOno6Mhms9E1gTCh1qJ6ebtWSKQuzkkhvZQXhkHvnr7Zs2df+0/XvuUtb6Ulg5DapJXLmCszVyYOr81uqKuZFHKyIbf/4bYvX//llStXtHd0dHS0BUFYUWGhReg6L2cuzvvloLint++MM8+67jOfOfHEE2vW7HFWBE1BhkYntaDdwhGUsMA+UHWOjidMY2vV9NDg0Je//KVv3fDtTCaTzWa5Bug0PezomxXYshIZ0Zk05Q/2D/ip1Nsuf9u11147ceJEzgir0jIuYzdyV937fWBJmqorAqqySqONjLbZ271797du+OYPvv+j3t7eCRMmcCNPjfpVPeScK4TxpDRU4UwumetiqRiWg1IYlEolpVRnZ+ekSZMmE0fOndvd1d09oWfKpMkdXZ09PT2pVCqTzaT9tDFaUsU+n88PDQ0NDAwMDQ3t2LGjt7d306ZNO3bu3LV9x9ZtW4eGhnUQ+ulUNp1J+b6fSqU8T1c3K1UUu5w11Si/XFpDMDo0/f0D3d3d7/jHy6947/umTJ5CHV0sshG3H1ctZ8W/rNrYu18ZcE31xXXPbdy48SvXf/mWW24RQnR0dETXT55PQ9wheWd6LDwUe+m7Jwlpzb7Q1p675gMfeO+V72WHCb7ma9ipvpdvLwAqSB1bi9tPDCcTAk1YYF+oOkvSH6eI5M6My5cv/9S/fOKZZ56lU7ZNLOzpjA0xo4fTmHAUCbSIQoxXLBbz+fzcuXM/8KEPXnLJJU7xkTwhpHVyFdKMa8bF1LYnidpTvxtfiausvOKqTKyCJIRYu3btN7/5zeXLl+fz+UkTJ6U8L3DZMFtA0ZhzdY5laP9o2ScxthTFY6VpBEhJpY1RQTkIgqBYKoUmVEaGOlCen/J86Xtp3/eUl/JT7PwYJZ1hUC6XdaiL5VIYhGEQ8HFIUahNp9O8DmpXdiv7Fh3k6AjTZQGVEKwJaYPDE8NZ4NDQkFLq3HPf8N73v++EE07kmjMV6mUjhQ1ju0uMlbwVVTok9nKLn348Abjq1xZ/0y337tmz53vf+96PfvDD559/3l0eRfck00hbLJEyTsVZrpkmuT3PhOHu3j3HHHvsddddt3jx4oTrM5qcQfNhJazEyGCsmg8lLLD/1Jwund3b1q1bP/6xj/1q+a+4ChoHaQ7bVrHDRgmb2grWQBgaGiqMFM47//xrr732uOOOqxdels1sr+U5Ge4Z9qzwJl2SctmZT/crVqy48cYbb/v9H/r7+traO7K5jI7d6+K0MnpcdNYneSg69/MCtqif4eJQxIExCnV0WHixWSqpdaBDoUNNbRuc/cZZoBRWuUQpX/m23Jsoydb2M1euMKwVZMOQW9GNMrwcHg4NDeay7a973d+858orTzn1lOhXUA6UklH00ibucTeiqqym48hqhPDqBpD4S15Rr/+pEKJaRrvmR9QVxb+LQqGwdNmyb3z964899tiEngnZbDagaxH+VWhBZvvUnCCoDZDWqDWV3lUqlRocHDTGvOvd7/7Qhz/U1dVVv9iBMyFoLrJkQo8d5mz5zp4kXBDG2w6Mn4aroTW5iw5Dj06X3/vejV/4wueHh4Z7untsJcbJLQsXpHgqXbMJMRvh7dy9o6Oj453vevd73vOeqVOnujTFxL4zorqDt+4NXLMgae84SgZsLwR4ekZUzwo7scZH1qz58Y9/snTp0t17drdlc+0dHVaikku5ZINUeQkuDJCcRzyA5GayODEVUSxT5AMRvbyyG8VxyFRcHqs3xpD/jkoe/2RnVt04k7AL55VKrHBqlPZpjfajHZRDw0OF4UJXT9cbzn392/7P5a85/XTWPhNGeNKLN6omPtYcZ37O+hy3kgEboyorGHXvH2m3ki6FWFlDh9Qep0ql0s033/yjH/5w1arVKd/r7OwKgqDuSQRNUUkauHTLI1p5qlwKe3t7jzv++M9+7rNnn302X9hxUD9gul3grxDbhGU1/6K3p5EKTVigaTTsCqZoIT1PPfromo98+Np77rtn+tTpyvOCIIwls2Qc76zMgxsf1dooT4RB2Nc3cMSsWe97/3svueSSjo5OVrBSFA7VC6wG61gIwiVbDTOzFyiNmgR8sn7qqadu+vlNS5f+8omNG7LpbEd7uwtK1BCe6KuIL26j3YmHYfgntL8qviHdYlBlh6pbi+NHyWSyWzM9XNOhHXdF2eqCqKxBVaI1lRWUMeFQPl8cKb7iFa944xvfeOmll84jLZEwDKNE3VPxdU9lHb469927AFwj0FG75ipisUttomMUJf2eEObWW3/zta99ddXKVdlMlvSctTGVGbDK9Ri3Pct4ojm6j6+UHBjsy2baLr/88g9+6INd3d22D3yUVmcAmoi0qyN2CNjE1q1owgLNp+YUz5XDQj7/ta9/9dvfvGGkONLVMyEMNPn6W1d/aZQ9G9uTJ/cxRRmlp1QhXxweHj7hxBPef9VVF110kY1AnFzWUHUSrSmN8ndeMDDY9dGG+hhuVEkI0dvb+9vf/vaXv7j5wT8/mC8UsulMe3u7VELrQEpPCcXzVnFy787yQjSShq5xXkrMB9vtUUo5F6Pk7tZ/2XhoOFF7pqhk2z4KxXxxpJRNZ49dcOxb3/qW8y64YPr06U5N04WoOLO0gZH2wEtc4iQvdOQLHufRMmAh3XgYXbJEkT/66Z/+9Mdvf/s7f7zjj0KY7u5uflMlJ4vqrz9Ia8io6DnE0FB0bXH++ed/9GP/fMwxxzg/LmQd4MAgg1CzBjRr3Giq/Ek7YS+RAYP9p8GYZjwSaoztzHrgvnuuu+6zDzz4557unmw2G4ZlqsAoYZRtSVakGGOUbdSVrGoZJWnDQ/lyubzo7EVXXX3NokWLqDSqudPJyPh8Hy9zJt7P+xKAG+9X7BLBLU68R8aYlStX/mr58v/+039veuapQn4kk01nszlfpQyP/RhN2+Nxfqx46We03ieipoZcP/zq+t3qf2p1rCqNT3QVI0P7Q0omw9CURkaKIwUh/cMOP2zx3/yvCy+84NWvPiWTybDhVSxiXSliuysB9oUS1mrRHbrxBGCWBQ/pLeHFSrhVu8y3qSys+IjdddddN37nu7ff8YcwNBN6JiT9H2svXIQ1KTasW21CJUVQCgeHh488at4HrvnAJZf8rYwFKRV5XeGMBw4McRNWzepN4qIYARg0HReAOelipcDBwcH/+NrXb7zxxsJIvru7Kzpthtz0VGnQt8Ua4wQceM1SGaH7+/s8lXrt/3rtNVdfc+ppp/H52hZSK5pRTHNK0MnSqEtn2RWIW775JQf6+x955OHf/e4Pf7zzzi1bnh0eLqTSfntbu+d5hhuDyN2frSniXaui4ZBxvRhFcqtqIrRbQ6XaLVeYFUuPKSWNDkdGyoWR4XQ6O3PmoWeddeZ551147LFHT5k6zXU5sR9UvDRe3+vm2qySx2rsEnTN3FGDMSQrShVdT9njaYS5/977b7jhhjvvvLNcKvVM7FLSDwOd9IisKbZX+vDpKkcpmR8eTmeyf/8P/3D1NVdPmTJF6yDUwve9ypUaTnrggCADEyiTEL4hSfV4HkTQhSPGkEBzSHTT2I5nKghSX3As77dyxYovfOFzd999d3tbezqdseVHYUdXqAyd8FGiszl5CBrf83SoB4b6c9m2C9/4pne+610nnHB8FIbDKLWy7cJ1W5RI4+oXMcc1HiMS6aTr3bbf1tGm+SnbyzM0OPjAgw/ce8+999xz14YnNg4MDChPZdLZXC6bTmdEfD3Mh6basJi/o2s8A+qryvWF63hHeCnZBnhtdFAOykGpXAq01rn23JFzjlx46sKzF732jDNO7+rq5ofWdJjHR0NWVxFq+pzHH4AT21fxMUwuTVdJagghfv/73//kxz/57z/9KZ/P9/T0kOdSIESlIG/nmIWpFhu3B8r3/ZFCoTBSPHvxa//pIx85eeHJUd4cBMpT/CbjdgOFcx04UMhAl6VJLufwB0CiCQu0jqoM2DW4mtAY4XleUC5///s3/uu//mt/f//ECZMMmyVVBtVF9fyMEHHtV0nppb2gHO7e09vV1XnB+Re858orFixYYOuTxkRJnIu4vCXxbVuedc9rw8YLjMeI2LsksUmxYUE8+EqtQ8ZZEAoh8vmhZ555ZsVfVqz4y8p16x/bsuXZ3t7+UrGUSnvZbNb3U57y2CPAcDjhoV1t5U1oMkkYYWoq0m6jEhNQLKuho6Q8NOWgWCoGQVBWnpzQM3n6odPnHTnvpJNOWvjqhXPnzu3s7IzzXc0aKdUO88naeE3C6r5Zc6zGHkOqSnnjWoLgCxde7OXWtnw+f9ttt33vxu89+OcHpTbO2DG5vzVHoKZg4HlesVgcGhqaP3/+VVdfveTiJalUKgzLVHKu/Z3W9XID0CpiKUo7BGyku7xFExY4MMQZKMVQWhVWQkhv48YNX/ny9Ut/eUsqlWpv72DbLqWE5PXCKCh59nqxrvoqpCmXy4P9Qx3dnW9+00VXXPGeY449pqaBiEnKhtiQZluhXCaWyICr5COIZGQR1SKL7p5xMh0bERnf99yzlEsjTz311JMbn1qzZs1j69ZufmbL1q1bh4aGSuViEAaCGnV9X9H8VSoKF75UQtGcsOAVS9cCZp+dKsZhEP3RJjBaKk+mUn5HR/fUqZNmz547f/5RRx/9qnlHvnLWnFnt7e1uV3i4mdPNRvVtw03RstJmJRLx1Yt3dTx9zm7B2HVdyeS1lNuMgYGBX9x003/9v/9auXKl53nd3d3sJ1jXZiXi1Lfya9XaaBN6ymNv5ilTprz9He9457veOWnSJGF0qEPP842Jc3oR68C01iYRgCokr+7Y1ScjjYQSFjhAmFpNKjd7Wmll+uXNN3/ly19Zt27tpMmTPc/TOlTSj++skmkZnYg11xHp58rzvWKpUMgXOzraz3n9OW+//B0LT3k137levsOeyrXhWqRNt2tSPSNrpRLrNbWcAKerIDuJ18QUcdI2OPnhKhQKW7Zs2bRp09atW7ds2fLcc88NDQ3t2rVr9+7dQ4NDpXKpWCyOjIxwuNVh6Pup6DZN3Sil0ul0Jp3O5nIdHR0TJ06cOnXqoYceOm3atMMOO2zWrFmHHXYYtwq7l0sGs8ajt2N98E1dBly/rFur7xGX1oWomD2EQnh0GRH4fprvtmnTpqVLl958881PrH88lUp1dHQks16ZcHGg6XDpGsijU5hij3MV6LC/vz+Xy1100ZL3v/+qV77yKPckKtFZ4GowlUIIznbggBBlwMquX0EJC7yYJNpWWaDIylzs3rXrP77x9Z/8+P8OD+e7urqk8LUOY1MdTWpQqsYo372B+dQ8Uizk88NtuY6zFi162+VvO+uss7itNwgCCj0qMWEbD+jyYqCLHc2wEKiUu+Po7EaYXNpX718khCgWi8PDwyMjI6WSDcBBEGitefv5UX4qlfL9dCb6k8lk2oj6Z+P82BjBfgnx4WowVTyKesk4VsRf+DiYxBgVt4UZ1mYpFourV6/+xU2/+O1vfrP1ua1dXd1tbTlh7IVCMvRWt1uxUKhg8WltQqPN8FBBefJ157z+6quvXrhwoft1J/c6vgo0NS3iABwYZNmEipswbS+DU8LiNTqJAAwOGMl3GuteOUGi1atXXv+l6++88450OpPNZDzPF3Y0tLYx2GohxuILVKE0nueXS6WBwQHfTy189cJLL7309W84d2K8mmgMNRtKp3clqruCXGoXN/++YAbc+OPiFjtFUrK6ft43eTnygh4PY+BCu4xxB6quh0u6a4LRnqxuRVyOnQEn9qu2Ih+vOGi3dzt37vz1r399089/vmrV6jAo53Jt6XTaTVfXyblUno3PV3xUqWZnBoeGyqVg4cKTr7rqfRe+8SIeoKLn8arnsqxnV7X/+X5dYwGwV/AasKmo7dAbmkt43A+KAAwODA3V9p3Fje/7Wutf/eqWb3/rWytXrkinM7lsO0fgmhN0dQesFIIWWXQUWz1PaiMGBvq1Do+YNef8885980VLjj/+eH5gEAZRbGHLYSMSnREUmpuRBDfc37FUMqqPwxh3SOx+rX5T/Ye3WmGqYVPVaJ5F482Aq/fLuOxfa7bfkM406eGHH176y6W/uvVXz25+Np1KdXZ2unGp+gbvxhIl9B5QUhVLI8P54QULFrzrH9/95osvymVzrk06IYwFQwXwUkE6i2mLEW4mA5eB4MXCDoDGE5xJMyVS/b3p61/72vr1j3d19uRyWddaVX+ypi84S1ak+0gjSzIKAMWRcj6f7+zuPu20U5csWXLuuedW2oC1Vi5fjEUWk0/W3M/GaOpae/UkozkFNYzZ1S+nqRNYJ1qU93oqun6bq3eEE2Ju8krxHXbu3P673/1u+fJb//Ln/xkeHs5ms+3t7UYKHYRJubSG+XriybVSnlQynx8aGswfffTR73z3u9/y1os7OzpjWSu/plEAGS546eCasKjrSliJcmuUKRSuE8GLRcN8xYXhPXt23XjjjT/6/v/dtn3rhJ6J6XSGV/jqH1JzwwV1WneUQVkPDAxIT8ydPe91r3/d+eeffwrZ/nPId+LSbjmGRo5UvJR7oBOphtqKDWU69uqxrvVqzLGrUQNww46tpNWx08oWQvQP9P/lf/5y222333HHbVu2PCu07Ohs8/20DkLjZnkb/fadH77NjLVRXnQ5NDxUKI6MHH3Mqy677H9f8r8vmzhxYmzOIRMyn7WbijMbeCkg2c+EXEYat97jbQoOPDXpTk0NVscTok9ufPKGG76xfPmtfb19XV1dfkrpUFbrVzSAZSypbSeUwmM9y0JhpL9vsKen58STTjr/gvNf/4Y3HH7YYXx/lvJg42ElpE6MyLdu3/chSOx3XBlVkSrxUxuAx87a7SAvlSWsf/Dw8EOrH/r1rbf+8c47n970pNayo6M9m01L4YW6bLTzv7IvVnOpVHtwSGqjkB8Jw3D+/KMuveyyt7zlrZMmT6Kst0yv6dU4O+E8Bl6CyIDsvt0oJMkMQQkLvCSoOcsnf0R5leZWrIdWr/7Ot7/929/+Jl8odHV1Kek5SX1tQjoJy7hbp9IWxHZ2pjL2Kzzlh0F5YGCgHATTZxx69tlnn3POOWecccaUqVPiFw3ZvKjhxyFuARNOFEvY+eLElEF1oJKJAZjxrC6PmcW+wCd09MfqxDKw65NS8cKtu6fVsqT7eUklyppkV0jpx9pVQojVK1fddvvtt91++6Nr1gRBOZ3OZHMZJZQmXW9JBoqxJIutuo224B0lwb6SRhRGRoKgfNRRR73tbW+/5O/+tiMuOFPolRVFlOYu2gPQbGI7QhE7lAmSoox7QXHlCF5S1GtukAhlFIZXrfzLt7/13dvv+P3w8Eh3VzediI2m8aR4wN2tJlZNAdFzCfInkh61TCvPL42UB4cGPM+bOu2QU0479cILLjzpxBNmxjkxV6eTLcp1H5NqpeTRd+dFHYBJJrVJp4TK8mtip+ydE6PNdpudC4Kb8BkcHHzkkUfu/OOdd9919xPrHu/r72tra+/s7BA8qm27oKtmpMkBmetwVnSzpive87wgDIfzg9KoBccd9/eXXfbmi97c2dXlNsB5YMRtXzwUjtMXeOli14BlZarATaOLBro/ALzYVAVO7tbR2rr5CLFq1cqf/Pgnv/n1b3p7d2dz2WymLdnJxSRy0KoRIqk0G/CSIHOU7ZVLwUixVBgpSCEOP+zwkxcuXHT2orPOOmv27Nku2Lg2xupXiUeRKx78Y2XAsRL7gfu41RUDRCXNFVVaGdXHyh4xrcmQyphksrtr16677777rrvuWvmXFWvXPTY4ONjV2ZXNRMST1cKTSkiPm07c8VHu/FPZEivsxaG9XC7nC3k/lVq8eNH/+fu3nf3a12ayWZ4vUsqrXzZGxAUvC6QOQqFiJSwhrUw++XpyTxZK0OAlSGWYNA4XSUfe9ese+9GPfrhs2bJt27Z1d3dnszkdstKFStQnFa1lWj8+8nmglUtFDT5S2ZVOFQUAYWRhZLhULGstJkzoWbBgwelnnrlw4YmvfOWrJk6c5LZKax0arQytLMc+BEp4mq5o3Ziyie177ceK+rqEPPAiiEnTQEPtVypZfK65dnHd5kl9j8LIyBPr169+6KH777tvxYoVz27ZEpTL2Uy2vb3d9/0gCKubkIW1hqhoblcLf0rjHKLIrykczucLw4Upk6cuft3/uvTSSxctOosfwaE3ftr6TB2DveBlgNShNnGLp70ONUIq9ioVrDdrtJENvWQAOOA44WA3qZssTMfGedGp+emnN/7X//uvpb+8ZfPmzSQo3a7I9rWqa5d9ANn5XxpDVkvkGpsQeaZoLUl9mcrawfBwvlgstnd0zJx52KuOPnrhySctXPjqI454xaRJk0ViS5IqlMplyaN0iFknn2pNrtZ96uoGh5zNkagRA+HtTgbd/v7+57Y99/Dqh+67//5HH31001NP7+7dk06lujq7/FRKUeSOjrOopP5KVKyi4hVnzvvjpXK3IUor4Qkhh4eHy0F51uxZSy66eMmSJUfOO5KPatLqESkveFkjAxOSC79gp7dYlY0bsAyfORCAwUuNSpoTr+gmwwk34XIY3rF9+7Jly37205+tWfOwkCKbyeTa2vjdbR/nhGfYlkQrN/srnWEivyg9RNM8q1IyCHSxWAx1oEORSvuzjjjimGOPPemkk185P/ozderUmm3WJEIh2Lo2afLnlF85XI1r+bjuOCRFthJ19nGMFBvnIEz7qJX0lFebiff19z2+7vG1ax99+OE1q1atfGbz5vzQcBiG6Uw6l835qRSLSvIrs4inSaS4bJigjVGJRV8rmB2fbCjKq2JppDA8LD3/pJNOvOSSv73wjW/kGgNd0Bjy81VxslzbKg/AywsZhAE7q8RXoXRC4sEk+2nCBSZ4GVAr60hhxw0sDQ0O3XnHHcuXL3vggQe2Pb+ts7Ozra2DC9d2PZJqnlIoWtw0iS4tUdc2TFVkE0rh5DpUoMvlYpAvkNl+tq1n0oRXHH74UfPnH3PMMfPmHjVl2uTJk6ckrYeY2HZP18hW7fXHTVZL6DjpqUayYu5v3nLXRpa8Z344v2v3rh3btz/55JPr1q9/4oknNm7cuHvnrkK+EIZBJpvNZDJ8fcP9aO4iwNSZ9SfKziI2YLCV+OhCPwqpkluXh4aHRgojMw6dsejMsy5+61tOPe20XC4nhCgHZSXtn1rrKpyawMsZySqpcd5rojSYLk7pWhzvbPDyoF6JKZkeuTAshHh8/eO33LJs2bJfbtiwwRjT2dmZyWTILZjUoKkL15ZLlY6yQaGM0DXxyTlnO83LKI02nlRRTC+XwlJ5pFQqF4sjUsr2ts7u7o4jZs85/BWHHz7zsHnzjjpi1qzDDz98woSJmUx6jD3ijrCmHB5KSkd9spGRkV27dj315IZnNj+7adMzz2x6ZtPTT2/btq23t3d4eFhKmSY8z0un006myh5hm79WlMuiG2xsSk/udMSSC7OVwCnlyMjwyEipLZc7/oTj3vymJYv/ZvERR8ziuwVB6HmqgRgWjBPAQYEMtfMyNdJIbdWwYEcIXq40HHh19n+x0Wzf7bfdccstSx+4/8+7du3KkXdQKiVJGd3Nw1QYTRHCPnOUx2lpPGPsgE0Ug+JF3zAMS6VSGIZBuSyF8vxUyk9PnNQzafLkmTMPmzHz0MMOmzlj2vQJRGd3V3dXV3dPj5foLm4K5XJ5gNi9Z09/X19vb++zz27Ztu25zZuf3f789p07tu/Zs6dcLtPeCD/lZzIZ3/ddmlvzbGNUtp33Q3T9IpTU0VnFKOpx04LKzKpcLrOzkzFi7pzZf3POORdceMGJJ56YTqfdQr51/I2rzVjuBQcfbEdIsVcqYUciuZTlrAnxpgcvG+q9DZzHvvN1SCojrl+3/vbb/vD73/72scfWDRfy2Wxbti0d3TmsFDkT4baqicrVqI3QsZiEcovHpuKqJD0ltREspKg118bDYrEUlKPYLKRMpVIpz8/ksn461dnR0d3d0xH93dXR1dmWzVHBPCKbzabT6VQ6JYXkZNRZ8YdaB+VyEATFYnFoaKhQKOTz+f6+3sJIoThSHBwcHhjo7+8fKBQKw0ODxdJIEIRG65SfFlKmM5kUm/5TPZqnfd0li7v+4DA82qmgouJRI5/CNXz7xNF1QLFYUlIeMmPaqaeeeu4bzl94ysIpU6aI2CqxonOSWNqXMCwCByMkxCHscq9gqdso5eUYDDtC8PKjgUqiMzeya48mlgu2EzVGBw8+8D+33X77HXfc/vjj64026XQmQ0KJRpBMIrdEx1JNCcc9bi0S5OMpEzVjxZVa2+8rtTDKaoIIoQX3PdrnNJo1tsIw+s+EYRAGYewkKOJalK7p+6Utd0IiLhuNbmkdxm1emjs8ZJR5Uj7rK4+QSvhRqPPoAsLmmTXyzu6QJi0xRjsVVGaA4hq1vdyhxq7iSHFkpKiUnDJlyqmnnXLeeecvXrx4Styn5px6yRdSJm2Tx5bHBOBlDSthybgH1EiJJixwsFHJgOOkSlQlxNrzolBEAza9//Pgn3/9m98++OADzzzzdBjo9o72TCbDkdhQgKTeK2eIK6w4FEtEO3M/oVjzyU4x2UqqjBW5aKnHxjz62ClR07oUhVc2cUqYEydbnCpXGNzDVedsKKoao7UwHj3etnwbNl0wcZ0gXupuqPA8mtNR5QhT4CSvKeplU8pTyghTKpaGh4e0EDMPnblgwbGve905r33t2TMPmxnP8oZ8tGRixqnmQPCYGOtvJ1rNEIfBwUCdHaGIZ/Ti/kUADjLqFxRZTsvEK8TstvTQ6tX33nvvH//0xw2Pb8gXRjLpTFtbWyrt07yNdHKKznGP+rAU/cjEzdIVGSleFaYCk6rROubJnNpGr1qjf9Gw12zs3ZSyQUBNPL9wL9FQHcw9xFlR1T9h8hV518IwHCkVy6WyUmLGjOmLFy9+7eLFJ5xw/KGHHu42jK2NGl7WsyxILENA/2kakXTfUgdesQSAllCnhMXvcShhgYOXMfxxXZU1KWu8YsWKe+6558H7H1j76No9fbs938/mctl0KuWnudWIE0rKjElGirScZJzb1plJJOdzEptU/RkbzW6oxhN37H4oTiZrktea56+ZQYrSecXykA20Mysvxyo98cYEYZgfyRfyBaN1d1f3nCPnnnzyyaeddtrpp58+Y8YMfmB8rU/5fuOtle44JPfHGjojAwYHHayERWcKGp43TpEVSljgr4a6MCO1toJZyvP4rV8ulzc8sWHlqhX33nvvw6sffu65Z0ul6D6sRKF8RVmbNGQqwMoSZP5v4rleuyJreElYqFgKqqqQ6243DLf7sB7EuaZLOhsG4zqkVYFOFLeFHUykUUWloqv2UJdKpZHiiDAik8vOOGT6kUfNO+PMMxaevHD27NldZJNQGRS21zRSCFzNA2BhO0KhrVyciM8RInblRwAGBxuNS7JuoVXWmhBzCEmlUu4ZhoeGV61esebhR1asWPnQ6tVbt24NycBBKZlOpdPpjEdd1kZHgY9rrSyeRXM4go16KmqXSV/CWFtz/KXmhtRYPtTPESXvFo8asjKGcWtQvHDNjVRKyiAMg6BcKpWDIBBSZNOZyVOmHLvg2IUnn3z8CSecfNLJHZ0d7pm5r8rhjvwLXkOYuuK7dC1i9mtkwOAgIfpQKWn7nak/xTVhVVzHUYIGfz00NCzi29yxRcmc71Sod+7YseGJDevWPbbm0TVr1659butze3bvGS4UMul0W1tbOpXyPF+pSoexjifva8OrXZOVQtRWmyt3GV9Udg/XsetDzffdniZf3YlV0489KUVotNG6SATlsvK9SZMnTz9k+qxZs044/viTTzp5ztw5k6dMcT6ANGGlpZX9kPUbjJMJAEnQhAXAWDQUHI5nhGqtgcIw2PLslqefeuqxdWvXrl27YcOG5557rq93oBwUjZZKyVTK9/10Ks2+fJKll60aJS/XajapbzAClMzdX3CbG96uievUriyddGUYJbgRWmvOX3nyeMaMGfPnz59/9Pwj58yZPWfO4Ye/IpvNJV/ODRE1vERAuAVgNHgO2F1709oUlLAAIMbZ68Rd0MJI5clkbbQ4UtixY8eTG55ct379li1bNmzYsHnz5v7+/t7e3ny+YEToeZ4Snud7vu/TcnOUd7LVD8czDvD1ClxVlwXWwaFuctcYKowLk4AkO8IwCNkVikaPy5S2eumcP7Fn8qRJE2ccOv2II46Y9Yo5848+6ogjXjF9+oy2to7k/vL1h8tx6xeqW/PbAOBgg5Swkn7gUMICIMm47WVZCivKZSlgsntA8oMTBOW+vr7+voEdO5/f9tz2rVuefX7789uf37Fj186Bvv6BwcFCoTBSKARhaLQOwsDEGsuSVomi+KyU9SqOs2QyzeURKqGNlrEsRxiGWgoRalltucCqzrlcrrOzs7OrfdLEKdOmTZs589CZMw+bPmPatEMOmThxYldXl+9XiVRzSxp3KDuzIw6/SZ2T+rlhnDcAGAPOgLlVMz65KFmZ90cABqCOxk1ScYdvUvnSaUIZIbxRJl+FECOFkd6+3l27du3etXtgcGBocHDPnt7+gf6+vr6hoaEgCFhdslwul4rFcrnMhrukLx34aV9ImU6lo9RZSHIqyqRSqVxbW45kLLu6OrsJUpueOCliYk/PhEw2M9reuaat2Mg43sVkPUAIkVD6BADsLayEJZIOX+R4iiYsAIiGGXD9DadNaUQ8WFttoEQjPS6wxQOvknqnx7IqYnhdNggCjo78POxOzw93q9GUJyv2URi72VhUCz67hNXZEzqTYjcu5bS1nVgVWpIB2GdsCbrmm2jCAsBSo8w0PhmmOOl9AaxXbrVNb7KQ61ZY9+0KuJJ/V0fZ+mdzpoHuPpUngR8RAK3BdkEnlLD4vIEmLACI8WfAyQdZ2enKg+sXRxvOO9ULXTWd8XyWx17HhUECAE1Bhlpz6cxOIzglrPizj2teAPafhlqSDe8wthGCo2FEb3iH0bYEn24AXlxk2YSe4QlEoRSp5UWJbyUI4yMK/qoZdxd01YPGnQHv40YhiALw8ifRhGXHkKIAzI0i/IHGxxsAAABoOr5MurBQSyZfuyebsBB9AQAAgOZiuzmNFd4gJXqeAxTajQO2qBkEAAAA+KvFFwmvMf6WFEIqZVxdGgAAAADNRhlJ44YmNtnmRV+NIWAAAACghSiuOLPSDRuJs747ys4AAABA6/Dtv7KiBY0mLAAAAKDVKO6CtkmwkUahCQsAAABoObYJS8q4DcugCQsAAABoOUqzCKWx+vHSaGHdkBB6AQAAgFahJJeeeeWX81/CjMPLBQAAAAD7hl/xRLNqlJLM+bkXC01YAAAAQEuIMmCWn4yCLhkjaWrCov/QhAUAAAC0BF/wyC/rYEkdZbtaSqms2SdyXwAAAKAFKG2o6Gys4ZqkjmiDJiwAAACglSgaAqamqyjikjcDyXAYwZ6EKD4DAAAAzUeGRktaAbbyz4YFsUKqSHvjMh8HAAAAwF6ipCYXfiFcE5YhWSxtBbEMmrAAAACApuPTkq8QLMfBgtBSSAklLAAAAKCFkBKW1CKeOEooYWEBGAAAAGgVpIQV6z8LIaCEBQAAABwAWAlL2a+kEEbFSlgSSlgAAABAiyAlLO6/EpqasAyUsAAAAIBW45SwrCEhKWEJKGEBAAAALYWVsCT7IXETFpSwAAAAgFbDSlgmaUQYK2Gh7AwAAAC0ChlqLUWshBV/U0QhGEkwAAAA0CoSSlimIolljNaCmrAMmrAAAACA5qMElaDZkVAIzbVnpZQ0SmgSyIIgBwAAANBsWAlLRmE4ynRJl9I2YQkhEXgBAACAlqAo4zVkiGSVN1gcmr9ll4WxFgwAAAA0FRmGmrSghRFKKmdHiCYsAAAAoIUo8uBXlAAb0sJiP0KtoYQFAAAAtAyf4qumArQSQitBa8KSS9NQwgIAAABagjLUeSWNIA/CuBELsRcAAABoJZzpKm64MkbFrvxQwgIAAABaiAx0qGq7nNGEBQAAALQWX9LIL1efJXvzkyIWfVsZY5TiMjWCMQAAANA0fMMLwNLVnlmWQ3HjM+IuAAAA0AqUEdJILWjoKAq41BGtYUcIAAAAtBIlhWYbButLGIddNGEBAAAArcOn0KtYbZIKzmTOT/HX1Z9RiAYAAACai5LGM0JSAdpw4qvJjtBACQsAAABoGb6QRlLl2TZhSSk1NWEJgdwXAAAAaBHK2DkkO4okjZbOjhDLwAAAAEBrUGxGyDbANuLacIzoCwAAALQKn9ufOfYqKY2RrAYdBWU0YQEAAACtwZdUbmYlLK2luwklLAAAAKB1+ILHf60Clh0+ghIWAAAA0FKUZv9fIU1l/VcYI9GEBQAAALQOJbVdAKbmZyNoHjj6B7O/AAAAQMvwqcRsqN/KSGmEUfSFpsFgNGEBAAAALUEJ4ZEVPye8ipaBo6+0gBIWAAAA0Cp8EWW9mhqwFH9LCiEVmrAAAACAFqLIdYGiMC8AC1LC0lDCAgAAAFoIKWGRK3BFj8OQEhbKzgAAAEDLkIEOlUjUma0Vv2F7whdz0wAAAICDF2WVr2IHfrJGou4rNGEBAAAALYOVsCjy2nwXSlgAAABAy2E7QpLeIEUOabTkZiyEXgAAAKBlKFr0NdKZIkmSoxSwIwQAAABaiB/9JVUUiDntlZ4xFI8TSTAK0QAAAEBz8SU58XPuK43QXI4WmrJj2BECAAAALcHaEZL7vh0/onK0EmjCAgAAAFqGMmxHaCSv+ko2BEYTFgAAANBKlNFsxGA43vJUcGUs+MXePgAAAOCgxCfbQcH++7EoNJWiTaX+jEI0AAAA0FyUEpLCrm29Mqx7FeXFlBpDCQsAAABoAb6xya+ppMFQwgIAAABajOKklzqhrRKWkEIb2BECAAAALURJqjCzBkf0b5zxQgkLAAAAaB0y0KGsHjniTizUnwEAAIDWQUpY0kpfSaHYiEEYTT+FEhYAAADQEqwSFpsQCk57NTVhCYEkGAAAAGgRTglLuCasKCGGEhYAAADQSpQw1U1YVImOhSkBAAAA0BJ8/seWmiWv+0Y5MOwIAQAAgNbBTViSJaAl/aHUlyeB0YQFAAAAtATXhOUEoSWUsAAAAIBWo7S0YtB1SlgIvQAAAECrUEIb14Ql7DKwgBIWAAAA0FJ8Ka3/grCN0EqSHaGAHSEAAADQMnxFNkjWh19HUTeKvVDCAgAAAFqJH7c9Wy0stkGCEhYAAADQUqIEWAhpXBMWmzGgCQsAAABoJT6lv4ZGj+KIG1ekEYEBAACAFuG7OjMv9LIgZZQHQwkLAAAAaBlKkge/FtE/xhhNlWhjtKb5JE3dWFYjGgAAAABNwiphuRK0bcJSpISlY5dCU5kPBgAAAMD+o7R15JdGUxXa6Cgh1lR/Vi/21gEAAAAHKX5FhFJSBzTVolkJSzphLKS/AAAAQFPxZTwEHDdheWjCAgAAAFqNb4RUPHdE2a8WsCMEAAAAWo7PcdUISn9p+NcqYcGOEAAAAGgZykh2IZSc+kqjpRRQwgIAAABairISHNRy5eZ9DewIAQAAgFYiwzCsKzRLq0OJJBgAAABoDVYJy6W8RrAvA5SwAAAAgBaihBXAiv33pZRGKKWkUULT2jCXpAEAAADQPJRWhgOsoTy3qgkLFWgAAACgNSguOlMhWtpSM40FG65NQwkLAAAAaAEyDMO6YV80YQEAAACtRRnOd+MmLNtzZbQmZ0Kt0YQFAAAANB9FFvxRkOU8WLEYllJSKgMlLAAAAKA1KEMehFLLKOONu6ENYi8AAADQSqzlr2ERaHJgkPFA8Iu9bQAAAMBBiwx0KDnsJr6JJiwAAACgpfiSa86CloGFNsKjKeCkHaE0hmMxS2YhKgMAAAD7i2+kFEZTCiyF8LgbK2FHyDVqg4lgAAAAoIkoI6n/meQ4yJfQVNsRavrP8IASRCkBAACApqCk4QVfGkcyRsTaz9VNWIi7AAAAQDPxpbFru8ZoisIe2QMLaZPgmmYslKABAACAJuALUrqKm6y0tlNI2kghjTJCK0mO/ZJXiLWbXAIAAADAPuNT75W2SljGEyZuwqKVXylVde8zMmAAAACgCSgtNc8hGc22SEZKHkLi2WBe/XVxFwEYAAAAaAI+WS2oeuVJU/Hil0K6JizMAQMAAABNoGJHaGzxWTrvI6hBAwAAAC3C539oDFiL6H9FLdBGSGGMMiZUSlFoNmyVhAwYAAAA2H98CrSGJoKVMLEotJQsxyGlR3fDGjAAAADQTJS2UVUaI2kYiR34heElYOvWr+NuLChyAAAAAE1AccKbUNuwE8FGQgkLAAAAaBW2CctZ8fM3YUcIAAAAtBTfFpzjWWDDQddACQsAAABoIT4luiYxdBR9JRWUsAAAAIAWogxZEQpSwqLoq6n/WUIJCwAAAGgdPkVYu+BLpWZhB5OkaKSEBQAAAIAmIMNQS6kpBifTXFOthOVq0BDiAAAAAJoAd1QpI6gRWmhNfxuj6W+jtSaRSi1ESPc08X+YTQIAAAD2Hd9II7QRyi75Rv/GXxluwhIi0fkssR4MAAAA7D8qSnyVYjFoXv6lISRSpYzQ1ZkuS2Vpuo0MGAAAANhHlBRSGmMXfA0NBQvpBKITURbhFgAAAGgasR2h4YxXOkmOUZSw0IQFAAAANAFa3DVG20mjUEihqQlL2yasIO7JQhMWAAAA0DR8mv/VUhqqO3uCbkgluQId2xG6lis0YQEAAABNQBlpKAa7zJaasLSrP9ekuWjCAgAAAJoA9T8LiqxS2mHgsZqwEHQBAACAJuC7ziua+hVSeMYYFfsykANSbBEsUIIGAAAAmoNP2a4wNv3V2n6fwnGU7oZSKnIjNNSx5TJgiY5oAAAAYJ8hP2AZ152lkqKS/BLchNUw60X0BQAAAPYRpaWRhuaAjTHRDU033JKv67fCGjAAAADQNJSMVbCiv132Gy/1IskFAAAAWoFPtWctBDdhGSk9lsHiwGwaNGEBAAAAYH9RSghhFDnwh8KYUGrDfgzCkCcDKWFpp4SlX+wNBgAAAA4GfGpl1tSC5QlplJCu+EyNzsiAAQAAgOZDdoRRxiutBEeiCYu+p9F5BQAAADQdRU3QKk5xrShWoyYsRGEAAACgafhGCta9MlE0lsZwD7ThJqxGSlgAAAAA2F98Sf1XJHVltCbNK4q+WrJDIf/ExWCoXwEAAABNgO0IubzsxoHpr/ib1XEX0RcAAABoAqSERQO/FHAp7hpjtEkoYZk4BkMPCwAAAGgOrIRFs8A2xTUsC91ICUsiDwYAAACagooXeqO0l1JgxQFWGdam9LgenUiCkQEDAAAA+wu5IRlphJJRhA1DGYVkKYzmsEsyHYZ7s4TLfQ0asgAAAID9weeF37j5iqKvlNH/dvFXCqmqQy0M+QEAAID9RfE/ZMbPE78mSok1GTJEkZcloJNBV6MQDQAAAOwn5MVAcpR2HMkNH/G/lc4sBF0AAACgafhxjxU5HUlF3dC8vmuL0vRvGLdAa5c0owQNAAAA7DOKV3oNdz2b0EijhaaOaG2M0TqkGxx00QsNAAAANAffSHb5JQloyfmu4D+kQOnR3RqO/yIDBgAAAPYRmvo1SSUsHd0W0hg3cSSQ9QIAAADNhYU4BE362rkjCrY0GyzYKxgAAAAATcYn4UnNvdDkwS/tMBI1RpMSFhebIUIJAAAANA2fjYBZDloabaKc2HDiK4UyJpRK2T6tKBgnu6ABAAAAsI/4ZIBkhKwMHUWBN3YjjJuwkAEDAAAAzUSxBAdlv7YJS4goEY41OLAGDAAAADQfJY20q79ScCrMmNq2Z0RiAAAAoGn4Tm/SWMcFJSkUS5sEe/EdUIIGAAAAmoYfK2HROLDQRigqQHM4VsaESilyIzRONuvF3mYAAADgZQ83YZESlpLCxEpYtU1YsCAEAAAAmonSglqghTB0S5LboNY1TVhQwgIAAACaieJis7XetxGWhaARcQEAAIBWIcMglCz8bJSQRgnleqGlRMEZAAAAaAl+nPlSE5bRmm1/DU8luSYsloiGEhYAAADQHHwhBbsuUPMVu/FbS8KEHSHGkAAAAIBmoqzghqFFXyMpxxXacENWPKCEsAsAAAA0FVdPNlKqiv4VNWHZyBwbJb2IWwkAAAAcZCjSoNQ24kqqOZsoLNuJYOPRvzIO1ciFAQAAgCbgk/UvBVdppJFa6so0kqmRvpJQwgIAAACags/9zdaKP27AElEsrjReJaIuoi8AAADQBJQROgqwWpILP7kS2uXf2CgYURcAAABoEhxijTFKapaZ1Jz4WgmO+C4S/VcAAABA82CRKymlEtIzksvMFIOjnFixOIeJx48QgAEAAICm4DJgn9Z7haYmZylC6oAWbNLv1n5RggYAAACai23CkqTBIYVn4qkjq41lrFcSAAAAAJqIsiobUc5LkdhoaTSJUErDktAoQQMAAABNotKEZTj22gXgKBGmQCyMzYkNStAAAABAs6g0YXlCCWmo3VmT/rOSwrNuDPY/AAAAAOwvsdG+vaG0XQJWwkhlffiNgBs/AAAA0FScy74koQ2ptY51N2TNPQAAAADQLKKgyxpXXILmr5OpMTJfAAAAoLm43iu+obX2G94PKTAAAADQRJLVZb6tXDqsVCUbBgAAAEATMTGVJqzkz1wwfilQk62/dDYMAAAA2Fvqg5pf35f1Ym+kxU1KJb8EAAAAXo5IKWsimnTRePxP4Tq4XsRonXx128+d+Hu0O4/x01bsznhetxVbtf+P3betGs9P93Z79v+Z92qb699XYzx/645z66h5uX3Y33F+2fCx49zZvfp0j+epmvKebOJ7Y39+Op6N3J9Pyv589vd2a/efl1TSuLdoraXWuv4HL4V0s/60+FLYKgAAAGAfqLmCib4Mw5C/GC28JS8/9/bKfR+iZk3crX8evs33Sd5uVgY8zm1ueLe9Olb7tlX1r1v/o/E/dt+2ueaZx/m6yV/ZOI/z/jzz2I8d7af176vkoRjn/o5/j8bzOWrWpWfDYzWe/R3PY+u/3Kv3Rs2m7tWnu+Fu7ud7cpxbtbfvjf356Xj2uimflH07I+3Dp3s/OTCv2/CZ9zZG1Dw2eZ//PwAA//8DXWqmE/+hEAAAAABJRU5ErkJggg==') !important;
    }
    
    div.stButton > button.btn-download-att {
        background-image: url('data:image/svg+xml;base64,PHN2ZyB3aWR0aD0iMTYiIGhlaWdodD0iMTciIHZpZXdCb3g9IjAgMCAxNiAxNyIgZmlsbD0ibm9uZSIgeG1sbnM9Imh0dHA6Ly93d3cudzMub3JnLzIwMDAvc3ZnIj4KPHBhdGggZD0iTTMgMTEuNzVDMyAxMi4yOTE5IDMuMzQzMTIgMTIuNzUgMy43NSAxMi43NUMzLjgzNjkgMTIuNzQ3MiAzLjkyMjMyIDEyLjcyNjcgNC4wMDEwNyAxMi42ODk4QzQuMDc5ODIgMTIuNjUzIDQuMTUwMjYgMTIuNjAwNSA0LjIwODEzIDEyLjUzNTZDNC4zNDU3MSAxMi4zOTE5IDQuNTM0NzQgMTIuMzA4NyA0LjczMzY1IDEyLjMwNDRDNC45MzI1NSAxMi4zMDAxIDUuMTI1MDQgMTIuMzc0OSA1LjI2ODc1IDEyLjUxMjVDNS40MTI0NiAxMi42NTAxIDUuNDk1NjQgMTIuODM5MSA1LjQ5OTk4IDEzLjAzOEM1LjUwNDMxIDEzLjIzNjkgNS40Mjk0NiAxMy40Mjk0IDUuMjkxODcgMTMuNTczMUM1LjA5NDU1IDEzLjc4NSA0Ljg1NTk5IDEzLjk1NDQgNC41OTA4NCAxNC4wNzA4QzQuMzI1NyAxNC4xODcyIDQuMDM5NTYgMTQuMjQ4MiAzLjc1IDE0LjI1QzIuNTA5MzcgMTQuMjUgMS41IDEzLjEyNSAxLjUgMTEuNzVDMS41IDEwLjM3NSAyLjUwOTM3IDkuMjUgMy43NSA5LjI1QzQuMDM5NTYgOS4yNTE4MiA0LjMyNTcgOS4zMTI4IDQuNTkwODQgOS40MjkxOUM0Ljg1NTk5IDkuNTQ1NTkgNS4wOTQ1NSA5LjcxNDk1IDUuMjkxODcgOS45MjY4OEM1LjQyOTQ2IDEwLjA3MDYgNS41MDQzMSAxMC4yNjMxIDUuNDk5OTggMTAuNDYyQzUuNDk1NjQgMTAuNjYwOSA1LjQxMjQ2IDEwLjg0OTkgNS4yNjg3NSAxMC45ODc1QzUuMTI1MDQgMTEuMTI1MSA0LjkzMjU1IDExLjE5OTkgNC43MzM2NSAxMS4xOTU2QzQuNTM0NzQgMTEuMTkxMyA0LjM0NTcxIDExLjEwODEgNC4yMDgxMyAxMC45NjQ0QzQuMTUwMjYgMTAuODk5NSA0LjA3OTgyIDEwLjg0NyA0LjAwMTA3IDEwLjgxMDJDMy45MjIzMiAxMC43NzMzIDMuODM2OSAxMC43NTI4IDMuNzUgMTAuNzVDMy4zNDMxMiAxMC43NSAzIDExLjIwNjkgMyAxMS43NVpNOS4wOTQzNyAxMS4zOTMxQzguNzczNzUgMTEuMTc3NSA4LjM4NjI1IDExLjA2NTYgOC4wNDQzOCAxMC45NjY5QzcuODc2ODcgMTAuOTIxNiA3LjcxMTg1IDEwLjg2NzYgNy41NSAxMC44MDVDNy43MDMxMiAxMC43MzEzIDguMTU2ODcgMTAuNzIzOCA4LjU1NDM3IDEwLjgyNTZDOC43NDY3NCAxMC44NzY0IDguOTUxNDEgMTAuODQ4NyA5LjEyMzM2IDEwLjc0ODZDOS4yOTUzIDEwLjY0ODUgOS40MjA0NCAxMC40ODQyIDkuNDcxMjUgMTAuMjkxOUM5LjUyMjA1IDEwLjA5OTUgOS40OTQzNiA5Ljg5NDg0IDkuMzk0MjcgOS43MjI4OUM5LjI5NDE3IDkuNTUwOTUgOS4xMjk4NiA5LjQyNTgxIDguOTM3NSA5LjM3NUM4LjY5MTM5IDkuMzE0MzYgOC40NDA0MiA5LjI3NTQ2IDguMTg3NSA5LjI1ODc1QzcuNTY2ODcgOS4yMTc1IDcuMDYyNSA5LjMyNjI1IDYuNjgxMjUgOS41ODE4N0M2LjQ5NTM3IDkuNzA2NjYgNi4zMzkxOSA5Ljg3MDc2IDYuMjIzNzMgMTAuMDYyNkM2LjEwODI4IDEwLjI1NDQgNi4wMzYzOCAxMC40NjkyIDYuMDEzMTIgMTAuNjkxOUM1Ljk0NDM4IDExLjIzODEgNi4xNjgxMyAxMS43MDg3IDYuNjQ1IDEyLjAxNjJDNi45NDM3NSAxMi4yMDk0IDcuMjkyNSAxMi4zMSA3LjYyOTM3IDEyLjQwNzVDNy44MTY4NyAxMi40NjMxIDguMTI1NjIgMTIuNTUxMyA4LjI0Njg3IDEyLjYyODFDOC4yNDUxIDEyLjY0MzYgOC4yNDA0NCAxMi42NTg3IDguMjMzMTIgMTIuNjcyNUM4LjE0ODEyIDEyLjc2OTQgNy42MzUgMTIuNzg0NCA3LjIwODc1IDEyLjY2ODdDNy4wMTg5NCAxMi42MjA2IDYuODE3ODIgMTIuNjQ4NiA2LjY0ODM4IDEyLjc0NjhDNi40Nzg5MyAxMi44NDQ5IDYuMzU0NjEgMTMuMDA1NSA2LjMwMTk3IDEzLjE5NDFDNi4yNDkzMiAxMy4zODI3IDYuMjcyNTQgMTMuNTg0NSA2LjM2NjY1IDEzLjc1NjJDNi40NjA3NyAxMy45Mjc5IDYuNjE4MzEgMTQuMDU2IDYuODA1NjIgMTQuMTEzMUM3LjEzNTg1IDE0LjIwMiA3LjQ3NjE1IDE0LjI0OCA3LjgxODEyIDE0LjI1QzguMjIyNSAxNC4yNSA4LjY3Njg4IDE0LjE3NjkgOS4wNTE4OCAxMy45MjgxQzkuMjQxNzUgMTMuODAyNiA5LjQwMTU2IDEzLjYzNjYgOS41MTk4NiAxMy40NDIyQzkuNjM4MTcgMTMuMjQ3NyA5LjcxMjA2IDEzLjAyOTUgOS43MzYyNSAxMi44MDMxQzkuODEyNSAxMi4yMjA2IDkuNTgyNSAxMS43MjA2IDkuMDk0MzcgMTEuMzkxOVYxMS4zOTMxWk0xMy41IDkuMjkyNUMxMy4zMTI1IDkuMjI2NDggMTMuMTA2NSA5LjIzNzU4IDEyLjkyNzIgOS4zMjMzNkMxMi43NDc5IDkuNDA5MTUgMTIuNjEgOS41NjI2IDEyLjU0MzcgOS43NUwxMiAxMS4yNjg4TDExLjQ1NjMgOS43NUMxMS40MjY5IDkuNjUyOTIgMTEuMzc4MSA5LjU2MjgxIDExLjMxMjkgOS40ODUxMkMxMS4yNDc3IDkuNDA3NDIgMTEuMTY3NSA5LjM0Mzc2IDExLjA3NyA5LjI5Nzk3QzEwLjk4NjUgOS4yNTIxOSAxMC44ODc2IDkuMjI1MjMgMTAuNzg2NCA5LjIxODc0QzEwLjY4NTIgOS4yMTIyNCAxMC41ODM3IDkuMjI2MzQgMTAuNDg4MSA5LjI2MDE4QzEwLjM5MjUgOS4yOTQwMyAxMC4zMDQ4IDkuMzQ2OTEgMTAuMjMwMiA5LjQxNTY0QzEwLjE1NTYgOS40ODQzNiAxMC4wOTU3IDkuNTY3NSAxMC4wNTQyIDkuNjYwMDNDMTAuMDEyNiA5Ljc1MjU1IDkuOTkwMyA5Ljg1MjU0IDkuOTg4NTEgOS45NTM5NEM5Ljk4NjcxIDEwLjA1NTMgMTAuMDA1NSAxMC4xNTYxIDEwLjA0MzcgMTAuMjVMMTEuMjkzNyAxMy43NUMxMS4zNDU4IDEzLjg5NTYgMTEuNDQxNiAxNC4wMjE2IDExLjU2OCAxNC4xMTA3QzExLjY5NDUgMTQuMTk5OCAxMS44NDUzIDE0LjI0NzYgMTIgMTQuMjQ3NkMxMi4xNTQ3IDE0LjI0NzYgMTIuMzA1NSAxNC4xOTk4IDEyLjQzMiAxNC4xMTA3QzEyLjU1ODQgMTQuMDIxNiAxMi42NTQyIDEzLjg5NTYgMTIuNzA2MyAxMy43NUwxMy45NTYzIDEwLjI1QzE0LjAyMjcgMTAuMDYyNSAxNC4wMTIgOS44NTYzMyAxMy45MjY0IDkuNjc2NzZDMTMuODQwOSA5LjQ5NzE5IDEzLjY4NzUgOS4zNTg5NyAxMy41IDkuMjkyNVpNMi4yNSA3LjI1VjNDMi4yNSAyLjY2ODQ4IDIuMzgxNyAyLjM1MDU0IDIuNjE2MTIgMi4xMTYxMkMyLjg1MDU0IDEuODgxNyAzLjE2ODQ4IDEuNzUgMy41IDEuNzVIOS41QzkuNTk4NTIgMS43NDk5MiA5LjY5NjA5IDEuNzY5MjYgOS43ODcxNCAxLjgwNjlDOS44NzgxOCAxLjg0NDU0IDkuOTYwOTIgMS44OTk3NSAxMC4wMzA2IDEuOTY5MzhMMTMuNTMwNiA1LjQ2OTM4QzEzLjYwMDMgNS41MzkwOCAxMy42NTU1IDUuNjIxODIgMTMuNjkzMSA1LjcxMjg2QzEzLjczMDcgNS44MDM5MSAxMy43NTAxIDUuOTAxNDggMTMuNzUgNlY3LjI1QzEzLjc1IDcuNDQ4OTEgMTMuNjcxIDcuNjM5NjggMTMuNTMwMyA3Ljc4MDMzQzEzLjM4OTcgNy45MjA5OCAxMy4xOTg5IDggMTMgOEMxMi44MDExIDggMTIuNjEwMyA3LjkyMDk4IDEyLjQ2OTcgNy43ODAzM0MxMi4zMjkgNy42Mzk2OCAxMi4yNSA3LjQ0ODkxIDEyLjI1IDcuMjVWN0g5LjI1QzkuMDUxMDkgNyA4Ljg2MDMyIDYuOTIwOTggOC43MTk2NyA2Ljc4MDMzQzguNTc5MDIgNi42Mzk2OCA4LjUgNi40NDg5MSA4LjUgNi4yNVYzLjI1SDMuNzVWNy4yNUMzLjc1IDcuNDQ4OTEgMy42NzA5OCA3LjYzOTY4IDMuNTMwMzMgNy43ODAzM0MzLjM4OTY4IDcuOTIwOTggMy4xOTg5MSA4IDMgOEMyLjgwMTA5IDggMi42MTAzMiA3LjkyMDk4IDIuNDY5NjcgNy43ODAzM0MyLjMyOTAyIDcuNjM5NjggMi4yNSA3LjQ0ODkxIDIuMjUgNy4yNVpNMTAgNS41SDExLjQzNzVMMTAgNC4wNjI1VjUuNVoiIGZpbGw9IiM3MzczNzMiLz4KPC9zdmc+Cg==') !important;
    }
    
    div.stButton > button.btn-delete-att {
        background-image: url('data:image/png;base64,iVBORw0KGgoAAAANSUhEUgAAAoAAAAKACAIAAACDr150AACAAElEQVR4nOy9CZRsR3mgGX/cm5mVtWTty1v19J4kD2qBEQ1CuM1gwG4bu20MBiMjxu5z3HPG2znGPmMbxuu0T5uDPd1tenqmOYbxaWwW2+ANIwwIY5AMpjFgxCKE9Ba996RX+6t9y7w3Ys6NPyJu3C0rqyrrVWXV/wnqZd68S9wt/vj/+BdfNBoMOAMADlLI6K8CAPAvY0JKBuAxJs0SZj/Yz3a5lJIxxjkXQtgPnHNcjru1H9y92eMyg90zbmL3nD160ba5uyraln6lX+lX+pV+bf4rM7T967atEkLYX92fspvnttmVPlZaZSVaauc7anMReFwXIYSUkkumDq8aw5I7MTsF9T/JmMRWMSY5T6+GzXVb714y+xNKZfcCFbXP3XPquqQ2yd22+VXI3fbw/Npi4w/VcXfBIT+jI3Pco8eBv6HH+df9Y9tW5YrSrNBt3mYrXFMLcSdZobPTNre4Puc8Unu1eE0K8KQ8564kNu3Te8s9PXsmltw9N2lxE3ncIjtaeS80P9D+NWMvx91Lq47bGR294x49DuqZPKg9H5Jncv++btuqlFaWu7z1bV1FtmjPO21zbjNS0jCSwUEQAOeQXNVtCgBTJui06p0dgLS+QhHNN9zFDgmCIAiiCbdAsrimcrQE41euf00ePimurdCWxhAN+Ndd2SUl7Vu0j6es+e6G9gJld2i3tVaFrC1lP37dCwd13L20ai/bduIZHb3jHr1f9+9qHM4974WDeiY7gtZF1V4OkZ3SjghFiPeBM8iuWqSJ7k7B3RbSgAmCIIhjAgSRLsxARrotJ9lGEARBEPuJVTU5sEj6ojeVNk3gR+1gRWYKgiAIgmgb1uLNmRa60qq9gCswiL8RBEEQBNFWEqHNVtsFdMvS4b8kgwmCIAiizfjWNU4H+JrAIwkSmKc0YxLABEEQBNFmuASQTCgjtFKFpU23wZlWggmCIAiCaDNcCgFMB/YqJyxp0k0yJ+MVQRAEQRDthKtpXhDRB51eA8UxQzkMFK9NEARBEO3HlyoKGD2f1ewvSlwANflLYcEEQRAEsR9wkBIEqLDf6D+B8ldKIYWTsYyUYIIgCIJoJ6oSMEiuFV7GWaQSM85BhwNTNDBBEARBtB9uMnMrJRgkk8oZS0imCzTp0gsH3U6CIAiCOFLoVJRYDUklxALJUBnGuWBmimcQBEEQBNE2fCV/hfqs0kMziJReFMXqXzI/EwRBEETb4QwkAJcAgumk0EwKJoVUTlhSCqocSRAEQRBtx2fAmRRgkl4BMzX4MREHBSERBEEQxD7ApQxV4mftaAVSqFTQKgc0QNIXmiAIgiCI9sBV0iudh0OqTNDW3Gwtz2SCJgiCIIj2woF5ElAFRk2XKycsFZgU6cKmOjBBEARBEO0Dw5BQAksmmHLFEmpWOFShwcpBmjRggiAIgmgrvoo9EsoFi6u/OOFri/GT7ksQBEEQ7QedsHQRQiV4TTZoZurzEwRBEATRbnwVasS1smtNzaDNzqT/EgRBEMR+AKEQDARI0M7QEmOAJUUfEQRBEMT+wZmUILmQ2t4sQQomdElCqaoyUBgSQRAEQbQbX+m9AtAJi1m9F7QVmnRggiAIgtgHfKlrMYCaBVZVCbE8ktRzwkBCmCAIgiDaDcdJX0y24bhdAcPUHCR9CYIgCGIf4GByXqH1GRiXwFV8EjDgTMtmEsMEQRAE0U64yv7MpTI6M8HC6K8AJqRyxmJClSMU5IRFEARBEO3EV8ZmIZUTlnLFUvqudsTS6ThIASYIgiCI9sJBxxyBlKoIoUTnK5BgSxRiEBIpwQRBEATRNriUWHRf5YRWqaAlxgPHLlmo/5IWTBAEQRBtgzMOSt0F9HwG4MxJiaULEwKQBkwQBEEQbQRCKaxsBckESM6AqfpIKkgpBK4qBINkOlMHQRAEQRB7hau8G1IVAGa2GKEKQFLiFjxlfAYyQRMEQRBEG+EglAuWLX+ks0AzKdEOTZZngiAIgmg/HNVeZX4GqV2eMRoJnbBIABMEQRBE+4FAhADAI1UXJJMm9ldSBiyCIAiC2D8407UXsBRhqDyyBEjBmGBSShEwJpRxWhx0UwmCIAji6OBH2q8QyuoMDDgkyhGiExZCyjBBEARBtI1I+qK3Fc7+goz0XcFAkBMWQRAEQewbPM71bCsiKd8rTk5YBEEQBLFv6EQcHJhAT2hdFhjVYTI7EwRBEMS+wKVyehYSaw6KUAqBtQijv+SERRAEQRD7gh9puarmL1e2Z/1/csIiCIIgiP2EK1uzVNkomQpJEuh5RZmwCIIgCGL/4Gh+Vt5WOiclJqKkTFgEQRAEsX/4yvGZqyBgqUKQOMpg7ZDFPGaN0gRBEARBtAmuag4y5XiFBYFR5ZXkhEUQBEEQ+4d2wpKqIgMwmweLkxMWQRAEQewfvtJusQqDErhSxwOjF5b6gerwEwRBEESb4Vj9iKMEdl2ulEQmtZcgCIIg9gNfuVgJZtysADz0iZYYGKwgJyyCIAiCaC8+qEBgoaQsSBHqYvwohDmTIXCuhLFgZIsmCIIgiDbhMwApJWjnKw6SyhESBEEQxL7DBYpcKZU3FriZsMBkwiLZSxAEQRDthUcy1pQ+kliZQWXCMt9k0jWLIAiCIIg2wDH2V6ecBCWSY+cr5ZNlY4MJgiAIgmgTqiC/UoGxKINgcSYsxpigTFgEQRAEsQ9wXQsJpApGAiWQI7GMajAAN7ovacAEQRAE0TZ8lesKzDxw9E1VJgSStwRBEASxf3BVAxgijRc4mqKxKL/UIlgoc7QkTyyCIAiCaCMqtwZo32fljqXjjrjSggGMkxaZoAmCIAiiffhK4KoKDGouWGiHaCFVTSQpgXNVHpjkL0EQBEG0D5UJiwnOtLarkmIxZY5mTEcg6ZxYB91UgiAIgjg6cIEqsGQqBRYAUzqwZI4blnT+EgRBEATRBngkfyXHkkdSYhZKTEVpJS6FIREEQRBEm/FBKiM00yUZmFQLsBKwUYKBZoAJgiAIoq0oL2gJArNAK0VXMiGlFKo+gxCSJbRhgiAIgiDaAEfvZwCBma+YysvBuS7DQKovQRAEQewHXDIlelUCLGYScWBYklqBdF+CIAiCaD/c1PzlKgeHccICkJJ8rwiCIAhiv4BQCJTB2tNKByAlnLAIgiAIgmgvnAsJDEzwkZBcleOXQqArlhTkhEUQBEEQbYcra3MIWuEFZX1mnHOQnAkG0qjFBEEQBEG0Dy4iEQxS5eBQIcHKCUuVKFSKMUEQBEEQ7cdnUud/Nrk4UApLo/wqaC6YIAiCINoKhKFgLDQ1BwVjnpK3AlR5hoNuHkEQBEEcTbjKwuFJbW72dOIr5X6lfLHICYsgCIIg2o8PAFII5X/FseqgVOUI0SuaskATBEEQxH7AJWbckKDyPkf/mExYOhslQRAEQRBtRxVjwCpIwLH6PuiSDCR9CYIgCGK/gFAIYIIBSMGA6xIMlAmLIAiCIPYVXxVhAC12hSoHHH0RajqYSyk559IEKhEEQRAE0RZ8I1qtiI3+BeDo+ExylyAIgiD2Ay6UygsSmGCYjIMB6r8kegmCIAhiv+CgXJ8ZEzrdlUk/SU5YBEEQBLF/+MAAVNZJwTABNFffnAKFZIgmCIIgiHbjMykZ40JluwJlgFZl+ZU9mpywCIIgCGJ/8HXpBSZ1FLAyQgNwNECT3CUIgiCI/YBLVQHYpr0CpjJhkRMWQRxCpPkrna8EQXQmmPMZPaFN0QWVDpqcsAji0AHmLzhfCYLoTCCQIciUtgtaIpMSTBCHCol56tRnMF/3spPC/bTlSARBNIOD0K+XzoDFpATBhGCMyhESRPvBt0nV/Nz5a7VDE3R2dakzvRs1GuIlErsAqWagEiuRrk0Q+wIEIlQ1gTEGySzFF9AkwyIvaILIxyiK0qZQl9p9okXlMSVDIbkEHFUUGDpnxLtWQfzNtmV5R5fYQmk+NNOAdTQEacAEsR/4aroXc0BjxyFUHkr78hIEUYzRD00aG2nCCrDGmLRFtVvUd2Xe14S6K+PfZHKD3G1TH/SLLZ2fZLw3SJ6ZCYXA9SWNxQmivfhg7FPYXRjBi15ZBEE0Ja0Bg3WewMh6pbZKuwQBhXWySEm+rBabPeCOvrq6eEIpt5q6apVQDeNM5YLnOIzgulsw7SfpSxBtBMIwtC+VykoJaJ+icoQEsT1CGWgFE1JyAJ1Z3ZQ1wRGtMONazvkBt7Y10PlD1QcHOwmV+kBGaYLYO1oAq6ExTgWr90oKZVWjTFgE0ZSUBhxPnUbvSxiGvu/bdefn5y9evHhZMTk5uby8HIZhJNI44GunrNZMcuDSxAGqWV79UkpVLFTppvhK2g8oFIU7daR+ZkLNSnMAY+ACPSRgOGSQUvrg9db6RoaHz5+7/cKFC3fcecfpM2dsm4Mg8DwvPl3VIcQTyQRB7IECJywz1CUnLILYFqm9r7TrsBAoqACAB0Hwz1/558/+w2c/97nPPfbYY5OTk6urqwCAgtkKUWbkY8KMDHZ6yAyRGeDK+MG1bKffUEjPCVvTd2plfNODIGBSVqvVkbGx5zznf7r/xfe/9KXf+eIX39/b24u/KrmrD2g+UJ9AEHsiEsDAjMmZ6WG4FNGonDl2JxLABJEXQKvdF6UU6JUshPC8SLhubm78xV/85Xv+6D2f+cwjW5ubnuf5vl+pVDzPy5Y5Sb1l2WnXljy5jOjOeWHt9G9qC7tDtX4YBPV6fau+FQZhuVz6ju/4Vz/xE//2R37ktX19fUIIaw9Tf7WWj7o1qcMEsQuUAHbDjdQLhcEHtkA/CWCCyEH5OEsTHiSlEKHwS34YBn/6Z3/2+//pHV/+ypc9j9f6ar7v26h6KyDtX3chfk5pqK34QLmrZX8teotVu1VEEjpf4fmoXoBzvry8vLm58e3ffu+b3/wLDzzw+lKpjKpwYkzAEsYzgiBaR5mg3fdHRyORExZBpMmRbAIdn5kUUjLped5jjz32S7/8Sw8//HClXOnrqymxK1z3K1cWpiZ6ciWljTAGky0297VMyemivRWRsljj5ihrV1dX19c3XvnKl//O77ztvvvuw5GEPSOdNo8KtxDEzkk6YenhPCMnLILIYmZmoteDm3lfNekrPI8zBh/84Ad/5S2/cv3atcHBAca4kILFJt6Ej0VKVd3m/QIWi7oWwomzAjirbeN7bUWpZPkFWHBNAFhaWqrV+t76lrf+/M+/GXjUb0QbSsY4A6rdQhC7AkIhrLi1w20Tes/IBE0QllRqC0y7EYRByfODIPyN3/j13/3d3+1R1OubTh5HTVYGp5TgQknsOES3Yu7NasPZOebUhyIAAIW053mbm5uLi4tvevBN//n3//PIyIiWwXlnRxBEK0BDhlyYL9ErhJmw4sR3JIAJQmPDgYwMDoXwPW96Zvqn/7ef/su/+suRkRHMaaESWYATtBPLOVQ97Vf7fmXDbfcPK1Z1yC8mBimwY+NnnMD2S6WZ6en77n/xn37gT86dOxeGAeeeHbczKuFCEDsBGmHAXW9GJgFQImNMIQlggmBGiBrJq6ZrMMz32Rs3XvPa13zxC1+YmDjRaDSUlRjTc8Q096JKJri4daeTiiQusl2bIGOB0Yqlcml6evp5z33eR/76w6fPnlGBwr4dc+iIZoIgWgDngHUSOqv0khMWQWSR2ms4EkWhlD7n8/PzP/zqV//DZz974sSJer2emt9NxdO7wbhW6Ww+wE15bO1CQmf9vOzyOAo5WZ0pK31TJmvf92dmpr/zX33nX/31hwcHB0zSTadyBOXKIogWACEEMBYyoV4Vrkf45IRFEEmyfkxBEDz4pjd96IMfnJiYaDQa7soppRbFG75KW4rU+gdOxYA+Zbnvu3aNVvVK/ZI3PT37hjf82Hv/+I+ioTsHbl3FqH4wQbSGj+VQ9GhcpdBTcFuL8KBbSBCHADQ8G+mr3J69t73tbR/64AfHJsabSF+cOkVHYpS7tVrtwoULZ8+eHR4eLpfLboXgJm+c1ZjT7Wo6W+xmz8h6X+OwoNFoLCwsPP3001NTU2tra57ndXV1uWI4MQqXUqiySEEoxsbG/vRPP3DPc//Fr/3qr4ZhIIFjKk1cgYQwQWwLBELNAUtTWhRkJH0Fo0xYBOEiTRk/nPr95MOf/Dc/+G+6e3pQ4+PAE7khmeSMS/XR43x9fT0Iwxe+8IWvfe1rX/ayl91zzz2oax4eVlZWvvrVrz766KN//dd//ZWvfMX3/e7u7iAIrLU8rRaryxFKsba2/omPffxl3/WyIFTpKjEsWtVRon6DIJoDgWiA5MYJS9cDpkxYBJHAROCi+rixsfFdL/+ur331a/39/ZH6y8ED7XWl3Zo4MMGAg+d5c7Ozz7/33l/+5V/+gR/4gZ6eHlwtDEOrIh/Iy5Wan7YBRSsrKx/5yEd+7/d+77HHHhsfHwuCwE3+gdpwpNari+GVSouLC/e96EWf+MTDlUqX6wVN/QZBbAtgR2A9J3S2d04BBQThIBgDKSKlT3i+9453vOPNb37z+Ph4o15H6zCKJfyLKV0939vc3Gw0Gv+7olar2aoGbtgP7t4N1b01oivt5KxAByvG2OLi4tve9rb/+l//n0qlUi57QSBS0VN2c9/3p6en/9v/+99+6qd/ypZOwlJM1H8QRHNABCHjqvoZkyCFVAN5QGdPZUNzbWsEcdywlToZsFAIDrC4tPgd3/mdVy5dqvXVsJ4gAKDgsbO5nPONjQ3G2Dvf+c4HH3wQVd5cf8ZdpI1sz3klfcrcJUIIFMN/9Ed/9PM///NCiK6uLkyZl1sbeG1t/fbz5z79qU8PjwxH5w5cmmQ+1G0QRBM40y+SevNBVRzThUl5pARjMPD+ZgUgiMMLYA5IDNhR5uIP/MmfPPH4N/tr/VYmpeJ5APjW1hYAf+973/fggw8GQWAzPkKGlNbrOmTt73mZIKhsWJTneWEYBkHw4z/+4+95z3t83280Gqglp8OW1ML+Wt/jT3zzAx/4E7OOwPPQKW4JgiiACy7xbRFCSWEpdDZocmAkCHwXpKqYrWotbGxuvOtd7+oql62MtLmurHAKw2BjY+Od73znq1/9Q41Gw/O8VAWheOfJbBip4oP7e17J8F916Pi4mHuyXq//0A/90Dve8Y6VlZVsJkuTugdCIbpKlT9+7x9vbGx4nse0+6akoTtBNIerUmRa52UQ56+NOhyby5ZkMXGcUeknw1Awxj71qb//xte/0d3T40bpuOt6nrewsPCzP/uzb3zjj+GcaErEJnbsqKG3OOove2jTNqFnvNX8bqPReNOb3vRzP/dz8zdvlkolnCdO2qKZEKK3t/fxx7/+mUceBQYiFPFBbtn5EEQHwsHG7EkdPaGGwozLg+kXCOJQAegADRJdpz7+tx9r1OvgcS2BICGNOOcrKyt33333b/7mb+JMsJr37YxXyWlhpNuiS6bneVLKf//bv/38e5+/tLyM09jJkGU8cX9tbf2hhz6ix/LWuexgT4kgDjeRnDUjX1WKH1ASCyGFLSF+C2akCOKwApjfyfO8lZWVz/+Pz1e7qyG6OitfXzedMia6+qVf+qWBgQFd6BPHs50Dek7pGSgedQphGPb19v7a//GrQkVMpHoDVZkCpAy6qtXPffZzqysr3POEsVXjbJaWw9SLEEQSzkAH/eLwlyu/K+ActCtjR3UeBNFu3LnSSxcvXbx4sbvaDcJM2SarlWxsrN977/N/5Ed+xPW6OugzaB3Ua6UeNOhC+wxjq37wB3/w/vvvX1lZSRVVVDPHTAhZrVYvX7505coVMxkOuqILzWQRRAHGMxOrjTP1yoCpQ3jQjSOIA0fHASgx/NWvfXVZmWFtQBGYkmHKjVFsbKy/+tU/1NfXZ32bD7r5O8Km0AD0O2Pmu5SyXC6//vWv39raQsOYW7wBzQC+562trX398cd1Pj1luAfSgAmiGI4zXNKE7GGtU2M8Iojjjg2lAYBvfOMb2dQZ5tfoJSqVyt/93d/Tern7Qwa++lzlkQRp1Vtzeq94xSsGBwcxJMkNvjJzvpHK/MQTT+jeQyX1iXoRQQKYIPJRLxsDQEGsPpETFkHEONL02WefzVTrswKYNRr1U6dO33HHnW587YE2fadYOanEqgmtsudy9uzZO+64c6u+ZfNWGjHMADz8fPXatdg5XCYCK8gETRApOOhwPZzJMS8fOWERBOKkfFpZWSuVyqn8FbgW534QhLfffvvAwIDrF91RuGFIxr1Zh0FH9PT0nL3tbKQBs3hcrgowSM4lB+Dcm5ycFEIVRrKnT3KXIArgGOOIvoy6oCdg9bSOcyEhiH0BUz03Go3FxZtusirUD+06jUbjxIkTpVLJ9YvuKBImaD0ZDPovzvpOjI+7WeJdFVky6Ze8hcVFlQUMdHgFacAEUQwXgD2FLlwK2gkLx7gdN4QniH0BAOr1uhItzM0fabMtormoVqvhhw6UvsyRkzjIAGmirIAxqfoHLClhJ8b1aUo7iOdBI4JLkLbAGjlhEUQBHIwzI8cCwFLGblkEQShwViaVKtKt04eUSqUOnPq1xHPAmJwndsIyZvdSqWTdntMOImoTqaoqSe0GTXPABNEMH3TOdDCzXcr7igksyYArdWyHQhDtAa3QAF6RGIknRDsY7AdwAtjDr9KU1o+T9WQ6BKsHg7FUY/ZOkrgE0Ryu3jQdrSdwShhTYTFywiIIFxUc4FQryg5Mj96bAvZk1VdpFrqYxBvaIGAtACR/CaI5vprowQTrHEfA0SukEtiS7ksQDgJA2MwbKTFcJJI7CuuEFX+1ljFIrxkPNazABZuS0zFnd/QVIYh9Bb0lFDgBzFT8n2DkhEUQLmo+NA78dev45thmm3oeZUsB7mvLpTmo1OG50rpGJX2ksC8A6czZSh3Py9ymCjVkx81k4gMmwNLTviR9CaI5vg320zO+UjDgyutR0stDEA6g3pT4tbCiNyf1VVPPo5S0ztWbs+p16ljZfaZWiwcHKAvdesO66oK1KavIB/0Lvvh6I6niEaXpHyBx9Eg6Y7SiZJIzLOQC5sRJ/hLENvi6ErB1a9ROWJKcsAjCRckawKQUqUBYq8uizwSzSiFmY3RENiuyK0EzXbnoa9FPKcXaasDMld8YuqundlG+qkQ82gmLay9oY33W+zQWeKFG6sDcsEWTHFvqk1epfUgIE0QhPuaRk6A7gEj+ykgPVn0Hx5qmnZlVgCB2hnTkYNaCyrnkHN8VnQMZfZ4x6jf1ghgXCkwyh2NavXvp5pFOCUXM3ujsaffpYKWMs0pai7fRzKWJX7YpJbmUGLzMmafqMCQyw6c0ecDLYKzTUVehVGXJpIgvmzQuJbtBJsckZNAmjiR+nEkds1JiITJywiKOHymbcUopxUgk17SbzXgVqYZCYCXgSCyF2pXCVBeL5JkKZ9qNTEUVOzTJMZhWT4GzSB8FcwAPPKyKD8YnucX9K2EsQxEqPxCtytvzxfPSYRGRziuZngiOTdk4oJBGeEf9yI7O0CFrv6eeiDh6+AKkTkepSpyCFIxzIZkxJdFjTxw6EhOcACkpmKlTlF6haCdFx8p+SE3fmqlW4IC1CpvJnSAI1jc3Ntc3VlZX6vX6+tr6zcWFer2xtra2sbG2ub65tLIsQrGysrK0tBQG4frGen1ra3NrUzIWNoJ62GBC+UkCiFCoI3pB2PA8L1I3OSv7ZQmyUqmUvLLn86GBAb9c6u3r6ypXKl1d/bVauVLp6e3pKleq1Z6+vt5SudTT3d3X11ft7q52VcvlUpPGYyIOK1+Zkb5aJEusahrqa1VwTbOXPZ60dq672TNm1NIJMO3IY9tbX3TfCeLw4KPB2VihpOPFSU5YxMGT6lJxodu3ZjvZlH7prplaLSVchbK4Cpt8EYDHETbgeZ5rE87ZlSqIND8/f/PmzaXl5Zvqw8LCwtzc3M2bN5eXlxcWFjY3N5eWllZWVxv1+tr6mhCiUW+sb24wwYIwCIJAChnKkIMXjYgjdVMVO8DGcB5bYrXjhsqSYVoh9cwRUwoxx9xdABBp5BxwhO17vgRWKpc449zzK5UKgCyXyz3VHr/k9fcP1Gq1nt7uWl9tZGRkcHBwYGBgZGRkaGiov79/aGhoZWXVuHFptdcUHdSZo73oQvn2KrkT5Fm9P3UTjRuK9qd2ps/1tzjGuMA9JZWnLPvkEMShAp2w7EutTFYS321o7qVJEPuKG+eTmWFtSXlFilaIHabMap7SXLPaq5QyCIKlpaWtra3UJjbtMyaC/vBH/uYz//Dozbn5jY2NuoLzSBCGIlJYpRCYqzJSWj3P1lPqKlUi3ZVXsanchNKicM0fQJgC3umTjf2ac8YcOqcmyksRyekwqDPG6ltby0tLkV799NOhCKWyROPZoUJfKpV83+/u7g7DcGBgIAgCz/P0BXQHOh5sbG5OTU13dVU4577vF90UYQYHhfWV9SAIbwfKYkgVKXedvbN3mTRg4vADOF/l+myqNHKxTkxOWMStJKEMFegxWbmS0qvScavGT8rOv+ayurp68+bNubm5hYWFqampZyafnZqcmpycnJudW1C67NLSklsKN+V4DABbW1tBEJRKJU+Rali2kVnBk11hp6QuYIpmFlpH3UfpaxVovIZhGPq+jx2Cex3c6XAhRLXaMzBQGxocGhweGh0dnRgfP3369MmJEyOjo8OjwyPDI8PDw03aHwjBhNA+YsrZi8dKN4Q4LEum4U5ZO3ImC6j/Ig4lEI3NpWA4fYRCWDCdSB3XoIEkcRDkStlcsWTtnMk6RflFETY2NtbW1lZWVmZnZ2/cuPHMM89cvXp1cnJyZmZmfn5+eXl5aXlpfX1ja2MjVGI7EqS+73nc4x7qr1kbOCqegqmauIwLZ/qmSE3P6m25U9c7zLRlShiZPWTjiVMtaa4+plbgwCPZKITHPav9Z7cNw7DRaAgh8C/K765Kpata7e3tRZv2xMTEiRMnTp06dfr06VOnTo2Pj/f393crsjPo6AJmJ4ndUUI6ANo4fuMowTkd8uIiDiMQyJBL49ugYgGBKRdoXujpQBD7Qa54yI2FtdptE6W20WjMz89PTU3duHFjcnJyenp6amrq+vXr+GFxcXF1dRVlA2PMM/i+7/nqE3ATIY+50dEiKous3xiKo3NSgJNCOU+SNUmakXtNcgV5Rp+2vkr5c6upC4vm5SL3tBxXKbMjnPV1d2V1Xyv2UIimZoKDIEB53Gg0QgWu7HleT09Pf3//2NjY6OjouNKYTxhOnjw5MjLS3d2d+8yo2ktCXw2sh2GEtGqYjfxKgidhY8VUHLM0Fu5Uj0eOXcT+AYEIIfHYSQAuJLOvGT1zxE5p0f+lic6XSnDRRNAKIZaWltDX6dlnn7106dK1a9euXr1648aNubm5paWlzc3Nra0tO+NrZzQ9z7Mey64sl0Imsl2xhPa0be+8jZm3+CI4+jQwnDN2KpU1NSwngmZzxE3BfHnu4KDQ0oDZN5SYA5mI793Whzw1+EB11r2DYRgGQYCy2S4vlUrd3d21Wg015rNnz55TnDp1amxsbGBwcGhw0Hplu9dCShEKCTJqLWfAcdDAki5gsQk7PZ1uRlAEse9oAWxSqOOTZ9LRkdAlWiMrSpuvYFcryt+E4jC1kyAIVldXJycnb9y4gVL22rVrN27cmFEsLS2tr69b8yNKWfwbC1dHqGcP7WbYSAtaZvTbAo+w3A/Zcy+SVQlbtHIEwyBbq5flzovbQbN2WNKeW9sK7MRBUXPN9Sd3248mABRMkQDmzNUXU5cxdTFTEwS5K1uN2W4ohECpHCjs+pVKpa+vb3h4eEJxVnHmzJlTp06dPXdbX1+fz73UKWD4tE3JmZ6bZ9rPK3EvjNmaNGBi/0g6YUltRrOzwpJJDuSERewMN1jT5BlOdNNWBDaJmp2dnZ2ampqcnLx27drTTz+Nqu309PT8/Pza2lpKVSqVSpVKBZVaa2LNTq9mZdj2PoaY38Jkh8rtjrNKfOprE10zcdFiucU4h9RAIZv0o8nRc2eOi6zZydMtPLvs3Gp2q9yvuX5SuTb57My33RAPjbe1oXANG6VSqaura2JiYnx8/NSpU+fPnz937tzZs2dPnDgxPjY+OjaKPnEpRLQ7gdPbalDhhH5ga1Ono+3VTDqPN8lkYtdEGjDXmi9Yw4x+uoROLiCFVGlfCULTZNrSXceVH9hH5brYrKysLC4uPvvss1euXLl06dKVK1eeffbZycnJBcXm5qaU0vO8crmMpmNXo7VOOm4Xn6vSbSuudtqXFsndJlfMXeK6jGWX5wqh5pc9VwBnXcezEj07NGm+8+wp5F7qJtcqdxZ821vmttO1kVhf9yAI6vU6TjAzxrq6umq12tDQ0Ojo6MmTJ2+77bY77rjj/Pnzp0+dHhoa7K3VyhkLtpA6UAuzWCvzNcQJxeysgLVbF8SDEUQrQEOG0djPDO+i8V0kcVH7lZiMgATwEWPX9rSiTtP+ivJQqJ/8PJ1jdW1tWjlGoRnZMj09vba2hjoNhsniZK2N57GdrD1WUY+fq1y6tk17CrhD13XIdU1qLgyaT2+7wtVVHO3IQOd0lBJFBXokWdMrLndje7LDi6Ib5B7XdRVGYwMusVHIuMTNjukOm3JltisCi9qQK32zA4ttp59bsf2mvN9d5y+UxI1GwzViM8b6evvGxkZPnjqF88q33XbbGcXJEyf7an3ZQ+C9QD8vd7gUe77lPRKpdNYsYxZirn2IvLSPJRCIABhnznySccLCSsH5mRCIo0STPq4ocCUV/JM7ZdtoNGZnZ6dnZq5eu3rpqYuXLl26eOnS1NTk/Nz86urq5uZmvV7HNcsKdKhx5UF2mjbbjNRCV0i4s4/ZM81umP3aRNfEDygarbnbXY6ZkxH8FWUAfsaZ6UqlgiderVbxQ5ehqsAPvu/jByvObYoMK92xJVLKtbW1IAgYYzgj3mg0VldXt7a21tfXV1dX6/X6+vr61tYW5gkJggBjl12LPV5/6xlu74iVPq75wco824Cia1V0K4vWL5pcb+7ylvuQ2NlulKZMRvdiq76FF4pz3tXV1dfXNzo8cvLkyTvuuuPChTsuXLhw9uzZiYmJ4eHhcrmcaqSxXUefPQDG0woKCmbHN53ZdCKJS0Fy93gDYRigicU6cajPwi1HSHQ6TXygmps3sy42VralDoG5op555pmnnnrq0sWLly9fvnrt+rVrV6enp5eXl8MwtHE+COq4rmOUm5qqSElyPYaaTz26e0CJlTXMFim17qRpyjqNAjWMCVDO2kvkeV6pVOKcl8vlqop87e/vryl6enoGBwdHFIODg72K7u7urq6unp6erq6ucrnc1dWFUrkdtz2BlBIdwlXS6YhNxerq6vr6+tLSEibOXFhYWF5eXllZualYXl7GrF4YO4TRRFaZtrYKvKE29VVqlJZqRiuj+SJz945OGe9X6r67D7B9ltDbKwyCeqMRRrI1esxqtdrY2Njp06fPnrvt/Lnbb7/9dpTKg4ODlUoldazs0+u2A0vN5STzcko104TyMQSCUJUmU2M0HXkkdU1QyoR1VGkivbIKbq6fVKPRmJycfPbZZ59++umLFy9+61vfeuqpp27cuLGwsICKF3Be8Ut+qVSulH3fx2rtam5Noq+ulDJ3/F80kWmNw9nG54YLW7lbpM66plTXQ8fqrFrAKu1QLWFK1eHd3d3VarWnp2dgYKCvrzY4ODA6OjwyMtavGBgYGBwcxOTJtVqtt7cXvcN2cY+y3sWteOTmzqcWpSVpghACpfXq6uqi4ebNmxj0hXJ6cnJycXFxa2trfn5+XYF6v8216YZ7uQ+StbSnbkTR3XTt6nZhE5u2q5S7P7nPQ2qkpR7a6D+7fhiGW4pGowEex6FVf61/Ynz89ttvv+OOO+68884L5y+cPXPm5OlT2UhlIaVQ8ws6mZd2q8amcPcu2qZQT3vcgDAUqgSSekiEiUNy5jVoUHYEaO5/lIq4zXbW9Xp9aWlpamrq4sWLX//615944onLly/Pzs6ihoTaLVpQscPVCoGaDNZ9ril86+4c8srFNrcY50rfovGE+zl1RlaRxQnCRqOB26K0sLZfq7COj4+PjY2gk+3Y2Fit1t/TE+mv1Wo1qwwV3YJcpTCrc2e3be4elb16TTTF3GndbEtaFNhBEKAyvby8vLq6Oj09jdHYU1NT8/Pz09PTqFWvrq5ubGygMLPX2bWF2Dnp3GuV9f0u+uDWZk75muWO6lJRZykfK/eNwMFcqB4adMAOggAAenp6hoaGxsfHz507d5fi277t206fPj0wMNDT05O48raYJDPRyclkR9khBXHkgVCEcbemBmgQPWpYakzSHPDRIBWH4/ZHaDpLqYBTU1OXL1/GQNvLiqeffnpmZmZjY0MIgaoAgnOZ1kPKWvxsDim3m3P9VgDr2xRIiqw67pqRm3dSqj/V7quxddEkYLJXo1wuY+85MjIyNjZ24sSJsbExnPMbHh4eGRkZHh4eGBjo6ura9tpm54CbC9QWKVJ5WzRXNp8xbZ3sHIQ7UGt+go1GY3l5eU4xOzs7Nzc3MzOD+chmZmYw7TZOWuPNRdXZOt+hYLZPV6oB2euQdQzMnkXuVTLqqY78MLuOPfXUGoA1amIfBWAiCBsGFMmlUmloaOjMmTOoJZ8/f/7MmTPoel2uxDMLuHKuxx91uccHTMShekt9wwU5YR1JXAcQq/JauTs9Pf3EE0/80z/902OPPfbkk0/euHEDkzXazXFu0i3Jl/KYtXvOPipZydFEEhR5ujaJ5XXVdyHCRiNoNOrK44l5Hu/u7lGxKIMjI6OjoyMnTpycmJi47bZzZ86cnpiYGBgYwBTERdfNuCuzTHn+/Nbmnmxn0dwbIPVT1l/a1aSLHLallOgUtrCwMD09fePGjStXruCkhqs3b25u4oDPasxukkv32rqJqXMzVGd9+HPuDqbgFRjpG8+QSKUQc8xGYlONmj6TA0fDO+4Wh3oYDYV7rVarAwMD42Pj33bXXc9/wb0vetGL/sVz7p44eQJ/tZI4dbVbny8nOhcIhYBkFTMz+GPkhHUEiDsdxwEEK/YwxlZWVz/1d3/30EMP/eM//uOTTz5Zr9dR/8AcUrb2gPXmdfdsHVhsjE0TmVokjVJxPrkbpjxvsY8LgmBzcxPT/ePKpVKpp6dneHhwdDRSZE+fPn3u3O2nT5+cmDg5NjaqNNqRJgE8qGBlY2zM10T8SGxEPx5+rEWad3MvcVcqu9O9uU58lnq9jqnNZmdnMQ3LM888c/36dSyrPDc3t7GxYdOw4FNqA8Tdw6UmmHNPxLY2PebAJeYDZ6ikYEJspfiyaKGas5Op8DbE+o2jgQRFMkrlcrl8/vz5++6771WvetX3f//312o1NBXkls8y7dfd8zF53o4PcSYsLERKTlhHEmkS8TXCoOxHoveppy6+//3v/bMPfejyxUtbm5vdPT3VatUVckwFWbh5ELPyMjcvUjYUtXnDcsNVrUQMkggh0Hm4VquNj49jgoUzZ86cPn365MmTmJuwt7c3txKtY3+2ZvJWNdQdPf+drgG3zk4Fc+4eUqK6KDlaEARra2sYQf6M4erVqzjfvLi4iC5grku2dQFLzb7n2vAT7lraFl1ssU9kwmomE+3jbceRUsrV1dW1tTWPexfuOP/AA2/4X9704xfuuAMf0azjvQ0hdqQvieEjgilHqMRtvJScsDoTN/Yfkv0gyjPf92fm5v7v//Jf3vPf33P9+rXevr5qpSJMeE96Ig0DJ/I8m4pEcpG3S6pTdoUfqrDoCYV/7VxvpVKpVqt9fX025e/Zs2dvv/32s2fPDil6e3tzr4PagwxD4VbmybTf5jKix3sbmsw9Z2cEdnExczNepaRX0dwEumTPzs4+/fTTV65cuXr1Ksrm2dnZlZWVzc1NDL/GwDA3wUvWTTq163wToKmkBObpYcXT6u4h4kg2JmWowrg9vra+vrayeubMmZ/4tz/xsz/7cxPj4zhMTJ2v1oyMikSq8JEBGiJQeYa4urHScVClcoQdifbhTMYX4kQaY+xDf/7nv/7rv/bEN5+o9fV1d/cEYZDY1r3RADITxZH63MSjJyuJI+3W45FUV/Ies0DY3BS+71er1ZGREawRa9Prn1SMjY3lakVupdisRpvbPEeQaJti7m4pNLMV9n6hsqK3yX6yiiz6AGbX3NrampmZQWF8/fr1q1evXr9+fX5+fnJy8ubNm+vr6ygabRCzdSfM9cHOjX1iTvXH3GCw+It9pXA4i8YhIZiKauPANzY2lpaW7rjzzt/6rd968I1vxPEo9zwcBauQJeHMepD0PTpAI2zw6A6DKZypMmHhb/iXOqAOQbpFWzHiUC1oNBqlUmlxafGtb3nru9797kq53NfXhzmAdnpzUzI1+6vrBYMiMxTRfzqTQxCGIvQ8r7u7G92MT506deedd96hOH369NDQUK1Wy008lDIe5v5tMgOd+rvzq0vcCnI9ty3ZydoiF7Dc4pXr6+srKyszMzPPPPPMRcUzzzzz9NNPT01NYb4R1FYrlQrKY3dGJqvxN/dgcJ3/leJqxsRqAJpSr/EJ3tzcWF/f+LEfe+N//L9+b2x8vBE0fM/DwrBmRWGnTux2REeDccCYU42DyctC5Qg7COkMiXUgt+oMOAMhRSBE2fe/9rWv/eRP/uQXv/TFocEhDhBKyU2YEPYybkrklJTKphhMdW2ujRf9TTDFklqVeZ5f6+sbGhqaOHHizjvuuOuuu86fP3/y5MlTp06dOHGiWq1mz8j1h0rEDTdtxjZXiaTvIabFu1Pknp3rbOUqsk2KbmE6kampqWefffby5cuPP/44+mOjoowGYc/zurq6UCr7vm/37L4aWWEcL1eVrRgDwVVoiTClLUFP8VgjkJRydm727ufc/e53v/slL3kJJgDhSuvlHISQnKtXHOtVQqEjG9Ep6FRtKns+M1qwU46QnLA6B2mTXagPONXk+/4jj/zDGx9849TU5MjwaBDU3WTCRXfWVTVSMs+mILbi1k7fYhAk6rVnzpy57bbbbr/99jvvvPPcuXPoG5XtBNG52rakyWOWDfdsYvkkG3KH0vzGbRsV3eIh3Nxq2WcyDEOUylevXr106RJGw1+5cmVqampubg4nTWywgOvn5cbB5wwplJurksN2higObbNZM33fX1pZrvZU3/0H737tD78mCAPOPWYShDAW/6EyTEcA44SlxG28lJywOgerAUur16rlQRCUS6W///Tf/+jrH1jfWO/t7a5vhVhYyJleYNs661oxbMu91et13Ao9pEZHR8+ePXvhwoV77rnnrrvuOnXq1Pj4eH9/f3ZXOqNjMpND9uhNMiySTCWQPZo0Ug5f9qsNH0qtv7CwMDc3d+3atUuXLj3xxBPf+ta3rl+/PjMzs7y87JbLLJfLOEjFML+s9cjOLbuB5InGACuXSisrK6GQ73/v+1/zmlc3Go1IOUa1V6KhUkGPf+cDgQh4pPLieEoyVQ1JCpZKk0ad3aFFOq6YHIWlevN9z/vSl7/0fd//qo3VjWq1GkZDaY6hZWqmH9KeIrgHRyfA1FE2Cz8AYPAPZqU/d+7cBcWZM2ey3sgpM3I8WZuc2sh1ctn3S0Z0JnschGW99HOVabuyu2ZWUV5fX79+/fqVK1cuX7781FNPXbp06dq1a7Ozs4uLiza3hs4WV/IxxTRqz4k26OEwSCmYNixFY1QOvF5vBEH4F3/+oe/9vu9FGeweHcw7THQ0EIQNAM6s+w45YXUmTlpHQEenK5cvf8/3/usbN27012qbG5sYd4HDaGPI1T4dbhYhzO6LGTn6+/sxsvb8+fN33333Pffcc+HCBSwzkHoeUon1m0zcutBDRRwgRVHLuQHNVlG2v6YKbAghlpeXZ2dnr169+vjjj3/zm9+8fPny1NTU5OTk0spyGIQlVVMy0oyBhYGN9wXJJJcQ7YBzJkUkilVlHN/3l1eWh4aGPvGxT9x9z91B0PA83zaxiQM20UFAIAQ3QeeMkRNW56E1YJPqSqgMPkGj8X2vetWnP/Pp0ZFRHI9nR/r4F4f2WJauXC6fPXv2hS984b333nvhwoVTqmL5iRMnsi856sSuu2l2Ni43fqPIcZTGecStpMl0cjaILruanbLFoSdOA2ePgvL46WtXv/bVr37xn774jW9848aNG40wGKj1VyoVFYkQacosGgEzO4BFU5YQgnt8ZW3lxf/yxR/96ENd1Sr6cqkDURjSEQG9oJkOAlYDKvUvOWF1EjbRlVQFiHzPe/vb3/6Wt7xlfHzchhvlpsKwqQyGh4dfo7jvvvuGhobcndvMAEV+UkXhti01O7kTesyIQ0sqmC33yW9uuJ6enn7kkUf+6q8+/Ld/+9DS0tLw8Ih6X2XKkTsS7Tr/pPRLpZmp6d/+P//9r/3Gr2OeL3VQYVM10CvT0UAgQi7B5OnX4WnkhNVBWBuZTbjxla985eUvf7kQApM5W1U1ZU8rl8sbGxvr6+tveMMb3vrWt9599934K9bpsz1IKqnkTuOGyXOKOPwUZQ8tUpSLYp9SQ1vXz9+Vx1/4whfe/va3f+QjD5XLpZ6eHmujcpJRR//hhkHQCEPxiY9/4v6X3C9EiDOGDCSLPtAL1dlAQwquCqWbEG+h/HQYZcI6/Ni8tdJ8lYyJULzu9a/7m7/58PDgUKgTYOG9k/YvJuhYWFiYmJj4/d///de97nXWquyalN0DNbEw56oFpAETR5vmL0XRjI+NNWKMve9973vLW94yPT09MDCA2m0swl2fDoCFxYX/+aUvfegjH610lVXGBk+qjvpALwDRBrgqxu88OiYBf4tBdcTBAiY5LEpQj/OPfvQjDz300PDwcKhDIIRx19B+dVLKSqWyuLj4/Oc//+GHH37d616HxYWwFKvds3SwIrnJh+zXVhrfJBiJIA4zzV+KrLqMGrDneR73ojcuDB988MFPfepT999//82bN0ulkq3EgDNKsUU6DIeHhx/5zCN/+Rd/ybmHkXxKBSY6HghDtH7YaHFywuokpJGuOGKWUn7Xy7/rf3z+84ODQ40ggLggK4YRMilFqVSam5u79957P/zhv5mYGFd10PyiW01CkSD2SNqnGl9YyUIRlnx/cXHxgQce+PjHPz42NmZt0Sk8z19eWXrePc/99Kc/3VWtOik5yGjU2XDGPBxMocVDZVESjAnBRKp6F3HY0PdFpbbA6aKHH/7k5z73j321WvQm66zvYA3VKqU7LC8vj4+Pv/e9752YGA/DUOWgZykdlFRSgmgX7oQx6Aii6IX0OG+EwcDAwHve854XvOAFS0vLjgkqkXE6DIP+Wv+Xv/LPH33oo+iAbTK+E50Nx1ykTGo3LOVaByoNB00wHHZsMh1mCq28733vFWHocY+hM7sp++OUeYlWe9e7/uCuu+4KgoZJ45yuo5CbdJcgiF0D2s1VJf01Xq++59fr9fHx8Xe96w/6+nqCIMSXlDGerfoghfyzD/6ZCEO+c3dI4nCisoODeiSE1oGZNWgyaR0BiEOI1YCFikS8fv3633/676vVKjp6ZFM6e563tLT4Uz/106961fcHQaAszzbbiq6iZCZ9GQ2xCaKt2JcRazHgUJeVSqV6vfGCF/zLX/iFX1hYWEDZGqebNLaoMBQ9Pd2f/dxnr1y+ot21JA2ROx5uBS5wHUnKWKKA5cG2j9gW6yr16KOPPvvMs9VqVQ2hTQU0x4tqa2vrttvO/OIv/iLaq/OiL6xLFCQ8AwiC2CvWyCSV66S0zlaeB2EY/rt/978+5znP2dzc9H1tfUwUJQvDSqUyNTP9yKOPoNESGJmpOh6OHa9VhND4rDIWAkBORApxqHAdiR9++GFbqiiujGSsVb7vLy8vv/a1rzl9+lSqvqmzPxuqJJ2vBEHsHWtbCs2w2CZLj97ZsbGxN73pwa2tuknYztzKhsCj/zh4n/j4J7TyDGSF7nh0oUqJaQy1E5aqfKOcsNCYSYOsQ4tNsrG+vv7FL36xq6sr60uFq4VhWKvV3vjGN6XCivKSW0kSvQTRbsDky/GYLgNsflAhSlLKH/3RHx0YHGio7HUpooE1Y9Vq11e/+tWlxUVuhDR1zh0N16ZKiaWhtc2Zc+2ERcOrwwwYSxYAPPnUU1NTU+Vy2ZWvuBqqv6urq/fee+9zn/tc1y0z8w6TBkwQ+0TKvREwk4ZWgFQB7wsXLrzk/pesr62hO3QqtlhKWSqVJicnv/71r5Px+WjARdSHCzMhIa0TFg7XDrp5xDaASQf/1FNPraysYIYdNxWA/bq5ufmKl7/c87wwDN2KvMlhFpi/4HwlCGLvWG8Mad48m8lOlUNS+V9f+cpXbm1tpSOXzOeSX1pZXfnmE0+kkuQQHQoH1wyC5RgYUCasjkA6xe2vXL6CNQRj1yrr9MFYKESlUrn3+S/QSq/ysqJbTBC3EEilVY/LcpspI8bYtz/v2/v6+twKKKkyiEEYXr16NVWCjOhQOJdoCdH5zxjjWI+BnLA6AzNAnpqaTOeLl+ZVBwgajf5a/5mzZ2MnZxpjEcQtB8e/NthXLdNiGcuenDl7ZmR4pNFopBw13GIPU1NTmLadOudOx2TbQDcsjFUjJ6xOQCbzVS0vL+PyeGiMGrDKwhGEQW+tb3h4KPpJlRMV9OoSxEFg0ucIAGHS4IAdSQ8ODvbV+jAvtLuV1Zs55ypcOF3rkOhEuNSZsADscIycsDqBlAobBGFqqKTfWFODpbtarXZ3qyFWPPdLrlYEcSvB1LAm9Iib0mTaK4sxVumqVLq6hK6kErtJxjPBHDY3N93Cw6QgdS5cGCulcX1V9fkFVWI47KTeuWw4rzVER8JaMr9ULpVKDJOt4FIz92SrFZL3cztxPcoP5WWVTqqWw9mPp1rY4mr7fSI7Olw6rsAk30BXDFOiDDCXrKpUWCr5fsrBKlHukDHBpM12R35YHY0fT+PHMlj55FnPAOKwYiuXZYqBa0/L6C3lJvkdOOFJUoV9G/2YmW0wTNGkJM2PBs7NFN2kaPnxJetIbkuOtVA+uajosrtCUUlmfbTtCjan9tlkJ225la1Xlc4tCtKkNn5uKZEmR9/24mRPv+iabJs13ZbrZuZtw5QLkKyNxPB9ZMxTvjexB6V6i1NnBCy/4iHRcfg25aBWeiU6YSXKEdLdPYQ4kUPgpnpXy9DRA2Wtfrv10Mqm4MG+w3QQ0LJ6ljvqzybuOZwa1eEh94rtdLXd/Zr6murfc3fSrluZtZo2f3KKRGNq5VYGNNkJmiYXp/np5x638JQzn8wrasP07V9pjVe2dLc9hMn/nA49oreso/FxGlCZM1SvzVHzESoaidsbTzL4sCFj03F0rzj3nKGSZJHmyzDnMw6xOOflchkAdIy/sytIf/RyfiGOKIf51W4xCqMVDXg/aOOB7K5KpRJI7eGRU5pM2aBh35pB3GJ8ocyR8V2UureW7kLi8OFowJKZwvtmOAzuNLAOBQ6CycnJvlqvCFUtNBAMq6JJLcxxGwnAJRNqgA4S4sBwu9fiYoVYBMKOC1R5j06cozIWAYlqCpjaUPYqFG/kbB2vmPmc0vCYccVxdhfHiTqKnQB0YdfbppuQsZQyk3GJWXMImktsRghcMf4an/oBYa9D6oP+y1jC0hNvFV0ZSF3rZtjpl4QhV78G0l0vftr1xXHeFtOYWPY7yTVM2/RllfGcjwD93oK9mULq7YIwbAQN9266kjieSE6cfce9YoTG17qSzt6vHhMOQidyaPVpJg4UQENiVgAjYRh2dXVdf+aZ7/nuf+1VOEiVfQWiV16AAAmqTn90x4Wd+NfiXIBy1HSeAnwqOMPHw3baYC1pibkx3b5E77SXJ8r2eWZQIPVIwvzuDCh2s63Q4xndUiElZ5KjxT7+AZLbqgsiQWCROakPIG0lMfySbIa0JkfnZlkBztxshdFdsnMFzrS91F09l1JgRZXURIJtSUI7TN5Lt6nSzFJoUadlh9yJP6aML2+L91rqQY4UbvlLMNMo0pwxSsHEoCQxwQoQOzWl7m9qigWY9XrSp46PNBPRH5msACdZtiU54yUZNwLvFRZ41Xk2WDxOVl8F6EEDMMkFhNHYKn645PLKak9PD2a4S83yQt4sdct3hzh0+PGoGN9BsMM8csLqFMAJLnS1MGfeSfUIc/NzRk6oVYUUIExZMy5lyBytCKJVBWBlf6OGYdcGKskHWNniHNNRF3QX5Ujd/RDAnEFocmfKHQrg7LYi8ZXri2NeCpO3N29bqQL6rFEhFkPmIrm7sEpw3sylzkuoZY0W7WDeShNZpm+iUAKY2bDwZuJyu2sfq2j23jXfYbOjyOQgxDkGXiKhgm/0AEGax8Ow02dE2mEOmPuo7hEPGY4go4Nw0yShJ9sgVKMrvH/FTw4YMQzmBibOR8lcPJ1YurP4ToPO8atMx9xcGXOBJYj4yZGYi9D3/SZGddfvjKW9L4kOwzd252gcHY22I/VIuWRJuPVzKsSuUMqaECzhBc0SNkn1gvtljzNf6jhEZkboVlv1WzlQ8oP7o+5UpDQyaB8fG6vUSNPs1kW7u623w2HBXraNd9C6y5veKLYY7+hYMmuG3slh9534OjR9slxyTSypVdQjUfRsuL/u6g4mD4WjUzCjpJR1WgtIDnqcoaW9GjsZewlaG+3sDhqy9HNS4Bvvnj5lo+xofBS0erYrGowZu5Ny5CEnrE4AWLG7ZqxjKb1AMKF1KWEtoRJj/pkV1dxoetrqqQMhcBrSCoOcp8IeDlr1Ed3t+bohUsIevuU9OBN01inV2iUzfVzyTAu2RfsAaDtAKnegu0NlbWwWfVQ4v27656ZAwrIt3QljmTNvXLiH1IT27nG1/NTFFG5dajscFDJ9KkmyHsupiCCIDQzpRMrmpJj5IJNLdo50JhXsa4OztdyaE7Wfjf6gTE4MXZqVq6R++JTZA8WxjgmV2r3DeqrjaxiGIfpX0gRwp+PrUVfU9XAMEAWBM0+MdN9DjuujwbIzV3poFcNVmhXQFShBxpLDWOggtrdqKc0SMaNZs1huwrwW1JQ9nXXWIteabtRsn4xZfxmNa+hrHtyiRa+u6Wn9qnRgH7Paj3N93Atl8/q6/WyRgGnx4uQGqm4bM4oDsuQIYE93MJuwyW1VaqiRNa42P7R7Md1DZB9UZz+wxzNKNMBOnyvnCcm1WNWTQsbaAc7MO9jLYhwIzK/SjqR184TMdWMEgCAI7FlTHHBHw6VydWUYYKbq8et5ILqjh56kCHQ7MuurIZPaR3S/QZsk1fQmB2YqJuFMmLTzjTgQL34MUh1oUVqD/Thxp8/NzxbUfGvzN1aA0okOHOOea+uz66OgcuUoGnpRezNfZfKmpPW21BlZNzq3b93JeeXsM/thm6wReYK/qM2tkJ3LRN+i7D7xcXIvQqoxOzoXgPROTEvsfZSp+9jKPUo/7VJ7RwNaDnVkPVh3MlvfFfB/5hMwPSeMTluYhsN6TugDQc5NweekUql4nrd/Y1ziluFb37y4HzcukHRXDzn4TqIxqre3Ry1DdYrZsbVWhKPXXUiTiBZ97KQxTTpuVLHDlXXIsq+3NYVltauUsphV2vbSTWQ1AHd51iPUXT/vuIlpcnvRslp17mm6TbC9vLZdA0uPBmTasp2ru7vWxWyyiB0JvzwX9BwTRdHmRdaFIlP5tguzLUmVEEhdMZ1uonjYkWpe0enb4LGUlpxRjov89nMuUfYhxPldKeK7zOzrJM2Lp6UzA/eJtb2ttjib1jgelaknwX37JiYmXLs0yeDOxcdOmqvCHICZsKTOVUiZsDoCfP3uvfde9Vk4M4UcTZ9c+VupuxlKtVDHuPA4rsKMxIG5sbyOcLVOXrYvcI9erIjkZ/jbEU22dXW13PT0rR831c5clajJaepfjRC1s+apddxWpcrMpeYUsqbjVs4ie2V29FP2hm57Z7dd6J6CG62ezaToLtn2xFMjpFaGGqmbm1q4m0cUzAyECnuPHbBE7AxupzVkxj6k28m2mepOnyyTz3ve83DkbfVgokMBrHulvCG0IzxGQai7Tk5Yhx0ptOPG7Ozsy172skuXLg0PD9fr9YQ+AeBGyOS6fBb1PrkLc22Dt4Zcj57c+c490vqE605NxFkNb4+y9hbQ5GrkCpUmO2llznKPV2B3TgDZTbISPXWDsraZokNnr48e5wpR6Mwdj0s4SnK/7C8vr4wOj3zmkc+cOXMGDd6swERBdAQQCuHkgDHhhU6WfprkP7RoX0vGQjUW/tjHPvb617++Xq8PDg7adRI3Dv3dbQCE2kEYhvGPmS4GHLLTctvKnuY6yi662tTTWORhtHcZ1vyZz44DcpuXauQu2rZreZy6iU3GKE2M7Tkrszi7h2vGZbHnsVkTR/IyZyI597GRJq9LtmEtOI7lDGiKVs69MrmDodz5jpSa3rwobxiGuRu6JgHgHApHADiDJJSp21tcWpScvef/+++v/9HXh2HAuYdplKh/7lygIUMu3Nk+oRRfE/fPWhqxEgeFNahJJYM/+clPvuWXf+Urj30lFNvGq0R4nmeldXYCknMeBMHCwsK+NZ8gjjJ9tVqlXGY6yoDbLlUKyXgknpeXlu0IuAkAjHPvrrvu+g+/8x9e88OvwQkgzjmlKux0IBAh1+MvjuG/ABwTyWhHExLAhxh3TI0ic2Vp+dFHH/3yP395dXXN47y7t8dTCVZYUu0IgsDzvM9//vN/93d/h3nvUs6WGG5YqVR+5md+plKpBEGQclLNnX8t0jByXUxT7lopE+W2CmiRduLu3L1Q7uZNbKfb+XClV8vaIVMT0tvuKnvFik4894ya77n5ttseNNdoYdpqYmxsM5w813a+Iz2ty7RaLHmcyFRKoSOxnQx8rfc5Wcd7d349NaxMKfepYlApP6zshXKvmH3m7RNrFWIp5dr6+h+/94/rm1u+79vE4ja9Ovf5wsLCd7/yu1/xilc0Gg3cIWrM2CoE37ve3t677777pS996fDwME4aguOxRf1z5wIiFImJQRO8aGcXiMOM0ElsVToeJTKx2FGLfPCDH3zggQdGR0exC0jJ8iAIyuXyF77whdtuu20/T4IgjiBT09Mvuu9Fq6srPveVg41JTK3wfX96evoD73//Gx54oPV9hmGIMj7lTU10KL7OeGVCUJjOhEXlCDsD7a1ssgzizYreTMFswmYVHqOn9wWmNwNoNBrlUqlarVYqlaIpQM/zNjY2Ll++fObMGXzzdzq1tn8nXjT7eyC0JdrqoDhsV/LWPz9NzCFFk/35GN+ZQIiS5126dHFlZSV+azD3nPaVZkKKUqXc19cnhGg0Gjb/s21e9uioYXfuk0Zk8W2ND7A+eRz0NDA51x16TLiQKThngz2AWXMgNyklbYQiMOZ7Hud8aGioUqmgUSvlKoX72drampubw7ijJgJ4L7QovLM25EPCtp44x5xCRzmbwdHx+bRxSq4tFz9kH4xbMOxLhU41v9HR0TkXUnjRv/zmzYV6vd5T7dbnzrmtwg8AIpRd5crA4CDn3FMvY+4znz3iIRkqEW0BM2EJvK04HWO9rw66bcT2SPNXJlVYU+RZvag6O4+uV2R9K6WUAwMD3d3dNsDfnQPGXTUajZmZGX2s7UJL3bnh1n133Uk7KI6pcP2xU0c8qJioxAU/0Gbshf2+kqlJ2Xga1fwKyVn5rFgtejBafHK2Pevsk59a7nonNEELRabHE3Nzs/WtOgceNwzio4dh0NPb09/fn/I5yLYq1bxdnCxxaPFtmLhywtIRwNLoScQhB5y/CbcmqYsDGE8N4xHD4trmADAwMNDX17eysoIWsNzECFNTU614RaV6hP3uHdK9+QF1RoekGXvhoE4hXaVXFyG4Rc0oemKbH7f5r9KZmmWMzUzPiDBMZIJ29hOKsLevz40YbN4GErpHEpULmHGO1RfUNCI+Pjxv0Ep0DFYyQ0JKg05+JzFP38DAQH9/PwZCFKVTmJ6ebi59CWIXWOONaENNosNB0hg+PT2NCnE27lmZoEV/X62/v59eruMMhh4xYZ4eZbpUbvBMuAniD7qdRNvQuVY4E0J0d3ePjIxgqINdwe0mAODmzZs020S0HTss5HnpuTsR+47ghO78/DxwXXMsm9gkCIKhoSE3ApA4hnAdHmqtjkw70wLjpPseRbBDUPl1oiGWtJGF2ek3KWWpVJqZnbG5LQ+68cTRocBG08G489YiFPPz85x7KQc0O5sjhBgeHrZKDnE84QIr8jM3dYCUNpU4cdRAj2mc6I3u7/j4uO0XUtYOIYTn+3Mzc5sbG0dMCS7ycDnodh0jYv/BI2OCVuBDtbm5OT8/7/teKkeNfYmEEKOjo657GnEM4aZahxbDuqg4dUZHFydxb3TPT5w4YZbLbHfged7i4uLG+kZbNOCso+lBib1c11zqB28lh1kDLvIMb+4x7j5Um5sbCwsLmBUn9XTZzcfHx3OTphHHB851kgar9YJxwgKMCKaO6YiBAtjKnhMnTnielyrsY9aUvuetrK7cXFrMZvXbBam4kd09XW0JmyEN+KBIXflbfLgWD1oYOpVZnn1+8O/qyurCwoLv+9kMrPZdw7EvacDHGa6dcrBkO7rJkhPWkcYtdYUCuIz54h2tFJFSep63vr4+PzfXlp6iLRpwUedIpDgkodIpUrfvFh+uxYO2rgFnd4s/3VxcWF1d9biXvfi4Oed8YmKCtZhjizii+CbvBoauCfOsUiasIwsAE4KpGmgRw8PD1WpVKI+sbNodzIE1MzuTFc+7I7WHA5QN2V716PWDuedItELRpWt+Se3C2dnZra2trq4uXJ6KMsAyJ+Pj4+7Uz606M+IQ4WOcuPKDxlzQEr+mFSXiqGDngFWhbzY4ONjb23vz5k3f980aTlVX1S/Mzs5hnRZ3IXHIOVRJno8YRWWV0btKSjk5NYXZW3PD6xuNRm9v7/DwMPkfHHN8M/2LmbBMNDBlwjq6oAZse4zu7u6BgYGZmZlYAKsb7zptbqyvx78SBFEMOl4tLy9ngwswPhg14P7+/t7e3lSI8EG3nbjV+CYdq6qVFanAHqYsBBang6Yn44gBWsQC51CtVgcHBxuNRrVa1T9Lk0RaOQFUq9X3ve99Tz755MbGhu0siiKXmnQlpniatI1o8lSlkmK6bXdrZVr1Tkih7DeJmjY8eqpzxpF76ux0YYv8ZIGpSgPxEQu+ppfnVRew0/Y6UF/Hcqe1W7t97E9Z3Lwm5OrNqRvt7iddRwgS4UTRaA9Na8ljsF051UtTsU1XFM7swVkIcbVh0BqFfjySpwnJB0maKHlzwXOellRmK1sf0O6qu7v7S1/6Es7s2PKg7laNRqO/v7+rq8tO61Afezzx8bZjBX7GQGDmFipHeNSxhuhKpTI4NNRoNMD0k7brwXQBPb093/zWE1/60pcYj6uZRo8JZ/lfTelDnQQo8SvW+3C3VVUT7a8mi3Xer0b4OU1MyAlMfQ3btGq3v2KXr/v06CsHXczHXBawnbVdrcVrZQrRZ4W0dA6k94wXVh/XDGrwVyGlvUfRmDr6ale2h4iumWkAJgDQ23rxykyI+HwBuGSCm4J6wNLHta0ypT7Mr0YM6nkuycxO8gclWI0/dRe8qG3M5lKVTbflZmUcfsWPYOt3P/OIxk+dHs3l3lAwewHJhBTd3d1YajCbR91qwJVKxfWmpj72GOLrRzrShgBD4lWeQnLCOuLgjcUR+vDwEGPuvC9D2cYh6lekkD3d3bXePv2cGC3BrZcF0mqacU9mRbAtTOyk3k+mXYDkpjm/xj/orpzl60A520rGEkGmcUfZZJYl86uRu7FC6rQ5uVZORokmrdpWD3RPv8lxM3vWp5CSMelti65G4nDmV7Mkey7xgSB9gLzrnD/gyKy5h20lJIS1zH+ucsUz4BObOAtdNd1cjZAz7gRuRgI4ofdHItbNMYkfMKiEcx6G4dDQkA1SIul7bPEFk56K/Y2UXjWaVf2aHmYfpuB4op2g+osMD48wR0Rh76bqGOqOToRChtI6zO/IQS9R8caKZEd/SFTEMc9c3q8tHaa4jwVZvKtW0oy03JRmzci0quk+2+QIaWRxSh/c/d5Yplk7GlEcCNlmF8jfwjMInc9Sado6rToA/mat3vFqjonbDYiSUo6MjOAUDxWTPs742n7E9OQRmNEkOWEdYewkoxDC9/2RkZHkz471VCUKRwc9nHd0V0yEPzr6lf7VztqZRVIJAjOLa5dpOS+NwuV0YfGvWa0veUbaBJ3oYR2tyNqEwdkkLeGSp5df/t2aARKKFTM9t25s2gQNRa1KLEyfHU5bbqdj50xys4Ss0em/mZ2ET7Q5teeUJSLva2KSVZ9vUtLk6Od5y3Npcb68lW3d5bJYL5fOXxbfwWbH1cU9laHeCm1pxpbuaC8rg5HR0dFslSTiuOGrB9F0gBDpw6pb0i4tuBI9H0cMlQkLu8Tozo6MDKd+1R90NJoRLUKic5bNY5nYKvOBOQJYV0BUs4mMu+vE5VKzAogli6k2eQoBIGcFSB4obxNW7IPq5g50+0phL1ymzRALNihoRv55FizMOSnIO5lUL4+qVo47VfGhsnew+Ktrv04fxd0kVvsKzqV5M3JbtaNtk8tT1u2ildNDk1gMa1UFh3rS9cUzxu7o3RBmDJlx74pFMuahtP6M1MEeWzhjXmxulp7JhCUpE9YRxnYF+M/w0HCpVLJVWWwmcD3jFTu4gnTqubT0VIDWOiWPPx8SchNypTIWuetDR41HXS/fW+DjU+QW7rZn/45+CwDzMGPKfFd8Gz9mZu0v1ipg74IrfaWUJc8fGhxilIfy2MO1ldH696tHjcoRHm0cURJ9GhoaqlQq6LHJzLQWKkzZbAPgur8aCrtXK4DhFglgY2zN+ZBYLaWx5QnXJhmA9/ccdk7uibi9//5Vk2zjxdlLWu9bRnagprwlGPahemyq/4XUGAhXDoWolMs470MC+JjjCya5Nqign6MyEwoMVCAnrKOJta1ijNnw8HClUmk0Gim3TLfvyEb95qYBasKt6SbB9X9NfkjMm+ady/6JKPd6tl4PoC2JP90z3cUOmzQjG4icM2Wely05a9VPxeoUR4E3S0G10yu8a9J2YxlHoqspFpPcyFnfba0QYVd39+joaPZd29dmE4cQH3SYn3kITIdFTlhHFMn0vB2TJpixv7+/u7t7YWEhmwXCfvU8rxPdNQu9clJOQwXWZmKnFKdhaTUHSLuOeCvQlcWYsiV61r4TCiHdbHMK+/qEQvT19g4PDdGTRvgYkh+PUiWWI5TkhHVEAcdsiI7AUKvV+vr65lTJIyQlkDjn9Xp9dW2VA4/zdDi4nWC2Q9xLp7xT3JBQUOkjrAsYZA6d9VAlRWSPdLQAbqJzN1nHuGPh68SYkD09PeVy2R2/on6PQUdhGNZqtf6BAdc2QE/d8cRXKShZbKjjlAnrOBAPrTDZ5NDQ0MWLF+OfHXsgANTr9bGxse/9l9/XqDeUawBXhRxEanKruZGzyELYQY9W6y9CyvSazS657bYpv9nmx23FULw73JK3zf25WpFedmGR4XpHNBnqZU85Ze+1Z5d7vqkBZeo+Zoeb7nErlcpjjz32zDPPoF+FfZukyV0jwnBwcLDaXbVuj2SCPrb4Rg8yKVNVmCZlwjq6gPM37k3GxsbcxD0unPONzY1ve85z/vQDf3KrG0sQHciDDz745JNPVqvVMAxRgVGLdYLpMAyHh4dsVWDSgI8zvkq1q9IR6Gy3QktfICesI0k6kYONSiwSwFLKUqk8PTW5Va+XS6UW9d1tyXWl2aMP7U61vW0Vx1yvojaqLEVtbu/VuAUKVvZwqRWKfs1OBBS1ue0OVjvdW5N5ClwipODAhRDXrl3zVVxfwg1N5/wEIcTI6BhjqSh34jjiS4Fi1uQ6Vc+D9cQijhyxBiylsPXRJiYmrEHMYg1uJc+/OX9zbX293N/PkmFMe2pKKpVHO/a20500WT+3ee2NBm7S5vZejf1WsFo5XNGv2XCvJqu1cYi2U/fjJmeHSziL3qa11bWZ2VlbvhMPEenBxgFSSDExfkIpOzKbXY44Vvh6ckL1qdoJS+I/5IR1xHE7o7GxMRtf4a6AC7nvLS2vLt1cGBoYsPHBu9CAc9Waop92yn5rwPvB/mnAt5iO0IDRGuw681sfl13v0z0L/LywuLC8tOR7XjrKCNUclSxr4sS4CfnsmFtM7Ac+l7qKaiRyBUhdAk3ggI6csI4J4+Pj9kZng5E87m1sbMzOzdx+/vbcrrZFgdG8hz3kGvB+GHX3TwO+xSboTtSA23Vl3GRqnPPZ2dn19XXrBW39n4ED5hb0uYd5KBmw3BTgxPEhkr9S+TwrPyyhRmXRUJEDZcI64riRviMjI+i0meqhjKsIbNU3Jyen9IZslz2jXbMTHU86q/Gd1dpbg3tN2ntl3J3MzMzUG3VwqvS7LvFSynK5a3R0NLshcQzhOkm+NNW0JVarYdlYT+KI4aYfGhoa6unpCcOwKAdyKMKZ6Rk3KsPtU255RkBp/sp0MZu9bOv+cstyd+07za/VXq4k4WDGsjPT041GwDOpQJnxwOrp6cEK3Kls5MQxhNunAzONO6+jpHfxaONK0Fqt1t/fHwRBrgkaOAghZ2ZncKFkiRoGB6FjWVcyKKgPtKtt3V92tMtDTfNrtZcrScRIMySdmp4WaiCbTbEJAEEQ9PX11Wq1thjViU6Hx0nDGXo+c6lzQpP96rgAAH19fQMDA6gBp/Ra9RAAE2JqaipWdu0jcTDPBmnArUMa8P6CnaepFCZnZmayeTzsgDUMw4GBgb6+vls2PU8cZjhIrJNusmHhEyOFjP4nJZUjPNLY3qGvr29wcDAIApapFGS9S+bm5tzZMxvdeCANN3/DPQjgzLb4Chj7j5RHQyI1v1Z7uZIE04NUpQFjUN/MzEyqDhKuhh8ajcbAwEBvb69raqIO9tji67oLYLtTzC/I95gljjjkuL2DEMLzvJGRkTAMs9klcWXP8+bn50OVzFZrwDjFdTAy2BpLveSSvW2r/sWThoxlumNpfq32ciUJ5mrADEAIOTd30wYU5KWxZCMjuvw21SIkuHTqzTH9GDFBTlhHnXiiV4VGSCnHx8dFMnd8KmBjfn5+bW0Nl4u2qEu7N3/u4xwwRgTIaCjqNDBpSEyV8d8Re9yDNDuRre6E5oB3TtPJCPcO6ltgPK3W1tfmb970PA9709QmqhgDnDhxAuOR3J+I4wkHIR0/PcmA6zRJB90yYl+Jsx8Yo/Lo6Cj3eOj0C64G7Pv+wtLiysqKrqDEZK4J2u2RtDm3Cbvv/G/1HHA2qmfX7hG720M8JW+90DNRLgVdOc0B75xidzyZLJUNjjOElHJldXVhIRLAjLGkp4S+QSoR9GhC5yFV5xjjMw5MCm7tbhIwGYetUUiPyFFFmlE63t3x8XGfe2EQ4FRWKnjR9/3l5eXlpaUTExNhGPrMEyZZaV6HLV2b2z41P/N5pzI4b1vjCWHngw+PUNLXM++auukVczfNfM4uKfqVSJOyOuCbEoah53nLy0srq8u+7ysjkVDjJScXB4DneePjY+4IjDjO+CCknvE1mbBURDCqQZQJ68gCphowEwIzgJ84cYKb/HlxEkrOcYaYe97m+gbjwDkvd2BlfoLYVzDDZSMI1ldWS+WyFDaIICGqOecjIyOk1RCIU45QRwILEMoJi5F55CgTO4mYezw2NuaXfBS3qfS8QgjOIuH7h3/4h3feeacUwuc+cI9pW4nUCjOW0JKCgSmKL7WPihraocU7/o9pZY4LCLnEbPWmgIypEAKJavp6S6adwJzzcQzieg29AdMWHX2a7hmbS2Fq+KuVmT00+peBTlGjV7Mf9MrKYITr6H9VA0R0PqAbEJ+R2k4yoRcrFckcSLdV6prc+n3Eo+sTY5yBwKAFY6IQKipbX4KE6hrtDlfWtVZwuiG1hj5JgPhA+uJJXQ8tcb7aMoanIJ2yltxYzvSeAdx8amAyDODu7XHNtQIZ3xKzst6b5AAi3pWwz4VqldBPBl5cKZwablFrOT4uzqVzHiTmqRtgDxXfP8bVnvWzIqWx9QAzT515/LRsFSKUgnkePHXxSamEsdCHteECOgiYc45psKhrJUw5Qgk69iJ6BM27w4EMUUcYtzw+zkeNjIxUu7rW1zay+fOU8wj3ff8//d5/FNItmmT6XNAdtltpOHm87DK7iOt+9VY/cKkjZhrQ/Pf9a++O9rxNq9zbYSVMwWESC2797Sg48v7dhaJtt9mnfc5zVvI83tvbF/s/65I20rxHoedpDditA0HRwMcWXw1IlWcHoF1R1WbQn4gjjrQqIUCtVqtWq6srqzmZsECP5wcGB4ymobQBIYErXU5Ird4J0CqoGsyhYqP1Pv1gWQ3FPl+gBDDHtaTWnq0eow6GbvlKp2M4VNSlM3FbMFqaWm6OqzUj012i1g1x1wlG8Jta2FJicBUeQHBm9UVXiWcm9tN+hbzgabum0ZxkKrLL7EAZoNQfvETquLEyqurWiUgR1Kthjx5vDzpqWVqfLqPBAi4za4DkIm6zNBohGK09unnc7Jurr0aT1HYEY1AQCXFhNFFXMZXOgfCr1X+NfmzOV3I76+4aS4Talmd2lasBO79Kxzzilvowy6W9zuBYOJhRa42VwuivDO8dxLeaJV0U0Q5hd8UwpkD5KTKWFKxBEAwoWGvlK4gjjw/GiibM22DKETIqR3i00UZmI2V7e3tr/f1TU1Opmqm6gpsKBdYbcl3rlNm5YBtEGs8OJ2IwuPIzwOJrsTHSWinBY9E6ultG467tET13Z9zuOrEI3P4SErJdah0ElB0S+1MwcgrbbRUakCw6U65ikADX19Yg6bk2XrdntwVfk9bJ2MLvOM3aU3aq4AG3iQxV3+/rZbrrZ8m023boE6tq2E6nrB5urGyheGhurKXcvtNSSh9Q9KELCPOZZ8rTRoMhX00K6MOh3RvL+SmzcPJJYr4eFdmv8RxWdJVjQ68ZAdnLwsEMRrTl3WeoNeK1x2YYdz/1qzs7xo2rMe6Ky/gJ1KLW+pUB2JM31xmsrd/cLA5mwJR+TQo0VFzO9a3XmQXtbcIwe3MPIQzF/8/e38Q492T3fXhV3cvuJtl8f+t+fvOXpX80EkbRQoEtaRMlkWxBjmDI2TiJBsnMRICMJAshApKshHjrrBwtYjhKJCOrgUdjbWRk4UDWjAFBDiAI2tgTBJYRQT/9nofvZLOb/cJ7qwLWqapbvJdks1+eh93N70c/PcPmvWTfyy7WqXPqnO8plyskg7V2mZt9f/C2CWmFbvwK6bbPdCUwkrDeNL5V4JwfHR21m63/+199b8VL8+pbEuEem4i7qT9r9vlsuaryL8Mc8t7fvtMTqyT9pF4nq+m9p1eA515hM2eSXVtpZ/mdfmU6Idk9SDWb8tYfPPPJ6M9BkiO1cr6/Kexd9Eq8yr+9zJVwd+Ope1fJH1Z5H10StTalq+v+6Cr940qGcOoc6V+PMuuM5CXevSgTHVAqWdCwVLcunvow1fqr4vaekyG94a1Wnln798qclnqJ9yr6AyqXS9FptY+Pj01StAvHYHI9VLwkLPM90VOUW9LB7r5d/PAXrbQ6nQ6JYfl9y129RCoz6163YNPza7stpXorpdrIPJHUThtf9V/XfizZ2OAmN8UvDN10j9lfsUsv5Oyv3rLc2fJZpa5w+2/c3gY/dUnZxdamJsT+4/R9sfWfcKoQLvs5bBpOW25h7TVs+kB2eSZLVv3KPZBxXK/XvSRHnioUBodGKLkStM43/oE01jedZgreMjSDdzqdTVkh22exFOk041UDsIudSE3uax/seF+uCpPqmx/6Dms/gdTN7tj9PvW5uSqvTeuPtaZo04WtddF2+di3v5V/YXS1/gm+IXTvcO8iJnVhm/4c2eWCfy+pj90L6aeN35bBnPqlWz6iTdGg7DX7IhupZV8sZbvTVkrFUgZCmFgMd/kSsMEHR8ht+wWqBk52TTIBGfAm8f2n8/PzlEVxc+7uk/jaX3GvG0ekrEVK68B/twf96k2e671R9Oycu8Wy7u4Br/XnstHONalwWy91u1O7SzA/+1ZrPeC15tDPGMjGA9ZeQyp46/9x1xr1VAOD7Spgm25h7Ue96U22fDJbxmHWrktpSqnOz9/RvrgyOdKmIgvW9zAJudneEna/gp4wqYJ0EgbH28Yt2DudTjb094x//V08YOdj+S/xHa8dryoVq6SIemoW3jFa/iDbv/YeswYpa8O239T2j277hT1o8bTWKcyuLVLWce0yJWUCUzeeCnikAumbrnlTmCRlU7NDKOUBpxx6d+YWQ/7oL4J/ee22lsEyDZRcAiE4UGifT7gO/JTxqNsQSqcY/sREGPCS8WerVqtFfVr2eD3ZCXGTm5hik0/DdZ84OkQPPv2CMmt9t1igROL/JZHN58r6r743nI1ju3dIham3xBXWCibfm993L9lKd/dWKZ9+x8jB/b9RB5lFINrtFkN8EVhClkxtRmdGgySsw4L+0JVKJZ/Px3FMUcS1XsLHvgbf1rr90bW+19qXb4ob+6HRXQLaTwm5Z0m5gJsc4pf8dcuaQD8q6/dl2vETzprnTWGDTYGKx8UD1t7O2nj1c0WAlC6RkkoeH580ag1GivsMGdCAkQrc0gGWksrtrALdC54LwDPizzL1er1YLDoPeG32iv/aXXY6d++7l8ro8Sf37alG/stT52fTf3TOP6eyF9PRz5bdqIwbvfbi761Lyd7Rpo3kLXY3uSTvv9QzpmIpuZ6VVz872ZQrt0HgzvH9+JRPvzZKkU2qcu+z8iezf4vsM/6htb9r+72krnNLrGWL958dEqnNjuVqMoqLhUKlVjFvzuAEA1LCImdXCJqB7JdfokjtEPAnoGq1enp6OpvNcrmcO2FLMsuO09zuccKsF+XbXcpkpgebNm5Tl5r1qrXptbZMmbPZenu28eIfaoMf+ipzWsaSpp7Zow32n8zuE691au/9GLO5UWs/q7W2fJfrXHtOyhffciVbHqfCM9k9/uWgjeVp8bRcqZiW08pXVgMHSpgMGjOMhE7CivVXKqkE3fd1go+FH0isVqu1Wu3zzz9fOwE9e07WFlLOUDatZnsml3u8o7sJHscWS/zod9hjQU7WcD7ihdmUaaVV5KI4qlTK9WrdHrGz7se8I/DCCWkvgnq7Smab8ZuOJbQLiBq1tw9NGfl8vl6vR1HkbwC7DKZs1O5Bv+IRk5r/OJvpussLnetsvHxTc2mkkBO1Zyc3vepDb7h4277ovjtaW8Ozy+dgQpTaSZJapNE8r9IK1VapmNs3ps4W6Yl91bTYV68r9dlyF/emi9ObU8VN1ptciWdYaaq1pUT2gXI97bM9rFLXsClV+0FZ7qn4durlfmrCulwzknIz8p/ZMHsURY1mq1ytuA0Qjn2+g8ckYXGriqaMSqvYnvAC3h5kaxuNRjZFiE6gH6MoorziTe5Cts0LPZ+tb9nRkKfEH3YZk94lUWTH2Rxm2woYCWU6rN96ZU9ubTWO+2ntOVnWeoQ7Xr+LjPNUzFl6fxfOvRC07QWZVqjMBnXdJ5M2Emv/Fqm/VPavoI8K/auZ7a+YrJmka3+Q7MEza3/SG/beL1oJc7gHW1IBskkA5gI8qZPU+e6ErA1O3a/Dney/yl7SikxnavDHRgZr+UFtCeGAgyKUXAXMTkDLISOZacyAdoQHhJssSAwrmzbsZhkhxGKx8BNfHanJMWtfn5hfmn3JFjdrTQ6tdTvIVHHbnImMhG10a7sdpJcXyncrnWWx7dZ56pK2eF273ilzvY/N0pg7u0Xn2N/G1vR34isdJpR0qxDrAZsGwU5te1XBeP0WgPenVPYdmLeGYV4fJv9m7A34bbH4ivu+bpsjmxLFti/p/E/Pd4gz+xfMGngXBmDeZ7UmmOz/cf1f5L4F+pvCc7lcEGz80Kjjp47HuNSEXcYCeMuEzOYMGiUsLpiZqrA5cSg4/4Nz7tQoU9MNHV0sFkdHR7/5m7/ZarVubm58VyM72WXnPmcAXPs/+85M93BT2WzYLX529p1TM7ibT90v9V0Vk4dq3ExuW/at2g7tidqAkHSdDDgPlIr0v4kp2tG+bokEbIoHJAZ2Y+4sXZ6w/6bSJ8kAC88e0y2oVCmRb1Syueir1lHaGYMCZrH9lEx7JZfC6frku1uQzKyHNkwxbukgvX7GnPowra1Gy7rF2aPexZOnLhiL9aVKWlimUva2Z/n576yUOjo6mkwufvVX/5vZ7CKXy6WKsrQLzml16+S8ML+CpB2h7dBKa8KkESmdhGjJ28afcd69e5eql02lHN/e3v7Yj/3Y93//9+/1kgF4WXzxxRfX1/PUejQV2f7ss3f2dMypgLkkLOEiR0YSS1FhH9oRHhB2jvgsDEOXZpKaRIQQl5eX79+//77v+z4/PwuAg8J9L6j9fhAE799/MZ/PT05O3DdlNYKtgoB/9u7MvpziKJhUD53QNv127cwllzqmxJCEdSj4bm673XZiWNmwKuf87u5uOBy6+WXHfbjVGCbtvTEXWHXZxzY6vfHd1uYVZ072c3bcla9/5+wvWptMu+P597z73ljJxnpc8DP7yTx0I3/tO2y+2vvfJxtzXvujn77uXup2Zt3xTSM29bb02KUZBkFAK9HBYHR7e0c6Nn4o22wPSHlyfNLpnNn3YX6S4I4fIHh7hM7ttTthyuQ7CCRhHQr+jFar1arVar/f1xklgZtKaFqheWc4HJpXrQars+kqaxOj7MYq8+c3MpDK5PSmr9DfmFy7Gbz6q1fsLr2h+xVr07W2/7i685dsr3LukoPWZO6szSR/7lLXlbdfNS3Z37X+Bh/6a/iqwsm959x7dF0eFuPc5S65P6hZQ6yt7MomXq21ypz7PzJ7MsuevDYFOpWk5npcMsa63V4cR6n7pcdCn5kvFBqNhucfK5suBw6X0EuytGOCu9pIcBC4Ug3GWKFQqFarX3zxxdHRUeocN6H0ej3zjHUq19bb+M/wdaelcgy224JNGQkbSn3Y2h/tL1pzVff+aEl7k35CTfYiN1zec3231rx95t75prN3v4xdXrj9nC1Hs/dAT/i341635c0f++Oau9iF1PqMvhRuYKzkGDJ2Ey3q9VqpdOqtP6AGDZZ+7nI8mJKM5VARFCPUglh808QK3hguCzqfz9dqNSo0ckdTbkG323XKufdW8b45lP039h6DgyMVBh8MevZp7rvgpnr+blGv1Qr5gk3JdurZGDwHjSB9IEkhaGnl6ZWUCu0IDwgXfzs+Om63284hTgVgyUj3ul2eit8dENz+G3iPwYHivjgXFxd+ydNKTJvzOIparfbJScF6xgKDB9AeMLfyQNwqcKAd4cGR7HgJ1rQ7VWs36oQQvV73ETk4b4XsvHlonwBIcBZ3MBgEQZCKFbk16yJatFqtIBfar4wbMxg8B42QXKdAKyv/pjOipRW73fflgU+Eq6lgirXbbf/JVD5LEIbj8eTmau7ST1K1j3u9j0+ACzsrhKABLUnjOO73+2SA6UvhigjMd0exVqvtYooYPIAQ3IjO2uRnuyBTGBkHhqlG47zdarsmrymzyjkPg2A6vZhMp350+pDSBZwHzBFFPGT8wT+bzabTqSudTyls01ep1Wpx5nImMHjAEkGNYFSSHSq4TsLSiVniYGZVYNbpjLFao56dR5wlDoLg8upyNB6nVO8PxgMGgDGvNE4pNZ1OLy8vSYSSrVbNmRI+xprNJuOYS8EK2gNmQtnu5Mp0W5EKSViHh6/F4aeT+IuwMAyv59eT8Ti7TwzA4eAXQ4/H46urqzAMU+V2rlw4lzvyUyv2fe3gpSB0vxWp25No0XbdgEUsfV8B3/fQoJYEjUbj+Pg46aG72nEhCIL59XxsDbCv5o8BAw4HXz9kPB5fXl464ZpUcbOU8iiXqzfqK/X0sMLANBTTYlgm7dnrG7PvawOfFDNlKF6v109OTkjk1p9KHHEU9/t99yqs6MEB4i9P+/3+zc2NL5vvr0SljI+Pj+u1+kpUCfMrsG3UhK0fN/ZYcSRhHRxmRc9VpVIplUpxHGdPoPROKeNut5sKtWEPGBwgVoeyS/Lp/saw+3YspCwUC9VajVF2qxFc3felgxcAxZmZEcjnDEpYBwtNCFKpghatdZrya4wr5ySGlXJ/YX3BIZDqzaCUcuKs7nm/BknGcblcqVQqthcwQtDAQElYellmkrC0cD2UsA4PM3lIJYRotVpODMuPM9N0EwTB++4Hv+jCD1Dv7w4A+BT45e+63pcPBgOKP6/tTRLHcaNRz+f1to7e5mEIQQONsNobth2hHl06loIkrAPDE5dvt9txHGe7+tCJQRD0ej06wV/sY6EG3jpunPtxZjUajZwKh98oiU6OpWy3O8p1ecIeMLAIyRnnkruiTyWZghLWQeLlM5+ducalK30AaUIJw7Db7V5eXqY1b2GDwZuGBBNs11abkxjLfr9P5tb/ClAMaXmyjDudjl7DSp1igz1gYBA29mzDjFahBUlYh8z5+bmvsJEKhORyufF4PLViWHxDkzgA3hjUNM42EzRhodnsstfr+0LQlKuYfB0k67TPOanuu68LvivA7AFzoZKAM5SwDhG/4ogxtlywe2HnlCiHEOLq8mo0GLqKRuTrgUPAy4SQbnk6HA7H41EYhnTOaioWuczs7LyjX7/SPRoAQWNhOZ6YJMcXSlgHiF8+QWJYTosjmwsthLi5vaHMT+0CQ4oSHByuPH40Gl5dXVESljvqgkexkmEozs71ng5HXBGsIJYrNPJvjCgHhxLWAZJqL1ir1U5PT10itJ/kbBqMRxGVAqdcZwwY8KZxYR5q6EsqHANS4fCXqtYYKxnL/EmhpRucmB0bWGFgMUbXdcfiSkIJ6zDxnd1qtVqpVKIo8hOs/Fi0lLJrPWBuo9DosQbeNr4qHLOtjbrdD4vFwveA3VJVCC5lXCqVa7UqTbSuOgmxIuCUsGhdZ5+DEtbh4ZKcae4olUq1Wi2Kouxp7t++0+JgaqVHHxJMwBtlVRfdzJC9XtcJp6cCRYzxKIqr1Wq5XNLPrNeqBAeLMA2yKKsPSlgHjPtbSylPT0/r9fpisXBHs4pX/cHAzjXcCOzBAwZvmrV1d4PBINu10/0Yx3GtViuXy75uKzxgQAiulJ0wFZM0kUIJ6xBxe1dSyiAIGo1GttkRQT/2hwNKyOLcs7swwOAA8GPRrjFJKhmCvh1xHJfL5aS/5/6uGbxABOmzmPI0DiWsw8XXkiQxLGqIlEqQdgv80Wh0e3vryseXL0QIGhwGvo9LOpR+pYDfnkRK2Wl3hBDGpdH7wIwhyQYwUsJSaSUsDiWsQ8S3tZzzVqtFATR/e9hFz4IgmEwml5eXyraxZNgDBgcDfU2EELe3t+PxOAxD17zELw2INf+/zz6jymFOCa42C2vP9wBeAIJT0NlXwtIgCesA8aPNzWbTrev9o84Az2az+dWVtcoce8DgcHBL1dlsNh6PfRms1U7AS1qdjokPKSWQWAM8SL3FU03TSlja/nIUdx4UqaYL7Xab1vUk8eNnl1A/htlsdjG98CNvfiI9AG8ezvnFxcV4PPZrkFLNCoUQzXZTB4m4EkYIGoJYgBC6OQdXTEoVU0kJlLAOk9Rf+fz8/Pj4OI5jf2fLnSmEuL6+nkwmZvM4oxcNwCFAGzFBEPj9B/1vyvHJcavdZnZVmng1+7pi8JIQpjJcccED/QyUsA4a3wPO5/O0B5w6gQzw3d0diWGRxi1WaeBA8HMSR6PR9fU1ecB+oRH9K6Us5AvtZsuvU3IK6gCIJNZsgohQwjpQ/MoKpVStViuVSn4Slr9JTDOOMcBm5GDAgIPArxfo9XrU+jebqEhJWIVCgSr6jBS00zjC1wXodoQ2dYYzt5MHJawDJGVlC4WCM8Cp6kZnhr/44gvt/up2DBzyAuAg8JU0XBDIV+dwnm4cx6TqSnV6uo0h10k28IAB0wbYtmBwIuNQwjpMUlb26OioXq+7JAB/W4tmGSFEr9cz0TZuphUMFXAg0Nfkw4cPm6wvGeBavXZSyHPbMoxJVAqABMGlaepBLi89gBLWQbGqcMtcmhVpcayVwSIDPOj1dGkjT7kFABwI3W7X9UHydnmZsnXAzXozFIFk0nxvBGrlQQIlYUnGqIqcVnBQwjosUtOHM8ZnZ2ebGv0qpQIR9IcD5TaJOUwvODiGw2GqBkl/HZipLlGq1W6t1OlBrAZ4kBLW0olRWv6KazkFiSSsQ8I3ui57k3N+dnaW0tXzCy1EKMbjyeXlpXYAuNcqFYA3jtN5HgwGVIPkfY90lNmuR886HcZ1cSelS2xY0YLDhJSwdMjEKJSapRmSsA4WF2RutVopeSx/oysQwWw2GwxHK/3HMa2AA4C+ERcXF9PplDxgrxch9f01waR2u82X1ljLG3GFxBrgQ+0IZWJ9mc6oYSaEgrFyCGR2sMyfu9lspkRu/SEhhLiaX/X7ffMq760AeNuQxR0MBrPZjDxgb5EacB7oSKLSXyJdBKx7dnIGDxisIHTNr1CUiaXbEUIJ69DIhqDp+Uajkc/nqSdSegxoOejbm1tqhgrAQUHflMFgMJ/PnQH2wkBK99VWRyfH9Uadms0xxpH6DFIIU/6rbDtCDiWsg8P3gH0nuFKpFAqFVCmw3xDpLloMR8N9Xz4AnxS3Hh0MBnd3d9SJIauEJaUs5gu1atWJP6diSPu+D7B/hOKrRWwKSlgHjd/yyBfDoqPOQitb5jgajRAjAQeFM6JkgG0Zkt7k5fQNWj4TR4vS6WmtVlNKCdIcVGhHCFYQVPbL/RUZlLAOGN8bLmviOHZHk2QT/aOUstvtKnRiAAcGrTj7/b7TaqX+RnYlKmh5Wi6XatUa9SLkZH9Njg1T+MoAakdI86nUTi9ngW4KrIUF0Y7wUKEAWqFQqFQqKQ/YT4QmNcqUUgcAbx4yui4Bwm4AC/p+6NSZQEpZqVbK1Ypbs0q7VqXdYHxngKCtCckYJWFJ4+RACeugcS2Pms2mnwWdUqwUXIx0GRK17gfgQKCwc7/fd491GNGk03AuGVNxHDebrSAIKI2R6e8VhLCADylhMUrC0oocCkpYwHm6zWaTWgK7Q/4mcZALRqPR7fWN36ANgDcPRZhJhcPuyzh9Z6N3pZRsNptJcYGNNcIAAwcpYUnj82ilFsZJCQt64YeLy+TsdDrZo2RopZRhGA4Gg5ubGz8zBYA3jFtr3tzcDIfDMAy9g+ZLoL8gSsrl+jXdTAy9GICHVsLSj3iixcGghHXgODf33bt3Tms+q4oVBuF0Ormaz325SgDeNvQVuL6+Ho1G5AGv7ZnNGG+3z/zERpPuSm+y33sALwOthMWE4EKZQSGsEhanZHoEog8Tp0bpa82nJpogCGaz2XQyQaIAOBxoqM/n8+l0mioC9r8FnPPz87OUgvr+rhq8RISpJ2GS/oESFvANbafTOT4+Xquft5x0hLi5uaFUFEwu4KCYTqdXV1dBEGQP0fclDEPawcFXA2xC6DJgKjkSUMICfrEvY6xWq5EapVvgr2ZB81jKbq/nXrjXCwfgU0DjfDgcXl9fu2yJ1NpUSpnP5107k71eL3i5CEk7Flp6A0pYwM/AIjXK09PTKIr8fSx7ojmn2++6FyJzABwIpMJha5CUv0JVSkVRdKplsJCfCLYg7JSqkqwAKGEdNonepFL5fL5Wq0VR5B9NHupB0+v2vBd/8ssF4FPhJyF2u13/SXpMXxwheBzHlUq1UChklq0AJIhEgt8+QyVtAu0IDxVf6KpQKLRarcVi4edeJQ/0416vl6RAY54Bbxc/258McCr531nou7uoWq3m83l/QweAFELX/FolLEbK0FDCOmj8oqOjo6N6vb5YLLyj7oFiWv9qPBp67dgwVMCbxV99kgFOhZftt4BLGdXrtePjY3cCZlGQJWQ0YkwzQjP9ci5otGDhdshIKYUQ7XZ7dY2faIQrxcJADIeDZI2vOKLQ4A3jfNyMEHQC5yyKokajkSrh28f1ghcNKWFx8nz1Fp5LwkI88dChIeEyOTNRaK7VKHOj0fhydmnbFMIJBm8ZzrkQQkpJKhwUI3ShaddtQQtBN5VSrpkYPGCQRSth6cihWccpI4KF4QJoZqnX636mid+0XykZhOHF7HIyHpvZB11ewNvFWdmrq6vpdBqGoTO9LpnGndxqNaEQB7YjyAG2cysz2VfUpQFKWIeKX1nRaDRope+2sqz1ZUqxIAjmV1fTiwsz15hOawC8Qdz3Yjwez2YzksHyjyYqHCJoNVv+JjFmUZBFUGNKxUgTmjOaV5WUDElYh4uvLdBoNHwxLK/iQjDGwyC8vrmZaDXK1Z7kALxZJpMJGWA/JkSHqJd2LpdrNBrufPjBYC1LR1erT8okCYs6XDIoYR06Tg76+PjYb8vv4s+MqSAQt7e3o+GQXqEYKi7AW4bs6Hg8vr6+DsPQ3/r1kySOjo7IAHOPfV87eHEIaiBtHBepGCVhSShhHTS+7lWtVisUCv4S3gbiFCWkRFHU12qUeiRRefAa7WgA3gA0pKkL51ohaDrn5OSk0Wg6w4wvAliLLuS0FUdUCMyUVsLCiDlgfC2O09NTXw7aO4fptHku4/hDr2t7knPse4E3ib8F8+HDhyiKqMrIT0ukM+M4LhaL1Wp135cMXjraALskLMG9JCxMo4cOrdxLpVK1WvUNcEp7j3HeG/T9IBs8YPD28MNC79+/T4m2uhA05zyO43KlXKlUMH+C7QiuRTgUY5JmTK4YkrCAnTWklEEQnJ2dSSnXdgVmjAVB8OH9B+jtgbeKv5oko9vr9Zz7654n/XzO+UJG1Ur16PiIlq34aoBNCKOEpZTuDGxSsZCEdeCkGry0Wi0/CYtw6SdhGHa7XTohVRCJ8QPeAFmxyW63SxvAq634jXqRjGW1VjNxIIVmDGAjQuq0VVt+ZLb2JNoRAg3NL6Tpk7WmruJiMBhcX1/7KScIQYM3gz+eOedRFPV6vTAMUzVITstXxtLUICH3GWxF7wEzo4SlKzuV0fTF1HnA+LpXVArsnvd1f2iQCCGm0+nl5aUfjqPzMf2At4HfeP/y8nI0GoVh6Hbo/AQIWoF22u1EWxBfBLABoWPOSiURFK2NRQcQSDxgfEeW5KDX9n6hPeD5fD4ej/0nMWbAm8G3vpzzC40QIhWCVkYLjismO2dnxhhzqHCAjQhKgzaZWJJJpiSTiumUG4lA4oGSMp+dTsfpzvtxZtt+XFzf3PRMKTDGDHhrpNqQDIfD+XxOKhz+/q6wEedQiHanzay6LxajYBPCZe5pNQ65dHwVF+QZY9gcKqlIcr1eJzEs/xxnhoUQsd4VY2igBd4iqcBPr9e7u7vzUw7t18HUIB0dHbVabfdaZEGDTQjFdQ9XZZry6xwsJc3OBVOYUA+S1CZurVY7PT2lmgp/3nEPYimHuj2q4npAYcYBbws/otPtdqMo8ouAvbNkHMeFwmmz2bKvQkAIbESQ7hXnJL+hJ01u/2Pwgg8Xf9+rXC5Xq1WadFIhaDo1lnG/11/tlYQZB7wdvFoj1u/3U8Jw1gwr3Yo/LpfLlUrZfokgBA02QsW+WsPXiBoFnDlBLCRhHS5+InSpVKrVahSCdsbVc3MVU6rf7+uMk+VgggcM3hjKo9frpQywn/8vZVyr1crlkv46ALCNULCl7dXhZyWlzsdiFIRmnAna4cN8epg4T7dQKJTLZZLacIHoJMKmuOBiMBhgqQ/eNkIIznm/P1h9WtI2MRXlRVFUKpUKxVNOng2+D2AzQmfJS62BlbTl1+MMSliHjgsmB0HQbDad1pUfatZ7FVKxCPoAAEm1SURBVMtV2mA4iOMYyzXwJnH6bkqp0WhM41wfWRpZKc1eL+dCSqVbaBsdSmTRgC0ISXsUips9DCWZghIWcBI/jFSgz87OtBi034ZBWFsswzAcjsdXV1cpJSwA3gYu4X+xWIzHYyoC1phdXnpMrZDOz864FtKnrp3MpbQCsIrQOatc2TR6l36F/OcDx+5pGaPbajX91Crr+xr50iAISAzL5EVDgg+8LdzK8urqajQaBUHgfRfIGHO9ixcrJTuttjKHmHJTKoLRIINLwnLJVkYJiyMJ67AxCc/WlLYarTAIZEyLenIIYrK0SnEhxHQymc1mxkijnzR4W7gcq9nscjqd+gaY3GN3Zi6Xa2odStPolcEDBhsROumKoiTUj19JRiJYaEd46BhtAb1y75yf5XI5rydSrINvjIZHGAbz+eV0OmGuUz8Ab5GLi+lsNqVODATpb9gwNQuCXK1eMxnRir4pMMBgPUIv0fTur6n85VwxwZGEBRj5skSn0yH1eVqZ2U78xtpyHtzd3Q36A5OUghwC8LZwaQ3j8fj6+jqjiE6Tp4rkIgyDdrudHFEBQtBgE0IxXa2mKIyiuAkwckyggKC5ptlsHp+cxNLqpVl5H9oaE4LHsRr0+/ZFCn4weDP4kpP9fv/u7s5lQdu94Vj/KOIoyuVy7XYn2drTezT6bTCjgjQ6fcA1zKIgNGeKS2WlKfd8gWB/pMSwCsVCHEX2kMk6oeR58pW7Wg7a5vIB8KYgi9vtdlNyNPR14JwLzmUkC/l8o9FI+vMrNHgFG6FmSH4LOUE+sFjVGd/3dYL94Gzw6elp5bQkpXQFGEolxZEUjtb9GJQyeVgYM+DV41tNmgbfv3+fEktXtgkh5yyKF5VyuVQ6dbvCnKWbZAPgELYToUvCMqkEUgeldYE51m4HDU0xuVyu3mhoAxzokJqgmccm0rMwyH14/4HW/ZwjFQu8BVJWU2nJ1dQzevnJlGRcBDJWZ52zfD6fdDBkK60MAfAhJSwvCQtKWMCDZg0pZRAEZ2dnrgmME01z6SlBEAxHA3KG9Q4wRg549aSsJud8Mpm4TsBuhuSmEwOPoqhz1vF0sgDYhjDpVspuZjBpUls5NoAPHb+vkVLqs88+8xUJbBkSbVuwMMyNRiMZR8I0RsXgAW8HksGSUg4GAzLAviar4oK+KlKpdrvthFpXV6tYkoI0ug54OZKkmWptMBorOEC4beDz8/NURqgvUEBaHBfTqTmK8QNeP/5eL+d8Pp+Px+MgCJJ+2EaphtutF9bWKhxqFYSgwVqEdoEF50GShKXHkTBpfVi7HS4pK0vVjStdCE3+p+JcBYG4vJz3+yM756AlMHhTKKWm0+lkMiEDbEUTSEtfcWtoW62Wv06FBwy2IJjJvbJ1R1r7VykpGZSwAJUamaqier2ey+VWV/SU8qzIA57fXPeHfRu1Rk8k8OpxUx+1NhrrjiNeJwaWCFHqTp2BCJrNpnkeW3jgPoTiZorlzEiKc2p7yZCEdeiYLQlu1Cgr1erJyYlTo6RWp8raYCHY7e3NaDRKbR4D8AagwTwcDm9vb6krsH+U3Fwp5fHJca1mdCjRChjci1Dc7PtazXCdhCUhJQiYH2RjjDWqtUKh4IIinAu97WX6vgjBF4uFM8D7vnAAngG3jqSvwGg0urm5cR6wmyOpDj6WspAvlEpl3wDjmwC2IMiFYU4ojZKw0M0GOOtrBX1q1VqpVHKVSCbxRAldE7z8L45lr9f387b2fQcAPIlULKfb7cZx7KqMvKNSKabiOF9YGmDXqx/6k2A7glSwBDkynHlJWGhHCBgNBhoJlXq1UqnEceyqL5ZOsBK6M7/iPFBKffjwwaWfYA0HXju+GqtSqtfr0RaMHds8aRrG+SJaNOrVVrtlxj88YHAfgutcAZ1/pdd0AklYIIHbXGel1MnJSaPR8PaASXIyqbUQYTAY9CkhC+s28Jagfd9+r8+Zp7+RlAksLW4cx/VGo1QuKa8QHlMn2IKwikVWCUshCQskuCRo3YhQ1et1Sgf1NsZcoYUKc2Gv14sWC0+sA4DXjb/cHPT7XCShndQIl1LWajVakjK0QAI7IKTe9XVFbFxJxpmEEhbQ0EBw67Bms+lUfjKn8jBYGuDr6xtsAIM3gxvwd3d3vUE/CIK1K0vKgq7XGqQqyF3waB/XDF4LwrWUsxFFcwDuC/CnD5qDOp3OpriIYioIgslkenNzgw1g8Jag8Tyfz/v9fi6X85/3p0qlWNNuAGtBBTgx4B6oHaHfekFwikJzASUsQDiP9t27d/6TvpVVammApxfT2eVsVTALgNcNDearq6vxeExC0O75VXV0dn52Zl9k1FjxHQBboHaERgmLMy6ZkvqxVLGSUDEFzK+D/NK7zwJdg5FShKbTQhHMr66Gw6FfQAnA22Ayubi4mHkhaOXVARtpmvNzZ4ChpADuRzDbg8Fm9zGuuGA6CQsD6DBx6ZuK6pASU9rpdE5OTigPK+sECyHu7u78hqmwweBVY7qk62E8Ho+ur6+tCofuBmZL5LW7q46OjmiPBhEgsCPCiImTBrTOxmJMSeUmYUygh4dL3zQtXrita2S1ev30tHS3WNCJ/ixDk1QURYPBwE9d2eN9APBEuHJ9r9lgMFjoka9Mu1bFbXCIKx7HcSFfaGkhaAQOwY4IlkrC8mZexhjkTA+RVQ+Yea0XioVCtVqJ9DTkd3rxEwXIAEOLA7wBjAesp0Gt8qY8GTipuKTTuOCxlIVCsXRazn41ANiEYGnFK517ZQSxuF91Dg4FZ4Bl0myBlvkn+XytWiUtjlQKqBsnSwNs3V8YYPCqoQAhDeNer7dyhHFb67ucNaMorlSqxWIxmx4BwCYE6bWoZK1HqVdWCYtBCevwcCFoYQMhejaRXOXzJ+12J4oikh1ILc5omPS6PZqXjFYlAK8WGwJczomDQW91MlwObvpJKBktblutZqFYhAoN2B1hJdOsEpbfjlCaLRBsBB8Wq3vAyqt3FEHQarfjOF67v6ubEoper8uViz9j6IDXjFXV4Jy75AZuW7g6lWjGhJSq1W7ljnJIfQC7I6RufUTjjJSwknaE8F4Ok1QWtFv169hIq9OymZ8rIWialYJAjMejeLEg7Vwk8YHXjY0AaQ944Hm3+mthzDGXSioZn511SA8L23ZgR4QTa/GVsIxN9p0hcDiksqDtip4mm3a7bXv1q0wIWgVBMJvNZpeXJgsaQwe8cmjYX19fT6fTIAi87Ae/V5KUUp51zvz8Z0Shwb0IoUwWlkvCUqSEpdCOEGhW//rNRtNvduSrAmkDHF5eXk0vLpCBBd4MnPPpdDqbzXK53GrIxy+4C1qttt++ENMmuBchzWJO2SQsZtoRKrQjBMxfztOE0mw2j46OXJJJaj4KAjGfzyfj8X6vGYDnggzseDy+uLiwKhyJBA05KVJGYchbrUbqhXu6ZPBqoCxo+scVdFLzS7QjBAl+Q6Tj4+M4jrn3vDPSQRDO5/OxNsAIxIFXjWt0zTkfj8fz+dwJQbvQICElC4Ncvd7ItgoGYAtGb5LrXV+z8UcCLxhAQONLbVBLYFKjdCf4MbcgCG5v75wBRjUweL34YjKDweD6+joMQ/+QOzGWUe7oqFZr+DsysMHgXoSiscJUMlhIjhL5q0DjB5k559VqtVgsUiVSNgQthIjj+P379+5J5eVw7e8mAHgwvrPb7/cXOrc/JYFOw1vKuFgo1OtV+pagGQnYkUQJy7XOQhIW8PHVnpVSpVKpasWwskfpQa/by8pEYxSB14W/xOz3+1JKKq5LEqHtzp2MWKlUrlQqXJfA+/I1AGzBKGHpeKJikqR/kYQF1kCL+iAI6vW6a4iUOmH5/4J/0X3vh++wGQxeOq75jB2rbgUpdP/NwWDg3F/vNVoVmqtIxvV6PZ8vSKmr8+Q+bwW8IqgdoVXCMgs3JGGB9dDE1Gq14jh2z6wMEm2h379/Txbar8rY20UDcC+29t3PrqJid3rc7/fJEvsbKzpxRuks6LjVanFBz1s1BYx5cB9CLsecZC6rnmklLFJ6wTYwWA0g07zjmp5ms6AZY2EY9nq9u7s7Vy6M5jDgpbPBA1Y6qCOl7PV6bjzbNaWZNZUSMpadM63CwRXpsO77fsDrQJC+oJlJVdKWAcs34OMH5c7OzlISHH60OQzDyWRydXXl7wpDnh68Ujjnt7e3w+EwCAJ6xo1k2hKmx53OmRGqtN8FmGFwL1RYbgLO2usVJPIraDcPSVgHjx+Uo2fIABO+IoHrxzCZTIbDYapVMEYReLlsCEFz/d9kMhmPx1QE7K8jpZTux3anTfWcprsNADsgbAq0VJQ5wK0SFkMSFmBrU6ja7bZTpU+FoKkU+PLyst81vdu4lnlJskYBeIFsCkHrPd7RaHRxcREEQWom9BadzOzLcNOzP+loAsBmhI02672O5VypknaE8FrAqv9KE1Cj0cjn81SJlBIGIg94sVi8//CBXq5tsN1I3ve9ALCedR6wcYJ1DdLNzQ2FoN3zbuTHcXx8fNxqtdxbKabQxgbsgqCcAaZsby2d1GfaEQKwCk09tVrt9PTUlQIT/k5wHMddbYCpv1ayW7y/Kwdgd1Kebrfbvb290VnQbDXln1Q45GnxtFGvJ66zggcMdkI4ySszmVISFkfYGayBlv/lcrler0dR5J73Q9A0JfX7fXOIKUnZpJiOwKvCubndbnexWFCJJjkqfvqVlLJcrpbLZbMjQ3XAMMBgBwSnlCvuNupcEhaUsMAalFKVSqVery8WC19pMpVy1R/0TYxOr+YQkAOvBRfLcU5Ir9ej7r9U9evvvFAIulqrlsolM2FypjhC0GAnBMlwkBKWLjunTWEkYYEMeg6SUp6cnLQaTdePIdWvkCam4WDo5iMSENrrpQOwK4nOhs176PV6dmVJ/3lFAZxFUVSrVE+Lp9kvAgDbEcoJvjBu/kUSFljHcnGmqG5cdc7fOa0rc9Sr+g2CYDgexXFM4kFJKGXftwDAvfgesC+D5Y6nzpSxbDQaYS6U9nwkYYEdEZKbvAJFCTPaD5bGEMNrAYZkXa+tbLvVTFUA+2nSVAo8n8/dRGZs877vAoDd4ZxTd6/JZOI6AevCYOXX4MUyfvfuHSnoryTTYMSD+9AhaG14bUskGlUK7QiBjx4dnOo0OOd+zqfv+1KMOgzD2Ww2n8999xf+AHhd0OidzWbj8dgVAVO/IzuklYylUqpFlfHeq+ABg10Q1M1DNwEmIUqhM1apOglKWMBA1RXux1bnLAxDVwqcUp0UQozH49ls5iva7/XyAXgYblSnDLANUNNmsIrjiHPRbDWNCAdaj4CHILTwhqtbEzqFngYekrBAimQ19tln5ycnJ9SWPxuCJjGs6XRqn8GUBF4ZblTPZrPpdOrJYNE4j6kdoVIyl8uROKvfghOAXaB2hJIpKVycUDHB0Y4QrOCCbqYfQ/u8UCg4D9gvzCAP+Pb2djQY2iYf8IDB64OGNGUzuCQsPZyVjQ4uXZR8Pt9ut+1RNP4CD0AoqmwzyTXKa0eIMQQSlPFk9Y6FUvVGrVgs+mL0qcaFcRz3BwNzAEMJvEJoJA8Gg8ViQfn8fiYhba1IyYrFYrPV1BMpzwqnA7AFUsISSmfXuBCL4gxJWMDHRJgpP4+z4ulppVIhLQ53QqpDw/v3763lxmgCrwnfgroiYHJsrWCRJEUOGctarVqr17UjQ/8gcAh2RTDqQ5iMG6GUE8SCEhZIIN0BSofOH+dJjdIXw0qJGHS7H0i7g5uWSAC8AlJbud1ud2mShW0rYoNBdOZisaiUy6fFU52Y5bogAbATgkTTbBt+04pfKSmRhAVW0YNgORSkkiLkrVaLtDjM0VSnNsF73R6VuOl5C2s48DpIehFqKJDDSQZOSxVZT5hxrqL4rtFs5nI5qRg3hSQA7IqwG3TOiZGcJ1Lj8H2Bg5s+HWaGSmV+uixo+jEIgmG/T1qUCms48HpIqZr3ej2K4vCVfRaaHgMpWbPZpKeQNwMeipDGsWFKknT+0mWRSMICazFe8NIPIAOctBr0sqCVUmEQDkejeBHRjIXRBF4Xbk3Z7/fDMNTpVtRbU2pTm1jcTrujg4d6AkXCIXgIwjZLtyl83PjEcFhAlqTYl7OkA/nqUfpXCDGeTibTifaBObKwwCvCVRPd3NyQCgeNbPrHjGhrbNsdLYNl5J9hgMEDECZJhrvBY4MtSMICGfzx0Gg0SHgyVQpMJwRBMJtdDocjq5KFIQReDW4beDKZXFxcCCGYErpxjbG+lGhIA55WoiYqjYUmeAiCS0Upz0rLT1olLAklLJDFHwmNRiOfz8dx7Cc/u6NCiNubm+FwaA7t6YIBeASunHcwGFxdXQVBoLjuwrCcHrUx1tUjsZJBEDTqDcaMoi88YPAghN7VsO0Il56v5EoJqECDdfijolwukxiWO5TygG/vbifTCRZw4NXhFpT9fv/29jYIAq6Wdpca1mitfH1OLAuFQqVa8aXiANgdoVyXaWWSsKCEBe5FKVWr1UqlEqlRZo9yzu8Wi9FohGUceF34alaj0ej29tZ1syYxOEWV7ZzJWJ6enlYq1STGg7UmeAhCCxvp8jUr86J3gbGVATZCyQHlcrlUKtEmhXue5h/ygKMoMiIGymTaA/C66Pf7URQJIRSXSkcHjYOi217HMi4Wi1XrATNUIoEHImjaFEKn2C+NbwAlLLAdCtDptX8lpcXhjxSl1IcPH5K9YQwi8Brwkwr7/b4Z4ZTnzG1LbH1CLKNSafktsEKtHB4weBBC17PxmFSwJHe9CKGEBTbhmv42m00/BO0UgjTLaWvQH7hlHEYReBV4S0be6/Ws+oYRorEZ0svjMpb1eq1QKCznSU4+MAY5eACCFnXclCCZULRWwkI7QrAep7bRbDZdK353yD4UQRAMhwMZS7uFhrEEXg3LOVB7wDoh1cZ4GDVcWM6RnDMpZaPRMl8BxYyLDMDOCEm5q8oqGSXtCLFrB9bjYnStVis143jJKCoMc73e4ObmxhhsDCfwGnDDe3G3GA2GgQiSQ94SU9eOLNegngIrDDB4GLoZAzdKWCa5QDGI94ItuBhdp9PJSFG6YaPCMJxMxtfX175S9B4vG4BdcBvA19fXw9EwDAJv3Cqby2AclWaz4RoVIgQNHooQRlXcdCU02VcmNo1qYLAeGhLtdpv6dngqHIl/EARiPB5fXs5SfYIBeOHQWJ3P56PRKJfLeUdcsbvpYdNut/z6eIxw8CAE/at7zBkhFx1ZkZIhCQusx80ynU4nDEOnRumfI6UMguDqajYeD32vYk+XDMCDmc4uLmYzIYTfndAd1S2/wvPzL9kZEsMbPBjaA17aXLGcVCUFUpZuDUMSFliDsjDGarXa8fFxKgpNcM6FCG5ub3u9nq2S3N9FA/AQaCQPh8P5fC6EcCtIt4g0Lb/CXKfTdvn/mC3BQyErK7ltys9Mw1coYYH1+AHnarVaLBZdJVJqAuKcR1Hc6/XpR8UxQ4HXAQ3UwWBwd3fnG10/2iylLBYL7XYTAR7waIRRTzN5zxxJWGA7bgJyWhxRFNGh1JihhKyuNcAcMTrwquh2u1Tv7ve6drUhURRpLbiyrwEHwIMQWthFMK6TsTjnPLBJWFDCAmtwi33OeT6fbzYbi8XC9ZP280WpJKlvDTBTjKMrEnjxOF+WhFRZeia0RUqLRbVay+cLNjqNeRI8GMHlcqqUNgAtuUISFtiC7wGfnJw0GvU4XuhWbcrPAqX2WlrOfmA2jXVDtz1fPQD34aa7fr/vP+mSH/Tgj6WUrVarUCikYtQA7I6gQCF3QkVa6wVJWOBedBZo0Gg041jaXWHldSSUFJUbDIwaJVcYTOCl4+rllFKUP+h5wOZ/6RQtg9XI5XIuJQKzJXgoQlJqTKKEReXkHEpYYAsUHdFKQA3aGLOVGO4EyoRm/UFfxlII3XkaapTgZcM9hsOhE5BRSiZ6lPosMsAIE4KnILhUKykGJGrKFMYT2IJb8jebbRLFtZUYzgMWUsowzE2nF5ezmVWwx6ACLxrnAS8Wi8lkQjVIemxLWyiiqMydc95ut91L9n3h4FUiyOU11pdLzgV1vBRMQAkL3IvW4jhyIWgXNeF6ZRcE4dXVfDqdQg4avArcbu50Op1ZFQ7baJ/0ikwUOgiCTqcDDSzwFATlOytTJhJIppTOy5Isdis7rO/AJprNZhCEOhpNgTiR7FwoFYbB1dXVdDK1p2OSAi8al+E/nU7n83kQBKleC+SuSBlxxpqNpnshJknwCHQS1nLqlJyZ/9PbvyYJC4BN0JK/2WzmcqHXFZhS+pRSQstBB9fXV+PJGDMUeF1MJpP5fB6GoZdjleRCS8VyuRy1QnJ7xvu+ZPD6EIraW1IS1nKKlDRV0mSJmCHI4rwEpVS9Xj85OdZrOL1u4zSEdJiOs0CI29u74XBIL8RgAi+WlKLqaDRyBpie1gOe9llYJGUud9zS/bCxuASPRuiEK07BFebmSLuYQxN1kMWvuyiXS/n8iZeHkvQD5kyJIFwsIicHjcEEXixOrpxGcq/Xu729DXQvQpvhz93wlnFcKJzUajWkyICnILiguLMeQ9wEn6GEBe6FPODKkqqXBS1oG1jrbiw9YCXlYDDwvF/4CuAlorNQmZvuut1uHMe0B0zJDXr319YgxVGpfFqulBB/Bk9BMBtVoWizEkpCCQtsxU03Uspc7qjVakdR5Kn0MdPZ0qSsiG5/oBIRSkxV4KVix61Sqtvr2qCNtKPa7stpM1w+rRwdn+hqeMyQ4JFQMwadhGVC0BxKWGA7qa6onU6HFmr2CdKkNM8EQdDrfeC2FRKmKvAycXMdPeh1e4HpBCy8BEOScxOxlJVKxRYNo80XeCSC+vBzxZmeMLlySlgYUuAeyJp2Oh33jDeLGbWsMAw/fPggZSyE8Y9hhsGLImnxq5XwTQj6Q1fHn82+r6IeNTRBcqbiqNnSNUg8afOFUQ0eiqAgYZLzzKngHNouYCO+r6DFsJoucOdCzVQ0yZjK5cLBYDCf35Bj7No27PUOAEhIxjMjb5bf3d11e90wDJ28OVlfesAVl4qRCocyJhu5MuAxkAFWuiOh0mFDq4TFkYQF1pNanJEgnx0nTste0JZwGIrp9GI2u7TtCtG6HLws/AIkvT5Ul5eXo9Eol8tJadq0ssTNpUJ31umcU+yQXgIPGDwCweRyBaeW/8OYFHI5wJZP6RQsDCywHt+IUgiatoFN3yNFoZTl/woRXl7Ohv2+l6KFQQVeKjpyMxlPLq8uRSAoK8brw++aYbM27bygwyZ4Arohv5I6+EzVSOT2UnHSvq8OvEh8IQ4Sw8rn81SJ5HxcbuNyQojb29uu6+xm3wFhFfBCcOOZ23zmfr9/c3MjhCkS8c9ljEspj45yNvDDESkEj0YorVhE7opez8nV/V+s70AaZ31pxqnVaqenp3Ec01Ezj+liJK1ZL6JF5HqbK6t1AA8YvBCS1aRt19Xtdu/uFkIIlZTN2bbp2gCfnOQp9SFp04+gDng4QuddCSsBzcx6j5sHaOAKsiQeg7a1xWKxXC6THLRzbM3Cji+HVyzj4WCQ1ToA4CXgYjmu7dGHD704ivym6LbXyHKijOP49LRUqZRNTtZq/RIAuyP0kk4vAI1YoK0AVkjCAvfDFT89PS2VSuQBW19Ajx99nCsuper3+9ZRUFYzC4MKvCC4t9fb7XXdloo+uDJcoyiqVCql0ilVCadEpAHYHUF6CVLXmEup068Y5WDpFCwoYYF1uPlIcVUqlRq1uhXDosPe5gUXnLPBsMdMfpZuFKw4otDgRWG3VJYjs9f7sFosR+Vz5ocoiqrVaqlUQj4/eCLaA1bc1MBxk3IvhOAcSlhgI7Y4kiupwjBstppkgM18pKivpdkGFkL0el2+fMA5l9zlaGF0gZeBk2mjYTkaTYRIJsAkRq0Hs5SyXq8XCgV3vv8mAOyOkJwzTolX1JVfclfrhgwssA7nHJDWlVKqc3ZGh/z5SNkpSQje7fUXC53VohhD0gp4YThTqu2rGo1Gqdiyb4njOCYVDhemhswqeBzCFpHbbly2LQMGE9iEm4zcg1ar5U9AZlbSj5VSuVw4mVzMLi91HM+m+QHwYvAr6Obz+Wg0Jg84ifSs6p+3222X/+yOwgMGD0XoZCsTcNazoiDRU7EcTgJJWCBLop1r1/7NZpMHQiVFGYlPoCuRjiaTycXFhU71U0juAy8Nf0hfXV1Np5NcLnTG1R/P5CWTB+xXBOz7DsCrRND/aPUr6sZAQtBoRwjux807n3322dHRke0KnArZsSAQM43d2EAIGrxQOOeXl7PpdEx7ve5J94C2gakI2IGRDB7H0tE1euIkM64zstCOEGzB91/p33fv3h0fHclYZtSeFecyCIKrq8vpeGJsM8TrwQvDebqMsQtNEASpDCznIodhSDJYGMbgiQiVNGrVk6M2xUoytCMEm8j6r81ms3CSj+PYOcHk+1JOgRDi7m4xHI2SVAMAXgZuMLudkclkYnQoNf4WL2VgHR8fuxacCEGDpyAo5SpZx1ESFkdQBWwk5QErpcqlclGrUSYp0KbvAlOK01zWt3LQSMACLwd/MNOkNxwOKc7s22ZnoaMoKhQK9Xo9K4oOwEMhEUomrBKWl4QFJSxwH8pMT4V8vlwuO/fXButW5rWeNsD6acxW4KXgizmbVvzdrjvkWWjzXBRF5XK5Uqn4TjNmSPA4BGmfxqSExaTWe0ESFtgJZZ2A8PioUqn4BtibmMzg+eL9e1O8xFA0CV4Qvs6Gb4BXzSp1axBSyVazlc/n4fiCpyOYjhNyUsJinhIWkrDAfZhdC6UE581mc7FYpAywnqG0XkHI+4OuidcxhW0z8ELwN3Hp8WAwIE1AfdwJpkqmT4uiuNVshqEtUtrz5YPXjZCmpQ3LKGFhfgQ7QTNRp9ORUtLOmY3dkRoHj+P4OJebjCd+kRK8B/AS8BOsKKTc7faCIGCM2mvyZKzq6M3i7q7ZbCaDnN5kv/cAXi1CR5pZEoeBEhbYGeVxfn5OT5o2hcokFdC8lgty0+nFzc2NdxSA/eMGJD1YLBaDQT8IAjdC7XpR6M05RkV3S4/Y9U5CWiF4LEKk2w4iCQvsij95nWk5aD+2vJLDIoLxZDIej6HbB14afp7z5eWlFoLmnpSqMp2vraN7fnbmxFZNtfs+Lx+8YoQrPpLGEYYSFtgdxe3qv9NqC76+ySCp911dXXldgTGowMuCRuZ4PJ7NZlqFwzzv+gSTVhFjrEEyWIkDDAMMHomQVPKr82joKShhgV3xEliqlepJPk+J0F5ii3Mg+Hx+Tb6FfSkSWMBLwUVlhsPBfH4dBDkbV3aDWXAupFK5XK5Sqdi5kaO1CHgKImlrY9dzuskrkrDA/ZhKXz1UypVyPp8nLQ4vEG3ieFoM62YyGa80TdrntQPAUr01GWODwfDm5joMAyqgU8oFqJfjXMYyn8+XSiV6rW7fr7tfIxINHoVtuWWVe03+FZSwwA4kO2FKVaqV0+JpHFPuqJm5nIauNsC3g37fHlWYr8ALwd83GQwGd3d3QujFIZVmmlCN4kpJGRfy+bI2wCS2qtP8GXpsgsch9CDTYRZufGEkYYEd4cvFvxknlWq1VC6RAXZd+ZfDiktSo4wi+f59V7n0aMawxgN7hzv0knAw6MdxbHQolbG+dmrkcRxXKpV6o+k2VuABg6cguO64pf9ZDh8JJSywM4olCc+np6e1Ws2FoLXXa3SvtBnmYRj2+j3OeZDI3O/16gHw3F+l9RD6vb6f3i9dVy+93RIvolqtVqvXrE+iGOwueAKuHaFNZ1VIwgK7wr1mMkKIRqNBSVj6SeFqOHSytAqDoNfr+iocHAs7sG/cFCe00e33B1wIvz8STzwQFctFrVYNba9+5lJn4AGDRyGkbUdIzi7XupTSmGMMKHAPzldQSpFCkAvpuTxofaIUQnS7vdvbWxfww6YZeDlwwWUUDwYD4Xoc0Q6vXTKq5SSpavU6hQaN62y26+gt9n0P4LUhuHF/l+PP381AyxqwIy7tmcSwnA1WJq1g+YRULMyF1Gk1Uc/CGAMvg+UYVXx+Pe/3+6TzbJ73vBDOhFKs0znjHiSjDwMMHocwQWeTM6O0SeacUZ9+gSQssB2ypDRCPvvsM88hlpSgQuZYKRUEwWg0mF/N0MMcvBwSLWiurubzwXBABtgM3HTLfe4kV239CELQ4PEIrgRj3MhgSaGo64cy+QdIwgKbMCq43hLtS1/6khDC0+JwJpYMcDibXY7HY5ae1ADYG24DhTM+m83G43EYhu6oV6FkxLDenZ/5R+EBg6cgdB6B1FVIJpzCFRNa9gW+L9iCEwpy67N2u318fOwCzJy7LTKmmBJBcHNz0x+O3DsgBA32jp8VOJ1Or66uqAbJ79BgSn6VCsNcx+6zYHoET0corkzsmfROl8YY7QjB/bjyDJdU1Ww0Sqella7AVCa8HGBCcL5YRKPB0Lxq+STGGNgz/gAejUbUsCvZA06CfyqW8fHxUbPeSAV+AHg0YumGcGFEX5jRf1Ec3gm4B2OArYyfYqxYKFbKZZNOb6uBmesOzFgcx91ezwlBY2sD7B3f3Pb7fSpkT9lXWi9KKQv5fLlcYomSDAYweBKCM57IrXHtEyvXlRBKWGAjbtvL5VkdHR8VT4skhmXFJk2mKKkKMca63W4S90M7BvAyoOFKg9N/0k2AnAsZy3K5UipVkix+AJ4GlSFZJSxpHBalpEQSFtgN4w1wXiyetlrtyPoQRsTAO4ecDGa1OziyVsALwDnBg8HAPenvDVO8ZxFF1Wq1UCislCEB8ASETpWRjEmThKWnSiGQhAXuga8+UkrljnKtVjOOY38LzVfLWhrgXl/xlbgfAPvF2VrfA04FmYUQTKp2u50v5NftEAPwGIRkFD80gkWeEhZML9iGX/rotIGazZbUBtivD3ZzWSDEcDjgruvlfm8AAI3Lw+rpBIWVjpneNCilpEK7tUcBeARCm15OfV11WRuUsMBOJKWP1DRGSs54p9P26zeSk/WTIgyn0+n1/Nr1MQRg79BYjaJoMplkOwQTSqlYyrOzM7f7C+sLno7WujJy0Ipz6SlhIQkL3E+itaGTrdrtthMS8qcweiYIggsNRhR4OZDFnc1mk8mERm9WKEYqFXDearWxewKeEUF9kKyYWgAlLPAg1KqsVaPZPDo6omGTXboFQlxdzSfTKZJIwcuBRuN0Or28vCQZrHQGlp4Kw1yu1WquvnI/FwzeDEJxliRhMShhgYfBvZAdY6zVaoVhGEVRavDQDkcQBPPr+WQ8drMbzDDYI/5u7mQymc/nQRCkfF9njHNh2Gg0VwY2ZkfwNMjKMr0TDCUs8Hhokmq1WrlcLoqiVKSOHN4wDG9vb6cXF6kuDgDsBScEzRgbj8dXV1euFVJKjVJqD7jRqCf9+eEBgydDraeN6oaZMKGEBR6CP09Vq9XT01OqREoZV2rav7hbDPsD/4V7umoAVkLNo9Ho7u5OV2By3/rS0TiOi/l8o9FgmQRDAB5NqJOslCsB5lzoQiQtUuQUfTHawFbcPFXRkJ6fH2R2Q0hK2e11/dokjC6wL/zhNxgMpJRBEPhrSsWYkpILsTTAp6fVSnUlLxUjFzwNoXvwUxLWcqKUUMICj4LmrDAMa7UaObtJp1UPxtlfvP8CIwq8BHxr2uv1aNxmZzzO2CLWMljFgvR1KDGKwdMQTItw6BokEjRiUMICj4Mmplar5Qf3/AeMsSAI3r9/zzl3Td/2fdXgoHHjs9frrYkt24BzHMeNRoPbQWtUVjmMMHgSQnFSwuKkB811KwYr04uxBXbCFy44Oztb26+NHudyuW63K6UkV8PveQ63GHxK1Gq7hV6vRynQJDxp9lBcQ6RYtlotpUi3iLZX0IUfPBVhsp+ZVb/i1MQGSljgwZBncK47lqcMqttUC8NwNBpdX1+vtpqB3gv4pJiO1HrICb3F2+/3aQPYF6fk3vhsd85MLIevyLDu8zbAK0doHSxlB5qy7QiVWI5AgZkR7ILvwrbb7ZRf61cckQGeTqd+lik8YPCJsVLkZhDO59ej0ch5wKkhTdDeiumuyWIYYPB0BFMiScKSgh5IpWIkYYFH0el0fG2slKaBCILpZDrQlUja8GJogT1g5zRjZceTyXQ6dTJY2dO1B9y2j7lOkUEIGjwVwcyGhqSlHS3xhBACSVhgZ/yK3kajcXx8LKX0k7DcQAqEmF/P33/xhXkhJMfBJ0atPCBzOuwPLi8vhZdj5SOlCkNBHrAWC+QUwDY/AvBYhO66QC6wLhph0ghDIwkL7IxvZSuVSrFYjG1bfr8UeOkBc75YLL54/956wFz6wWp6t33fDnjLGDeD1A/MwrHb/XB9fe3rUNpBa1ptFvL5Rr1OadH2HIxW8FSETrgS3Pq/ZIRdWTAGF7gXP9qslCqVStVqNY7jlOl1ygZSyp7ufE4haKqBY0klHHwK8DExGafKtqDRBvhDb7GIaMm4WjtndChPT0ulcokZdSJkQYPnwWhBM0pJWBphQX1aOZKwwG6kao2cAfbFsFI9zEn0gGY77tld+BTgY5NoaEjyZhUNyChapPIVTMqVLgIuV2vlchkxGvC8CG5Go14NSu0PK0rDkqSIhSQssB3fuCqlisVivV53IejUGo7Ehj58+KBlT4W/e6xggMHHJyl+C7hxOhjr9j+k6uJ8qdSlB1wsnpycIF8BPC9aCUtJI71hhxaUsMAjWC7YYhkEQbPZdC2BU1rQ1Ja/3+9LKQUlS7uehgjqgY+P3VxTfqi51+ullGH8TgxRFNVqtUKhkNohxkoRPBEhdR4qU3aiZK4dIQYYeBhLX0EHTlrNJnnA7nl6QGMsDMPJdDqfz62iva1ChwEGnwQ/ZkOiV4PBgFKg/fWiM7dRFJ2fneVyOSmlWt00AeApCNeHwSz37I8IO4MdWZmt9J5Zu9ny9339eY084IvZxdXVldeiwaRAIwQNPjbZCPP19fV4PKZOwC4Py6+dk7Fst9rOLSYjjIEKno4gB9hzRQTpn+rdOSRhgfvxjSvXGXytdtv3J1InLw3wxcXl5WUSzTPRF3jA4FOQkmmbzWZkgFMV7clpTHXa7WyKFgwweCK07cEZF0aTiOvpUEnJoIQFHohiXCwnqfPz8+OjY7/Zkf8gCILLy8uLiwvzjODocA4+GelWg5zTaHRC0P5opAdBELQ7HfvypFE6w8QIngYlYSmmpJY4NYNLCMEZkrDAA+FmwJydnZ3kT1K9CN2DIAiur6/Hw5Etx+RrfWUAPgYuBO2GHMVjXAdrwo1YqdTx8TEZYG2ZOV/X7AuARyBsZw9tfRVny/9XSnqrPAC2ki00araahUIhjmO/DYMf34uiaDgc+s9gLgOfhmyS83g8ns/nJIOVKliXUso4LhQKHSMETf3TzaYJPGDwRITRvjI15zSwuDbCGFtgJzwhSWNry5VKuVyOoihret24+qDFsJgyXsU+bwAcGKkIc7/fT5Wt+8vKOI5LpVKr3bZphhir4NnQIpRcGoPrkrD0jIiqc7ALiXyBbe5WKpXq9fpisfAjzylvuPvhg5afJEHofd8DOBh82VR6pktrwdXtYTdcoygqnZaq1Wo2QRoTI3gipIQlFLm/NCEyyaRkSMICD4dkg3K5HBng1L6aP5YGg4EW3zAGek/XCw4OF4lxNrjX66WKg80DPSdGMm61W3mtwpGUzKGDNXgOtBDHyiYcLe+ghAUeg/MbWq2W0+JIbbkppYIwGPT7lLSlGELQ4NORki4nGSz/ULITzJngfHF3d3Z2FtBS0onGwAMGz8HS91WkhKWlAzklYjErEQ3AQ3AB53fv3q3VFaITQhEORkPXlwZt+cGnxw1OUib3n+F6SmRalDeO4vP2mU6HltreYqyCZ0NwLQTtOl8aXQR0IwQPxw/utdvttS4CnSNCMZlM5pdXqYIQAD42KR3KOI4Hg4HfCdicRv2CNefvzo1JZigPAc+J0P3gaKJUOsNekC3WkZYAkRawO/44abVafsw5FbULguBiejEcDV1Wy14vHBwQqRD0dDqdTCZkgOlJs2PCknUhrSZt6RHGKng2hP7PtqZmgpm+DApKWOBx0ATXaDSOjo5cdUdKGjoQwfXN9WAwggQH+MS48UadFcbj8Ww2C4LAL0Byi0WpJOe81qhTPJA6pe/18sGbQmh3VzImXbwFSljgcfjubKVSyefzcRy7Q/5pQojbu7vpdOLvvQHwyXCJgf1+//r62vUi9JsgccallMcnx+VymSH4DD4CQnFFaVdW3kUai4zRBh6I3y21Wq0Wi0XygLP2lXO+iKLxZOwXhKAbEvgE+F33qRzu9vY2Wy9HPRiklMVCsVqt6u06bnKjAXgmtBKW3vcliVPTjhBKWODh+Lu85XK5WCzSFkY2jiKEWNzdOQEE83J0QwIfn9TM1u/37+7uKATt9/VyMljFQrFcKlMNkuvnD8CzQEpYVI5JM59rR4hyN/AYaNiUNdm2/G5qi6Lo888/d+4vjUB4wOBj4+cGKqX6/T5ppvpNgt3JSwN8WiyXS14TfkyG4NkQLGlsqR8ItCMET4KGzcnJSa1WlXJFi8M3t0EQ9Ho9vzUNDDD4NPh+Rb/b89MD3QlmrEpZKp0WiqeuMTCGJ3hGbDtCZoveFNoRgidB05ZSqlwuSSWd3fWldOnfyXiy0isJIWjw8TFakjZU0x/0Uw0KfcFnuTTApaOjnHkGqm3gWRGSK6oxN4NPW2BlzDHWeuDBuI4x9XpDSeXLY/mJWiSCf3d7l8x3MMDg40NjkSslhFgsFt1+j1Kg/RRCOpNiOc1mw2smwqACDZ4RwT29QBpfpIOFEQYeh3N5O50zIYLUUWeMj46OPnz4cHl5iSxo8Ckh31bpITebzd5/8T7M5RhfoxFNz7RabVMAbFxgpMWAZ2NpgO14UpxLm4SlBKOWDBht4MHQgDl/d55yF/yBlMvlPv/88z//8z9P5O+VggcMPhJ+IyNuO+r/xV/8xZ/92Z8dHx/J1Y6Znm45++yzL5Hd9f1j+CfgWRC6vM0koCoFJSzwVKibtFLqKz/8w8d688wd8kWIwjC8vrn+zne/szrfYaSBjwJ3jYxIdFcPue9+97uTySQMQ9PpaDVNIY5jIcQP/dAPuXD0vm8CvDUoCcsoYVn9FyRhgcdidXQ551/+8g+enZ1FiygV3LNnsjCX+0ff+tbt7a2LS5ujMMPgufE84KVvwTm/u7v71re+tbS+kvJgVCpTQUpZrVZ/5Ed+xM8iRFAQPCNC6W6EJgkLSljgiXgOxNn5ux/90R+d31yt6vwxI3Eg40ql/H/9i3/xzW9+Uwjhx1qUtcdJsO9p+8N+3efTIzpb3i191c/6e5/O2ut5xsvb1/0q5tq3Ke+BSXj2Lkk/zZWUS9f2d3/3d//wD/+wUqlIJen1qfYh8/nVV77ylR/4gR/w488v5E8J3gZCj0koYYHnwaU4R3EUhrmf/MmfXCwWpHpvg3sr46pYLP7ar/3an/7pn4ZhGEeRWwi6aCBzYZkn7A+n1Aef6L5k321T06d7f+8mi/i4uT71wh0/h2f06vw3dH/x7Bpl+0U+5vfq+UvRLKaoYwInjV1/g4NrbyOWMgjCyWTyP/ydv5PL5bgQtGxkqx+4jOVicffTP/0fHB8fx3G80ioYHjB4JhIlLO4rYekqJIw28FA4TbvLgSOYUn/tr/5ssVBcLBbWPjGl28vYcmFVLBbfv3//9a9/fTqdhrmca97g8mH83BnF1Aq7/Mh2O/lp+PYmzdbfm9r/Tv2bftv77nflhU+535UPjW18q61Hs3f04Ivc5Ra8J00Vhysz8vY1lC3wiGWcC0Km1C//8i//q3/5L0ulkoxjbhst+B/73WJRLJ7+3M/9XMr0AvCMCC4p6Oeif0oxyaRklISlkIQFHoL2cLmUgRCxUj/+Ez/+l/+dv3xzc8O5kFJ33PLaqVIhZqvV+oM/+INf/MX/tNfrhblQKRXJOPEp6V0ZT/5L9uL4Cmt/TNaWK4qYXPFH4y9JXQ29n0G24l+mfu/qRfrGaa3FWnHg7rvflReuOzl95et+XHlnsxbf8HszR/379aMC/i2sXKS6545WPyy+BsadPIa2uIwp80cxKz5roWUsI219b25u/sv/+r/69re/3Wq1nAglXaorCBZC3Nxef+WHv/JX/sqP00DFNAg+BqHiZmTRMCQhaDsi9deJu06YANwDzeJSp9dLKXNh+PVvfO33//l3RaAdkOXEyJlnrkgUutVq/dN/+n/+tZ/92b/7P/7dn//r/6FgQikVLw02i+w7KxuHlt5gzP6oVs/Xv5PTvqAwDzjtRAvGpVkOrFwSM4ZEySS2mbwJvXz1R71Y9b4q3Nuw5qtXlXnt8sflZdDrKPdH2bQM+072ndm61xoFO7oLendmI+PSy3nbegv+W2U/WO9k+1n4b+NfhnKfs/0kuP85r/tx21Xp/3O3vvra5G/kTpb65OUFcHsB+m8qhAjDkDH2J3/yJ//tf//f/d7v/V6r1cqWeNhFgxQiXCwWX/3qV8MwjOIoFKFKdlcAeDZ4pGKhlt956rfFlt9ZoZJvE+OCK6m4wOAD96OY8XGYlGR+ZrPZv/vv/3v/z/e+V6nWojjyrZ2b0ZRSuVxuMp0yJv/W3/qPv/61r//UT/3U8fHxvu8GvBGiKPrjP/7j3/md3/mHv/UPh6NhvVGPo5gMM9UXCWE0D6SUQRBcXl790Jf//9/5zj+vVKtSSWEVEQB4XngUx4JxlQSPlgaY5k5uE1b9uRKA+7HhR6lUEATf/OY3v/rVr5LPoaxHQrOeL/sXBoFSbDAc5PP5r3zlKz/xEz/x5S9/OcyFemmoz+SM6wfWp9R+jkhypBVz/pmyCdk2isNMyy9u1c8557GSnmfnVg/0gEkjkG760NlvhHaItSNoumgz5jxD4/MpLvWXyrmIikkjtWn2yJ0esfEWyb9bem986cZxG6jXr0yWw8nvNUeVtEHg5Sdhfi+9L+OJX57645ivs/W89f1KlaSZ8yRWYEK8y3dW3IXY3edMMWjl3or+IOTTK/MB0C2YP5j1jt0tMH2/9hY5hU/IceVMKB5zJXyHmiU+r7kqaT9/xfVgUBSFph20aBF9/vnnf/RHf/S9733v4uKiUqnkcjlKSvDnNF+HMgiC4Wj0v//Wb/1nX/salQJjDgQfCe4axpkRlgSOUIkEHokzquRSCCF+8av/ybf+0W+32527uzu/KskGSilivbSvYRhEUTSfz29ub0hK+pXBn1tL89nfcO988jvK5XLFYjGVz5zd06VITK/X+4Vf+Jvf/vZvr+xtK1JGAOA54VEsjWNgN1H0Spf6qAvKR8DqD+yI3XnVmSzauMZSBUJ8/vnnP/PTP/35X3xeqZQWi4jzwLkgUpqdSqU3QZZ2WPBg6XdYiXzPi1szdQvOZHpKdUkLxi1VpjCUcesBm7Wm2XRdn+XA7Qaw9kjJezM+qL0U96M7akr5vJMf981x/u7qO7tdYmar+JVzmJffW8F44h/v9Nuzv8h/1YbLWP+j+70mAHDPUUV/CuV9knZPQvvRJBK0HEaP/GD9hC8XfUmdQMEYPQJlEOQury7rtcbv//7v/+AP/ltL95frfTmzAb3jnw6AXeGRjM3XmKYkL43EpJAg/AIeDlk32stQel/t9//ZP/uFv/kfBYE4Pj7W2afCN6YUH2ZJGpQfpbTBGMWEjmfa+Ce735tKHX4+34tSuYw9YC5Q/PHZfgsfz7nc/kk+6Edr1+0K6SFZnk++Qdv2zSZq6d9NCYOLOLq5uf32b//jn//5v65bWQsvCc/2h8NECJ4PHsnIVa/rwS3Z0uWlVH6FPWDwIFRqhrRpqlEc58Lw27/921/72n+eOzo6OclH0ULobAMdkObOETKVI/YBW/p70tSb+A6oxSTKrmQIJwm1zC/xZPZH6wEbJ8xuLGYndtpb9fZkGZeJe5fk43r3bl3v55upV97K2wTm7qr5hpOf8ovuO/q0H/VfZ3VT29t4tt6mS2Pf+SJTl+zt7OrAgRRMSKaSwWA3+4MguLm9uZ5f/8b/+r/9F9/4Om3MmSor8/fHBAieHx7JBVfCRrCUNsVCMiRhgcdgCpDsPCmolk3LIMRRnMvl/vG3v/X1b/xSEATFYuH27o70OlaGmZ2dlam+UckkSCdzL7tqy4Rsg5Y0st3jlMm812ImMViT+JN22ja97e5B4I2X7602Urew8v5uKbC6OnloCHrtq7ZfxparuueaUxew+sD97Ta91SM+WJMTJmlQKubVbx0fHV1eXkYy+vv/89//xte/sVgswjBkSnHtirglHkdBJnhudAg6iefRcyYACKMLHkqyB8yMeVRMbwZLI1oUBMF3v/Odb/zSL/3553/eaDTiOGYy0a8w0kqchDKMCozJQmCcC+MoM68x3FNCkru8lq9uYSrfUXvI+4D94npOC5ujQGu4IAjGk3Gz2fqNf/Abf+Nv/LyUVJ4U2JpgbgrGmQ7V0AITgGeCL2dAm5Vi6tx1AYOePpGEBR5GxjKZxBeXDkONCP/1v/7TX/3VX/kn/+T/OD0tnp6WIhnLWNq0KSlYsLaFsHtGJU2sWUo9KnXao8ftSvSS8U0O4scgFTj9mL9q5Teq9VtOW6LJz+kPrr2GHT+BVGJz6tPLjg0KowSBmM8vZ7Orn/mZv/rrv/7rP/qj/3YURUEQmJMo/Ky82Dg8YPDc8FhKLbTB/Cx7IzODEDR4JvxJMI7jMAwXi8X/8hv/4H/6e3/v3/yb/7dUKuXzeTLPkkpbk33gFTbZ401z96PlA1Nv+Nj7fh2/2v+OZz63LflUz+n5p67Bf3LH125/0ilNcs6FELe3txcXF3/pL33/r/zKr/ztv/3L+Xw+iiJqDGx7YsLWgo8OX6g4MEkJpipEVx/Z7TcYYPAQ1sRm/Vn1/2PvjnIjKZIGjmeEC4E4ACB9SLyAPm7CBXhjz8Aehj0Bz+zTchV20YKEBFrQ7j7zgMaVsaqIzOzqqpqatitrDDP/H8Ljttvltt3d0ZEVGVHPbmTfpCQiP//8ry+++MuXX375yy+/qOpbb7315ptvRhYSnQLXsaF9r/Uz7O/gXiqzM8uHgtNuUHzJWg/QRq8vPuqgq25oi37a83R2P6NdH3A9IzJW8p49e/bbb7/d3z979913P/vsT59//uf33/8/v7NFzXNpXlmK78iAcbLYhpTaLExvaa7x2CrbHQnAOOz6SbPsy43l6JTSTz/99Nevvvrb11//49tv//vv/4w+rnUYhojEi9FAm/fJI4vPO6udi6x6HTOel35tjmdYX2H9LW68zY+wmVgv3o8XPde3uZbTlYg7v9gtHG3+Jm/8Ba4/so6+MRPzjTfeeOeddz7++ONPPvnk008//eCDD6JLpbr6J6P/EF4eGccspZV62XFBJyycpw4lTO0ZPyqz4rPff/f937/55p/ff/fDDz/8+OOPv/76a1mXXnVRWAyaXeRJ60R5Ha3XOdZmJIgb2aLvZjBIl4EmcQWN+QyPC6ib57MX11lH/fU6fPv45uuDVGcfPe/0eekIIK1FSmpFb/O/4M5P8bxf8ubPtTkGqi0dzy/Ov3Zx8NbcNH60eBn39ttvv/feex9++OH/u48++iiu7HuNLDb7tsSXJBcvk7eijMkncqm3T5a9BoEiLPTURhHFvSkaOlk8m5tNT4R6VWM6n+S/ufD4vAC8vyK9f3/ePMINhy2tR/yTmtIoctd6G7/wW7zwe+2sQq/j7vw3thObNw+1+In2f/xHrDSsb/ktR94/Pfy8P01c1Os7VXtFcr2kQdzFE4girGXT9vYKlyVo9GXXg/9iJnvrKTGf1rBIZIFHsNqQvAXdVopVy6H3zkEAp5J7u1fTdBl+4kVYl4kkrSoh8QoRPbUJtpcBRjJ/Nnzq24dX0KpmnpVnPKXBV2RyrMf4s2GKbgPp0q2NeydOILX0T2Yb32xjjfQPkp3YVgk0j50ns781vN2d2r9PdDPxWotOWFHeoElM/bxvdIupLwx5HsEJNrtJcRcD8NpQT3m19dvNKfs0wuwlMuaN2ayWRLIqiH6kvpXZ+wDw2hjKebjYOpHK2kz85+uBd7PnRZ4gAQDoQ0spgu888hYJ2bskSD0FbLMd9y0JJhUGAOCQqH+25O0RUl0S9Mstyi46wZIKAwBw1CBl7Ee5HO0Q5DKOMCoGbZYEUy0DAMBRg9WAat4f3xNfkzJGPXas59kmECEAAwBw3CBl7OUUfT3ASv23dl8XXQVcoi8AAIdoLruMJE71SsrRhCPaA9plH3BivwgAAL1o7cjWuiLECMyyCh3TXZK04meqoAEA6CA6YUmbptYGTzOOEACA82ht/hyF0KOpT4aznEsnrHtviZXqCWEAANDBkETMcu1NrpH6+gzg5DnwXb0m2TAAAN1oFouVZ299JZKyj+iSuv7M6V4AAPobfCJ6qqMxW/9Jb8SRElt+AQA4g4zj2KZTJ0llHKHH3N/9BFYAAP6otLwVjW5YMY7QR/RThAUAwFlqEVYyFU1JS6MNmY8jDGTDAAB0o+ZVVxKdsCzO/1pup4ApwgIA4ATqLTjUknrXK4u0t3bmIAADAHAKGS3HrqPygZhC6GVZFGEBAHASTXmKvlbmLmQTr8PKiyIsq0VYJMQAAHQwmA8/mm86KvMIRWZFWC0VJicGAKADtRZSIwpbNOHQS0+OxBwkAAA609RO9l7N/DUj4gIAcJohmj6LSI6JwKbq/9aYfDcbF8wSNAAAfQyx5SjHCWBvBp1TTCRMYppkFFFPks2Dca7NswAAwOMNZRhw8p1IIipTkBWPt57uLjphkQEDANCBtk1I9Zxvnv4vZ4DN2H0EAMAJvBNWRGCpgTbqsPyDs4SXGAwAQDcyWpY0heCy8GziO4BjRfqpbx0AAK+oIWVLXgI95bg5Z1HJJmkKy2LqJ4ijI0ekxsznBwCgg+iEFcP4ReROIgUWqUvOch13ib4AAHSgHm+zWDbv+Oyrz5Zza8ORa/SlLwcAAN2oWOl6Fad+IxWu/bDWCS8ZMAAAHQx+lrf21pAkcucVWElKB6zohNWWoDkHDABAB4N4gVUUYYnJmLLUleh4xzPj6JR1RwAGAKCLIQYgRbVzkqSl3Ll1wopkuLWfJPoCANCBmtR2VxbNr3Ky1glrXoQVCMAAAHSgdQeSRUfodOmElYi3AACcZEgxDtg7UUrE4emit6ekCAsAgHMM7WSvhQjCXnYlkQ+XIqw4E0wABgCggyHGAXuc9ZorbwKd1N8pM5K0hl7WpAEA6ENNYtE5WaxFS04+oL8WYVnNeucNKemHBQDAIVoG//oJX7MaWS+dOMRsXQJNHgwAwCGSx5wurTfaJuAyjrDOQZq3g2YmEgAARw1ZfRm6prV5Cr1RhBXjCP0McTkzvB6OBAAAHiOqoC2ZiWqypLH0rKU7pZdCqywDLtEXAIBD1JNeX4LOEYlzKoMJY29Sjl3BfmVhKjAAAF2oDyBUaVHVO0KbT2bwgFumBPvnWvEzVdAAABwylNO79U0MBl50wrregCRkwAAAHKRtFLCZ5ZRzshxvamusUocV54TJfQEA6GHwhlfZWieskgtLKc6iExYAACfQXLtvWGxFsuiEleiEBQDAeTTS3ai+shiKdN0Ja6sBFnkwAACHSB6zSfagqpHl+vgjK/G3oBMWAAA9Db7MHJ2wTLKZiGWLBWlJapbUV6lF6IQFAEA3g8U4/lZ+FSORpiQ4lQtJV19F9AUA4BAtfa+sROFoiJWtrT/H7iM6YQEA0JNKpLxefjWF4jL7KOqfW8Fze7toiQUAAB5j8NO5ZSeSLzhLdIhOs05Yssx4yYABADhEPfhqEvUgnGMDsF06Yd2nFFuF6YQFAEA3Q1t5Vs9zp7c5lfEMUxIcvaDZAQwAQE9qfhJYahGWJ8RTRC6XLokvZ38BAOgmthh5fC2ngm3e/4qEFwCAM8h4P4r49MEpF1YV34MU4xdkVX0FAAB6UM93xUy9+8aYYyyD5WzZa7GiCCunNPr181PfYAAAXgVD2W8ksddIU+k6WVJfKUVYiRYcAAB0VNpMivkcwrLjNyYTMnkQAICzRMsNs5SlFWE5W5Y9E4kBAOhmKMvK4ud3RX0n8BSKZdkJiyVoAAC6GcRqJ+hkksacVLJ5BBbvzzGKqJ8XNmIwAAC9DLXbRmz+1eiHlaTEZM+AaYMFAEBnQxYTK0P2S+BNyXKZDmxlTiFz+AEA6MmroEXjhK+vM0cgnnfCogklAACdDeLndy2y2ykBlpr6XqYT0hALAIC+1EcviPlO4JRTTpZL66tsZjmPKd37RdpgAQDQzeAne7OXXGn5v/5nUxKssyuTCAMA0IdmifYb8U+aEmCvijZru3+NJWgAAPrS8tanAsdk/lAich0PDAAAOhpSabJhdfygV0NLHQx8VYRFIgwAQB8ap3pLEZbH4TylvGOOf301mp1IAAD0NXi6m30TkpYNSNNHYn+wXDYDJzJgAAC60VS2GIlPRDKzPD8VbARdAAD6sVJhZdpy3DIESUh0AQA4Sywui4ia9+IoRVg+jyEGIYlJyYsBAEAnswzYfODCrAjLpkQ4ZxljM9JT31QAAF5BmlQtNh3FBiTPi2UKzOoZMOvRAAD0pzmNpeXVlBXXIqxL8ysCMAAA3VyWoFPOdRSwRtfJWJAuqS9L0AAA9NOKsAaRu9pqsgzeb7P4/a2QAwMA0J3WaDwPsx6RpU7mBwAAvWmM478mbRASAAA4g3rB1VWsZfgRAABnK/P2hZG/AACcrJVAtyVoixgc2TDBGACAM7QS6OntOI7zj67fj1jd4vT6WOtF7BvND3vkOAAA/M4tIql5z8ncPk0UBADgPFeZp7leh15ktL0Ou/ONjn/TxVc96MjtJcvTfvZBP/j+V93yfW/8Rg/62iO36sjNOO+O+tIeCzfejM2Lm1c+469/xI2Phed9+aMfKbfcqoO3+ZZbdeQZaf+vvz7yzs0+7/H7Ep5Vjh+qr8uvep4Bz2/K5rU3r8kSNAAA+1rIuwTgzXPA66+Zv7Nj53XQ4jjtaovP3vgzzH+M9Tfdv82Lz84P1Q6yf+TFlZ/8s4sfYf/P96DXqvu/qxv/ajtfe/BWPcgL/9yPPvL6u9z4WDjV7T/v/r3u4F//4O2//bGwf5zNR0rHO9JCr8fvwWek/b/++sgvzID7Pn5PelZ56JFPuntvHmrxy/lfAAAA//+0DJ6Mddh3agAAAABJRU5ErkJggg==') !important;
    }

    [data-testid="stFileUploadDropzone"] svg {
        display: none !important;
    }
    [data-testid="stFileUploadDropzone"] div::before {
        content: '';
        display: block;
        width: 48px;
        height: 48px;
        background-image: url('data:image/svg+xml;base64,PHN2ZyB3aWR0aD0iNDgiIGhlaWdodD0iNDkiIHZpZXdCb3g9IjAgMCA0OCA0OSIgZmlsbD0ibm9uZSIgeG1sbnM9Imh0dHA6Ly93d3cudzMub3JnLzIwMDAvc3ZnIj4NCjxjaXJjbGUgY3g9IjI0IiBjeT0iMjQuNSIgcj0iMjQiIGZpbGw9IiNCNUU4RjciLz4NCjxwYXRoIGQ9Ik0zMS41IDE2LjM3NUgxNi41QzE1LjgwMzggMTYuMzc1IDE1LjEzNjEgMTYuNjUxNiAxNC42NDM4IDE3LjE0MzhDMTQuMTUxNiAxNy42MzYxIDEzLjg3NSAxOC4zMDM4IDEzLjg3NSAxOVYyOS41QzEzLjg3NSAzMC4xOTYyIDE0LjE1MTYgMzAuODYzOSAxNC42NDM4IDMxLjM1NjJDMTUuMTM2MSAzMS44NDg0IDE1LjgwMzggMzIuMTI1IDE2LjUgMzIuMTI1SDMxLjVDMzEuODQ0NyAzMi4xMjUgMzIuMTg2MSAzMi4wNTcxIDMyLjUwNDUgMzEuOTI1MkMzMi44MjMgMzEuNzkzMyAzMy4xMTI0IDMxLjU5OTkgMzMuMzU2MiAzMS4zNTYyQzMzLjU5OTkgMzEuMTEyNCAzMy43OTMzIDMwLjgyMyAzMy45MjUyIDMwLjUwNDVDMzQuMDU3MSAzMC4xODYxIDM0LjEyNSAyOS44NDQ3IDM0LjEyNSAyOS41VjE5QzM0LjEyNSAxOC42NTUzIDM0LjA1NzEgMTguMzEzOSAzMy45MjUyIDE3Ljk5NTVDMzMuNzkzMyAxNy42NzcgMzMuNTk5OSAxNy4zODc2IDMzLjM1NjIgMTcuMTQzOEMzMy4xMTI0IDE2LjkwMDEgMzIuODIzIDE2LjcwNjcgMzIuNTA0NSAxNi41NzQ4QzMyLjE4NjEgMTYuNDQyOSAzMS44NDQ3IDE2LjM3NSAzMS41IDE2LjM3NVpNMzEuODc1IDI5LjVDMzEuODc1IDI5LjU5OTUgMzEuODM1NSAyOS42OTQ4IDMxLjc2NTIgMjkuNzY1MkMzMS42OTQ4IDI5LjgzNTUgMzEuNTk5NSAyOS44NzUgMzEuNSAyOS44NzVIMTYuNUMxNi40MDA1IDI5Ljg3NSAxNi4zMDUyIDI5LjgzNTUgMTYuMjM0OCAyOS43NjUyQzE2LjE2NDUgMjkuNjk0OCAxNi4xMjUgMjkuNTk5NSAxNi4xMjUgMjkuNVYxOUMxNi4xMjUgMTguOTAwNSAxNi4xNjQ1IDE4LjgwNTIgMTYuMjM0OCAxOC43MzQ4QzE2LjMwNTIgMTguNjY0NSAxNi40MDA1IDE4LjYyNSAxNi41IDE4LjYyNUgzMS41QzMxLjU5OTUgMTguNjI1IDMxLjY5NDggMTguNjY0NSAzMS43NjUyIDE4LjczNDhDMzEuODM1NSAxOC44MDUyIDMxLjg3NSAxOC45MDA1IDMxLjg3NSAxOVYyOS41Wk0yOC4xMjUgMzQuMzc1QzI4LjEyNSAzNC42NzM0IDI4LjAwNjUgMzQuOTU5NSAyNy43OTU1IDM1LjE3MDVDMjcuNTg0NSAzNS4zODE1IDI3LjI5ODQgMzUuNSAyNyAzNS41SDIxQzIwLjcwMTYgMzUuNSAyMC40MTU1IDM1LjM4MTUgMjAuMjA0NSAzNS4xNzA1QzE5Ljk5MzUgMzQuOTU5NSAxOS44NzUgMzQuNjczNCAxOS44NzUgMzQuMzc1QzE5Ljg3NSAzNC4wNzY2IDE5Ljk5MzUgMzMuNzkwNSAyMC4yMDQ1IDMzLjU3OTVDMjAuNDE1NSAzMy4zNjg1IDIwLjcwMTYgMzMuMjUgMjEgMzMuMjVIMjdDMjcuMjk4NCAzMy4yNSAyNy41ODQ1IDMzLjM2ODUgMjcuNzk1NSAzMy41Nzk1QzI4LjAwNjUgMzMuNzkwNSAyOC4xMjUgMzQuMDc2NiAyOC4xMjUgMzQuMzc1Wk0yNy4wNDU5IDIyLjcwNDFDMjcuMjU3MyAyMi45MTU0IDI3LjM3NiAyMy4yMDIxIDI3LjM3NiAyMy41MDA5QzI3LjM3NiAyMy43OTk4IDI3LjI1NzMgMjQuMDg2NSAyNy4wNDU5IDI0LjI5NzhDMjYuODM0NiAyNC41MDkyIDI2LjU0NzkgMjQuNjI3OSAyNi4yNDkxIDI0LjYyNzlDMjUuOTUwMiAyNC42Mjc5IDI1LjY2MzUgMjQuNTA5MiAyNS40NTIyIDI0LjI5NzhMMjUuMTI1IDIzLjk2ODhWMjcuMjVDMjUuMTI1IDI3LjU0ODQgMjUuMDA2NSAyNy44MzQ1IDI0Ljc5NTUgMjguMDQ1NUMyNC41ODQ1IDI4LjI1NjUgMjQuMjk4NCAyOC4zNzUgMjQgMjguMzc1QzIzLjcwMTYgMjguMzc1IDIzLjQxNTUgMjguMjU2NSAyMy4yMDQ1IDI4LjA0NTVDMjIuOTkzNSAyNy44MzQ1IDIyLjg3NSAyNy41NDg0IDIyLjg3NSAyNy4yNVYyMy45Njg4TDIyLjU0NTkgMjQuMjk4N0MyMi4zMzQ2IDI0LjUxMDEgMjIuMDQ3OSAyNC42Mjg4IDIxLjc0OTEgMjQuNjI4OEMyMS40NTAyIDI0LjYyODggMjEuMTYzNSAyNC41MTAxIDIwLjk1MjIgMjQuMjk4N0MyMC43NDA4IDI0LjA4NzQgMjAuNjIyMSAyMy44MDA4IDIwLjYyMjEgMjMuNTAxOUMyMC42MjIxIDIzLjIwMyAyMC43NDA4IDIyLjkxNjMgMjAuOTUyMiAyMi43MDVMMjMuMjAyMiAyMC40NTVDMjMuMzA2NyAyMC4zNTAxIDIzLjQzMDkgMjAuMjY2OSAyMy41Njc2IDIwLjIxMDFDMjMuNzA0NCAyMC4xNTMzIDIzLjg1MSAyMC4xMjQxIDIzLjk5OTEgMjAuMTI0MUMyNC4xNDcxIDIwLjEyNDEgMjQuMjkzNyAyMC4xNTMzIDI0LjQzMDUgMjAuMjEwMUMyNC41NjcyIDIwLjI2NjkgMjQuNjkxNCAyMC4zNTAxIDI0Ljc5NTkgMjAuNDU1TDI3LjA0NTkgMjIuNzA0MVoiIGZpbGw9IiMwMDRBNjgiLz4NCjwvc3ZnPg0K');
        background-size: contain;
        background-repeat: no-repeat;
        margin: 0 auto 10px;
    }

</style>
"""


def clear_modal_state():
    """Wipe modal state variables without triggering a rerun."""
    st.session_state.pop('active_modal', None)
    st.session_state.show_delete_modal = False
    st.session_state.modal_open = False
    st.session_state.modal_action = None
    st.session_state.modal_lead_id = None
    st.session_state.modal_lead_name = None
    st.session_state.modal_data = {}
    # Also clear any edit state
    for key in list(st.session_state.keys()):
        if key.startswith('edit_'):
            del st.session_state[key]

def close_modal():
    """Clear any active modal from session state and rerun.
    CRITICAL: All state must be cleared BEFORE st.rerun() to prevent race conditions."""
    clear_modal_state()
    st.rerun()


# --- PERFORMANCE OPTIMIZATION LAYER ---
# Note: We don't cache SQLAlchemy ORM objects as they don't serialize well with st.cache_data
# Instead, we use eager loading (joinedload) for performance optimization

def get_leads_cached(include_deleted=False):
    """Optimized retrieval of leads with eager-loaded relationships.
    Uses joinedload for performance instead of caching."""
    from app.crud import crud_leads
    db = SessionLocal()
    try:
        if include_deleted:
            return crud_leads.list_deleted_leads(db, limit=1000)
        else:
            return crud_leads.list_leads(db, limit=1000)
    finally:
        db.close()

def render_download_csv(df: pd.DataFrame, filename="data.csv", label="Download data as CSV"):
    """
    Renders a standard Streamlit download button for a DataFrame.
    """
    if df is None or df.empty:
        return
        
    csv = df.to_csv(index=False).encode('utf-8')
    st.download_button(
        label=label,
        data=csv,
        file_name=filename,
        mime='text/csv',
        key=f"dl_btn_{filename}_{hash(str(df.columns))}"
    )

def clear_leads_cache():
    """Placeholder for cache invalidation (no-op since we removed caching)"""
    pass

def get_status_display_text(lead, db):
    """
    Returns a simple formatted status string for the dropdown.
    """
    status = lead.priority or "Not Called"
    emoji = "🔴"
    if status == "Pending": emoji = "🟡"
    elif status == "Called": emoji = "🟢"
    
    return f"{emoji} {status}"

def get_status_pill_class(status):
    """Returns the CSS class for the status pill based on the status string."""
    s = (status or "").strip().lower()
    if s == "called":
        return "status-called"
    if s == "pending":
        return "status-pending"
    return "status-not-called"

def get_updater_info(lead):
    """Returns styled HTML with updater name and localized time placed inside the selectbox visually."""
    if not lead.call_status_updated_by:
        return ""
    
    from app.utils.activity_logger import utc_to_local
    from datetime import datetime
    local_time = utc_to_local(lead.call_status_updated_at)
    time_str = local_time.strftime("%I:%M %p").upper()
    
    # Compare both in the same localized timezone
    now = utc_to_local(datetime.utcnow())
    if local_time.date() == now.date():
        date_str = "TODAY"
    else:
        # Check if yesterday
        yesterday = now.date() - datetime.now().date().resolution
        if local_time.date() == yesterday:
            date_str = "YESTERDAY"
        else:
            date_str = local_time.strftime("%m/%d")
        
    text_color = "#444" # Dark grey for all light pastel backgrounds
    
    # Absolute overlay with pointer-events none, matching the selectbox margin properties
    return f"""
    <div style="width: 155px; margin-top: -24px; margin-left: auto; margin-right: 0; text-align: center; pointer-events: none; position: relative; z-index: 10;">
        <p style="font-size: 10px; color: {text_color}; margin: 0; line-height: 1; font-weight: 800; text-transform: uppercase; white-space: nowrap; overflow: hidden; text-overflow: ellipsis; padding: 0 4px; text-align: center; opacity: 1.0;">
            {lead.call_status_updated_by} • {date_str} {time_str}
        </p>
    </div>
    """

@st.cache_data(ttl=300) # Stats can live longer
def get_stats_cached(func_name, *args, **kwargs):
    """Generic cached wrapper for services_stats functions"""
    from app import services_stats
    db = SessionLocal()
    try:
        func = getattr(services_stats, func_name)
        return func(db, *args, **kwargs)
    finally:
        db.close()


def init_session_state():
    """Initialize all session state variables with secure token-based persistence"""
    db = SessionLocal()
    
    # --- PAGE FILTER STATE (PERSISTENCE FIX) ---
    if 'main_navigation' not in st.session_state:
        st.session_state.main_navigation = "Dashboard"
    if 'username' not in st.session_state:
        st.session_state.username = None
    if 'user_role' not in st.session_state:
        st.session_state.user_role = None
    if 'db_user_id' not in st.session_state:
        st.session_state.db_user_id = None
    if 'employee_id' not in st.session_state:
        st.session_state.employee_id = None
    if 'show_signup' not in st.session_state:
        st.session_state.show_signup = False
    if 'show_forgot_password' not in st.session_state:
        st.session_state.show_forgot_password = False
    if 'stats_view_mode' not in st.session_state:
        st.session_state.stats_view_mode = 'individual'
    if 'show_user_dashboards' not in st.session_state:
        st.session_state.show_user_dashboards = False
    
    # --- MODAL STABILITY STATE (CHICAGO FIX) ---
    if 'modal_open' not in st.session_state:
        st.session_state.modal_open = False
    if 'modal_action' not in st.session_state:
        st.session_state.modal_action = None
    if 'modal_lead_id' not in st.session_state:
        st.session_state.modal_lead_id = None
    if 'modal_lead_name' not in st.session_state:
        st.session_state.modal_lead_name = None
    if 'modal_data' not in st.session_state:
        st.session_state.modal_data = {}
    
    if 'status_filter' not in st.session_state:
        st.session_state.status_filter = "All"
    if 'call_status_filter' not in st.session_state:
        st.session_state.call_status_filter = "All"
    if 'show_only_my_leads' not in st.session_state:
        st.session_state.show_only_my_leads = False
    if 'active_inactive_filter' not in st.session_state:
        st.session_state.active_inactive_filter = "Active"
    if 'show_deleted_leads' not in st.session_state:
        st.session_state.show_deleted_leads = False
    if 'lead_type_filter' not in st.session_state:
        st.session_state.lead_type_filter = "All"
    if 'leads_sort_by' not in st.session_state:
        st.session_state.leads_sort_by = "Newest Added"
    if 'tag_color_filter' not in st.session_state:
        st.session_state.tag_color_filter = "All"

    
    # Referrals Filters
    if 'referral_status_filter' not in st.session_state:
        st.session_state.referral_status_filter = "All"
    if 'show_only_my_referrals' not in st.session_state:
        st.session_state.show_only_my_referrals = True
    if 'referral_active_inactive_filter' not in st.session_state:
        st.session_state.referral_active_inactive_filter = "Active"
    if 'show_deleted_referrals' not in st.session_state:
        st.session_state.show_deleted_referrals = False
    if 'referral_type_filter' not in st.session_state:
        st.session_state.referral_type_filter = "All"
    if 'referral_call_status_filter' not in st.session_state:
        st.session_state.referral_call_status_filter = "All"
    if 'referral_tag_color_filter' not in st.session_state:
        st.session_state.referral_tag_color_filter = "All"

    if 'payor_filter' not in st.session_state:
        st.session_state.payor_filter = "All"
    if 'ccu_filter' not in st.session_state:
        st.session_state.ccu_filter = "All"
    if 'referral_auth_filter' not in st.session_state:
        st.session_state.referral_auth_filter = "Pending"
    if 'referral_lead_type_filter' not in st.session_state:
        st.session_state.referral_lead_type_filter = "All"
    if 'referrals_sort_by' not in st.session_state:
        st.session_state.referrals_sort_by = "Newest Added"

    # Confirmations Filters
    if 'confirm_payor_filter' not in st.session_state:
        st.session_state.confirm_payor_filter = "All"
    if 'confirm_ccu_filter' not in st.session_state:
        st.session_state.confirm_ccu_filter = "All"
    if 'confirm_care_filter' not in st.session_state:
        st.session_state.confirm_care_filter = "All"
    if 'confirm_lead_type_filter' not in st.session_state:
        st.session_state.confirm_lead_type_filter = "All"
    if 'confirmations_sort_by' not in st.session_state:
        st.session_state.confirmations_sort_by = "Newest Added"
    if 'confirm_status_filter' not in st.session_state:
        st.session_state.confirm_status_filter = "Active" # Default to see Active clients
    if 'confirm_tag_color_filter' not in st.session_state:
        st.session_state.confirm_tag_color_filter = "All"
    if 'confirm_show_deleted' not in st.session_state:
        st.session_state.confirm_show_deleted = False

    # New Caregiver Type Filters
    if 'referral_caregiver_type_filter' not in st.session_state:
        st.session_state.referral_caregiver_type_filter = "All"
    if 'confirm_caregiver_type_filter' not in st.session_state:
        st.session_state.confirm_caregiver_type_filter = "All"
    if 'referral_ccu_filter' not in st.session_state:
        st.session_state.referral_ccu_filter = "All"
    if 'confirm_ccu_filter' not in st.session_state:
        st.session_state.confirm_ccu_filter = "All"

    
    # Timezone Detection - Force Central Time as requested
    st.session_state.user_timezone = "America/Chicago"

    try:
        # Initialize basic auth state if missing
        if 'authenticated' not in st.session_state:
            st.session_state.authenticated = False

        # Secure Token-Based Session Check
        if not st.session_state.authenticated:
            token = get_session_token()
            
            if token:
                # Validate token against database
                user = crud_session_tokens.validate_token(db, token)
                
                if user:
                    # Token is valid - auto-login
                    st.session_state.authenticated = True
                    st.session_state.username = user.username
                    st.session_state.user_role = user.role
                    st.session_state.db_user_id = user.id
                    st.session_state.employee_id = user.user_id
                    st.rerun()
                else:
                    # Token is invalid or expired - clear it
                    clear_session_token()
    finally:
        db.close()
    
    # Start email scheduler (runs once per session)
    if 'email_scheduler_started' not in st.session_state:
        st.session_state.email_scheduler_started = False
        if st.session_state.authenticated:
            try:
                from app.email_scheduler import start_scheduler
                start_scheduler()
                st.session_state.email_scheduler_started = True
            except Exception as e:
                pass  # Scheduler error won't break the app
    
    db.close()



def inject_custom_css():
    """Inject global CSS styles and time fix JS"""
    st.markdown(GLOBAL_CSS, unsafe_allow_html=True)
    inject_time_fix_script()


def inject_time_fix_script():
    """Inject a professional, high-reliability script for timezone localization"""
    st.markdown("""
        <script>
        (function() {
            if (window._uiEngineRunning) return;
            window._uiEngineRunning = true;

            const formatLocal = (utc, style) => {
                const date = new Date(utc.includes('T') ? (utc.includes('Z') ? utc : utc + 'Z') : utc);
                if (isNaN(date.getTime())) return null;
                
                const tzOptions = { timeZone: 'America/Chicago' };
                
                if (style === 'ago') {
                    // Get current time in Central Time for accurate relative comparison
                    const nowInCST = new Date(new Date().toLocaleString('en-US', tzOptions));
                    const dateInCST = new Date(date.toLocaleString('en-US', tzOptions));
                    
                    const diff = Math.floor((new Date() - date) / 1000);
                    if (diff < 60) return "Just now";
                    if (diff < 3600) return `${Math.floor(diff/60)} minutes ago`;
                    
                    const today = new Date(nowInCST); today.setHours(0,0,0,0);
                    const timeStr = date.toLocaleString('en-US', {
                        ...tzOptions,
                        hour: '2-digit', minute: '2-digit', hour12: true
                    });
                    
                    if (dateInCST >= today) return `Today at ${timeStr}`;
                    return date.toLocaleString('en-US', {
                        ...tzOptions,
                        month: '2-digit', day: '2-digit', year: 'numeric',
                        hour: '2-digit', minute: '2-digit', hour12: true
                    });
                }
                
                return date.toLocaleString('en-US', {
                    ...tzOptions,
                    month: '2-digit', day: '2-digit', year: 'numeric',
                    hour: '2-digit', minute: '2-digit', hour12: true
                });
            };

            const runFix = () => {
                const elms = document.querySelectorAll('.local-time');
                elms.forEach(el => {
                    const utc = el.getAttribute('data-utc');
                    const style = el.getAttribute('data-style');
                    if (!utc) return;
                    
                    const newText = formatLocal(utc, style);
                    if (newText && (el.innerText !== newText || el.innerText.includes('UTC'))) {
                        el.innerText = newText;
                        el.dataset.processed = "true";
                    }
                });
            };

            // Global Singleton Processor
            window.processTimeFix = runFix;

            // Scroll Lock Logic
            window.setScrollLock = (locked) => {
                document.body.style.overflow = locked ? 'hidden' : '';
                const mainContainer = document.querySelector('[data-testid="stAppViewContainer"]');
                if (mainContainer) mainContainer.style.overflow = locked ? 'hidden' : 'auto';
            };
            
            // Aggressive Polling + Observer
            setInterval(runFix, 1000);
            new MutationObserver(runFix).observe(document.body, {childList:true, subtree:true});
            
            // Try to reach parent if in iframe (Streamlit component behavior)
            try {
                if (window.parent && window.parent.document !== document) {
                    new MutationObserver(runFix).observe(window.parent.document.body, {childList:true, subtree:true});
                }
            } catch(e) {}
            
            runFix();
        })();
        </script>
    """, unsafe_allow_html=True)


def render_time(dt, style='datetime', is_badge=False):
    """
    Returns HTML for a timestamp that will be automatically localized by JS.
    styles: 'datetime' (default), 'ago' (relative time)
    """
    if not dt:
        return "N/A"
    
    # Ensure it's treated as Central Time (CST/CDT) in Python-side fallback
    from app.utils.activity_logger import utc_to_local
    fallback_text = ""
    # Hardcode CST for this user as requested
    cst_dt = utc_to_local(dt, 'America/Chicago')
    if style == 'ago':
        from app.utils.activity_logger import format_time_ago
        fallback_text = format_time_ago(dt, 'America/Chicago')
    else:
        fallback_text = cst_dt.strftime("%m/%d/%Y %I:%M %p (CST)")
    
    cls = "local-time"
    if is_badge:
        cls += " badge"
        
    # Standard HTML part
    html = f'<span class="{cls}" data-utc="{dt.isoformat()}" data-style="{style}">{fallback_text}</span>'
    
    # Universal Trigger: A hidden element that executes the fix logic whenever it loads
    # This ensures conversion even if global scripts are delayed or throttled
    trigger = f'<img src="data:image/gif;base64,R0lGODlhAQABAIAAAAAAAP///yH5BAEAAAAALAAAAAABAAEAAAIBRAA7" style="display:none" onload="if(window.processTimeFix)window.processTimeFix()">'
    
    return html + trigger


def prepare_lead_data_for_email(lead, db):
    """Prepare comprehensive lead data dictionary for email reminders"""
    lead_data = {
        'id': lead.id,
        'first_name': lead.first_name,
        'last_name': lead.last_name,
        'phone': lead.phone,
        'email': 'N/A',
        'city': lead.city,
        'zip_code': lead.zip_code,
        'dob': str(lead.dob) if lead.dob else 'N/A',
        'medicaid_no': lead.medicaid_no,
        'source': lead.source,
        'event_name': lead.event_name,
        'word_of_mouth_type': lead.word_of_mouth_type,
        'other_source_type': lead.other_source_type,
        'staff_name': lead.staff_name,
        'created_by': lead.created_by,
        'last_contact_status': lead.last_contact_status,
        'last_contact_date': str(utc_to_local(lead.last_contact_date, 'America/Chicago')) if lead.last_contact_date else 'N/A',
        'active_client': lead.active_client,
        'referral_type': lead.referral_type,
        'e_contact_name': lead.e_contact_name,
        'e_contact_relation': lead.e_contact_relation,
        'e_contact_phone': lead.e_contact_phone,
        'comments': lead.comments,
    }
    
    # Add payor information
    if lead.agency:
        lead_data['agency_name'] = lead.agency.name
    else:
        lead_data['agency_name'] = 'N/A'
    
    # Add agency suboption information
    if lead.agency_suboption:
        lead_data['agency_suboption_name'] = lead.agency_suboption.name
    else:
        lead_data['agency_suboption_name'] = ''
    
    # Add CCU information
    if lead.ccu:
        lead_data['ccu_name'] = lead.ccu.name
        lead_data['ccu_address'] = lead.ccu.address or ''
        lead_data['ccu_phone'] = lead.ccu.phone or ''
        lead_data['ccu_fax'] = lead.ccu.fax or ''
        lead_data['ccu_email'] = lead.ccu.email or ''
        lead_data['ccu_coordinator'] = lead.ccu.care_coordinator_name or ''
    else:
        lead_data['ccu_name'] = 'N/A'
        lead_data['ccu_address'] = ''
        lead_data['ccu_phone'] = ''
        lead_data['ccu_fax'] = ''
        lead_data['ccu_email'] = ''
        lead_data['ccu_coordinator'] = ''
    
    return lead_data


def send_initial_lead_reminders(db, lead_id, username):
    """
    Consolidated logic to send the first notification email for a lead or referral.
    Respects lead.send_reminders preference.
    """
    from app.crud import crud_leads, crud_users, crud_agencies, crud_ccus, crud_agency_suboptions, crud_email_reminders
    from app.utils.email_service import send_referral_reminder_email, send_simple_lead_email
    import streamlit as st

    lead = crud_leads.get_lead(db, lead_id)
    if not lead or not getattr(lead, 'send_reminders', True):
        return False

    # Target the assigned staff (lead creator), not necessarily the logged-in admin
    target_username = lead.created_by or username
    user = crud_users.get_user_by_username(db, target_username)
    if not user or not user.email:
        st.error(f"**Email skipped:** Assigned user '{target_username}' has no email address in the database.")
        st.info("Please update the user's profile with a valid email address.")
        return False

    success = False
    subject = ""

    try:
        if lead.active_client:  # Is a referral
            # Get payor (agency) info
            agency_name = "N/A"
            agency_suboption = ""
            if lead.agency_id:
                agency = crud_agencies.get_agency(db, lead.agency_id)
                if agency: agency_name = agency.name
            if lead.agency_suboption_id:
                subopt = crud_agency_suboptions.get_suboption_by_id(db, lead.agency_suboption_id)
                if subopt: agency_suboption = subopt.name

            # CCU info
            ccu_name, ccu_phone, ccu_fax, ccu_email, ccu_address, ccu_coordinator = ["N/A"] * 6
            if lead.ccu_id:
                ccu = crud_ccus.get_ccu_by_id(db, lead.ccu_id)
                if ccu:
                    ccu_name = ccu.name
                    ccu_phone = ccu.phone or "N/A"
                    ccu_fax = ccu.fax or "N/A"
                    ccu_email = ccu.email or "N/A"
                    ccu_address = ccu.address or "N/A"
                    ccu_coordinator = ccu.care_coordinator_name or "N/A"

            referral_info = {
                'name': f"{lead.first_name} {lead.last_name}",
                'phone': lead.phone,
                'dob': str(lead.dob) if lead.dob else 'N/A',
                'creator': username,
                'created_date': datetime.now().strftime('%m/%d/%Y'),
                'status': lead.last_contact_status,
                'referral_type': lead.referral_type or 'Regular',
                'payor_name': agency_name,
                'payor_suboption': agency_suboption,
                'ccu_name': ccu_name,
                'ccu_phone': ccu_phone,
                'ccu_fax': ccu_fax,
                'ccu_email': ccu_email,
                'ccu_address': ccu_address,
                'ccu_coordinator': ccu_coordinator,
                'care_status': lead.care_status or 'N/A',
                'priority': lead.priority or 'Medium'
            }
            success = send_referral_reminder_email(referral_info, user.email)
            subject = f"New Referral [{referral_info['referral_type']}]: {lead.first_name} {lead.last_name}"
        
        else:  # Regular lead
            lead_info = {
                'name': f"{lead.first_name} {lead.last_name}",
                'phone': lead.phone,
                'creator': username,
                'dob': str(lead.dob) if lead.dob else 'N/A',
                'source': lead.source,
                'status': lead.last_contact_status,
                'created_date': datetime.now().strftime('%m/%d/%Y')
            }
            success = send_simple_lead_email(lead_info, user.email)
            subject = f"New Lead: {lead.first_name} {lead.last_name}"

        # Record attempt
        if success:
            crud_email_reminders.create_reminder(db, lead.id, user.email, subject, "system", "sent")
            st.info(f"Notification email sent to {user.email}")
        else:
            crud_email_reminders.create_reminder(db, lead.id, user.email, subject, "system", "failed", "Email service error")
            
        return success
    except Exception as e:
        print(f"[ERROR] send_initial_lead_reminders: {e}")
        return False



def get_call_status_tag(priority, updated_by=None, updated_at=None):
    """Returns HTML for a color-coded Call Status box with optional tracking info."""
    colors = {
        "Not Called": ("#EF4444", "white"),
        "Pending":    ("#EAB308", "#1a1a1a"),
        "Called":     ("#22C55E", "white"),
    }
    bg, text_color = colors.get(priority, ("#6b7280", "white"))
    label = priority or "Not Called"
    # Build tracking line
    tracker_html = ""
    if updated_by:
        time_str = ""
        if updated_at:
            from app.utils.activity_logger import format_time_ago
            try:
                time_str = f" &bull; {format_time_ago(updated_at)}"
            except Exception:
                pass
        tracker_html = f"<div style='font-size:0.68rem; margin-top:2px; opacity:0.92;'>by {updated_by}{time_str}</div>"
    return (
        f'<div style="display:inline-block; background-color:{bg}; color:{text_color}; '
        f'padding:4px 12px; border-radius:6px; font-size:0.78rem; font-weight:700; '
        f'text-transform:uppercase; min-width:90px; text-align:center; line-height:1.4;">'
        f'{label}{tracker_html}</div>'
    )


# Keep old name as alias for any missed references
def get_priority_tag(priority, updated_by=None, updated_at=None):
    return get_call_status_tag(priority, updated_by, updated_at)




def get_tag_color_dot(color_name):
    """Returns a small colored emoji dot for a given color name."""
    if not color_name or color_name == "None":
        return ""
    
    color_map = {
        "Blue": "🔵", "Purple": "🟣"
    }
    return color_map.get(color_name, "⚪")


def render_tag_color_picker(lead_id, current_color, db, page_type="leads"):
    """Renders an inline color picker for tags."""
    colors = ["None", "Blue", "Purple"]
    color_icons = {"None": "🚫 None", "Blue": "🔵 Blue", "Purple": "🟣 Purple"}
    
    current_val = current_color if current_color in colors else "None"
    
    selected_color = st.selectbox(
        "Assign Color Tag:",
        options=colors,
        format_func=lambda x: color_icons.get(x, x),
        index=colors.index(current_val),
        key=f"tag_select_{lead_id}_{page_type}"
    )
    
    if selected_color != current_val:
        from app.crud.crud_leads import update_lead
        from app.schemas import LeadUpdate
        
        new_tag = None if selected_color == "None" else selected_color
        update_lead(db, lead_id, LeadUpdate(tag_color=new_tag), st.session_state.username, st.session_state.get('db_user_id'))
        
        from frontend.common import clear_leads_cache
        clear_leads_cache()
        st.rerun()


def get_referral_status_tag(lead):

    """Returns HTML for color-coded referral status tags (multiple possible)"""
    if not lead.active_client:
        return ""
    
    tags = []
    
    # Base tag for any referral
    tags.append('<span class="referral-tag referral-sent">Initial Referral Sent</span>')
    
    # Specific stage tags
    if lead.last_contact_status == "Assessment Scheduled":
        tags.append('<span class="referral-tag referral-assessment">Assessment Scheduled</span>')
        
    if lead.last_contact_status == "Assessment Done":
        tags.append('<span class="referral-tag referral-assessment">Assessment Done</span>')
    
    if lead.authorization_received:
        tags.append('<span class="referral-tag referral-confirmed">Authorized</span>')
    
    if lead.last_contact_status == "Not Approved":
        tags.append('<span class="referral-tag referral-rejected">Referral Rejected</span>')
    
    if lead.last_contact_status == "Services Refused":
        tags.append('<span class="referral-tag referral-refused">Services Refused</span>')
    
    return " ".join(tags)


def get_status_emoji(status):
    """Maps status strings to simple emojis for headers"""
    status_map = {
        "Initial Referral Sent": "📞", "Initial Call": "📞", "Intro Call": "📞", "Follow Up": "📨",
        "Awaiting CCU": "🏢", "No Response": "🔇", "Inactive": "💤", "Not Interested": "🚫",
        "Care Start": "✅", "Not Start": "❌", "Assessment Scheduled": "🗓️", "Assessment Done": "📝",
        "Initial Referral Sent": "📤", "Not Approved": "🚫",
        "Services Refused": "🙅"
    }
    return status_map.get(status, "📄")

def get_referral_status_emoji(lead):
    """Returns emoji/text-based status tags for expander titles (no HTML)"""
    if not lead.active_client:
        return ""
    
    tags = []
    tags.append("🔵 Sent")
    
    if lead.last_contact_status == "Assessment Scheduled":
        tags.append("🟠 Assessment")
        
    if lead.last_contact_status == "Assessment Done":
        tags.append("🟠 Assessment Done")
    
    if lead.authorization_received:
        tags.append("🟢 Confirmed")
    
    if lead.last_contact_status == "Not Approved":
        tags.append("🔴 Rejected")
    
    return " | ".join(tags)




def open_modal(modal_type, target_id, title=None, message=None, **kwargs):
    """Set the active modal in session state and rerun"""
    # STEP 0: Reset ALL previous modal state to prevent "clash" or "ghost" popups
    clear_modal_state()
    
    st.session_state.modal_open = True
    st.session_state.modal_action = modal_type
    st.session_state.modal_lead_id = target_id
    st.session_state.modal_lead_name = title
    st.session_state.modal_data = {
        'title': title,
        'message': message,
        **kwargs
    }
    
    # 2. Maintain legacy dictionary for backward compatibility
    st.session_state['active_modal'] = {
        'modal_type': modal_type,
        'target_id': target_id,
        'title': title,
        'message': message,
        **kwargs
    }
    st.rerun()




@st.dialog("Action Required")
def confirmation_modal_dialog(m):
    """
    Native Streamlit dialog for general confirmation actions.
    """
    from app.db import SessionLocal
    db = SessionLocal()
    try:
        title = m.get('title', 'Confirm Action')
        message = m.get('message', 'Are you sure?')
        icon = m.get('icon', '🗑️')
        confirm_label = m.get('confirm_label', 'CONFIRM')
        indicator = m.get('indicator')
        
        indicator_html = f'<div style="margin-top:15px; font-size:0.9rem; color:#6b7280; font-weight:700;">💡 {indicator}</div>' if indicator else ""
        
        # Custom Header
        st.markdown(f"""
        <div class="modal-dialog-header">
          <div class="modal-icon">{icon}</div>
          {title}
        </div>
        <div class="modal-body-content">
          {message}
          {indicator_html}
        </div>
        """, unsafe_allow_html=True)
        
        # Optional fields for specific modals
        send_notif_val = True
        if m['modal_type'] == 'auth_received':
            from app.crud import crud_leads
            lead_obj = crud_leads.get_lead(db, m['target_id'])
            st.markdown("<h4 style='font-weight: bold; color: #00506b;'>Notifications & Tracking</h4>", unsafe_allow_html=True)
            send_notif_val = st.checkbox("Send Auto Email Reminders for this Lead", value=getattr(lead_obj, 'send_reminders', True), key=f"auth_notif_chk_{m['target_id']}")
            st.divider()
        elif m['modal_type'] == 'create_lead_confirm':
            st.markdown("<h4 style='font-weight: bold; color: #00506b;'>Notifications & Tracking</h4>", unsafe_allow_html=True)
            send_notif_val = st.checkbox("Send Auto Email Reminders for this Lead", value=True, key=f"create_notif_chk")
            st.divider()

        c1, c2 = st.columns(2)
        with c1:
            if st.button("CANCEL", use_container_width=True, key=f"dialog_cancel_{m['modal_type']}_{m['target_id']}"):
                close_modal()
        with c2:
            if st.button(confirm_label, type="primary", use_container_width=True, key=f"dialog_confirm_{m['modal_type']}_{m['target_id']}"):
                from app.crud import crud_leads, crud_users, crud_agencies, crud_ccus, crud_events
                
                success = False
                msg = ""
                
                if m['modal_type'] == 'perm_delete':
                    if crud_leads.delete_lead(db, m['target_id'], st.session_state.username, st.session_state.get('db_user_id'), permanent=True):
                        msg = "Success! Lead has been permanently removed."
                        success = True
                elif m['modal_type'] == 'soft_delete':
                    if crud_leads.delete_lead(db, m['target_id'], st.session_state.username, st.session_state.get('db_user_id'), permanent=False):
                        msg = "Success! Lead moved to Recycle Bin."
                        success = True
                elif m['modal_type'] == 'soft_delete_ref':
                    if crud_leads.delete_lead(db, m['target_id'], st.session_state.username, st.session_state.get('db_user_id'), permanent=False):
                        msg = "Success! Referral moved to Recycle Bin."
                        success = True
                elif m['modal_type'] == 'restore_ref':
                    if crud_leads.restore_lead(db, m['target_id'], st.session_state.username, st.session_state.get('db_user_id')):
                        msg = "Success! Referral has been restored."
                        success = True
                elif m['modal_type'] == 'perm_delete_ref':
                    if crud_leads.delete_lead(db, m['target_id'], st.session_state.username, st.session_state.get('db_user_id'), permanent=True):
                        msg = "Success! Referral has been permanently removed."
                        success = True
                elif m['modal_type'] == 'mark_ref_confirm':
                    clear_leads_cache()
                    st.session_state['mark_referral_lead_id'] = m['target_id']
                    st.session_state['current_page'] = 'Mark Referral Page'
                    st.toast("Heading to Mark Referral Page...")
                    success = True
                elif m['modal_type'] == 'create_lead_confirm':
                    from app.schemas import LeadCreate
                    from datetime import datetime
                    ld = m['lead_data']
                    dob_val = None
                    if ld.get('dob'):
                        dob_val = datetime.strptime(ld['dob'], '%Y-%m-%d').date()
                    
                    lead_in = LeadCreate(
                        **{k: v for k, v in ld.items() if k not in ['dob', 'send_reminders']},
                        dob=dob_val,
                        send_reminders=ld.get('send_reminders', True)
                    )
                    new_lead = crud_leads.create_lead(db, lead_in, st.session_state.username, st.session_state.get('db_user_id'))
                    if new_lead:
                        msg = f"Success! Lead '{new_lead.first_name} {new_lead.last_name}' created successfully!"
                        success = True
                elif m['modal_type'] == 'approve_user':
                    crud_users.approve_user(db, m['target_id'], st.session_state.username, st.session_state.db_user_id)
                    msg = "Success! User has been approved."
                    success = True
                elif m['modal_type'] == 'reject_user':
                    crud_users.reject_user(db, m['target_id'], st.session_state.username, st.session_state.db_user_id)
                    msg = "Success! User request has been rejected."
                    success = True
                elif m['modal_type'] == 'delete_agency':
                    crud_agencies.delete_agency(db, m['target_id'], st.session_state.username, st.session_state.db_user_id)
                    msg = "Success! Payor has been deleted."
                    success = True
                elif m['modal_type'] == 'delete_ccu':
                    crud_ccus.delete_ccu(db, m['target_id'], st.session_state.username, st.session_state.db_user_id)
                    msg = "Success! CCU has been deleted."
                    success = True
                elif m['modal_type'] == 'delete_event':
                    crud_events.delete_event(db, m['target_id'], st.session_state.username, st.session_state.db_user_id)
                    msg = "Success! Event has been deleted."
                    success = True
                elif m['modal_type'] == 'auth_received':
                    from app.schemas import LeadUpdate
                    update_data = LeadUpdate(authorization_received=True, last_contact_status="Confirmed", send_reminders=send_notif_val)
                    if crud_leads.update_lead(db, m['target_id'], update_data, st.session_state.username, st.session_state.get('db_user_id')):
                        msg = "Success! Authorization marked as received."
                        success = True
                elif m['modal_type'] == 'unmark_ref':
                    from app.schemas import LeadUpdate
                    update_data = LeadUpdate(active_client=False, referral_type=None)
                    if crud_leads.update_lead(db, m['target_id'], update_data, st.session_state.username, st.session_state.get('db_user_id')):
                        msg = "Success! Client has been unmarked as a referral."
                        success = True
                elif m['modal_type'] == 'undo_auth':
                    from app.schemas import LeadUpdate
                    # Set authorization_received=False and move status back out of "Confirmed" if needed
                    update_data = LeadUpdate(authorization_received=False, last_contact_status="Initial Referral Sent")
                    if crud_leads.update_lead(db, m['target_id'], update_data, st.session_state.username, st.session_state.get('db_user_id')):
                        msg = "Success! Authorization has been unmarked."
                        success = True
                elif m['modal_type'] == 'update_password':
                    new_pwd = st.session_state.get('pending_password_update', {}).get('new_password')
                    if new_pwd:
                        if crud_users.update_user_credentials(db, m['target_id'], new_pwd, st.session_state.username, st.session_state.db_user_id):
                            msg = "Success! Your password has been updated."
                            success = True
                            st.session_state.pop('pending_password_update', None)
                        else:
                            st.error("**Failed to update password**")

                if success:
                    if msg: st.session_state['success_msg'] = msg
                    clear_leads_cache()
                    close_modal()
                else:
                    st.error("Operation failed. Please try again.")
    finally:
        db.close()

@st.dialog("Delete Lead")
def show_delete_modal_dialog(lead_id, name):
    """
    Native Streamlit dialog for the special delete lead action from view_leads.
    """
    from app.db import SessionLocal
    db = SessionLocal()
    try:
        st.markdown(f"""
        <div class="modal-dialog-header">
          <div class="modal-icon">🗑️</div>
          DELETE LEAD?
        </div>
        <div class="modal-body-content">
          Are you sure you want to delete <b>{name}</b>?<br><br>
          💡 It will be moved to the Recycle Bin.
        </div>
        """, unsafe_allow_html=True)
        
        col1, col2 = st.columns(2)
        with col1:
            if st.button("CANCEL", use_container_width=True, key=f"dialog_cancel_delete_lead_{lead_id}"):
                st.session_state.show_delete_modal = False
                st.rerun()
        with col2:
            if st.button("DELETE", type="primary", use_container_width=True, key=f"dialog_confirm_delete_lead_{lead_id}"):
                from app.crud import crud_leads
                if crud_leads.delete_lead(db, lead_id, st.session_state.username, st.session_state.get('db_user_id'), permanent=False):
                    st.session_state['success_msg'] = f"Success! Lead '{name}' moved to Recycle Bin."
                    clear_leads_cache()
                st.session_state.show_delete_modal = False
                st.rerun()
    finally:
        db.close()

@st.dialog("Edit Lead", width="large")
def show_edit_modal_dialog(m):
    """
    Native Streamlit dialog for editing a lead.
    """
    from app.db import SessionLocal
    db = SessionLocal()
    try:
        from app.crud import crud_leads, crud_agencies, crud_ccus, crud_events
        from app.schemas import LeadUpdate
        from datetime import datetime
        
        lead = m.get('lead_data', {})

        st.markdown(f"""
        <div class="modal-dialog-header">
          <div class="modal-icon">📝</div>
          Edit Lead: {m["title"]}
        </div>
        """, unsafe_allow_html=True)
                
        lead = m['lead_data']
        is_referral = lead.get('active_client', False)
        col1, col2 = st.columns(2)
        with col1:
            new_first = st.text_input("First Name", value=str(lead.get('first_name') or ""), key=f"edit_first_{m['target_id']}")
            new_last = st.text_input("Last Name", value=str(lead.get('last_name') or ""), key=f"edit_last_{m['target_id']}")
            new_custom_user_id = st.text_input("Employee ID", value=str(lead.get('custom_user_id') or ""), key=f"edit_custom_user_id_{m['target_id']}")
            new_phone = st.text_input("Phone", value=str(lead.get('phone') or ""), key=f"edit_phone_{m['target_id']}")
            new_email = st.text_input("Email", value=str(lead.get('email') or ""), key=f"edit_email_{m['target_id']}")
            new_staff = st.text_input("Staff Name", value=str(lead.get('staff_name') or ""), key=f"edit_staff_{m['target_id']}")
            
            source_options = ["Home Health Notify", "Web", "Direct Through CCU", "Event", "Word of Mouth", "Transfer", "Other"]
            current_src = lead.get('source', 'Other')
            src_idx = source_options.index(current_src) if current_src in source_options else source_options.index("Other")
            new_source = st.selectbox("Source", source_options, index=src_idx, key=f"edit_source_{m['target_id']}")
            
            new_event_name = lead.get('event_name')
            new_soc_date = lead.get('soc_date')
            new_other_source = lead.get('other_source_type')
            new_word_of_mouth = lead.get('word_of_mouth_type')
            
            if new_source == "Event":
                events = crud_events.get_all_events(db)
                event_list = [e.event_name for e in events]
                curr_event = lead.get('event_name')
                e_idx = event_list.index(curr_event) if curr_event in event_list else 0
                new_event_name = st.selectbox("Select Event", event_list, index=e_idx, key=f"edit_event_{m['target_id']}")
            elif new_source == "Transfer" or is_referral:
                soc_val = lead.get('soc_date')
                if isinstance(soc_val, str) and soc_val:
                    try: soc_val = datetime.strptime(soc_val, '%Y-%m-%d').date()
                    except: soc_val = None
                new_soc_date = st.date_input("SOC Date", value=soc_val, key=f"edit_soc_{m['target_id']}", format="MM/DD/YYYY")
            elif new_source == "Other":
                new_other_source = st.text_input("Specify Source", value=str(lead.get('other_source_type') or ""), key=f"edit_other_src_{m['target_id']}")
            elif new_source == "Word of Mouth":
                wom_options = ["Caregiver", "Community", "Client", "Staff"]
                curr_wom = lead.get('word_of_mouth_type')
                w_idx = wom_options.index(curr_wom) if curr_wom in wom_options else 0
                new_word_of_mouth = st.selectbox("Type", wom_options, index=w_idx, key=f"edit_wom_{m['target_id']}")

            st.markdown('<div style="margin-top: 15px;"></div>', unsafe_allow_html=True)
            new_street = st.text_input("Street", value=str(lead.get('street') or ""), key=f"edit_street_{m['target_id']}")
            new_city = st.text_input("City", value=str(lead.get('city') or ""), key=f"edit_city_{m['target_id']}")
            new_state = st.text_input("State", value=str(lead.get('state') or "IL"), max_chars=2, key=f"edit_state_{m['target_id']}")
            new_zip = st.text_input("Zip Code", value=str(lead.get('zip_code') or ""), key=f"edit_zip_{m['target_id']}")
            
            # Caregiver Type
            current_caregiver = lead.get('caregiver_type', 'None')
            if current_caregiver not in CAREGIVER_TYPES: current_caregiver = "None"
            new_caregiver_type = st.selectbox("Caregiver Type", CAREGIVER_TYPES, index=CAREGIVER_TYPES.index(current_caregiver), key=f"edit_caregiver_{m['target_id']}")
            
        with col2:
            is_auth_received_page = (st.session_state.get('main_navigation') == "Authorizations")
            
            if is_auth_received_page:
                # Special Status Logic for Authorizations Received
                st.write("**Referral Status:**")
                
                # Determine current group
                current_care_status = lead.get('care_status')
                initial_group = "Active"
                if current_care_status in ["Hold", "Terminated", "Deceased"]:
                    initial_group = current_care_status
                elif current_care_status and "Transfer" in current_care_status:
                    initial_group = "Transfer Received"
                
                # Main Group selection
                options = ["Active", "Hold", "Terminated", "Deceased"]
                status_group = st.radio("Main Status", options, 
                                         index=options.index(initial_group) if initial_group in options else 0,
                                         horizontal=True, key=f"edit_status_group_{m['target_id']}")
                
                new_care_status = status_group
                new_status = lead.get('last_contact_status') # Keep existing contact status
                
                if status_group == "Active":
                    # Show sub-options
                    sub_options = ["None", "Care Start", "Care Not Start"]
                    initial_sub = "None"
                    if current_care_status == "Care Start": initial_sub = "Care Start"
                    elif current_care_status == "Not Start": initial_sub = "Care Not Start"
                    
                    selected_sub = st.selectbox("Care Sub-Status", sub_options, 
                                                index=sub_options.index(initial_sub),
                                                key=f"edit_care_sub_{m['target_id']}")
                    
                    if selected_sub == "Care Start": new_care_status = "Care Start"
                    elif selected_sub == "Care Not Start": new_care_status = "Not Start"
                    else: new_care_status = None
                else:
                    new_care_status = status_group # Hold, Terminated, or Deceased
            
            elif is_referral:
                status_options = ["Initial Referral Sent", "Assessment Scheduled", "Assessment Done", "Not Approved", "Services Refused"]
                current_status = lead.get('last_contact_status', 'Initial Referral Sent')
                status_idx = status_options.index(current_status) if current_status in status_options else 0
                new_status = st.selectbox("Status", status_options, index=status_idx, key=f"edit_status_{m['target_id']}")
                new_care_status = lead.get('care_status')
            else:
                status_options = ["Initial Referral Sent", "Not Interested", "No Response"]
                current_status = lead.get('last_contact_status', 'Initial Referral Sent')
                if current_status in ["Initial Call", "Active", "Intro Call", "Follow Up"]: current_status = "Initial Referral Sent"
                status_idx = status_options.index(current_status) if current_status in status_options else 0
                new_status = st.selectbox("Status", status_options, index=status_idx, key=f"edit_status_{m['target_id']}")
                new_care_status = lead.get('care_status')
            
            # Call Status is now changed inline on the lead header row (not here)
            
            dob_value = lead.get('dob')
            if isinstance(dob_value, str) and dob_value:
                try:
                    dob_value = datetime.strptime(dob_value, '%Y-%m-%d').date()
                except:
                    dob_value = None
            
            from datetime import date
            def on_edit_dob_change():
                dob_key = f"edit_dob_{m['target_id']}"
                age_key = f"edit_age_{m['target_id']}"
                if st.session_state.get(dob_key):
                    today = date.today()
                    dob = st.session_state[dob_key]
                    age = today.year - dob.year - ((today.month, today.day) < (dob.month, dob.day))
                    st.session_state[age_key] = age

            new_dob = st.date_input("Date of Birth", value=dob_value if dob_value else None, min_value=date(1900, 1, 1), max_value=date.today(), key=f"edit_dob_{m['target_id']}", on_change=on_edit_dob_change, format="MM/DD/YYYY")
            age_key = f"edit_age_{m['target_id']}"
            if age_key not in st.session_state:
                st.session_state[age_key] = int(lead.get('age') or 0)
            new_age = st.number_input("Age / Year", min_value=0, max_value=3000, step=1, key=age_key)
            new_ssn = st.text_input("SSN", value=str(lead.get('ssn') or ""), key=f"edit_ssn_{m['target_id']}")
            new_medicaid = st.text_input("Medicaid #", value=str(lead.get('medicaid_no') or ""), key=f"edit_medicaid_{m['target_id']}")
            new_e_name = st.text_input("Emergency Contact", value=str(lead.get('e_contact_name') or ""), key=f"edit_ename_{m['target_id']}")
            new_e_relation = st.text_input("Relation", value=str(lead.get('e_contact_relation') or ""), key=f"edit_erelation_{m['target_id']}")
            new_e_phone = st.text_input("Emergency Phone", value=str(lead.get('e_contact_phone') or ""), key=f"edit_ephone_{m['target_id']}")
            
        
        new_comments = st.text_area("Comments", value=str(lead.get('comments') or ""), height=100, key=f"edit_comments_{m['target_id']}")
        
        enable_global = st.checkbox("Edit CCU/Payor", value=False, key=f"enable_entity_mgmt_{m['target_id']}")
        
        if enable_global:
            st.divider()
            ent_col1, ent_col2 = st.columns(2)
            with ent_col1:
                agencies = crud_agencies.get_all_agencies(db)
                agency_map = {a.name: a.id for a in agencies}
                agency_list = ["None"] + list(agency_map.keys())
                curr_agency_id = lead.get('agency_id')
                curr_agency_name = "None"
                if curr_agency_id:
                    for name, aid in agency_map.items():
                        if aid == curr_agency_id:
                            curr_agency_name = name
                            break
                new_agency_name_sel = st.selectbox("Payor", agency_list, index=agency_list.index(curr_agency_name), key=f"edit_agency_sel_{m['target_id']}")
                new_agency_id = agency_map.get(new_agency_name_sel)
                if new_agency_id:
                    exp_a_key = f"expand_a_edit_{m['target_id']}_{new_agency_id}"
                    if exp_a_key not in st.session_state: st.session_state[exp_a_key] = False
                    with st.expander(f"Edit {new_agency_name_sel} (Globally)", expanded=st.session_state[exp_a_key]):
                        agency_obj = crud_agencies.get_agency(db, new_agency_id)
                        if agency_obj:
                            a_addr, a_phone, a_fax, a_email = getattr(agency_obj, 'address', '') or "", getattr(agency_obj, 'phone', '') or "", getattr(agency_obj, 'fax', '') or "", getattr(agency_obj, 'email', '') or ""
                            u_a_addr = st.text_input("Payor Address", value=a_addr, key=f"global_a_addr_{new_agency_id}")
                            u_a_phone = st.text_input("Payor Phone", value=a_phone, key=f"global_a_phone_{new_agency_id}")
                            u_a_fax = st.text_input("Payor Fax", value=a_fax, key=f"global_a_fax_{new_agency_id}")
                            u_a_email = st.text_input("Payor Email", value=a_email, key=f"global_a_email_{new_agency_id}")
                            if st.button("Update Payor Details", key=f"global_a_save_{new_agency_id}"):
                                crud_agencies.update_agency(db, new_agency_id, new_agency_name_sel, st.session_state.username, st.session_state.get('db_user_id'), address=u_a_addr, phone=u_a_phone, fax=u_a_fax, email=u_a_email)
                                st.session_state[exp_a_key] = False
                                st.success(f"**Global Update Successful!**")
                                st.toast(f"Payor Updated Globally!", icon="✅")
            with ent_col2:
                ccus = crud_ccus.get_all_ccus(db)
                ccu_map = {c.name: c.id for c in ccus}
                ccu_list = ["None"] + list(ccu_map.keys())
                curr_ccu_id = lead.get('ccu_id')
                curr_ccu_name = "None"
                if curr_ccu_id:
                    for name, cid in ccu_map.items():
                        if cid == curr_ccu_id:
                            curr_ccu_name = name
                            break
                new_ccu_name_sel = st.selectbox("CCU", ccu_list, index=ccu_list.index(curr_ccu_name), key=f"edit_ccu_sel_{m['target_id']}")
                new_ccu_id = ccu_map.get(new_ccu_name_sel)
                if new_ccu_id:
                    exp_c_key = f"expand_c_edit_{m['target_id']}_{new_ccu_id}"
                    if exp_c_key not in st.session_state: st.session_state[exp_c_key] = False
                    with st.expander(f"Edit {new_ccu_name_sel} (Globally)", expanded=st.session_state[exp_c_key]):
                        ccu_obj = crud_ccus.get_ccu_by_id(db, new_ccu_id)
                        if ccu_obj:
                            c_addr, c_phone, c_fax, c_email, c_coord = getattr(ccu_obj, 'address', '') or "", getattr(ccu_obj, 'phone', '') or "", getattr(ccu_obj, 'fax', '') or "", getattr(ccu_obj, 'email', '') or "", getattr(ccu_obj, 'care_coordinator_name', '') or ""
                            u_c_addr = st.text_input("CCU Address", value=c_addr, key=f"global_c_addr_{new_ccu_id}")
                            u_c_phone = st.text_input("CCU Phone", value=c_phone, key=f"global_c_phone_{new_ccu_id}")
                            u_c_fax = st.text_input("CCU Fax", value=c_fax, key=f"global_c_fax_{new_ccu_id}")
                            u_c_email = st.text_input("CCU Email", value=c_email, key=f"global_c_email_{new_ccu_id}")
                            u_c_coord = st.text_input("Coordinator", value=c_coord, key=f"global_c_coord_{new_ccu_id}")
                            if st.button("Update CCU Details", key=f"global_c_save_{new_ccu_id}"):
                                crud_ccus.update_ccu(db, new_ccu_id, new_ccu_name_sel, st.session_state.username, st.session_state.get('db_user_id'), address=u_c_addr, phone=u_c_phone, fax=u_c_fax, email=u_c_email, care_coordinator_name=u_c_coord)
                                st.session_state[exp_c_key] = False
                                st.success(f"**Global Update Successful!**")
                                st.toast(f"CCU Updated Globally!", icon="✅")
        else:
            new_agency_id, new_ccu_id = lead.get('agency_id'), lead.get('ccu_id')
        
        st.markdown("<h4 style='font-weight: bold; color: #00506b;'>Notifications & Tracking</h4>", unsafe_allow_html=True)
        new_send_reminders = st.checkbox("Send Auto Email Reminders for this Lead", value=lead.get('send_reminders', True), key=f"edit_send_reminders_{m['target_id']}")

        if is_referral:
            st.markdown('<div style="margin-top: 10px;"></div>', unsafe_allow_html=True)
            if st.button(" UNMARK AS REFERRAL", use_container_width=True, key=f"edit_unmark_{m['target_id']}"):
                # Close edit dialog and open unmark confirmation
                render_confirmation_modal(
                    modal_type='unmark_ref', 
                    target_id=m['target_id'], 
                    title='Unmark Referral?', 
                    message=f"Convert <b>{lead.get('first_name')} {lead.get('last_name')}</b> back to a standard Lead?", 
                    icon='🚫', 
                    type='warning', 
                    confirm_label='UNMARK'
                )
                st.rerun()

        st.divider()
        c1, c2 = st.columns(2)
        with c1:
            if st.button("CANCEL", use_container_width=True, key=f"edit_cancel_{m['target_id']}"):
                close_modal()
        with c2:
            if st.button("SAVE CHANGES", type="primary", use_container_width=True, key=f"edit_save_{m['target_id']}"):
                try:
                    update_dict = {
                        "first_name": new_first, "last_name": new_last, "phone": new_phone, "staff_name": new_staff, "source": new_source,
                        "event_name": new_event_name, "soc_date": new_soc_date, "other_source_type": new_other_source, "word_of_mouth_type": new_word_of_mouth,
                        "city": new_city, "street": new_street, "state": new_state, "zip_code": new_zip, "last_contact_status": new_status,
                        "dob": new_dob, "medicaid_no": new_medicaid, "e_contact_name": new_e_name, "e_contact_relation": new_e_relation,
                        "e_contact_phone": new_e_phone, "active_client": lead.get('active_client'), "comments": new_comments, "age": new_age if new_age > 0 else None,
                        "agency_id": new_agency_id, "ccu_id": new_ccu_id, "send_reminders": new_send_reminders,
                        "care_status": new_care_status,
                        "custom_user_id": new_custom_user_id,
                        "caregiver_type": new_caregiver_type,
                        "email": new_email,
                        "ssn": new_ssn
                    }
                    
                    if is_referral and new_status == "Not Approved":
                        update_dict["care_status"], update_dict["authorization_received"] = None, False
                    
                    schema_data = LeadUpdate(**update_dict)
                    res = crud_leads.update_lead(db, m['target_id'], schema_data, st.session_state.username, st.session_state.get('db_user_id'))
                    if res:
                        st.session_state['success_msg'] = f"Success! Lead '{new_first} {new_last}' updated successfully!"
                        clear_leads_cache()
                        close_modal()
                    else:
                        st.error("Save failed: update_lead returned None")
                except Exception as e:
                    import traceback
                    st.error(f"Error: {e}")
                    st.exception(e)
    finally:
        db.close()

@st.dialog("Add Comment", width="small")
def show_add_comment_dialog(lead_id, lead_name):
    """
    Native Streamlit dialog for adding a new comment to a lead.
    """
    from app.db import SessionLocal
    db = SessionLocal()
    try:
        from app.crud import crud_notes
        
        st.markdown(f"**Add a new update for {lead_name}**")
        content = st.text_area("Update Details", placeholder="Enter notes, calls, or other updates here...", height=150)
        
        col1, col2 = st.columns([1, 1])
        with col1:
            if st.button("Save Comment", type="primary", use_container_width=True):
                if content.strip():
                    crud_notes.add_new_comment(db, lead_id, st.session_state.username, content)
                    st.session_state['success_msg'] = "Comment added!"
                    st.toast("Comment Added", icon="💬")
                    close_modal() # Use centralized close helper
                else:
                    st.warning("Please enter some text.")
        with col2:
            if st.button("Cancel", use_container_width=True):
                close_modal()
    finally:
        db.close()

def render_comment_stack(lead_obj):
    """
    Shared component to render the chronological stack of comments for a lead.
    """
    # Fetch comments - lead_obj is an ORM object from joinedload
    comments = []
    if hasattr(lead_obj, 'lead_comments'):
        comments = lead_obj.lead_comments
    
    # Sort by created_at desc (newest first) in case not sorted by query
    comments = sorted(comments, key=lambda x: x.created_at, reverse=True)
    
    if not comments:
        st.caption("No updates yet for this lead.")
        return

    st.markdown("---")
    st.markdown("**Updates / Comments History:**")
    
    # Render with custom styling
    for comment in comments:
        # Format date MM/DD/YYYY
        local_time = utc_to_local(comment.created_at, st.session_state.get('user_timezone'))
        date_str = local_time.strftime('%m/%d/%Y %I:%M %p')
        
        st.markdown(f"""
        <div style="background-color: #f9fafb; border-radius: 0.5rem; padding: 0.75rem; border-left: 4px solid #3b82f6; margin-bottom: 0.75rem;">
            <div style="display: flex; justify-content: space-between; margin-bottom: 0.25rem;">
                <span style="font-weight: bold; color: #1f2937;">{comment.username}</span>
                <span style="font-size: 0.75rem; color: #6b7280;">{date_str}</span>
            </div>
            <div style="color: #4b5563; font-size: 0.875rem; white-space: pre-wrap;">{comment.content}</div>
        </div>
        """, unsafe_allow_html=True)

@st.dialog("Document Preview", width="large")
def show_file_preview_dialog(file_path, filename):
    """
    Renders a preview of the document. Supports PDF and common images.
    """
    import base64
    if not os.path.exists(file_path):
        st.error("File not found.")
        return

    ext = filename.split('.')[-1].lower()
    
    try:
        with open(file_path, "rb") as f:
            file_bytes = f.read()

        if ext == 'pdf':
            base64_pdf = base64.b64encode(file_bytes).decode('utf-8')
            # Using iframe for better cross-browser PDF embedding
            pdf_display = f'<iframe src="data:application/pdf;base64,{base64_pdf}" width="100%" height="800" type="application/pdf"></iframe>'
            st.markdown(pdf_display, unsafe_allow_html=True)
        elif ext in ['png', 'jpg', 'jpeg']:
            st.image(file_bytes, caption=filename, use_container_width=True)
        elif ext == 'csv':
            import pandas as pd
            import io
            df = pd.read_csv(io.BytesIO(file_bytes))
            st.dataframe(df, use_container_width=True)
        elif ext == 'txt':
            st.text_area("File Content", value=file_bytes.decode('utf-8', errors='ignore'), height=400)
        elif ext == 'docx':
            import docx
            import io
            import pandas as pd
            doc = docx.Document(io.BytesIO(file_bytes))
            
            # Structured view container
            with st.container(height=500, border=True):
                # 1. Render Paragraphs
                for para in doc.paragraphs:
                    if not para.text.strip():
                        continue
                    
                    # Detect Headers
                    if para.style.name.startswith('Heading'):
                        level = para.style.name.replace('Heading ', '')
                        try:
                            st.markdown(f"{'#' * int(level)} {para.text}")
                        except:
                            st.markdown(f"### {para.text}")
                    else:
                        # Detect Bold/Italic in runs
                        md_para = ""
                        for run in para.runs:
                            t = run.text
                            if run.bold: t = f"**{t}**"
                            if run.italic: t = f"*{t}*"
                            md_para += t
                        st.markdown(md_para)
                
                # 2. Render Tables (Structured request)
                for i, table in enumerate(doc.tables):
                    st.markdown(f"**Table {i+1}**")
                    table_data = []
                    for row in table.rows:
                        table_data.append([cell.text for cell in row.cells])
                    if table_data:
                        df_t = pd.DataFrame(table_data)
                        st.dataframe(df_t, use_container_width=True)

                # 3. Render Images / Graphs (User report)
                st.divider()
                st.markdown("### 📊 Graphs & Illustrations")
                found_images = False
                for rel in doc.part.rels.values():
                    if "image" in rel.target_ref:
                        try:
                            image_bytes = rel.target_part.blob
                            st.image(image_bytes, use_container_width=True)
                            found_images = True
                        except:
                            continue
                if not found_images:
                    st.info("No embedded graphs or images found in this document.")
        else:
            st.warning(f"Preview not available for .{ext} files. Please download to view.")
            st.info(f"File: {filename}")
    except Exception as e:
        st.error(f"Error loading preview: {str(e)}")
    
    if st.button("Close Preview", use_container_width=True):
        close_modal() # Use the proper modal closer to wipe all state

def handle_active_modal():
    """
    Centralized handler for all active modals in the application.
    (Updated with Stability Refactor logic + Ghost Popup Prevention)
    """
    has_modal_open = st.session_state.get('modal_open', False)
    has_modal_action = st.session_state.get('modal_action')
    has_active_modal_dict = 'active_modal' in st.session_state
    
    # 0. GHOST KILLER: Wipe stale state if open flag is True but no action is defined
    if has_modal_open and not has_modal_action and not has_active_modal_dict:
        close_modal() # Wipe everything
        return

    # 1. SPECIAL DELETE MODAL (legacy support for view_leads.py)
    if st.session_state.get('show_delete_modal', False):
        lead_id = st.session_state.get('delete_lead_id')
        name = st.session_state.get('delete_lead_name', 'Unknown')
        # GHOST FIX: Clear state WITHOUT rerun so we can actually reach the dialog call below
        clear_modal_state() 
        show_delete_modal_dialog(lead_id, name)
        return

    # 2. HANDLE GENERIC ACTIVE_MODAL
    m = None
    if has_modal_open:
        if has_modal_action:
            # Reconstruct dictionary from flat state
            m = {
                'modal_type': st.session_state.modal_action,
                'target_id': st.session_state.modal_lead_id,
                'title': st.session_state.modal_lead_name or st.session_state.modal_data.get('title'),
                'lead_data': st.session_state.modal_data.get('lead_data'),
                'icon': st.session_state.modal_data.get('icon'),
                'type': st.session_state.modal_data.get('type'),
                'confirm_label': st.session_state.modal_data.get('confirm_label'),
                'cancel_label': st.session_state.modal_data.get('cancel_label'),
                'indicator': st.session_state.modal_data.get('indicator'),
                'message': st.session_state.modal_data.get('message')
            }
        elif has_active_modal_dict:
            m = st.session_state['active_modal']
    
    if not m:
        return
        
    # GHOST FIX: Clear the trigger state BEFORE dispatching.
    # This ensures that subsequent reruns don't re-trigger the modal logic.
    clear_modal_state()
    
    # Dispatch to specific dialog functions
    if m['modal_type'] == 'save_edit_modal':
        show_edit_modal_dialog(m)
    elif m['modal_type'] == 'file_preview':
        show_file_preview_dialog(m['lead_data']['file_path'], m['lead_data']['filename'])
    else:
        confirmation_modal_dialog(m)


def render_confirmation_modal(title, message, icon="🗑️", type="info", confirm_label="DELETE", cancel_label="CANCEL", target_id="modal", indicator=None, modal_type='soft_delete'):
    """
    Triggers a confirmation modal by setting isolated session state variables.
    """
    # 1. HARD RESET ALL MODAL STATE
    st.session_state.modal_open = True
    st.session_state.modal_action = modal_type
    st.session_state.modal_lead_id = target_id
    st.session_state.modal_lead_name = title
    
    modal_data = {
        'title': title,
        'message': message,
        'icon': icon,
        'type': type,
        'confirm_label': confirm_label,
        'cancel_label': cancel_label,
        'indicator': indicator
    }
    st.session_state.modal_data = modal_data
    
    # 2. Maintain legacy dictionary for backward compatibility
    st.session_state['active_modal'] = {
        'modal_type': modal_type,
        'target_id': target_id,
        'title': title,
        'message': message,
        **modal_data
    }
    # No rerun here as this is often called inside a button click which will trigger its own rerun/processing
    return None


def get_pagination_params(key_prefix, default_limit=10):
    """
    Initializes and returns common pagination parameters.
    Should be called at the TOP of a page function.
    """
    page_key = f"{key_prefix}_page"
    rows_key = f"{key_prefix}_rows_per_page"
    
    if page_key not in st.session_state:
        st.session_state[page_key] = 0
    if rows_key not in st.session_state:
        st.session_state[rows_key] = default_limit
        
    page_index = st.session_state[page_key]
    rows_per_page = st.session_state[rows_key]
    
    skip = page_index * rows_per_page
    limit = rows_per_page
    
    return skip, limit, page_index, rows_per_page


def render_pagination(total_items, key_prefix):
    """
    Renders a unified Material-UI style pagination bar at the bottom of a list.
    Only renders the UI; state management should be handled by get_pagination_params.
    """
    # 1. Access existing state initialized by get_pagination_params
    page_key = f"{key_prefix}_page"
    rows_key = f"{key_prefix}_rows_per_page"
    
    # Defensive check
    if page_key not in st.session_state or rows_key not in st.session_state:
        get_pagination_params(key_prefix)
        
    page_index = st.session_state[page_key]
    rows_per_page = st.session_state[rows_key]
    
    # Calculate metadata
    num_pages = max(1, (total_items // rows_per_page) + (1 if total_items % rows_per_page > 0 else 0))
    
    # Adjust page index if it exceeds total pages (e.g. after a filter is applied)
    if page_index >= num_pages:
        page_index = max(0, num_pages - 1)
        st.session_state[page_key] = page_index
        # If we had to adjust, we should rerun to ensure consistency, 
        # but usually this is called after the query, so the query might have used the 'too large' index.
        # However, search_leads/count_search_leads handle this gracefully in the DB layer or return empty lists.
        # st.rerun() here might cause infinite loops if not careful.
        
    start_item = (page_index * rows_per_page) + 1 if total_items > 0 else 0
    end_item = min((page_index + 1) * rows_per_page, total_items)
    
    # 2. Render UI
    st.markdown("---")
    
    # MUI-like bar
    p_col1, p_col2, p_col3, p_col4, p_col5 = st.columns([1.5, 1, 2, 0.5, 0.5])
    
    with p_col1:
        st.markdown("<div style='margin-top: 5px; text-align: right; font-weight: 600; color: #6b7280;'>Rows per page:</div>", unsafe_allow_html=True)
    
    with p_col2:
        options = [10, 20, 50, 100]
        try:
            cur_idx = options.index(rows_per_page)
        except ValueError:
            cur_idx = 0
            
        new_rows = st.selectbox("Rows per page", options, index=cur_idx, 
                               label_visibility="collapsed", key=f"rows_sel_{key_prefix}")
        if new_rows != rows_per_page:
            st.session_state[rows_key] = new_rows
            st.session_state[page_key] = 0 # Reset to first page
            st.rerun()
            
    with p_col3:
        st.markdown(f"<div style='margin-top: 5px; text-align: center; color: #111827; font-weight: 700;'>{start_item}-{end_item} of {total_items}</div>", unsafe_allow_html=True)
        
    with p_col4:
        btn_prev = st.button("⟨", key=f"prev_{key_prefix}", use_container_width=True, disabled=(page_index == 0))
        if btn_prev:
            st.session_state[page_key] -= 1
            st.rerun()
            
    with p_col5:
        btn_next = st.button("⟩", key=f"next_{key_prefix}", use_container_width=True, disabled=(page_index >= num_pages - 1))
        if btn_next:
            st.session_state[page_key] += 1
            st.rerun()
            
    return page_index * rows_per_page, rows_per_page


def render_api_status():
    """Diagnostic tool to check if the FastAPI backend is running."""
    import urllib.request
    import os
    
    backend_url = os.environ.get("BACKEND_API_URL", "http://127.0.0.1:8000")
    st.sidebar.markdown("---")
    
    try:
        with urllib.request.urlopen(f"{backend_url}/health", timeout=3) as response:
            if response.getcode() == 200:
                st.sidebar.markdown("**System: Live**")
            else:
                st.sidebar.markdown("**System: Offline**")
    except:
        st.sidebar.markdown("**System: Offline**")



def export_leads_to_excel(leads):
    """
    Exports a list of lead objects to an Excel file (XLSX).
    Returns the binary content of the Excel file.
    """
    data = []
    for lead in leads:
        # Format CCU Information
        ccu_info = "N/A"
        if lead.ccu:
            ccu_parts = [f"Name: {lead.ccu.name}"]
            
            # Address parts
            addr_parts = [p for p in [lead.ccu.street, lead.ccu.city, lead.ccu.state, lead.ccu.zip_code] if p]
            if addr_parts:
                ccu_parts.append(f"Address: {', '.join(addr_parts)}")
            
            if lead.ccu.phone:
                ccu_parts.append(f"Phone: {lead.ccu.phone}")
            
            if lead.ccu.email:
                ccu_parts.append(f"Email: {lead.ccu.email}")
            
            ccu_info = "\n".join(ccu_parts)

        # Get Payor name safely
        payor_name = "N/A"
        if lead.agency:
            payor_name = lead.agency.name

        # Format SOC Date
        soc_str = lead.soc_date.strftime('%m/%d/%Y') if (hasattr(lead, 'soc_date') and lead.soc_date) else "N/A"

        # Build row
        row = {
            "ID": lead.id,
            "Name": f"{lead.first_name} {lead.last_name}",
            "Staff": lead.staff_name or "N/A",
            "Phone": lead.phone or "N/A",
            "Email": lead.email if lead.email else "N/A",
            "SSN": lead.ssn if lead.ssn else "N/A",
            "Emergency Contact": lead.e_contact_name if lead.e_contact_name else "N/A",
            "Relation": (lead.e_contact_relation or getattr(lead, 'relation_to_client', None)) or "N/A",
            "EC Phone": lead.e_contact_phone if lead.e_contact_phone else "N/A",
            "Call Status": lead.priority if lead.priority else "Not Called",
            "Contact Status": lead.last_contact_status or "N/A",
            "Referral": "Yes" if lead.active_client else "No",
            "Authorization": "Received" if lead.authorization_received else "Pending",
            "CCU Information": ccu_info,
            "Payor": payor_name,
            "Caregiver Type": lead.caregiver_type or "None",
            "Start of Care": soc_str
        }
        data.append(row)

    # Create DataFrame
    df = pd.DataFrame(data)

    # Buffer for Excel file
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name='Leads')
        
        # Auto-adjust columns width (optional but nice)
        worksheet = writer.sheets['Leads']
        for i, col in enumerate(df.columns):
            # Set width to max length of column name or values, capped at 50
            max_len = max(
                df[col].astype(str).map(len).max(),
                len(str(col))
            ) + 2
            worksheet.column_dimensions[chr(65 + i)].width = min(max_len, 50)
            
    return output.getvalue()
